# -*- coding: utf-8 -*-
"""V16 Step 1: 标注规则详细指南 (docx)。

5 秒能判定的标签指南 + 决策树 + 真实例子 + 边界对照。
用户打开 xlsx 同时打开本 docx 对照判断。
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent / "outputs" / "v16_标注规则指南.docx"

d = Document()

# Title
t = d.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("V16 标注 Review 指南")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x5D, 0x82)

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("配合 v16_REVIEW_ME.xlsx 使用 — 5 秒判定一条").italic = True

d.add_paragraph()

# 假想场景
d.add_heading("假想客户(判定基准)", level=1)
d.add_paragraph("当前要生成报告的客户 = 福建中锐网络股份有限公司")
d.add_paragraph("• 实控人:黄祖海 / 配偶康恩慧")
d.add_paragraph("• 行业:其他未列明信息技术服务业")
d.add_paragraph("• 主营:智慧水利 + 智慧教育")
d.add_paragraph("• 注册资本:4100 万元 / 实收资本:4100 万元")
d.add_paragraph("• 5 家关联企业:汉鼎/教育科技/青云/软件科技/海沃")
d.add_paragraph("• 2025 年营收 10010 万 / 净利 494 万 / 资产负债率 42.5%")
d.add_paragraph()
d.add_paragraph("→ 你判断每条 element 时,问自己:**这段内容如果原样保留在中锐报告里,合不合适?**", style=None)

# 决策树
d.add_heading("3 秒决策树(从上往下问一遍即可)", level=1)
steps = [
    ("Q1", "这条是 Word 样式标题(Heading 1-4)、或带章节编号(一、/(一)/1/(1)/①/※)?", "→ SCAFFOLD(章节骨架)"),
    ("Q2", "这条是 \"字段名:字段值\" 格式(冒号分隔)?", ""),
    ("    ", "  → 字段名属于企业固定信息(注册资本/法人/主营/股东/经营地址)?", "→ FILL(代码填中锐对应值)"),
    ("    ", "  → 字段名属于评级/审批类(PD评级/白名单/申报单位/绿色信贷/投向政策行业)?", "→ CLEAR(留空+pending)"),
    ("Q3", "这条是占位符或括号指引(下划线 ____ / XX 年 XX 月 / (面积等) / (天使轮/A/B/C…))?", "→ SLOT"),
    ("Q4", "这条含 \"请填写 / 如涉及 / 请说明 / (材料未提供) / 注:\" 这种纯指引短句?", "→ PRESERVE(保留指引)"),
    ("Q5", "这条是表格表头(年份/股东名称/资产名称 这种 column 名)?", "→ SCAFFOLD"),
    ("Q6", "这条是中长段(>40 字),含别家公司名/别家人名/别家年份/别家具体数字?", "→ REWRITE(LLM 整段改写为中锐内容)"),
    ("Q7", "都不是? 短文本默认 SCAFFOLD; 长文本默认 REWRITE", ""),
]
for q, body, decision in steps:
    p = d.add_paragraph()
    r1 = p.add_run(f"{q}  ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x2E, 0x5D, 0x82)
    p.add_run(body)
    if decision:
        p.add_run("  ")
        r2 = p.add_run(decision)
        r2.bold = True
        r2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

d.add_paragraph()

# 6 类详细
d.add_heading("6 类标签详解 + 真例子", level=1)

LABELS = [
    {
        "code": "SCAFFOLD",
        "ch": "骨架保留",
        "color": RGBColor(0x2E, 0x86, 0x4B),
        "what": "模板的固定结构,任何客户都通用,系统原样保留",
        "look": "章节标题 / 章节编号 / 字段标签 / 表格表头 / 单位说明 / 报告大标题",
        "action": "代码不动,原文呈现",
        "examples_yes": [
            "一、授信客户背景情况",
            "1. 授信客户基本情况",
            "(1) 主要股东的背景",
            "①基本情况:",
            "单位:万元",
            "提款条件及贷(投)后监控要求落实情况",
            "(在表格里)股东名称 | 持股比例 | 出资方式",
            "(在表格里)年份 | 总资产 | 净资产",
        ],
        "examples_no": [
            "授信客户全称:兴业资产管理有限公司  ← 这是 FILL,因为含别家公司名",
            "PD评级:4  ← 这是 CLEAR,字段值是别家评级",
            "公司设董事会,成员五人,由股东委派...  ← 这是 REWRITE,中长段含组织描述",
        ],
    },
    {
        "code": "FILL",
        "ch": "代码填",
        "color": RGBColor(0x2E, 0x5D, 0x82),
        "what": "字段值或表格 cell 是别家客户的具体值,系统应替换为中锐对应值",
        "look": "\"字段:别家公司名/别家数字/别家人名\" + 字段属于企业固定信息",
        "action": "代码从 中锐 client_profile 取对应值替换",
        "examples_yes": [
            "授信客户全称:兴业资产管理有限公司  → 改为 福建中锐网络股份有限公司",
            "注册资本:195000 万元  → 改为 4100 万元",
            "实际控制人:许翔  → 改为 黄祖海",
            "(表格 cell)兴业资产管理有限公司  → 改为 福建中锐网络股份有限公司",
            "(表格 cell)195,000 万元  → 改为 4100 万元",
            "(报告大标题)关于兴业资产管理有限公司5亿元综合授信额度授信报告  → 改为 中锐相应标题",
        ],
        "examples_no": [
            "PD评级:4  ← 这是 CLEAR,中锐没这个评级数据",
            "白名单类型:B类  ← 这是 CLEAR,需人工核",
            "我行投向政策对应行业:商业服务业  ← CLEAR,中锐行业不同需人工填",
        ],
    },
    {
        "code": "CLEAR",
        "ch": "清空+pending",
        "color": RGBColor(0xB7, 0x6B, 0x00),
        "what": "字段值是别家客户的评级/审批/政策类内容,中锐通常没现成数据,系统清空+加 pending tag 让客户经理人工补",
        "look": "\"评级 / 评估类 / 审批类 / 政策行业\"字段 + 别家具体值",
        "action": "代码把值清空,文本变成 \"PD评级:(待补充)\",加入 pending_tags",
        "examples_yes": [
            "PD评级:4",
            "申报单位:大客户一部",
            "白名单类型:B类",
            "绿色信贷标识:绿色四类",
            "环境和社会风险分类:√A □B □C",
            "客户环境和社会表现的动态评估结果:非常满意",
            "我行投向政策对应行业:商业服务业",
        ],
        "examples_no": [
            "注册资本:195000 万元  ← FILL,中锐有 4100 万",
            "实际控制人:许翔  ← FILL,中锐有黄祖海",
            "公司设董事会  ← REWRITE,需描述中锐组织",
        ],
    },
    {
        "code": "REWRITE",
        "ch": "LLM 整段重写",
        "color": RGBColor(0x8E, 0x44, 0xAD),
        "what": "正文段落,内容是别家客户的具体经营/财务/股东/管理描述。系统让 LLM 基于中锐材料整段改写",
        "look": "中长段(>40 字)+ 含别家公司名/人名/具体数字/具体年份",
        "action": "LLM 看大标题主旨 + 中锐材料 + 历史段落作参考,整段重写",
        "examples_yes": [
            "兴业资管控股股东是兴业国信资产管理有限公司...",
            "公司董事会由谢斌、赖富荣、吴红珍、倪勤、汪祖福等五人组成",
            "谢斌先生,1971 年出生;现任兴业资产管理有限公司党委书记...",
            "兴业资管 2022 年营收 X 亿元,同比增长 Y%...",
            "近三年及一期内公司实际控制人未发生变化",
        ],
        "examples_no": [
            "一、授信客户背景情况  ← SCAFFOLD,这是章节标题",
            "PD评级:4  ← CLEAR,字段-值",
            "______________  ← SLOT,占位符",
            "请填写下表  ← PRESERVE,纯指引",
        ],
    },
    {
        "code": "SLOT",
        "ch": "占位槽",
        "color": RGBColor(0x16, 0xA0, 0x85),
        "what": "占位符或简短括号指引,本身就是空白等待填",
        "look": "下划线 ____ / 大段空白 / XX 年 XX 月 XX 日 / 短括号说明 (面积等) (天使轮/A/B/C…)",
        "action": "如果有对应客户数据 → FILL;如果无数据 → 留空 + pending",
        "examples_yes": [
            "______________________(下划线)",
            "(面积等)",
            "(天使轮/A/B/C…)",
            "(及坐落位置)",
            "(请附 531 系统测算表截图)",
            "X年X月X日",
            "(若有,请说明)",
        ],
        "examples_no": [
            "请填写下表  ← PRESERVE,有动作指令的指引",
            "如未落实,请说明原因  ← PRESERVE,完整指引句",
            "注册资本:______  ← FILL(标签是企业固定信息,系统填客户值)",
        ],
    },
    {
        "code": "PRESERVE",
        "ch": "指引保留",
        "color": RGBColor(0x55, 0x55, 0x55),
        "what": "纯说明性指引文本,告诉客户经理\"应该怎么写\",本身不是内容",
        "look": "请填写 / 如涉及 / 请说明 / (材料未提供) / 注:/ 备注:/ 说明:",
        "action": "保留原文不动",
        "examples_yes": [
            "请填写下表",
            "如涉及,请填列下表",
            "如未落实,请说明原因",
            "注:本表数据取自 531 系统",
            "备注:具体情况见附件",
            "(具体情况如下表所示)",
        ],
        "examples_no": [
            "(面积等)  ← SLOT,纯括号短指引",
            "______________  ← SLOT,占位符",
            "兴业资管 2022 年营收  ← REWRITE,含具体客户数据",
        ],
    },
]

for L in LABELS:
    # 标签 header
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(f"【{L['code']}】 {L['ch']}")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = L["color"]

    p = d.add_paragraph(); p.add_run("是什么:").bold = True; p.add_run(" " + L["what"])
    p = d.add_paragraph(); p.add_run("视觉特征:").bold = True; p.add_run(" " + L["look"])
    p = d.add_paragraph(); p.add_run("系统动作:").bold = True; p.add_run(" " + L["action"])

    p = d.add_paragraph(); p.add_run("✓ 是这一类的(YES 例):").bold = True
    for ex in L["examples_yes"]:
        d.add_paragraph(f"   • {ex}")

    p = d.add_paragraph(); p.add_run("✗ 不是这一类的(易混淆):").bold = True
    for ex in L["examples_no"]:
        p2 = d.add_paragraph(f"   • {ex}")

# 边界对照
d.add_heading("常见疑难 5 例对照", level=1)

cases = [
    {
        "text": "实际控制人:许翔,中国国籍,本科学历...",
        "answer": "REWRITE",
        "why": "虽然以\"实际控制人:\"开头像字段,但后面是长段人物简介(中长 + 别家人名 + 学历)。整段需 LLM 替换为黄祖海简历。",
        "tip": "字段-值如果\"值\"是中长段简历/描述 → REWRITE;如果是单值\"许翔\" → FILL",
    },
    {
        "text": "(3) 实际控制人:福建省招标股份有限公司:",
        "answer": "SCAFFOLD",
        "why": "虽然含别家公司名,但这是章节小标题(\"(3) 实际控制人:\")。代码后置会自动把别家公司名清掉(留 \"(3) 实际控制人:\")。",
        "tip": "标题行(<25 字 + 编号开头)即使含别家公司名也算 SCAFFOLD,代码会清干净",
    },
    {
        "text": "本部聚焦于智慧水利与智慧教育两大核心业务线",
        "answer": "SCAFFOLD 或 PRESERVE(本身已是中锐内容)",
        "why": "如果整段已经是中锐自己的业务描述(无别家实体),系统不需要重写,保留即可。",
        "tip": "看 element 里的实体是否已是当前客户的;是 → 不动 / SCAFFOLD",
    },
    {
        "text": "近三年及一期内公司实际控制人未发生变化",
        "answer": "REWRITE",
        "why": "这是叙述性陈述句,需要 LLM 根据中锐实际情况重新评估(中锐近年实控人是否变化也要 LLM 写)。中长段无具体外来实体但属于报告叙事内容。",
        "tip": "中长段 + 是\"过去时陈述句\"(\"未发生变化\"\"实现\"\"取得\") → REWRITE",
    },
    {
        "text": "单位:万元",
        "answer": "SCAFFOLD",
        "why": "这是表格的单位说明,任何报告都通用,保留。",
        "tip": "短指引(<15 字)且是\"单位/附件/说明\"开头 → SCAFFOLD,不是 PRESERVE",
    },
]

for c in cases:
    p = d.add_paragraph()
    p.add_run("文本: ").bold = True
    p.add_run(c["text"])
    p = d.add_paragraph()
    p.add_run("应该是: ").bold = True
    r = p.add_run(c["answer"])
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p = d.add_paragraph()
    p.add_run("为什么: ").bold = True
    p.add_run(c["why"])
    p = d.add_paragraph()
    p.add_run("规律: ").bold = True
    p.add_run(c["tip"])
    d.add_paragraph()

# 操作回顾
d.add_heading("xlsx 操作回顾", level=1)
ops = [
    "1. 默认每行的\"你的判断\"列已填\"✓ 对\" — 跳过即可",
    "2. 看到我标错了 → 把\"你的判断\"改成\"✗ 错\",\"如错,正确标签\"列下拉选一个",
    "3. 如果样本本身坏(抽错了) → \"你的判断\"选\"⊘ SKIP\"",
    "4. 橙色行 = 我的规则置信 <0.7,重点看",
    "5. 全部过完保存,告诉我文件路径",
]
for op in ops:
    d.add_paragraph(op)

d.save(OUT)
print(f"[done] {OUT}")
print(f"  size: {OUT.stat().st_size} bytes")
