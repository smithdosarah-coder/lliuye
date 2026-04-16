# -*- coding: utf-8 -*-
"""Dump latest section_gen_test_*.docx into text for V14-v3 verification."""
import sys
import io
import os
import glob

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)

from docx import Document

candidates = sorted(glob.glob("outputs/section_gen_test_*.docx"), key=os.path.getmtime)
if not candidates:
    print("No docx found")
    sys.exit(1)
target = candidates[-1]
print(f"[Target] {target}")

doc = Document(target)

out = []

def walk_cell(cell, depth=0):
    pre = "  " * depth
    for p in cell.paragraphs:
        txt = p.text
        if txt.strip():
            bolds = [r.bold for r in p.runs]
            bold_flag = "[B]" if any(b is True for b in bolds) else ""
            out.append(f"{pre}{bold_flag}{txt}")
    for t in cell.tables:
        out.append(f"{pre}--TABLE({len(t.rows)}r x {len(t.rows[0].cells) if t.rows else 0}c)--")
        for r_idx, row in enumerate(t.rows):
            row_cells = []
            for c in row.cells:
                row_cells.append(c.text.strip().replace("\n", " / ")[:40])
            out.append(f"{pre}  R{r_idx}| " + " | ".join(row_cells))
        out.append(f"{pre}--/TABLE--")

for p in doc.paragraphs:
    if p.text.strip():
        bolds = [r.bold for r in p.runs]
        bold_flag = "[B]" if any(b is True for b in bolds) else ""
        out.append(f"{bold_flag}{p.text}")

for i, table in enumerate(doc.tables):
    out.append(f"\n===== TOP TABLE #{i} =====")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            out.append(f"\n[R{r_idx}C{c_idx}]")
            walk_cell(cell, 1)

dump_path = "outputs/_v14v3_dump.txt"
with open(dump_path, "w", encoding="utf-8") as f:
    f.write("\n".join(out))

print(f"[Written] {dump_path} ({len(out)} lines)")
