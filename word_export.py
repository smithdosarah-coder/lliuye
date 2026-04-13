# -*- coding: utf-8 -*-
"""
Word 报告导出模块 v7

v7 变更（对应7个问题）：
1. ★ 彻底剥离 Markdown 格式 — 正文段落永不加粗（仅子标题加粗）
2. ★ "需补充"内容→Word侧栏批注（不再标红），正文留空+旁边标签提示
3. ★ 封面表格自动填充（授信金额/行业/PD评级等从章节内容中提取）
4. ★ 清除英文词汇残留（likely → 可能 等）
5. 模版注入改为全文注入（不只是标题大纲）
"""

import re
import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from lxml import etree

# ====== 报告结构常量 ======
RPT_TITLE_PREFIX = "关于"
RPT_TITLE_SUFFIX = "授信调查报告"
RPT_COVER_LABELS = [
    "授信客户全称", "授信金额", "行业投向",
    "PD评级", "申报单位", "授信品种",
]
RPT_CHAPTERS = [
    "一、授信客户背景",
    "二、经营/管理风险分析",
    "三、财务风险分析",
    "四、结论与建议",
]

_CHAPTER_TITLE_PATTERNS = [
    r'^[一二三四五六七八九十]+[、.．]\s*\S+',
    r'^第[一二三四五六七八九十]+[章节]\s*',
    r'^[（(]?[一二三四五六七八九十]+[）)]\s*',
    r'^[IVX]+[、.．]\s*',
    r'^\d+[、.．]\s*\S+风险',
]

# ====== 需补充材料的检测模式 ======
_NEED_SUPPLEMENT_PATTERNS = [
    # ★ v7.3: 增加 LLM 实际输出的"未提供"变体
    r'材料未提供', r'数据材料未提供', r'信息材料未涉及',
    r'该数据材料未提供', r'该信息材料未涉及',
    r'未提供', r'未涉及',  # ★ v7.3: 兜底匹配（"余额未提供"、"总额未提供"等）
    r'\[待补充\]', r'待补充',  # ★ v7.3: 匹配 [待补充] 标记
    r'建议补充', r'需补充', r'需要补充', r'请补充',
    r'无法评估', r'无法分析', r'无法计算', r'无法进行',
    r'无法了解', r'无法判断', r'无法确定',
    r'信息缺失', r'数据缺失', r'资料缺失',
    r'缺少.*材料', r'材料中未.*涉及', r'提供的材料中未',
    r'尚未获取', r'暂无.*数据', r'暂未获取',
]
_SUPPLEMENT_RE = re.compile('|'.join(_NEED_SUPPLEMENT_PATTERNS))

# ★ v7: 纯"需补充"段落（整段都是关于缺失材料，正文应留空）
_PURE_SUPPLEMENT_PATTERNS = [
    r'^该(?:数据|信息|部分)?材料未(?:提供|涉及)',
    r'^材料未(?:提供|涉及)',
    r'^(?:该|此)(?:部分|项|方面).*材料未',
    r'^(?:该|此)(?:部分|项|方面).*(?:缺失|缺少|不足)',
    r'^(?:由于|因).*材料.*(?:未提供|缺失|不足).*(?:无法|暂无)',
    r'^暂无.*(?:数据|材料|信息)',
    r'^(?:尚|暂)未获取',
    r'^(?:该项|本项).*(?:无法|暂无)',
    # ★ v7.3: 匹配 LLM 生成的"XX余额未提供"等短句
    r'^.{0,8}(?:余额|总额|金额|数据|信息|比率|周转)未提供',
    r'^\[待补充\]$',  # 纯 [待补充] 占位行
]
_PURE_SUPPLEMENT_RE = re.compile('|'.join(_PURE_SUPPLEMENT_PATTERNS))

# ★ v7.3: 主观建议检测（LLM自己的判断/建议，应转为批注）
_SUBJECTIVE_SUGGESTION_PATTERNS = [
    r'建议采用', r'建议授予', r'建议将', r'建议评定',
    r'建议给予', r'建议设定', r'建议设置',
    r'建议(?:本行|我行|银行).*(?:给予|采取|实施|执行)',
    r'(?:综合考虑|综合以上).*建议',
    r'笔者认为', r'我们认为', r'分析师认为',
]
_SUBJECTIVE_RE = re.compile('|'.join(_SUBJECTIVE_SUGGESTION_PATTERNS))

# ★ v7: 英文词汇替换表
_ENGLISH_REPLACEMENTS = {
    'likely': '可能',
    'Likely': '可能',
    'LIKELY': '可能',
    'probably': '可能',
    'possibly': '可能',
    'approximately': '约',
    'about': '约',
    'significant': '显著',
    'significantly': '显著地',
    'slightly': '略',
    'overall': '总体',
    'stable': '稳定',
    'increase': '增长',
    'decrease': '下降',
    'risk': '风险',
    'relatively': '相对',
    'percent': '百分比',
    'however': '然而',
    'Therefore': '因此',
    'therefore': '因此',
    'Moreover': '此外',
    'moreover': '此外',
    'Furthermore': '此外',
    'furthermore': '此外',
    'Generally': '总体上',
    'generally': '总体上',
    'mainly': '主要',
    'mainly': '主要',
    'respectively': '分别',
    'total': '合计',
    'Total': '合计',
    'average': '平均',
    'Average': '平均',
    'Annual': '年度',
    'annual': '年度',
    'revenue': '营收',
    'Revenue': '营收',
    'profit': '利润',
    'Profit': '利润',
    'loss': '亏损',
    'Loss': '亏损',
    'growth': '增长',
    'Growth': '增长',
    'ratio': '比率',
    'Ratio': '比率',
    'market': '市场',
    'Market': '市场',
    'industry': '行业',
    'Industry': '行业',
    'company': '公司',
    'Company': '公司',
    'analysis': '分析',
    'Analysis': '分析',
    'based': '基于',
    'Based': '基于',
    'according': '根据',
    'According': '根据',
    'currently': '目前',
    'Currently': '目前',
    'due': '由于',
    'Due': '由于',
    'the': '',
    'The': '',
    'and': '和',
    'for': '',
    'from': '',
    'with': '',
    'this': '',
    'that': '',
    'which': '',
    'were': '',
    'was': '',
    'has': '',
    'have': '',
    'had': '',
    'are': '',
    'not': '不',
    'but': '但',
    'also': '也',
    'very': '非常',
    'more': '更',
    'most': '最',
    'less': '较少',
    'high': '高',
    'low': '低',
    'good': '良好',
    'bad': '不良',
    'strong': '强',
    'weak': '弱',
    'new': '新',
    'old': '旧',
    'large': '大',
    'small': '小',
    'long': '长期',
    'short': '短期',
    'N/A': '不适用',
    'n/a': '不适用',
    'YoY': '同比',
    'yoy': '同比',
    'QoQ': '环比',
    'qoq': '环比',
    'MoM': '环比',
    # ★ v7.3: 新增常见泄漏词
    'EOD': '',      # End of Document 标记
    'eod': '',
    'END': '',
    'end': '',
    'Note': '注',
    'note': '注',
    'Summary': '概要',
    'summary': '概要',
    'Chapter': '章节',
    'chapter': '章节',
    'Section': '节',
    'section': '节',
    'Conclusion': '结论',
    'conclusion': '结论',
    'Recommendation': '建议',
    'recommendation': '建议',
}
# 只替换长度>=3的词（避免误伤缩写）或特定短词
_ENGLISH_WORD_RE = re.compile(
    r'\b(' + '|'.join(re.escape(k) for k in _ENGLISH_REPLACEMENTS if len(k) >= 3) + r')\b'
    + r'|(?<!\w)(' + '|'.join(re.escape(k) for k in _ENGLISH_REPLACEMENTS if len(k) < 3 and k in ('N/A', 'n/a', 'YoY', 'QoQ', 'MoM')) + r')(?!\w)'
)


# ====== 字体设置 ======
def _set_font(run, name, size_pt, bold=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)


# ====== v7: Markdown 彻底剥离 ======
def _strip_markdown(text: str) -> str:
    """彻底去除所有 Markdown 格式符号，返回纯文本"""
    # 去掉标题标记
    text = re.sub(r'^#{1,6}\s+', '', text)
    # 去掉引用标记
    text = re.sub(r'^>\s*', '', text)
    # 去掉列表标记（但保留内容）
    text = re.sub(r'^[\-\*\+]\s+', '', text)
    text = re.sub(r'^\d+\.\s+', '', text)
    # ★ 去掉加粗/斜体标记（核心修复）
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)  # ***粗斜***
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)       # **粗**
    text = re.sub(r'\*(.+?)\*', r'\1', text)            # *斜*
    text = re.sub(r'__(.+?)__', r'\1', text)            # __粗__
    text = re.sub(r'_(.+?)_', r'\1', text)              # _斜_
    # 去掉行内代码
    text = re.sub(r'`(.+?)`', r'\1', text)
    # 去掉残留的*号（边界情况：不成对的*）
    text = text.replace('**', '')
    return text.strip()


def _clean_english_words(text: str) -> str:
    """替换中文报告中不应出现的英文词汇，保留专业缩写"""
    # 保留的专业缩写（不替换）
    _KEEP = {'GDP', 'PD', 'LGD', 'EAD', 'ROE', 'ROA', 'EBITDA', 'PPP',
             'COD', 'BOD', 'ISO', 'API', 'EPC', 'BT', 'BOT', 'TOT',
             'PMI', 'CPI', 'PPI', 'IPO', 'PE', 'VC', 'CEO', 'CFO',
             'SWOT', 'KPI', 'OA', 'IT', 'HR', 'ESG', 'ETF', 'WACC'}

    def _replace(m):
        word = m.group(0)
        if word.upper() in _KEEP:
            return word
        return _ENGLISH_REPLACEMENTS.get(word, word)

    result = _ENGLISH_WORD_RE.sub(_replace, text)
    # 额外清理：去除残留的孤立英文单词（超过3字母且不在保留列表）
    def _clean_residual(m):
        w = m.group(0)
        if w.upper() in _KEEP or w in _ENGLISH_REPLACEMENTS:
            return _ENGLISH_REPLACEMENTS.get(w, w)
        return ''  # 删除未知英文词
    result = re.sub(r'(?<=[\u4e00-\u9fff])\s*\b([A-Za-z]{4,})\b\s*(?=[\u4e00-\u9fff])', _clean_residual, result)
    # 清理多余空格
    result = re.sub(r'  +', ' ', result)
    return result.strip()


def _clean_bracket_tags(text):
    return re.sub(r'【[^】]*】', '', text)


def _is_subheading(line: str) -> bool:
    """判断是否是子标题（仅子标题允许加粗）"""
    patterns = [
        r'^[（(]\s*[一二三四五六七八九十]+\s*[）)]',
        r'^[一二三四五六七八九十]+[、.．]',
        r'^第[一二三四五六七八九十]+[节部分]',
        r'^\d+[\.\、]\s*\S+',
        r'^\d+\.\d+[\.\s]',
    ]
    stripped = line.strip()
    return any(re.match(p, stripped) for p in patterns)


def _is_pure_supplement(line: str) -> bool:
    """判断该行是否整段都是关于缺失材料（正文应留空）"""
    stripped = line.strip()
    if not stripped:
        return False
    # 短段落且包含补充关键词 → 大概率是纯补充信息
    if len(stripped) < 80 and _SUPPLEMENT_RE.search(stripped):
        return True
    # 长段落中也可能以"该数据材料未提供..."开头
    if _PURE_SUPPLEMENT_RE.match(stripped):
        return True
    return False


def _has_supplement_note(line: str) -> bool:
    """判断该行是否包含补充材料的注解（可能是混合内容）"""
    return bool(_SUPPLEMENT_RE.search(line))


def _extract_supplement_note(line: str) -> tuple[str, str]:
    """
    将混合行拆分为 (正文内容, 补充批注)
    如果无法拆分，返回 (原文, "")
    """
    # 尝试用句号/逗号分隔，找到包含"需补充"的分句
    # 常见模式："正文内容。该数据材料未提供，建议补充。"
    sentences = re.split(r'([。；;])', line)
    body_parts = []
    note_parts = []

    i = 0
    while i < len(sentences):
        sent = sentences[i]
        # 检查下一个是否是标点
        sep = sentences[i + 1] if i + 1 < len(sentences) else ""
        full_sent = sent + sep

        if _SUPPLEMENT_RE.search(sent):
            note_parts.append(full_sent.strip())
        else:
            body_parts.append(full_sent.strip())

        i += 2 if sep else 1

    body = "".join(body_parts).strip()
    note = "".join(note_parts).strip()

    if not body and note:
        # 整段都是补充内容
        return ("", note)
    if body and note:
        return (body, note)
    return (line, "")


def _add_plain_line(para, text, font_name, font_size, bold=False):
    """★ v7: 写入纯文本（无 Markdown 解析，不会意外加粗）"""
    text = _strip_markdown(text)
    text = _clean_english_words(text)
    if not text:
        return
    run = para.add_run(text)
    _set_font(run, font_name, font_size, bold=bold)


# ====== Word 批注系统（从 word_comments_v2 移植） ======

def _get_max_comment_id(doc: Document) -> int:
    """获取文档中现有最大 comment ID"""
    max_id = -1
    body = doc.element.body
    for crs in body.iter(qn('w:commentRangeStart')):
        cid = crs.get(qn('w:id'))
        if cid is not None:
            max_id = max(max_id, int(cid))
    for rel in doc.part.rels.values():
        if 'comments' in str(rel.reltype).lower() and 'Extended' not in str(rel.reltype):
            try:
                tree = etree.fromstring(rel.target_part.blob)
                for ce in tree.findall(qn('w:comment')):
                    cid = ce.get(qn('w:id'))
                    if cid is not None:
                        max_id = max(max_id, int(cid))
            except Exception:
                pass
    return max_id


def _get_or_create_comments_part(doc: Document):
    """获取或创建 comments.xml part"""
    for rel in doc.part.rels.values():
        if 'comments' in str(rel.reltype).lower() and 'Extended' not in str(rel.reltype):
            return rel.target_part
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    root = etree.Element(qn('w:comments'), nsmap=nsmap)
    blob = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
    part = Part(
        PackURI('/word/comments.xml'),
        'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
        blob,
        doc.part.package,
    )
    doc.part.relate_to(
        part,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments',
    )
    return part


def _insert_comment_on_para(doc, comments_tree, para_elem, comment_id, text, author="Agent"):
    """在指定段落元素上插入批注"""
    now = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    # 1. commentRangeStart
    range_start = etree.Element(qn('w:commentRangeStart'))
    range_start.set(qn('w:id'), str(comment_id))
    para_elem.insert(0, range_start)

    # 2. commentRangeEnd
    range_end = etree.SubElement(para_elem, qn('w:commentRangeEnd'))
    range_end.set(qn('w:id'), str(comment_id))

    # 3. commentReference run
    ref_run = etree.SubElement(para_elem, qn('w:r'))
    ref_rpr = etree.SubElement(ref_run, qn('w:rPr'))
    ref_style = etree.SubElement(ref_rpr, qn('w:rStyle'))
    ref_style.set(qn('w:val'), 'CommentReference')
    comment_ref = etree.SubElement(ref_run, qn('w:commentReference'))
    comment_ref.set(qn('w:id'), str(comment_id))

    # 4. 在 comments.xml 中添加评论内容
    comment_elem = etree.SubElement(comments_tree, qn('w:comment'))
    comment_elem.set(qn('w:id'), str(comment_id))
    comment_elem.set(qn('w:author'), author)
    comment_elem.set(qn('w:date'), now)
    comment_elem.set(qn('w:initials'), author[:2])

    for cline in text.split('\n'):
        cp = etree.SubElement(comment_elem, qn('w:p'))
        cr = etree.SubElement(cp, qn('w:r'))
        ct = etree.SubElement(cr, qn('w:t'))
        ct.text = cline
        ct.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


class _CommentManager:
    """★ v7.4 增强：批注管理器 — 收集批注并智能合并，每章最多8条"""

    MAX_COMMENTS_PER_CHAPTER = 8  # ★ v7.4: 限制每章批注数量

    def __init__(self):
        self.pending: list[tuple[int, str]] = []  # [(para_index, comment_text), ...]
        self._para_count = 0

    def next_para(self):
        self._para_count += 1

    @property
    def current_para_idx(self):
        return self._para_count - 1

    def add_comment(self, comment_text: str):
        """为当前段落添加一条批注"""
        self.pending.append((self.current_para_idx, comment_text))

    def _consolidate(self) -> list[tuple[int, str]]:
        """★ v7.4: 合并批注 — 连续的"⚠️ XX未提供"类批注合并为一条"""
        if len(self.pending) <= self.MAX_COMMENTS_PER_CHAPTER:
            return self.pending

        # 分类：⚠️ 缺失类 vs 💡 建议类 vs 其他
        missing_comments = []  # (para_idx, text)
        other_comments = []
        for para_idx, text in self.pending:
            if text.startswith("⚠️"):
                missing_comments.append((para_idx, text))
            else:
                other_comments.append((para_idx, text))

        # 如果缺失类批注太多 → 合并为几条
        if len(missing_comments) > 4:
            # 按段落位置分组，每组合并为一条
            group_size = max(1, len(missing_comments) // 3)
            merged = []
            for i in range(0, len(missing_comments), group_size):
                group = missing_comments[i:i + group_size]
                if len(group) == 1:
                    merged.append(group[0])
                else:
                    # 合并为一条，附到最后一个段落
                    items = [c[1].replace("⚠️ ", "").strip() for c in group]
                    # 截取每项关键词（最多10字）
                    short_items = [it[:15] for it in items]
                    combined = "⚠️ 以下信息需补充材料：" + "；".join(short_items)
                    if len(combined) > 200:
                        combined = combined[:197] + "..."
                    merged.append((group[-1][0], combined))
            missing_comments = merged

        result = other_comments + missing_comments
        # 按段落顺序排序
        result.sort(key=lambda x: x[0])

        # 最终限制总数
        if len(result) > self.MAX_COMMENTS_PER_CHAPTER:
            result = result[:self.MAX_COMMENTS_PER_CHAPTER]

        return result

    def apply_all(self, doc: Document):
        """将所有收集的批注写入文档（★ v7.4: 先合并再写入）"""
        if not self.pending:
            return

        # ★ v7.4: 先合并，减少批注数量
        consolidated = self._consolidate()

        comments_part = _get_or_create_comments_part(doc)
        tree = etree.fromstring(comments_part.blob)
        max_id = _get_max_comment_id(doc)
        cid = max_id + 1

        for para_idx, text in consolidated:
            if para_idx < len(doc.paragraphs):
                _insert_comment_on_para(
                    doc, tree,
                    doc.paragraphs[para_idx]._element,
                    cid, text, author="信贷报告Agent",
                )
                cid += 1

        comments_part._blob = etree.tostring(
            tree, xml_declaration=True, encoding='UTF-8', standalone=True,
        )


# ====== 封面表格自动填充 ======

def _extract_header_info(sections: dict) -> dict:
    """从章节内容中自动提取封面表格需要的信息"""
    info = {
        "credit_amount": "",
        "industry": "",
        "pd_rating": "",
        "apply_unit": "",
        "credit_type": "",
    }

    all_text = "\n".join(sections.get(f"ch{i}", "") for i in range(1, 5))

    # 授信金额
    for pat in [
        r'(?:本次)?(?:申请|拟)?授信(?:总)?(?:金额|额度)[为：:\s]*(\d[\d,，.]+\s*[万亿]?元)',
        r'(?:总)?(?:授信|敞口)(?:金额|额度)[为：:\s]*(\d[\d,，.]+\s*[万亿]?元)',
        r'(\d[\d,，.]+\s*万元).*授信',
    ]:
        m = re.search(pat, all_text)
        if m:
            info["credit_amount"] = m.group(1)
            break

    # 行业投向
    for pat in [
        r'(?:所属|属于|从事).*?行业[为是：:\s]*([\u4e00-\u9fff]+(?:业|服务|工程))',
        r'行业(?:投向|分类)[为是：:\s]*([\u4e00-\u9fff]{2,15})',
        r'(?:主营|经营).*?(?:业务|范围)[为包括：:\s]*([\u4e00-\u9fff]{4,20})',
    ]:
        m = re.search(pat, all_text)
        if m:
            info["industry"] = m.group(1)
            break

    # PD评级
    for pat in [
        r'PD\s*评级[为：:\s]*([A-Za-z0-9\+\-]+)',
        r'(?:客户|信用)评级[为：:\s]*([A-Za-z0-9\+\-]+)',
        r'风险等级[为：:\s]*([\u4e00-\u9fff]{1,4})',
    ]:
        m = re.search(pat, all_text)
        if m:
            info["pd_rating"] = m.group(1)
            break

    # 申报单位 — 通常是分支行名
    for pat in [
        r'(?:申报|经办|承办)[单支分]*[位行][为是：:\s]*([\u4e00-\u9fff]+(?:支行|分行|营业部))',
    ]:
        m = re.search(pat, all_text)
        if m:
            info["apply_unit"] = m.group(1)
            break

    # 授信品种
    for pat in [
        r'授信品种[为包括：:\s]*([\u4e00-\u9fff]+(?:贷款|保函|保理|承兑|信用证)[\u4e00-\u9fff、+/]*)',
        r'(?:流(?:动资金)?贷(?:款)?|保函|保理|承兑|信用证)',
    ]:
        m = re.search(pat, all_text)
        if m:
            info["credit_type"] = m.group(0) if m.lastindex is None else m.group(1)
            break

    return info


# ====== 子标题检测 ======
def strip_chapter_title(text):
    lines = text.split("\n")
    while lines and not lines[0].strip():
        lines.pop(0)
    if not lines:
        return text
    first = lines[0].strip()
    for ch_title in RPT_CHAPTERS:
        if first == ch_title or first.startswith(ch_title):
            lines.pop(0)
            while lines and not lines[0].strip():
                lines.pop(0)
            return "\n".join(lines)
    for pat in _CHAPTER_TITLE_PATTERNS:
        if re.match(pat, first):
            lines.pop(0)
            while lines and not lines[0].strip():
                lines.pop(0)
            return "\n".join(lines)
    return text


# ====== 报告生成核心函数 ======

def _write_chapter_body(doc: Document, body: str, body_font: str, body_size: float,
                        heading_font: str, cm: _CommentManager):
    """
    ★ v7 核心：将章节文本写入 Word，处理批注和格式

    - 纯补充段落 → 正文写 "（待补充）" + 侧栏批注
    - 混合段落 → 正文写有效内容 + 侧栏批注补充说明
    - 普通段落 → 纯文本写入，无加粗（子标题除外）
    """
    body = strip_chapter_title(body)
    body = _clean_bracket_tags(body)

    for line in body.split("\n"):
        line = line.strip()
        if not line:
            continue

        # ★ v7: 先清理 Markdown 和英文
        clean_line = _strip_markdown(line)
        clean_line = _clean_english_words(clean_line)

        if not clean_line:
            continue

        is_sub = _is_subheading(clean_line)

        # ★ v7: 检查是否是纯"需补充"段落
        if _is_pure_supplement(clean_line):
            # 正文写短占位符
            pp = doc.add_paragraph()
            cm.next_para()
            pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pp.paragraph_format.first_line_indent = Pt(body_size * 2)
            # 正文留空（只写"（待补充）"灰色占位）
            run = pp.add_run("（待补充）")
            _set_font(run, body_font, body_size, bold=False)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            # 侧栏批注
            cm.add_comment(f"⚠️ {clean_line}")
            continue

        # ★ v7: 检查是否是混合内容（有正文+有补充注解）
        if _has_supplement_note(clean_line) and not is_sub:
            real_body, note = _extract_supplement_note(clean_line)
            if note:
                pp = doc.add_paragraph()
                cm.next_para()
                pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pp.paragraph_format.first_line_indent = Pt(body_size * 2)
                if real_body:
                    _add_plain_line(pp, real_body, body_font, body_size, bold=False)
                else:
                    run = pp.add_run("（待补充）")
                    _set_font(run, body_font, body_size, bold=False)
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                cm.add_comment(f"⚠️ {note}")
                continue

        # ★ v7.3: 检查是否包含主观建议（应转为批注而非写入正文）
        if not is_sub and _SUBJECTIVE_RE.search(clean_line):
            # 整句都是建议→正文保留但加批注标记
            pp = doc.add_paragraph()
            cm.next_para()
            pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pp.paragraph_format.first_line_indent = Pt(body_size * 2)
            _add_plain_line(pp, clean_line, body_font, body_size, bold=False)
            cm.add_comment(f"💡 此为Agent基于分析的主观建议，非客观事实，请审核确认")
            continue

        # ★ v7: 正常段落 — 子标题加粗，正文不加粗
        pp = doc.add_paragraph()
        cm.next_para()

        if is_sub:
            pp.paragraph_format.first_line_indent = Pt(0)
            pp.paragraph_format.space_before = Pt(8)
            _add_plain_line(pp, clean_line, heading_font, body_size, bold=True)
        else:
            pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pp.paragraph_format.first_line_indent = Pt(body_size * 2)
            _add_plain_line(pp, clean_line, body_font, body_size, bold=False)


def make_word_report(client_name, credit_amount, industry, pd_rating,
                     apply_unit, credit_type, sections):
    """生成 Word 报告，返回文件路径"""
    doc = Document()
    cm = _CommentManager()

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)

    # ★ v7.1: 封面表格留空，由客户自行填写（不自动推测）

    # 封面标题
    cover = doc.add_paragraph()
    cm.next_para()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(72)
    cover.paragraph_format.space_after = Pt(24)
    title_text = RPT_TITLE_PREFIX + (client_name or "XXX")
    if credit_amount:
        title_text += "\n" + credit_amount + RPT_TITLE_SUFFIX
    else:
        title_text += "\n" + RPT_TITLE_SUFFIX
    _set_font(cover.add_run(title_text), "黑体", 22, bold=True)

    doc.add_paragraph()
    cm.next_para()

    # 封面表格 — ★ v7 自动填充
    values = [client_name, credit_amount, industry, pd_rating, apply_unit, credit_type]
    tbl = doc.add_table(rows=6, cols=2, style="Table Grid")
    tbl.columns[0].width = Inches(1.8)
    tbl.columns[1].width = Inches(4.2)
    for i, (label, val) in enumerate(zip(RPT_COVER_LABELS, values)):
        cells = tbl.rows[i].cells
        cells[0].text = label
        cells[1].text = val or ""
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_font(r, "仿宋", 16)
    doc.add_page_break()

    # 正文
    ch_keys = ["ch1", "ch2", "ch3", "ch4"]
    for title, key in zip(RPT_CHAPTERS, ch_keys):
        hp = doc.add_paragraph()
        cm.next_para()
        hp.paragraph_format.space_before = Pt(14)
        hp.paragraph_format.space_after = Pt(6)
        _set_font(hp.add_run(title), "黑体", 16, bold=True)

        body = sections.get(key, "")
        _write_chapter_body(doc, body, "仿宋", 14, "黑体", cm)

        sep = doc.add_paragraph()
        cm.next_para()

    # ★ v7: 统一写入所有批注
    cm.apply_all(doc)

    safe = "".join(c for c in (client_name or "report") if c.isalnum() or c in "_-")[:20]
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix=safe + "_v72_report_")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def make_word_from_template(template_path: str, sections: dict,
                            client_name: str = "") -> str:
    """用用户模版的页面设置和字体生成报告"""
    try:
        template = Document(template_path)
    except Exception:
        return make_word_report(client_name, "", "", "", "", "", sections)

    doc = Document()
    cm = _CommentManager()

    for src_sec in template.sections:
        for dst_sec in doc.sections:
            dst_sec.top_margin = src_sec.top_margin
            dst_sec.bottom_margin = src_sec.bottom_margin
            dst_sec.left_margin = src_sec.left_margin
            dst_sec.right_margin = src_sec.right_margin

    # 检测模版字体
    body_font, body_size, heading_font, heading_size = "仿宋", 14, "黑体", 16
    for para in template.paragraphs:
        for run in para.runs:
            if run.font.name and run.font.size:
                style = para.style.name if para.style else ""
                if "Heading" in style:
                    heading_font = run.font.name
                    heading_size = run.font.size.pt
                else:
                    body_font = run.font.name
                    body_size = run.font.size.pt
                break
        if body_font != "仿宋":
            break

    # ★ v7.1: 封面表格由客户填写，不自动推测

    ch_keys = ["ch1", "ch2", "ch3", "ch4"]
    for title, key in zip(RPT_CHAPTERS, ch_keys):
        hp = doc.add_paragraph()
        cm.next_para()
        hp.paragraph_format.space_before = Pt(14)
        _set_font(hp.add_run(title), heading_font, heading_size, bold=True)

        body = sections.get(key, "")
        _write_chapter_body(doc, body, body_font, body_size, heading_font, cm)

        sep = doc.add_paragraph()
        cm.next_para()

    # ★ v7: 统一写入所有批注
    cm.apply_all(doc)

    safe = "".join(c for c in (client_name or "report") if c.isalnum() or c in "_-")[:20]
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix=safe + "_v72_tpl_")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def parse_template_structure(template_path: str) -> str:
    """提取模版标题大纲"""
    try:
        doc = Document(template_path)
        headings = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ""
            if ("Heading" in style or
                (para.runs and any(r.bold for r in para.runs) and len(text) < 50) or
                _is_subheading(text)):
                headings.append(text)
        return "\n".join(headings) if headings else "[模版中未发现标题]"
    except Exception as e:
        return f"[模版解析失败: {e}]"


def parse_template_full_text(template_path: str) -> str:
    """★ v7: 提取模版完整正文（不只是标题），用于深度分析参考"""
    try:
        doc = Document(template_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        return "\n".join(parts)
    except Exception:
        return ""
