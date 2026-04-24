# -*- coding: utf-8 -*-
"""深柱 5 家材料包生成器 · v2 Task B 专用

运行方式：
    cd <repo_root>
    py data/mock/deep-pillar/_gen/generate_dp.py

产物：`data/mock/deep-pillar/DP001~DP005/` 各 20-40 份异构原始文件。
- PDF：营业执照、章程、身份证扫描件样式、审计报告、完税证明、申报表、租赁合同等
- XLSX：财报（资产负债表/利润表/现金流量表）、项目清单、专利证书明细、资质表
- XLS：银行流水（Excel 97-2003 格式，参照中锐续贷包形态）
- DOCX：银行授信补充问答、公司章程附件

零答案字段：不生成 difficulty / tags / benchmark_ref 等元数据。
命名规则：序号前缀混乱、日期拼接、扩展名混用、中文空格混用（仿真实客户提交）。
"""

from __future__ import annotations

import os
import random
import sys
from datetime import datetime
from pathlib import Path
from typing import Iterable

# 第三方库
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
import xlwt
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)

# 本地 profiles
sys.path.insert(0, str(Path(__file__).parent))
from profiles import all_profiles  # noqa: E402

# ----------------------------------------------------------------------------
# 字体（使用 reportlab 内置 CID 中文字体，不依赖本地字体文件）
# ----------------------------------------------------------------------------
pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

CN_FONT = "STSong-Light"
CN_FONT_BOLD = "STSong-Light"  # reportlab 内置 CID 仅 STSong-Light 可用 · 粗体用字号区分

# ----------------------------------------------------------------------------
# 随机种子（每家不同，保留跨 run 可复现）
# ----------------------------------------------------------------------------
SEED_MAP = {"DP001": 1001, "DP002": 1002, "DP003": 1003, "DP004": 1004, "DP005": 1005}


def _seed_for(dp_id: str) -> random.Random:
    return random.Random(SEED_MAP[dp_id])


# ----------------------------------------------------------------------------
# PDF 生成辅助
# ----------------------------------------------------------------------------
def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CNTitle",
            fontName=CN_FONT_BOLD,
            fontSize=16,
            leading=22,
            alignment=1,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CNHeading",
            fontName=CN_FONT_BOLD,
            fontSize=13,
            leading=18,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CNBody",
            fontName=CN_FONT,
            fontSize=10.5,
            leading=16,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CNMono",
            fontName=CN_FONT,
            fontSize=9.5,
            leading=14,
            textColor=colors.HexColor("#444444"),
        )
    )
    return styles


def _write_pdf(path: Path, title: str, body_blocks: list):
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = _pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=2 * 28.35,
        rightMargin=2 * 28.35,
        topMargin=2.5 * 28.35,
        bottomMargin=2 * 28.35,
    )
    story = [Paragraph(title, styles["CNTitle"])]
    for block in body_blocks:
        if isinstance(block, str):
            if block == "__pb__":
                story.append(PageBreak())
            elif block.startswith("## "):
                story.append(Paragraph(block[3:], styles["CNHeading"]))
            else:
                story.append(Paragraph(block, styles["CNBody"]))
        elif isinstance(block, tuple) and block[0] == "table":
            data = block[1]
            tbl = Table(data, hAlign="LEFT")
            tbl.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), CN_FONT),
                        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F0F0F0")),
                        ("FONTNAME", (0, 0), (-1, 0), CN_FONT_BOLD),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(tbl)
            story.append(Spacer(1, 6))
        elif isinstance(block, tuple) and block[0] == "spacer":
            story.append(Spacer(1, block[1]))
    doc.build(story)


# ----------------------------------------------------------------------------
# XLSX 生成辅助（openpyxl）
# ----------------------------------------------------------------------------
def _write_xlsx(path: Path, sheets: dict, freeze_header: bool = True):
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)
    for sheet_name, rows in sheets.items():
        ws = wb.create_sheet(sheet_name[:30])
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, val in enumerate(row, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=val)
                if r_idx == 1:
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal="center")
        # autofit columns（粗暴）
        for col_idx in range(1, (len(rows[0]) if rows else 1) + 1):
            ws.column_dimensions[get_column_letter(col_idx)].width = 18
        if freeze_header:
            ws.freeze_panes = "A2"
    wb.save(str(path))


# ----------------------------------------------------------------------------
# XLS 生成辅助（xlwt · Excel 97-2003）
# ----------------------------------------------------------------------------
def _write_xls(path: Path, sheet_name: str, header: list, rows: Iterable[list]):
    path.parent.mkdir(parents=True, exist_ok=True)
    book = xlwt.Workbook(encoding="utf-8")
    sheet = book.add_sheet(sheet_name[:30])
    header_style = xlwt.easyxf("font: bold on; align: horiz center; pattern: pattern solid, fore_colour gray25")
    for c_idx, h in enumerate(header):
        sheet.write(0, c_idx, h, header_style)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            sheet.write(r_idx, c_idx, val)
    for c_idx in range(len(header)):
        sheet.col(c_idx).width = 256 * 18
    book.save(str(path))


# ----------------------------------------------------------------------------
# DOCX 生成辅助（python-docx）
# ----------------------------------------------------------------------------
def _docx_new() -> Document:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "宋体"
    st.font.size = Pt(11)
    return doc


def _docx_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x30, 0x30, 0x30)


def _docx_save(doc: Document, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


# ============================================================================
# 6 大类生成器
# ============================================================================

# --------------- 1. 资质类 -----------------
def gen_qualifications(p: dict, outdir: Path, rng: random.Random):
    name = p["name_full"]
    dp = p["dp_id"]
    lp = p["legal_person"]

    # 营业执照副本（pdf · 多年版本 · 模拟"变更后最新+调档件"）
    issue_year = p["establish_year"] + rng.randint(8, 12)
    _write_pdf(
        outdir / f"1、{p['name_short']}营业执照副本{issue_year - 2000}.{rng.randint(4, 12)}.{rng.randint(1,28)}.pdf",
        "营业执照（副本）",
        [
            "## 基本信息",
            f"统一社会信用代码：{p['uscc']}",
            f"名称：{name}",
            f"类型：有限责任公司（自然人投资或控股）",
            f"住所：{p['region_full']} · {p['address_detail']}",
            f"法定代表人：{lp}",
            f"注册资本：人民币 {p['registered_capital_wan']} 万元",
            f"成立日期：{p['establish_year']}-0{rng.randint(3,9)}-{rng.randint(10,28)}",
            f"营业期限：{p['establish_year']}-0{rng.randint(3,9)}-{rng.randint(10,28)} 至 长期",
            "经营范围：一般项目："
            + "、".join(rng.sample(
                ["机械零部件加工", "工业自动化设备研发", "金属材料销售",
                 "机电产品销售", "家用电器及电子产品零售", "家居建材批发",
                 "装饰装修工程", "建筑工程设计与施工", "仓储服务",
                 "企业管理咨询", "财务咨询", "市场营销策划", "五金交电销售",
                 "新材料技术研发", "工程技术咨询"],
                k=rng.randint(5, 8),
            )),
            ("spacer", 18),
            "发照机关：" + p["region_full"].split("-")[-1] + "市场监督管理局",
            f"注册登记日期：{issue_year}-{rng.randint(2,10):02d}-{rng.randint(1,28):02d}",
        ],
    )

    # 章程调档
    ch_year = 2025
    _write_pdf(
        outdir / f"1、{p['name_short']}章程调档{ch_year - 2000}.{rng.randint(4,9)}.{rng.randint(1,28)}.pdf",
        f"{name} 公司章程（调档版本）",
        [
            "## 第一章 总则",
            f"第一条 为了规范 {name}（以下简称“公司”）的组织和行为，保护公司、股东、债权人的合法权益，依据《公司法》等相关法律法规，制定本章程。",
            f"第二条 公司名称：{name}",
            f"第三条 公司住所：{p['region_full']} · {p['address_detail']}",
            "第四条 公司为有限责任公司，股东以其出资额为限对公司承担责任。",
            f"第五条 公司成立于 {p['establish_year']} 年，法定代表人为 {lp}。",
            "## 第二章 经营范围",
            f"第六条 公司主营业务：{p['main_business']}",
            "公司依法自主经营，依法纳税，维护股东合法权益，接受政府有关部门的监督管理。",
            "## 第三章 注册资本与股东",
            f"第七条 公司注册资本人民币 {p['registered_capital_wan']} 万元。",
            f"第八条 公司股东及出资情况：{lp}（占 {rng.randint(60,100)}%）、其他自然人股东若干（合计占 {rng.randint(0,40)}%）。",
            "## 第四章 公司治理结构",
            "第九条 公司设执行董事一人，由股东会选举产生，任期三年。",
            "第十条 公司设监事一人，由股东会选举产生，任期三年。",
            "第十一条 公司设总经理一人，由执行董事任命或兼任。",
            "## 第五章 财务会计与利润分配",
            "第十二条 公司依法建立财务会计制度，依法纳税，接受审计。",
            "第十三条 公司弥补亏损、提取法定公积金后，剩余利润按股东出资比例分配。",
            "## 第六章 附则",
            "第十四条 本章程由全体股东共同制定，自股东签署之日起生效。",
            ("spacer", 20),
            f"调档日期：{ch_year}-{rng.randint(3,8):02d}-{rng.randint(1,28):02d}",
            f"发照机关存档：{p['region_full'].split('-')[-1]}市场监督管理局",
        ],
    )

    # 资质表 xls
    _write_xls(
        outdir / f"1、{p['name_short']}资质表.xls",
        "资质明细",
        ["证书名称", "签发机关", "有效期起始", "有效期截止", "证书编号", "备注"],
        [
            [q[0], f"{p['region_full'].split('-')[-1]}主管部门",
             f"{q[1]}-0{rng.randint(1,9)}-{rng.randint(1,28):02d}",
             f"{q[1] + 3}-0{rng.randint(1,9)}-{rng.randint(1,28):02d}",
             f"{p['region_full'].split('-')[-1][:2]}{rng.randint(10000,99999)}",
             ""]
            for q in p["qualifications"]
        ],
    )

    # 法人身份证（pdf · 扫描件样式）
    _write_pdf(
        outdir / f"1、{lp}身份证{2019 + rng.randint(0,5)}.pdf",
        f"{lp} 身份证扫描件",
        [
            "说明：本文件为身份证扫描件（正反面）。",
            "",
            "## 正面信息（已脱敏）",
            f"姓名：{lp}",
            "性别：男",
            "民族：汉",
            f"出生：1970 年代",
            f"住址：{p['region_full']}（具体门牌脱敏）",
            f"公民身份号码：{p['legal_person_id_mask']}",
            "",
            "## 反面信息",
            f"签发机关：{p['region_full'].split('-')[-1]}公安局",
            f"有效期限：长期",
            "",
            "_本扫描件由客户经理采集用于授信尽调 · 仅限信贷业务使用_",
        ],
    )

    # 开户许可证（pdf · 较老年份）
    open_year = p["establish_year"] + rng.randint(0, 2)
    _write_pdf(
        outdir / f"1、{p['name_short']}-开户许可证{open_year}.{rng.randint(3,8)}.{rng.randint(1,28)}.pdf",
        "人民币银行结算账户开户许可证",
        [
            f"核准号：J{rng.randint(1000,9999)}{rng.randint(10000,99999)}",
            f"账户名称：{name}",
            f"账户性质：基本存款账户",
            f"开户银行：{p['open_bank']}",
            f"账号：{p['open_bank_acct']}",
            f"核准日期：{open_year}-0{rng.randint(3,9)}-{rng.randint(1,28):02d}",
            ("spacer", 20),
            "核准机关：中国人民银行（脱敏版）",
        ],
    )

    # 专利证书明细（xlsx · 按 profile 的 patents_count）
    if p["patents_count"] > 0:
        patents = []
        for i in range(p["patents_count"]):
            patents.append([
                f"{p['industry']}相关专利 #{i + 1}（脱敏名）",
                rng.choice(["发明专利", "实用新型", "外观设计"]),
                f"ZL{rng.randint(2015, 2024)}1{rng.randint(10000000, 99999999)}",
                f"{rng.randint(2015, 2024)}-{rng.randint(1,12):02d}-{rng.randint(1,28):02d}",
                lp,
                "有效",
            ])
        _write_xlsx(
            outdir / f"{'00' if rng.random() < 0.5 else '1'}、{p['name_short']}专利证书明细.xlsx",
            {
                "专利明细": [
                    ["专利名称", "专利类型", "专利号", "授权公告日", "专利权人", "状态"],
                ] + patents
            },
        )

    # 高新认定 / 专精特新备案公告（pdf）
    for q, q_year in p["qualifications"]:
        if "高新" in q or "专精特新" in q:
            _write_pdf(
                outdir / f"1、对{p['region_full'].split('-')[0]}{q}备案公告_{q_year}第{rng.randint(1,3)}批.pdf",
                f"关于 {p['region_full'].split('-')[0]} 认定机构 {q_year} 年{q[:4]}企业备案公告",
                [
                    f"根据《{q}管理办法》及相关实施细则，经主管部门审核，现将{p['region_full'].split('-')[0]} {q_year} 年认定备案的{q[:4]}企业予以公示。",
                    ("spacer", 10),
                    "## 公示内容（摘录）",
                    f"企业名称：{name}",
                    f"统一社会信用代码：{p['uscc']}",
                    f"备案/认定类别：{q}",
                    f"有效期：{q_year} 年 1 月 1 日 至 {q_year + 3} 年 12 月 31 日",
                    f"主管机关：{p['region_full'].split('-')[0]}科技主管部门（脱敏）",
                    ("spacer", 12),
                    "公示期 10 个工作日，期间如有异议可向主管部门反映。",
                    f"公示日期：{q_year}-0{rng.randint(3,9)}-{rng.randint(1,28):02d}",
                ],
            )


# --------------- 2. 场所类 -----------------
def gen_venue(p: dict, outdir: Path, rng: random.Random):
    for idx, (start_yr, end_yr) in enumerate(p["lease_periods"], start=1):
        tag = "" if idx == 1 else f"({idx})"
        _write_pdf(
            outdir / f"1、{p['region_full'].split('-')[-1]}房产租赁协议{tag}.pdf",
            "房产租赁合同",
            [
                f"甲方（出租方）：{p['region_full'].split('-')[-1]}某物业资产管理有限公司（脱敏）",
                f"乙方（承租方）：{p['name_full']}",
                "## 第一条 租赁标的",
                f"甲方将其位于 {p['address_detail']} 的房产出租给乙方用于办公及经营。",
                f"建筑面积约 {rng.randint(300, 2800)} 平方米。",
                "## 第二条 租赁期限",
                f"租赁期限自 {start_yr}-{rng.randint(1, 12):02d}-01 起至 {end_yr}-{rng.randint(1, 12):02d}-{rng.randint(20,28):02d} 止。",
                "## 第三条 租金及支付",
                f"月租金人民币 {rng.randint(8, 25)} 元/平方米。付款方式：季付/半年付（双方确认）。",
                "## 第四条 用途",
                f"乙方承诺租赁房产用于与 {p['main_business']} 相关的办公与经营活动。",
                "## 第五条 违约责任",
                "任何一方违约，应按月租金的 30% 向对方支付违约金；连续 3 个月未付租金构成根本违约。",
                "## 第六条 其他",
                "本合同自双方签字盖章之日起生效，一式两份。",
                ("spacer", 24),
                f"甲方（签章）：脱敏   乙方（签章）：{p['name_short']}",
                f"签订日期：{start_yr}-{rng.randint(1,12):02d}-{rng.randint(1,28):02d}",
            ],
        )


# --------------- 3. 财务类 -----------------
def _balance_sheet(p: dict, year: int, rng: random.Random):
    # 依据 profile 的 total_asset_2024 / total_liab_2024 倒推年份比例（合理波动）
    base_ta = p["total_asset_2024_wan"]
    base_tl = p["total_liab_2024_wan"]
    yr_offset = year - 2024
    # 量级缩放因子（越老越小），并叠加小幅随机波动
    scale = max(0.5, 1.0 + yr_offset * 0.09 + rng.uniform(-0.05, 0.05))
    ta = int(base_ta * scale)
    tl = int(base_tl * scale)
    eq = ta - tl
    return [
        ["资产项目", "期末数（万元）", "", "负债及所有者权益项目", "期末数（万元）"],
        ["流动资产合计", int(ta * 0.62), "", "流动负债合计", int(tl * 0.55)],
        ["  其中：货币资金", int(ta * 0.08), "", "  其中：短期借款", int(tl * 0.30)],
        ["  应收账款", int(ta * 0.22), "", "  应付账款", int(tl * 0.18)],
        ["  存货", int(ta * 0.18), "", "  预收账款", int(tl * 0.06)],
        ["  其他流动资产", int(ta * 0.14), "", "  其他流动负债", int(tl * 0.01)],
        ["非流动资产合计", int(ta * 0.38), "", "非流动负债合计", int(tl * 0.45)],
        ["  固定资产", int(ta * 0.22), "", "  长期借款", int(tl * 0.30)],
        ["  无形资产", int(ta * 0.05), "", "  应付债券", 0],
        ["  递延所得税资产", int(ta * 0.02), "", "  其他非流动负债", int(tl * 0.15)],
        ["  其他非流动资产", int(ta * 0.09), "", "负债合计", tl],
        ["", "", "", "所有者权益合计", eq],
        ["", "", "", "  实收资本", p["registered_capital_wan"]],
        ["", "", "", "  未分配利润", max(0, eq - p["registered_capital_wan"])],
        ["资产总计", ta, "", "负债和所有者权益总计", ta],
    ]


def _income_statement(p: dict, year: int, rng: random.Random):
    rev_key = f"revenue_{year}_wan"
    if rev_key not in p:
        rev = int(p["revenue_2024_wan"] * (0.85 + rng.uniform(-0.05, 0.1)))
    else:
        rev = p[rev_key]
    cost = int(rev * rng.uniform(0.72, 0.85))
    gp = rev - cost
    op_profit = int(gp * rng.uniform(0.18, 0.34))
    net = int(op_profit * rng.uniform(0.74, 0.92))
    return [
        ["项目", "本期金额（万元）"],
        ["一、营业收入", rev],
        ["减：营业成本", cost],
        ["    税金及附加", int(rev * 0.008)],
        ["    销售费用", int(rev * 0.045)],
        ["    管理费用", int(rev * 0.068)],
        ["    研发费用", int(rev * 0.028)],
        ["    财务费用", int(rev * 0.015)],
        ["二、营业利润", op_profit],
        ["加：营业外收入", rng.randint(10, 80)],
        ["减：营业外支出", rng.randint(5, 40)],
        ["三、利润总额", op_profit + rng.randint(0, 50)],
        ["减：所得税费用", int((op_profit + 10) * 0.18)],
        ["四、净利润", net],
    ]


def _cashflow_statement(p: dict, year: int, rng: random.Random):
    rev_key = f"revenue_{year}_wan"
    rev = p.get(rev_key, p["revenue_2024_wan"])
    inflow_op = int(rev * rng.uniform(0.90, 1.08))
    outflow_op = int(rev * rng.uniform(0.78, 0.95))
    return [
        ["项目", "本期金额（万元）"],
        ["一、经营活动产生的现金流量：", ""],
        ["销售商品、提供劳务收到的现金", inflow_op],
        ["收到的税费返还", rng.randint(20, 150)],
        ["收到其他与经营活动有关的现金", rng.randint(80, 300)],
        ["经营活动现金流入小计", inflow_op + rng.randint(100, 450)],
        ["购买商品、接受劳务支付的现金", outflow_op],
        ["支付给职工以及为职工支付的现金", int(rev * 0.08)],
        ["支付的各项税费", int(rev * 0.028)],
        ["支付其他与经营活动有关的现金", rng.randint(60, 280)],
        ["经营活动产生的现金流量净额", int(rev * rng.uniform(0.05, 0.13))],
        ["二、投资活动产生的现金流量：", ""],
        ["投资活动现金流入小计", rng.randint(100, 500)],
        ["投资活动现金流出小计", rng.randint(200, 900)],
        ["投资活动产生的现金流量净额", -rng.randint(100, 500)],
        ["三、筹资活动产生的现金流量：", ""],
        ["取得借款收到的现金", int(rev * rng.uniform(0.10, 0.30))],
        ["偿还债务支付的现金", int(rev * rng.uniform(0.08, 0.22))],
        ["筹资活动产生的现金流量净额", rng.randint(-400, 400)],
    ]


def gen_financial(p: dict, outdir: Path, rng: random.Random):
    years_available = [2022, 2023, 2024, 2025]
    if p.get("quirks", {}).get("missing_2022_audit"):
        # 2022 的 xlsx 财报产出，但审计报告故意缺失
        pass
    for yr in years_available:
        # 每年一份财报 xlsx（含 3 张表）
        fname_variants = [
            f"2、{yr}年度财务报表-{p['name_short']}{yr + 1}0{rng.randint(1,3)}{rng.randint(1,28):02d}.xlsx",
            f"2、{yr}年度财务报表-{p['name_short']}.xlsx",
            f"2、财务报表{yr}年-{p['name_short']}.xlsx",
        ]
        fname = rng.choice(fname_variants)
        _write_xlsx(
            outdir / fname,
            {
                f"资产负债表{yr}": _balance_sheet(p, yr, rng),
                f"利润表{yr}": _income_statement(p, yr, rng),
                f"现金流量表{yr}": _cashflow_statement(p, yr, rng),
            },
        )

    # 审计报告（2023 / 2024 · pdf 扫描件样式）
    skip_audit_years = set()
    if p.get("quirks", {}).get("missing_2022_audit"):
        skip_audit_years.add(2022)
    for yr in [2023, 2024]:
        if yr in skip_audit_years:
            continue
        _write_pdf(
            outdir / f"2、审计报告{yr}年-{p['audit_firm_alias']}.pdf",
            f"{p['name_full']} {yr} 年度财务报表审计报告",
            [
                f"致：{p['name_full']} 股东",
                ("spacer", 10),
                "## 一、审计意见",
                f"我们审计了 {p['name_full']}（以下简称“公司”）的财务报表，"
                f"包括 {yr} 年 12 月 31 日的资产负债表、{yr} 年度的利润表、"
                f"现金流量表以及相关财务报表附注。"
                + (
                    " 我们认为，该财务报表在所有重大方面按照企业会计准则编制，"
                    "公允反映了公司的财务状况、经营成果和现金流量。"
                    if not p.get("quirks", {}).get("financial_vs_flow_large_gap_pct")
                    else " 我们已就部分往来款项及关联交易事项与管理层进行了沟通，"
                         "详见审计意见说明段（强调事项）。"
                ),
                "## 二、形成审计意见的基础",
                "我们按照中国注册会计师审计准则的规定执行了审计工作。审计准则要求我们遵守职业道德，"
                "策划并实施审计工作，以对财务报表是否不存在重大错报获取合理保证。",
                "## 三、关键审计事项",
                f"本期识别的关键审计事项包括：收入确认、应收账款坏账准备、存货减值测试、关联方交易披露等。"
                f"具体审计应对详见附注说明。",
                "## 四、管理层和治理层对财务报表的责任",
                "管理层负责按照企业会计准则的规定编制财务报表，使其实现公允反映。",
                "## 五、注册会计师对财务报表审计的责任",
                "我们的目标是对财务报表整体是否不存在重大错报获取合理保证，并出具包含审计意见的审计报告。",
                ("spacer", 18),
                f"会计师事务所：{p['audit_firm_alias']}会计师事务所（特殊普通合伙）",
                f"地址：某省某市（脱敏）",
                f"签字注册会计师：李某   王某",
                f"报告日期：{yr + 1}-0{rng.randint(4,6)}-{rng.randint(15,28):02d}",
            ],
        )


# --------------- 4. 纳税类 -----------------
def gen_tax(p: dict, outdir: Path, rng: random.Random):
    # 完税证明（按所属期 · 2023 / 2024 / 2025 Q1-Q2）
    periods = [
        ("2023", "按所属期"),
        ("2024年度", "（按所属期）"),
        ("202501-12所属期", ""),  # 2025 跨期汇总件
    ]
    for period, suffix in periods:
        _write_pdf(
            outdir / f"3、税收完税证明{period}{suffix}.pdf",
            f"税收完税证明（{period}）",
            [
                f"纳税人名称：{p['name_full']}",
                f"纳税人识别号：{p['uscc']}",
                f"证明所属期：{period}",
                ("spacer", 10),
                ("table", [
                    ["税种", "税目", "所属期", "实缴金额（元）", "征收机关"],
                    ["增值税", "一般纳税人", period, f"{rng.randint(350000, 1800000):,}", "税务一分局"],
                    ["城市维护建设税", "一般", period, f"{rng.randint(12000, 55000):,}", "税务一分局"],
                    ["教育费附加", "一般", period, f"{rng.randint(8000, 38000):,}", "税务一分局"],
                    ["印花税", "购销合同等", period, f"{rng.randint(2000, 18000):,}", "税务一分局"],
                    ["企业所得税", "年度申报", period, f"{rng.randint(120000, 800000):,}", "税务一分局"],
                ]),
                ("spacer", 16),
                f"本证明由 {p['region_full'].split('-')[-1]}税务局（脱敏）出具，仅证明上述税种实缴情况。",
                f"开具日期：{period.split('-')[0] if '-' not in period else period.split('-')[-1][:4]}年末次月",
            ],
        )

    # 增值税及附加税申报表（按月 · 2023-2025 若干月）
    months_to_gen = [
        ("202312", "2023 12月"),
        ("202412", "2024 12月"),
        ("202503", "2025 03月"),
        ("202506", "2025 06月"),
        ("202509", "2025 09月"),
        ("202512", "2025 12月"),
    ]
    # 根据 profile bank_flow_months 覆盖范围适当筛选
    months_to_gen = [m for m in months_to_gen if m[0] <= "202512"]
    for mcode, mhum in months_to_gen[: rng.randint(4, 6)]:
        var = rng.choice([
            f"3、增值税及附加税申报表-{mcode}.pdf",
            f"3、增值税申报表-{mcode[:4]}年{mcode[4:]}月.pdf",
        ])
        _write_pdf(
            outdir / var,
            f"增值税及附加税费申报表（{mhum}）",
            [
                f"纳税人名称：{p['name_full']}",
                f"纳税人识别号：{p['uscc']}",
                f"所属期：{mhum}",
                ("spacer", 10),
                ("table", [
                    ["项目", "本期数（元）"],
                    ["销售额（含税）", f"{rng.randint(2800000, 9600000):,}"],
                    ["销项税额", f"{rng.randint(360000, 1200000):,}"],
                    ["进项税额", f"{rng.randint(280000, 1100000):,}"],
                    ["应纳税额（增值税）", f"{rng.randint(50000, 280000):,}"],
                    ["城市维护建设税（7%）", f"{rng.randint(3500, 22000):,}"],
                    ["教育费附加（3%）", f"{rng.randint(1500, 8500):,}"],
                    ["地方教育附加（2%）", f"{rng.randint(1000, 6000):,}"],
                ]),
                ("spacer", 12),
                f"申报人（签字）：{p['legal_person']}  受理日期：{mcode[:4]}-{mcode[4:]}-{rng.randint(13,18):02d}",
            ],
        )


# --------------- 5. 银行流水 -----------------
def gen_bank_flow(p: dict, outdir: Path, rng: random.Random):
    root = outdir / "4、银行流水"
    root.mkdir(parents=True, exist_ok=True)
    rev_2024 = p["revenue_2024_wan"]
    # 日均流水目标 = 年营收 / 250（工作日）
    per_day_base = rev_2024 * 10000 / 250  # 单位元

    for bank_full, bank_role in p["banks"]:
        # 简化银行名（仅取前几字做目录名）
        short = bank_full.split("银行")[0] + "银行"
        bank_dir = root / short
        bank_dir.mkdir(parents=True, exist_ok=True)

        for mcode in p["bank_flow_months"]:
            # 决定格式：主结算用 xls（仿真实银行老格式），辅助账户部分用 xlsx
            use_xls = "主结算" in bank_role or rng.random() < 0.7
            fname = f"{mcode}.xls" if use_xls else f"{mcode}.xlsx"
            fpath = bank_dir / fname

            # 流水条数（一个月 15-60 笔）
            num_rows = rng.randint(15, 55)
            header = ["交易日期", "交易时间", "交易类型", "金额（元）",
                      "借/贷", "对方户名", "对方账号", "摘要", "余额"]
            rows = []
            balance = per_day_base * rng.uniform(3, 12)
            year = int(mcode[:4])
            month = int(mcode[4:])
            # 日期范围
            from calendar import monthrange
            last_day = monthrange(year, month)[1]

            # 关联方过账（仅 DP004 / DP005 quirks）
            related_party_share = p.get("quirks", {}).get("related_party_flow_pct", 0) / 100.0
            related_parties = []
            if p["dp_id"] == "DP004":
                related_parties = ["汇德顺达建设工程有限公司"]
            if p["dp_id"] == "DP005":
                related_parties = ["星胤商管运营（某省）", "星悦产业投资有限公司", "慎之工贸（某省）有限公司"]

            customers = [c[0] for c in p["main_customers"] if c[1] is not None]
            suppliers = [s[0] for s in p["main_suppliers"]]

            for _ in range(num_rows):
                day = rng.randint(1, last_day)
                date_str = f"{year}-{month:02d}-{day:02d}"
                time_str = f"{rng.randint(8,17):02d}:{rng.randint(0,59):02d}:{rng.randint(0,59):02d}"
                side = rng.choice(["贷", "借", "贷", "借", "贷"])  # 贷(入)稍多
                base_amt = rng.uniform(0.1, 3.5) * per_day_base / 4
                # 关联方过账：有概率把对手方改成关联方
                if related_parties and rng.random() < related_party_share * 1.5:
                    counterpart = rng.choice(related_parties)
                    # 关联方往往金额偏大
                    base_amt *= rng.uniform(1.8, 3.2)
                    memo = rng.choice(["往来款", "调拨", "合作项目款", "内部往来", "借款"])
                else:
                    if side == "贷":
                        counterpart = rng.choice(customers) if customers else "（零售客户合计）"
                        memo = rng.choice(["货款", "销售回款", "项目款", "零售款", "其他"])
                    else:
                        counterpart = rng.choice(suppliers) if suppliers else "（综合供应）"
                        memo = rng.choice(["采购款", "货款支付", "服务费", "预付款", "税款缴纳"])
                amt = round(base_amt, 2)
                # 账号脱敏
                opp_acct = f"{rng.choice(['6222', '6228', '6225', '4036'])}{'*' * 10}{rng.randint(1000,9999)}"
                tx_type = rng.choice(["网银转账", "同城票据", "ATM 存款", "柜面转账", "跨行汇款"])
                if side == "贷":
                    balance += amt
                else:
                    balance -= amt
                rows.append([
                    date_str, time_str, tx_type, amt, side,
                    counterpart, opp_acct, memo, round(balance, 2),
                ])
            # 按日期 + 时间排序
            rows.sort(key=lambda r: (r[0], r[1]))

            if use_xls:
                _write_xls(fpath, f"{mcode}流水", header, rows)
            else:
                _write_xlsx(fpath, {f"{mcode}流水": [header] + rows})


# --------------- 6. 补充材料 -----------------
def gen_supplement(p: dict, outdir: Path, rng: random.Random):
    # 银行授信补充问题及材料 docx
    main_bank = p["banks"][rng.randint(0, min(1, len(p["banks"]) - 1))][0]
    main_bank_short = main_bank.split("银行")[0] + "银行"
    doc = _docx_new()
    doc.add_paragraph(
        f"{main_bank_short}{rng.choice(['授信','续贷','流贷'])}补充问题及材料回复",
    ).runs[0].font.size = Pt(14)
    doc.add_paragraph(f"客户名称：{p['name_full']}")
    doc.add_paragraph(f"统一社会信用代码：{p['uscc']}")
    doc.add_paragraph(f"申请事项：{p['renew_purpose']}")

    _docx_heading(doc, "一、贵行前期调查中提出的补充问题", 2)
    q_items = [
        ("贵公司 2024 年营业收入同比变化原因？",
         f"2024 年度营业收入 {p['revenue_2024_wan']} 万元，较 2023 年 {p['revenue_2023_wan']} 万元"
         f"{ '增长' if p['revenue_2024_wan']>p['revenue_2023_wan'] else '下降' }，主要原因是{p['main_business']}"
         f"板块订单结构变化、原材料价格波动及部分客户回款节奏差异。"),
        ("主要供应商和客户集中度？",
         f"2024 年前 3 大客户合计占营业收入约 {rng.randint(28, 55)}%；前 3 大供应商占采购成本约 {rng.randint(32, 60)}%。"
         f"已与核心客户{p['main_customers'][0][0]}签订年度框架协议，账期 {rng.choice([45, 60, 90])} 天。"),
        ("公司近 12 个月最大单笔交易及其性质？",
         f"最大单笔为与 {p['main_customers'][0][0]} 签订的年度合作协议下的季度结算款，金额约 {rng.randint(120, 480)} 万元，属正常业务结算，具备合同、验收单、发票三单匹配。"),
        ("是否存在对外担保？",
         rng.choice([
             "截至目前无对外担保事项。",
             f"存在为关联公司 {p.get('main_customers', [('','')])[1][0] if len(p.get('main_customers',[]))>1 else '暂无'} 提供连带责任保证担保 1 笔，余额约 {rng.randint(200, 1200)} 万元，不超过 50% 净资产。",
         ])),
        ("贷款用途及还款来源？",
         f"本次申请用途为补充流动资金，用于日常{p['industry']}业务经营周转；"
         f"还款来源为企业经营性现金流入（月均经营性入账约 {rng.randint(400, 1800)} 万元）。"),
        ("是否存在未决诉讼或监管处罚？",
         rng.choice([
             "截至本函件签署日，公司无涉及金额超过 100 万元的未决诉讼事项，且近 3 年无监管处罚记录。",
             f"近 3 年存在 1 起应收账款催收诉讼（涉案金额约 {rng.randint(50, 240)} 万元），已进入执行阶段。除此外无其他重大诉讼。",
         ])),
    ]
    for q, a in q_items:
        p1 = doc.add_paragraph()
        p1.add_run("问：").bold = True
        p1.add_run(q)
        p2 = doc.add_paragraph()
        p2.add_run("答：").bold = True
        p2.add_run(a)

    _docx_heading(doc, "二、随函提交材料清单", 2)
    for i, item in enumerate(
        [
            "最近 3 年度审计报告（含 2023、2024 年度）",
            "2025 年最新一期财务报表",
            "前 3 大客户合作协议摘要（脱敏件）",
            "专利及资质证书复印件",
            "房产租赁续签协议",
            "主要对外担保清单（如有）",
        ],
        start=1,
    ):
        doc.add_paragraph(f"{i}、{item}", style="List Number")

    doc.add_paragraph("")
    p_sign = doc.add_paragraph(f"回函单位：{p['name_full']}（加盖公章）")
    p_sign.alignment = 2
    p_sign2 = doc.add_paragraph(f"法定代表人：{p['legal_person']}（签字）")
    p_sign2.alignment = 2
    p_date = doc.add_paragraph(f"日期：2026 年 3 月 {rng.randint(15,25)} 日")
    p_date.alignment = 2

    docx_name = f"{main_bank_short[:2]}行-{p['name_short']}授信补充问题及材料{rng.choice(['1','（回复）','',''])}.docx"
    _docx_save(doc, outdir / docx_name)

    # 在建项目清单 xlsx（如行业是工程/建材）
    if p["industry"] in ("家装工程连锁", "建材贸易") or "工程" in p["main_business"]:
        rows = [["项目名称", "甲方/客户", "合同金额（万元）", "开工日期",
                 "当前进度", "已结算金额（万元）", "应收账款（万元）", "预计完工"]]
        for i in range(rng.randint(5, 10)):
            amt = rng.randint(120, 2400)
            paid = int(amt * rng.uniform(0.25, 0.85))
            rows.append([
                f"{rng.choice(['某商业综合体','某住宅精装','某酒店改造','某办公楼装修','某物流园配套'])}{i + 1} 期",
                rng.choice([c[0] for c in p["main_customers"]]),
                amt,
                f"{rng.choice([2023, 2024, 2025])}-{rng.randint(1,12):02d}-{rng.randint(1,28):02d}",
                rng.choice(["60%", "40%", "85%", "完工验收中", "收尾"]),
                paid,
                amt - paid,
                f"{rng.choice([2025, 2026])}-{rng.randint(1,12):02d}",
            ])
        _write_xlsx(
            outdir / f"{main_bank_short[:2]}行-在建项目清单.xlsx",
            {"在建项目": rows},
        )

    # DP005：集团互保链摘要 docx
    if p.get("quirks", {}).get("cross_guarantee_chain"):
        doc2 = _docx_new()
        doc2.add_paragraph("集团内部互保情况说明").runs[0].font.size = Pt(14)
        doc2.add_paragraph(f"出具方：{p['name_full']}")
        doc2.add_paragraph(f"日期：2026 年 3 月 {rng.randint(5, 24)} 日")
        _docx_heading(doc2, "一、目前集团内主要互保关系", 2)
        lines = [
            f"1、{p['name_short']}（母公司）为 星胤商管运营（某省） 提供连带责任保证 约 4500 万元",
            f"2、星胤商管运营（某省） 为 星悦产业投资有限公司 提供连带责任保证 约 3200 万元",
            f"3、星悦产业投资有限公司 反向为母公司 {p['name_short']} 提供股权质押 相关授信 约 3800 万元",
            "4、慎之工贸（某省）有限公司 与母公司间存在相互往来款 暂未转为担保",
        ]
        for l in lines:
            doc2.add_paragraph(l)
        _docx_heading(doc2, "二、说明", 2)
        doc2.add_paragraph(
            "以上互保金额均在各主体净资产 50% 以内，且不涉及对外集团外担保。"
            "集团层面已制定互保上限管理办法，由集团财务总监统筹监控。"
        )
        _docx_save(doc2, outdir / f"集团内部互保情况说明.docx")

    # DP002: 客户在手订单汇总（替代"在建项目"）
    if p["dp_id"] == "DP002":
        rows = [["订单编号", "客户类别", "商品类目", "金额（元）", "下单日期", "状态"]]
        for i in range(rng.randint(12, 20)):
            rows.append([
                f"SO-{rng.randint(2024,2025)}{rng.randint(1000,9999)}",
                rng.choice(["团购渠道", "KA 客户", "酒店采购", "电商直播", "零售顾客"]),
                rng.choice(["冰箱", "洗衣机", "空调", "电视", "厨小电", "洗碗机"]),
                rng.randint(800, 68000),
                f"2025-{rng.randint(1, 9):02d}-{rng.randint(1,28):02d}",
                rng.choice(["已交付", "配送中", "待发货", "已结算"]),
            ])
        _write_xlsx(
            outdir / f"{main_bank_short[:2]}行-大额订单汇总.xlsx",
            {"大额订单": rows},
        )


# ============================================================================
# 主 orchestrator
# ============================================================================
def build_for_profile(p: dict, root: Path):
    dp_dir = root / f"{p['dp_id']}_{p['name_short']}"
    # 清空旧目录（如已存在）
    if dp_dir.exists():
        for item in sorted(dp_dir.rglob("*"), reverse=True):
            if item.is_file():
                item.unlink()
            elif item.is_dir():
                item.rmdir()
        dp_dir.rmdir()
    dp_dir.mkdir(parents=True, exist_ok=True)

    rng = _seed_for(p["dp_id"])

    gen_qualifications(p, dp_dir, rng)
    gen_venue(p, dp_dir, rng)
    gen_financial(p, dp_dir, rng)
    gen_tax(p, dp_dir, rng)
    gen_bank_flow(p, dp_dir, rng)
    gen_supplement(p, dp_dir, rng)

    # 统计
    files = [f for f in dp_dir.rglob("*") if f.is_file()]
    return dp_dir, len(files)


def main():
    repo_root = Path(__file__).resolve().parents[4]  # data/mock/deep-pillar/_gen -> repo root
    deep_pillar_root = repo_root / "data" / "mock" / "deep-pillar"
    print(f"Building deep-pillar at: {deep_pillar_root}")
    summary = []
    for p in all_profiles():
        dp_dir, n = build_for_profile(p, deep_pillar_root)
        summary.append((p["dp_id"], dp_dir.name, n))
        print(f"  {p['dp_id']} {p['name_short']}: {n} files -> {dp_dir.relative_to(repo_root)}")
    total = sum(n for _, _, n in summary)
    print(f"\nTotal: {total} files across {len(summary)} deep-pillar companies")


if __name__ == "__main__":
    main()
