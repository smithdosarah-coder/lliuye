# -*- coding: utf-8 -*-
"""Agent1 内部 KB 生成器 · v2 Task C

产出：
- data/mock/channel-kb/historical-clients/  12 家简要画像（md 和 docx 混合）
- data/mock/channel-kb/marketing-preferences/  4 份营销倾向 docx
- data/mock/channel-kb/product-catalog/  1 份产品目录 xlsx

环境边界（§3.5 #5）：
- 本目录**不**含外部候选企业池
- 只是银行已知事实：已成交画像 + 自家营销偏好 + 自家产品

零答案字段：
- 不写 match_score / recommend_weight / look_alike_tags 等答案

运行：
    py -X utf8 data/mock/channel-kb/_gen/generate_kb.py
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

REPO_ROOT = Path(__file__).resolve().parents[4]
KB_ROOT = REPO_ROOT / "data" / "mock" / "channel-kb"

# 12 家历史已成交客户（脱敏 · 与 deep-pillar DP001-005 企业名不重复）
HISTORICAL = [
    {
        "name": "晟洁精密五金（昆山）有限公司",
        "industry": "制造业-精密五金",
        "region": "江苏省昆山市",
        "scale": "小型（营收约 4800 万/员工 52 人）",
        "products_signed": "流贷 500 万 · 12 个月期",
        "credit_line_wan": 500,
        "signed_at": "2023-08",
        "business_highlights": [
            "苏州某头部消费电子金属结构件二级供应商",
            "获国家高新+市级专精特新",
            "应收账期约 60 天，核心客户结算稳定",
        ],
        "risk_comments": [
            "毛利率处行业中位，略受铝价波动影响",
            "单一大客户（占营收 42%）依赖度偏高，已做合同重组缓释",
        ],
        "rm_notes": "由杭州分行跨区发起 · 与集团股权关系核查通过 · 2024 年续贷正常",
    },
    {
        "name": "泓盈机械轴承（宁波）有限公司",
        "industry": "制造业-轴承",
        "region": "浙江省宁波市北仑区",
        "scale": "小型（营收约 5200 万/员工 48 人）",
        "products_signed": "流贷 600 万 + 票据池敞口",
        "credit_line_wan": 600,
        "signed_at": "2023-11",
        "business_highlights": [
            "主营汽车底盘轴承 · 主机厂三级供应商",
            "进入 IATF16949 体系",
            "自建试验检测中心",
        ],
        "risk_comments": [
            "新能源转型周期挑战，传统车客户占比仍 68%",
            "原材料轴承钢价格联动风险",
        ],
        "rm_notes": "宁波分行存量 · 2025Q1 进入续贷调查",
    },
    {
        "name": "锐迅电子（东莞）有限公司",
        "industry": "制造业-声学模组",
        "region": "广东省东莞市长安镇",
        "scale": "小型（营收约 5600 万/员工 64 人）",
        "products_signed": "供应链金融 800 万",
        "credit_line_wan": 800,
        "signed_at": "2024-05",
        "business_highlights": [
            "TWS 耳机声学模组为国内 3 家头部品牌二级供应",
            "发明专利 18 项",
            "珠三角隐形冠军名录在册",
        ],
        "risk_comments": [
            "行业价格竞争激烈，毛利率 22% 为中等偏下",
            "核心客户账期 90 天 · 供应链金融方案已对接",
        ],
        "rm_notes": "深圳分行推介 · 本行首家声学类供应链金融客户",
    },
    {
        "name": "鸿新食品工业（合肥）有限公司",
        "industry": "制造业-食品加工",
        "region": "安徽省合肥市蜀山区",
        "scale": "小型（营收约 2800 万/员工 42 人）",
        "products_signed": "经营性贷款 350 万",
        "credit_line_wan": 350,
        "signed_at": "2024-02",
        "business_highlights": [
            "冷冻烘焙半成品 · 华东便利店 + 连锁咖啡 B 端",
            "省级粮油加工示范企业",
            "通过 HACCP + ISO22000 双认证",
        ],
        "risk_comments": [
            "食品安全抽检曾有 1 次整改（已完成）",
            "原料面粉价格波动传导压力",
        ],
        "rm_notes": "合肥分行小微客群 · 纳入本行食品工业组合管理",
    },
    {
        "name": "启星软件技术（杭州）有限公司",
        "industry": "科技-行业 SaaS",
        "region": "浙江省杭州市余杭区",
        "scale": "小型（ARR 约 6500 万/员工 78 人）",
        "products_signed": "知识产权质押贷款 600 万",
        "credit_line_wan": 600,
        "signed_at": "2024-08",
        "business_highlights": [
            "制造业 MES/APS SaaS · 付费客户 380+",
            "C 轮融资 2 亿已到账",
            "NRR 110%（行业健康水平）",
        ],
        "risk_comments": [
            "商业模式 SaaS 订阅相对稳定",
            "团队变动需关注",
            "轻资产 · 本次以知识产权+应收做增信",
        ],
        "rm_notes": "杭州分行科创金融中心主推 · 本行科技贷标杆案例",
    },
    {
        "name": "雅宁家纺制造（南通）股份有限公司",
        "industry": "制造业-家纺",
        "region": "江苏省南通市海门区",
        "scale": "中型（营收约 4.2 亿/员工 328 人）",
        "products_signed": "综合授信 3000 万（流贷 + 国内证）",
        "credit_line_wan": 3000,
        "signed_at": "2022-12",
        "business_highlights": [
            "四件套/被芯出口 · 北美电商渠道 KA 供应",
            "自动化仓配中心升级完毕",
            "B2C 自有品牌市占排行前 20",
        ],
        "risk_comments": [
            "电商渠道依赖度 68%",
            "库存周转天数 118 天偏高",
            "近两年关注国内代工客户占比提升",
        ],
        "rm_notes": "南通分行对公大户 · 已签订年度框架贸融协议",
    },
    {
        "name": "博瀚物流运输（嘉兴）有限公司",
        "industry": "服务业-物流",
        "region": "浙江省嘉兴市南湖区",
        "scale": "小型（营收约 8600 万/员工 126 人）",
        "products_signed": "车辆融资租赁 1200 万",
        "credit_line_wan": 1200,
        "signed_at": "2023-06",
        "business_highlights": [
            "长三角冷链干线 + 城市配送一体化",
            "自有冷藏车 120 台",
            "省冷链标杆企业",
        ],
        "risk_comments": [
            "运价周期下行中",
            "客户主要为生鲜电商 + 连锁餐饮 · 账期 30-60 天",
        ],
        "rm_notes": "嘉兴支行存量 · 融资租赁 + 保险一体化服务",
    },
    {
        "name": "融创建材连锁（成都）有限公司",
        "industry": "零售商贸-建材",
        "region": "四川省成都市武侯区",
        "scale": "小型（营收约 9200 万/员工 58 人）",
        "products_signed": "经营性贷款 500 万 + 承兑 1000 万",
        "credit_line_wan": 500,
        "signed_at": "2023-04",
        "business_highlights": [
            "西南建材连锁 · 门店 18 家 + 加盟 36 家",
            "与 3 家主流涂料品牌签区域代理",
            "数字化订货系统上线",
        ],
        "risk_comments": [
            "地产下游传导影响较大",
            "应收账款账期 90-120 天",
            "2024 年已收紧新增地产项目客户",
        ],
        "rm_notes": "成都分行关注类观察中 · 2025Q2 重审",
    },
    {
        "name": "禾元职业培训（西安）学校有限公司",
        "industry": "服务业-教育",
        "region": "陕西省西安市高新区",
        "scale": "小型（营收约 3200 万/员工 48 人）",
        "products_signed": "税易贷 200 万",
        "credit_line_wan": 200,
        "signed_at": "2024-03",
        "business_highlights": [
            "IT/职业资格类培训 · 年入学 6000 人次",
            "省人社厅补贴目录在册",
            "与 8 家高校合作就业基地",
        ],
        "risk_comments": [
            "预付款规模相对可控",
            "政策敏感度高（需监控补贴政策变化）",
        ],
        "rm_notes": "西安分行税易贷小微 · 纳税 A 级企业",
    },
    {
        "name": "海通机电贸易（青岛）有限公司",
        "industry": "跨境外贸-机电出口",
        "region": "山东省青岛市市南区",
        "scale": "中型（营收约 1.8 亿/员工 76 人）",
        "products_signed": "国内信用证 + 福费廷 2000 万",
        "credit_line_wan": 2000,
        "signed_at": "2023-09",
        "business_highlights": [
            "机电产品综合商社 · 主力市场东南亚 + 南美",
            "海关 AEO 高级认证",
            "年度出口额 2800 万美元",
        ],
        "risk_comments": [
            "汇率波动敏感 · 已做期权锁汇",
            "目标市场关税政策风险",
            "应收账款以 L/C 为主 · 风险较低",
        ],
        "rm_notes": "青岛分行国际业务主推 · 本行出口贷标杆",
    },
    {
        "name": "颂安医疗器械（苏州）有限公司",
        "industry": "制造业-医疗影像零部件",
        "region": "江苏省苏州市吴中区",
        "scale": "小型（营收约 3800 万/员工 52 人）",
        "products_signed": "科技贷 400 万",
        "credit_line_wan": 400,
        "signed_at": "2024-06",
        "business_highlights": [
            "DR/CT 影像设备结构件 · 进入 3 家头部客户合格供方",
            "ISO13485 认证",
            "市级专精特新",
        ],
        "risk_comments": [
            "医疗器械行业监管严格",
            "出口欧洲样件订单需关注 CE 进度",
        ],
        "rm_notes": "苏州分行科创金融 · 2025Q1 营收同比 +28%",
    },
    {
        "name": "兴旺生鲜合作联社（寿光）",
        "industry": "农业-蔬菜合作社",
        "region": "山东省潍坊市寿光市",
        "scale": "中型（营收约 1.2 亿/联合农户 680 户）",
        "products_signed": "助农贷 1500 万（联保+保证）",
        "credit_line_wan": 1500,
        "signed_at": "2022-10",
        "business_highlights": [
            "设施蔬菜合作联社 · 国家级示范",
            "覆盖华北商超直采 + 长三角批发",
            "建成冷库 1.5 万立方米",
        ],
        "risk_comments": [
            "天气 + 疫情因子影响产量波动",
            "联保机制已覆盖 90% 额度",
        ],
        "rm_notes": "潍坊分行涉农金融 · 与当地农担公司合作",
    },
]


def gen_historical():
    out = KB_ROOT / "historical-clients"
    out.mkdir(parents=True, exist_ok=True)

    # md 格式（前 7 家）
    for i, c in enumerate(HISTORICAL[:7]):
        body = f"""# {c['name']} 客户画像（脱敏档案）

> 本档由对应分行客户经理维护 · 用于本行获客推荐与 look-alike 匹配的种子样本。
> 本档记录**本行已知**事实，不含外部候选企业。

## 基本信息

- **行业**：{c['industry']}
- **区域**：{c['region']}
- **规模**：{c['scale']}

## 成交情况

- **首次成交时间**：{c['signed_at']}
- **成交产品**：{c['products_signed']}
- **授信额度（万元）**：{c['credit_line_wan']}

## 业务亮点

"""
        for h in c["business_highlights"]:
            body += f"- {h}\n"
        body += "\n## 风险观察\n\n"
        for r in c["risk_comments"]:
            body += f"- {r}\n"
        body += f"\n## 客户经理备注\n\n{c['rm_notes']}\n"
        body += "\n---\n*档案版本：v2 · 维护主体：对应分行 · 档案密级：内部*\n"
        (out / f"{c['name']}.md").write_text(body, encoding="utf-8")

    # docx 格式（后 5 家）
    for c in HISTORICAL[7:]:
        doc = Document()
        st = doc.styles["Normal"]
        st.font.name = "宋体"
        st.font.size = Pt(11)

        t = doc.add_paragraph()
        r = t.add_run(f"{c['name']} 客户画像（脱敏档案）")
        r.bold = True
        r.font.size = Pt(14)

        p = doc.add_paragraph()
        p.add_run("本档由对应分行客户经理维护 · 用于本行获客推荐与 look-alike 匹配的种子样本。").italic = True

        doc.add_heading("基本信息", level=2)
        for k, v in [("行业", c["industry"]), ("区域", c["region"]), ("规模", c["scale"])]:
            pp = doc.add_paragraph()
            pp.add_run(f"{k}：").bold = True
            pp.add_run(v)

        doc.add_heading("成交情况", level=2)
        for k, v in [
            ("首次成交时间", c["signed_at"]),
            ("成交产品", c["products_signed"]),
            ("授信额度（万元）", str(c["credit_line_wan"])),
        ]:
            pp = doc.add_paragraph()
            pp.add_run(f"{k}：").bold = True
            pp.add_run(v)

        doc.add_heading("业务亮点", level=2)
        for h in c["business_highlights"]:
            doc.add_paragraph(h, style="List Bullet")

        doc.add_heading("风险观察", level=2)
        for r in c["risk_comments"]:
            doc.add_paragraph(r, style="List Bullet")

        doc.add_heading("客户经理备注", level=2)
        doc.add_paragraph(c["rm_notes"])

        doc.add_paragraph("")
        p_sign = doc.add_paragraph("档案版本：v2 · 维护主体：对应分行 · 档案密级：内部")
        p_sign.alignment = 2
        doc.save(str(out / f"{c['name']}.docx"))


# -------- marketing-preferences --------
def gen_preferences():
    out = KB_ROOT / "marketing-preferences"
    out.mkdir(parents=True, exist_ok=True)

    # 2026-Q1 重点拓展
    doc = Document()
    doc.add_paragraph("2026 年第一季度重点拓展客群指引").runs[0].font.size = Pt(14)
    p = doc.add_paragraph("文号：公司部 [2026] 第 03 号 · 发文日期：2025-12-18 · 密级：内部")
    p.runs[0].italic = True

    doc.add_heading("一、背景与目标", level=2)
    doc.add_paragraph(
        "结合 2026 年全行小微及对公授信组合战略，本季度将重点聚焦"
        "**精密制造、专精特新、科技型创新企业**客群，"
        "服务实体经济转型升级，配合“科创金融”主题年考核要求。"
    )
    doc.add_paragraph(
        "考核目标：本季度新增对公授信客户 280 户中，"
        "上述客群占比不低于 45%；新增授信总额不低于 18 亿元。"
    )

    doc.add_heading("二、重点拓展画像", level=2)
    for head, items in [
        ("2.1 精密制造类",
         ["年营收区间：5000 万 — 3 亿元",
          "员工数：30 人以上",
          "拥有发明专利 5 项及以上",
          "国家/省级高新技术企业资质",
          "近 3 年毛利率稳定在 18% 以上",
          "主要客户为上市公司或央国企合格供应商"]),
        ("2.2 专精特新类",
         ["获得国家/省级专精特新小巨人或备案企业",
          "年营收区间：3000 万 — 5 亿元",
          "研发投入占营收比例 ≥ 3%",
          "本行优先匹配专项授信 + 科创券业务"]),
        ("2.3 科技型创新企业",
         ["最近 1 年内完成 A 轮及以上股权融资",
          "ARR（年度经常性收入）≥ 3000 万元",
          "NRR（年度净留存率）≥ 100%",
          "优先行业：企业级 SaaS / AI+产业应用 / 工业软件",
          "产品可作为担保资产（通过知识产权质押）"]),
    ]:
        doc.add_heading(head, level=3)
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    doc.add_heading("三、配套产品", level=2)
    for line in [
        "针对精密制造：优先匹配「流动资金贷款」+「应收账款池融资」组合",
        "针对专精特新：优先匹配「专精特新贷」+「研发费用池融资」组合",
        "针对科技型：优先匹配「科技贷」+「知识产权质押」+「投贷联动」组合",
    ]:
        doc.add_paragraph(line, style="List Number")

    doc.add_heading("四、落地要求", level=2)
    doc.add_paragraph(
        "1、各分行公司部月度报送重点客户拓展进展；\n"
        "2、总行客群管理部每月评估进展并通报；\n"
        "3、季度末前 5 位分行获“重点拓展优胜”奖励。"
    )
    doc.save(str(out / "2026-Q1-重点拓展.docx"))

    # 2026-Q2 区域重点
    doc2 = Document()
    doc2.add_paragraph("2026 年第二季度区域拓展重点指引").runs[0].font.size = Pt(14)
    p2 = doc2.add_paragraph("文号：公司部 [2026] 第 14 号 · 发文日期：2026-03-20 · 密级：内部")
    p2.runs[0].italic = True

    doc2.add_heading("一、季度区域聚焦", level=2)
    doc2.add_paragraph(
        "2026 年 Q2 区域聚焦两大核心经济带：**长三角制造业走廊 + 珠三角智造带**。"
        "配合国家长三角一体化战略与大湾区制造业升级战略，"
        "强化本行区域客户组合密度。"
    )

    doc2.add_heading("二、长三角重点", level=2)
    for line in [
        "江苏苏州（精密制造、集成电路配套）",
        "浙江宁波（汽车零部件、家电配套）",
        "浙江杭州（科技 SaaS、电商服务）",
        "江苏无锡（半导体封测、光伏配套）",
        "上海周边（高端装备、生物医药）",
    ]:
        doc2.add_paragraph(line, style="List Bullet")

    doc2.add_heading("三、珠三角重点", level=2)
    for line in [
        "广东深圳（电子信息、AI 硬件）",
        "广东东莞（声学模组、智能终端）",
        "广东佛山（家具、家电、五金）",
        "广东中山（小家电、灯饰）",
    ]:
        doc2.add_paragraph(line, style="List Bullet")

    doc2.add_heading("四、非重点区域指引", level=2)
    doc2.add_paragraph(
        "本季度在长三角 + 珠三角外的客户拓展原则上需与目标画像严格匹配后再"
        "进入授信流程；对跨区域客户集团企业，优先由区域主办行牵头。"
    )
    doc2.save(str(out / "2026-Q2-区域重点.docx"))

    # 避开清单
    doc3 = Document()
    doc3.add_paragraph("2026 年度客户拓展避开清单（行业+客群维度）").runs[0].font.size = Pt(14)
    p3 = doc3.add_paragraph("文号：风险部 [2026] 第 02 号 · 发文日期：2025-12-28 · 密级：内部")
    p3.runs[0].italic = True

    doc3.add_heading("一、行业避开清单", level=2)
    for line in [
        "商业地产开发与运营（含商业综合体自持+租赁主体）",
        "强周期性基建（政府投资项目依赖度 > 70%）",
        "传统燃煤火电发电（碳中和受限）",
        "煤炭开采（非合规矿井）",
        "民间高利贷相关行业",
        "非持牌网络小贷平台",
        "P2P 残余主体",
    ]:
        doc3.add_paragraph(line, style="List Bullet")

    doc3.add_heading("二、客群避开清单", level=2)
    for line in [
        "年营业收入 < 1000 万元的纯小微（不具备对公授信资质）",
        "成立 < 18 个月的创业企业（如无股权融资或科创属性）",
        "法人涉及失信被执行记录",
        "最近 24 个月内被监管机构行政处罚金额 > 50 万元的企业",
        "5 户以下产销集中度 > 80%（单客户过度集中）的企业",
    ]:
        doc3.add_paragraph(line, style="List Bullet")

    doc3.add_heading("三、政策敏感提示", level=2)
    doc3.add_paragraph(
        "对于上述清单以外但处于监管政策过渡期的行业（如教育培训转型期、"
        "医美相关服务、短视频 MCN 等），各分行应与风险部前置沟通、一事一议。"
    )
    doc3.save(str(out / "避开清单.docx"))

    # 行业组合建议
    doc4 = Document()
    doc4.add_paragraph("2026 年对公授信行业组合建议（内部参考）").runs[0].font.size = Pt(14)
    p4 = doc4.add_paragraph("文号：组合管理 [2026] 第 01 号 · 发文日期：2026-01-10 · 密级：内部")
    p4.runs[0].italic = True

    doc4.add_heading("一、年度授信组合结构目标", level=2)
    for line in [
        "制造业合计：占比 40%（较 2025 年 +3pct）",
        "消费服务业：占比 20%",
        "科技信息：占比 18%",
        "涉农涉小：占比 12%",
        "基础设施（非高风险子领域）：占比 8%",
        "其他：2%",
    ]:
        doc4.add_paragraph(line, style="List Bullet")

    doc4.add_heading("二、行业配比边界", level=2)
    doc4.add_paragraph(
        "任一细分行业单季度新增授信占比不得超过本行对公新增授信的 12%；"
        "超过需总行组合管理委员会审批。"
    )

    doc4.add_heading("三、跨行业组合提示", level=2)
    doc4.add_paragraph(
        "对上下游集中度超过 35% 的客户，应考虑行业传导风险，"
        "在尽调阶段强化下游客户稳定性评估（如地产下游建材客户需特别关注地产违约传导）。"
    )
    doc4.save(str(out / "2026年度行业组合建议.docx"))


# -------- product-catalog --------
def gen_product_catalog():
    out = KB_ROOT / "product-catalog"
    out.mkdir(parents=True, exist_ok=True)
    path = out / "对公信贷产品目录-2026版.xlsx"

    wb = Workbook()
    wb.remove(wb.active)

    ws = wb.create_sheet("产品目录（总览）")
    header = ["产品名", "定位", "适用客群", "额度区间（万元）", "期限", "定价（参考）", "核心担保方式", "风控要点"]
    rows = [
        ["经营性流动资金贷款", "短期周转", "对公/小微 · 经营稳定 · 现金流覆盖", "50-5000", "12 个月内", "LPR+30-180BP", "保证 · 抵押 · 信用（视评级）", "对账单匹配 + 发票真实性核验"],
        ["专精特新贷", "助力国家专项客群", "专精特新/小巨人资质企业", "100-3000", "12-36 个月", "LPR+20-90BP", "专项担保基金 + 股权质押 · 信用", "资质在证 + 研发费用真实性 + 年营收 5000 万 — 5 亿"],
        ["科技贷", "支持创新型科技企业", "高新技术企业 / 具备 A 轮及以上融资", "100-5000", "12-36 个月", "LPR+50-200BP", "知识产权质押 · 投贷联动 · 担保基金", "NRR/ARR 稳定 · 团队稳定 · 融资到账流水可验"],
        ["国内信用证", "对公贸易结算", "对公贸易客户 · 具备国内证结算习惯", "100-30000", "开证 6 个月内", "开证费 + 手续费", "保证金 + 货押 · 授信", "基础交易真实性 + 货物权属"],
        ["福费廷", "应收票据买入", "大中型出口/对公企业", "100-20000", "与 L/C 票期匹配", "按天贴现", "L/C · 票据", "开证行资信 + 远期票据真实性"],
        ["应收账款池融资", "应收账款质押滚动", "对公 · 下游分散 · 应收账期可控", "300-10000", "12 个月", "LPR+80-200BP", "应收账款质押", "应收池 DPD 管理 · 账期集中度"],
        ["供应链金融（反向保理）", "核心企业上游供应商融资", "核心企业合格供应商", "50-3000 / 户", "180 天内", "年化 4.5%-7%", "应收账款+核心企业确权", "核心企业确权 + 应付真实性"],
        ["税易贷", "纳税数据小微信用贷", "连续纳税 ≥ 24 个月 · 纳税等级 A/B", "10-300", "12 个月", "年化 4.2%-6.8%", "纯信用", "纳税数据波动监控 + 法人征信"],
        ["助农贷", "涉农小微联保/保证", "农业合作社 / 涉农小微", "20-1500", "12 个月", "年化 4%-5.5%", "联保 · 农担公司保证", "农业周期 · 自然灾害预警 · 联保圈健康度"],
        ["知识产权质押贷款", "科创企业 IP 变现", "发明专利/商标等权属清晰", "50-2000", "12-36 个月", "LPR+100-250BP", "知识产权质押", "评估机构独立性 + 处置退出路径"],
    ]
    ws.append(header)
    for r in rows:
        ws.append(r)

    # 格式
    header_fill = PatternFill("solid", fgColor="D9E1F2")
    for col_idx, _ in enumerate(header, start=1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        ws.column_dimensions[get_column_letter(col_idx)].width = 22
    ws.freeze_panes = "A2"

    # 产品概述 sheet（说明）
    ws2 = wb.create_sheet("说明")
    intro = [
        ["项", "内容"],
        ["发布日期", "2026-01-05"],
        ["版本", "v2026.01"],
        ["维护部门", "产品部-对公产品中心"],
        ["密级", "内部-客户经理可见"],
        ["用途", "供各分行对公/小微客户经理用于产品推介、匹配、报价参考"],
        ["与营销偏好的关系", "详见 marketing-preferences/ 目录下当季营销指引"],
        ["备注", "本目录不含外部候选客户信息；客户线索来自分行自建渠道 + SearchProvider 实搜"],
    ]
    for row in intro:
        ws2.append(row)
    for col_idx in (1, 2):
        ws2.column_dimensions[get_column_letter(col_idx)].width = 32
    ws2.cell(row=1, column=1).font = Font(bold=True)
    ws2.cell(row=1, column=2).font = Font(bold=True)

    wb.save(str(path))


def main():
    print(f"Building channel-kb at: {KB_ROOT}")
    gen_historical()
    print("  historical-clients/ 12 files generated")
    gen_preferences()
    print("  marketing-preferences/ 4 files generated")
    gen_product_catalog()
    print("  product-catalog/ 1 file generated")


if __name__ == "__main__":
    main()
