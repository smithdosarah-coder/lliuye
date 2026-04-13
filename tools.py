# -*- coding: utf-8 -*-
"""
工具注册表 v7.3

v7.3 变更：
1. ★ .doc 文件二进制解析 — 彻底解决老版Word文件无法读取问题
2. ★ 增强表格截断保护 — 大型财务附注不再丢失数据
v7 变更：
1. ★ 新增 search_web — 网络搜索（DuckDuckGo，零配置）
2. ★ 新增 self_reflect — 自省工具（审视已写内容找遗漏）
3. 优化 read_files — docx 表格提取更完整
"""

import os
import re
import json
import subprocess
import tempfile
import hashlib
from datetime import datetime


TOOL_SCHEMAS = [
    {
        "type": "function",
        "function": {
            "name": "read_files",
            "description": "读取用户上传的文件。只需调用一次，结果会缓存。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_paths": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "文件路径列表",
                    }
                },
                "required": ["file_paths"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "classify_documents",
            "description": "分类文件（A背景/B财务/C附注/D其他）。只需调用一次。",
            "parameters": {
                "type": "object",
                "properties": {
                    "files_summary": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "filename": {"type": "string"},
                                "preview": {"type": "string"},
                            },
                        },
                    }
                },
                "required": ["files_summary"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "extract_financial_data",
            "description": "提取财务指标。只需调用一次。",
            "parameters": {
                "type": "object",
                "properties": {
                    "financial_text": {"type": "string"},
                },
                "required": ["financial_text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "audit_materials",
            "description": "盘点材料完整性。只需调用一次。",
            "parameters": {
                "type": "object",
                "properties": {
                    "all_materials_text": {"type": "string"},
                },
                "required": ["all_materials_text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_chapter",
            "description": "撰写单个章节（1-4）。如需写多章，推荐用 write_all_chapters。",
            "parameters": {
                "type": "object",
                "properties": {
                    "chapter_number": {"type": "integer", "enum": [1, 2, 3, 4]},
                    "basic_info": {"type": "string", "description": "客户基本信息"},
                    "materials": {"type": "string", "description": "参考材料"},
                    "previous_chapters": {"type": "string", "description": "已有章节"},
                    "special_instructions": {"type": "string"},
                },
                "required": ["chapter_number", "basic_info", "materials"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_all_chapters",
            "description": "★ 推荐！一次写完全部4章。写完后系统自动导出Word。",
            "parameters": {
                "type": "object",
                "properties": {
                    "basic_info": {"type": "string", "description": "客户基本信息"},
                    "materials": {"type": "string", "description": "全部材料"},
                    "special_instructions": {"type": "string"},
                },
                "required": ["basic_info", "materials"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "review_report",
            "description": "审查完整报告。",
            "parameters": {
                "type": "object",
                "properties": {
                    "full_report": {"type": "string"},
                },
                "required": ["full_report"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "revise_chapter",
            "description": "根据审查意见修正章节。",
            "parameters": {
                "type": "object",
                "properties": {
                    "chapter_number": {"type": "integer", "enum": [1, 2, 3, 4]},
                    "current_content": {"type": "string"},
                    "review_comments": {"type": "string"},
                    "full_report_context": {"type": "string"},
                },
                "required": ["chapter_number", "current_content", "review_comments"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "export_to_word",
            "description": "导出Word。通常不需手动调用，写完4章后系统自动导出。",
            "parameters": {
                "type": "object",
                "properties": {
                    "client_name": {"type": "string"},
                    "credit_amount": {"type": "string"},
                    "industry": {"type": "string"},
                    "pd_rating": {"type": "string"},
                    "apply_unit": {"type": "string"},
                    "credit_type": {"type": "string"},
                },
                "required": ["client_name"],
            },
        },
    },
    # ★ v7: 网络搜索工具
    {
        "type": "function",
        "function": {
            "name": "search_web",
            "description": "搜索网络获取行业动态、政策法规、企业公开信息。建议在写报告前搜索客户所在行业的最新信息。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "搜索关键词（中文），例如：'环保工程行业2024年政策' 或 '福建省环保设计院'",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最大返回结果数（默认5）",
                    },
                },
                "required": ["query"],
            },
        },
    },
    # ★ v7: 自省工具
    {
        "type": "function",
        "function": {
            "name": "self_reflect",
            "description": "★ 审视已写的章节内容，与原始材料对比，找出遗漏的数据、空泛的表述、缺失的分析。建议在 write_all_chapters 之后调用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "focus": {
                        "type": "string",
                        "description": "重点审视方向，例如：'上下游分析是否列出了具体供应商名称和金额'",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "recall_memory",
            "description": "查询历史记忆。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query_type": {
                        "type": "string",
                        "enum": ["similar_reports", "preferences", "lessons"],
                    },
                    "industry": {"type": "string"},
                    "client_name": {"type": "string"},
                },
                "required": ["query_type"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "save_lesson",
            "description": "保存经验教训。",
            "parameters": {
                "type": "object",
                "properties": {
                    "category": {
                        "type": "string",
                        "enum": ["prompt_improvement", "common_error", "user_feedback", "industry_knowledge"],
                    },
                    "content": {"type": "string"},
                },
                "required": ["category", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ask_user",
            "description": "向用户提问。仅在关键信息缺失时使用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "question": {"type": "string"},
                },
                "required": ["question"],
            },
        },
    },
]


# ---- 工具实现 ----


def _read_doc_binary(file_path: str) -> str:
    """
    ★ v7.3: 读取老版 .doc (OLE2/CFBF) 二进制文件。
    策略：
    1. 优先 antiword（如果安装了）
    2. 优先 soffice 转 docx（如果安装了 LibreOffice Writer）
    3. ★ 核心回退：UTF-16LE 二进制文本抽取 — 适用于绝大多数中文 .doc
    """
    # 方式1: antiword
    try:
        result = subprocess.run(
            ["antiword", file_path], capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # 方式2: LibreOffice soffice 转 docx
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx", "--outdir", tmpdir, file_path],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0:
                converted = [f for f in os.listdir(tmpdir) if f.endswith(".docx")]
                if converted:
                    docx_path = os.path.join(tmpdir, converted[0])
                    from docx import Document
                    doc = Document(docx_path)
                    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                    for t_idx, table in enumerate(doc.tables):
                        rows_data = []
                        for row in table.rows:
                            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                            deduped = [c for j, c in enumerate(cells) if j == 0 or c != cells[j - 1]]
                            rows_data.append("\t".join(deduped))
                        if rows_data:
                            parts.append(f"\n[表格{t_idx + 1}]")
                            parts.extend(rows_data)
                    text = "\n".join(parts)
                    if text.strip():
                        return text
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    except Exception:
        pass

    # 方式3: ★ v7.3 核心 — UTF-16LE 二进制文本抽取
    try:
        return _extract_doc_utf16le(file_path)
    except Exception as e:
        return f"[.doc文件读取失败: {e}]"


def _extract_doc_utf16le(file_path: str) -> str:
    """
    从 .doc 二进制文件中抽取中文文本。
    .doc (OLE2) 内部存储 Word Document Stream，文本以 UTF-16LE 编码。
    本函数直接解码二进制流，提取有意义的中文+数字文本块。
    """
    with open(file_path, "rb") as f:
        raw = f.read()

    # 解码为 UTF-16LE（忽略无法解码的字节）
    text = raw.decode("utf-16-le", errors="ignore")

    # ---- 清洗：只保留有意义的文本块 ----
    # 有意义 = 包含中文字符 或 包含数字+逗号的财务数据
    meaningful_lines = []
    # 按换行和特殊分隔符切割
    raw_lines = re.split(r'[\r\n\x00]+', text)

    # 中文字符正则
    _CN_RE = re.compile(r'[\u4e00-\u9fff]')
    # 财务数字正则（如 59,861,342.84）
    _NUM_RE = re.compile(r'\d{1,3}(?:,\d{3})+(?:\.\d+)?')
    # 垃圾字符比例检测（控制字符等）
    _CTRL_RE = re.compile(r'[\x00-\x08\x0e-\x1f\x7f-\x9f]')

    seen = set()
    for line in raw_lines:
        line = line.strip()
        if not line or len(line) < 2:
            continue
        # 跳过控制字符比例过高的行（>30% = 二进制垃圾）
        ctrl_ratio = len(_CTRL_RE.findall(line)) / len(line) if line else 1
        if ctrl_ratio > 0.3:
            continue
        # 清理残留控制字符
        clean = _CTRL_RE.sub('', line).strip()
        if not clean:
            continue
        # 必须包含中文 或 财务数字
        has_cn = bool(_CN_RE.search(clean))
        has_num = bool(_NUM_RE.search(clean))
        if not has_cn and not has_num:
            continue
        # 去重（精确重复行）
        sig = clean[:80]  # 用前80字符做签名
        if sig in seen:
            continue
        seen.add(sig)
        meaningful_lines.append(clean)

    result = "\n".join(meaningful_lines)

    if not result.strip():
        return "[.doc文件解析后无有效文本内容]"

    return result


def _read_pdf_with_ocr_fallback(file_path: str, max_ocr_pages: int = 50) -> str:
    """Read PDF: try pdfplumber first, fall back to OCR for scanned documents.

    OCR results are cached in <file>.ocr_cache.txt so re-runs are instant.
    """
    # Step 1: try pdfplumber (fast, for text-based PDFs)
    texts = []
    total_pages = 0
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texts.append(t)
                for table in page.extract_tables():
                    if table:
                        for row in table:
                            texts.append("\t".join(str(c) if c else "" for c in row))
    except Exception:
        pass

    text_result = "\n".join(texts)

    # If we got meaningful text (> 100 chars), return it
    if len(text_result.strip()) > 100:
        return text_result

    # Step 2: OCR fallback for scanned PDFs
    # ★ v7.11: Check cache first — avoids re-running expensive OCR every session
    # Cache is stored under ./memory_store/ocr_cache because input material folders may be read-only.
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        cache_dir = os.path.join(base_dir, "memory_store", "ocr_cache")
        os.makedirs(cache_dir, exist_ok=True)
        st = os.stat(file_path)
        key = f"{os.path.abspath(file_path)}|{st.st_size}|{int(st.st_mtime)}"
        cache_name = hashlib.md5(key.encode("utf-8")).hexdigest()
        cache_path = os.path.join(cache_dir, cache_name + ".txt")
    except Exception:
        cache_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "memory_store", "ocr_cache", (os.path.basename(file_path) + ".txt"))

    if os.path.exists(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as f:
                cached = f.read()
            if cached.strip():
                fname = os.path.basename(file_path)
                print(f"[OCR] {fname}: 使用缓存结果 ({len(cached)} chars)")
                return cached
        except Exception:
            pass  # cache read failed, fall through to re-OCR

    try:
        import fitz  # PyMuPDF
        import numpy as np
        from rapidocr_onnxruntime import RapidOCR
        from PIL import Image
        ocr_engine = RapidOCR()
        doc = fitz.open(file_path)
        total_pages = len(doc)

        fname = os.path.basename(file_path)

        # ---- Select pages to OCR (speed) ----
        # For scanned audit reports, OCR-ing every page is very slow. We first
        # do a low-DPI, header-cropped quick OCR to locate likely statement pages.
        KEYWORDS = [
            "资产负债表", "利润表", "现金流量表",
            "所有者权益变动表",
        ]

        def _quick_ocr_hit(page_index: int) -> bool:
            try:
                page = doc[page_index]
                # Low DPI + crop header area for speed
                mat_q = fitz.Matrix(96 / 72, 96 / 72)
                pix = page.get_pixmap(matrix=mat_q)
                img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                w, h = img.size
                crop_h = max(200, int(h * 0.28))
                img = img.crop((0, 0, w, crop_h))
                img_np = np.array(img)
                res, _ = ocr_engine(img_np)
                if not res:
                    return False
                head_text = "".join([ln[1] for ln in res if ln and len(ln) > 1])
                return any(k in head_text for k in KEYWORDS)
            except Exception:
                return False

        is_audit_report = ("审计报告" in fname) or ("audit" in fname.lower())
        if is_audit_report:
            # Audit reports are often scanned and long; for form filling we mainly need header/meta info.
            pages = list(range(min(total_pages, min(max_ocr_pages, 8))))
        else:
            base_pages = [0, 1] if total_pages >= 2 else [0]
            hits: list[int] = []
            for i in range(total_pages):
                if _quick_ocr_hit(i):
                    hits.append(i)
                    if i + 1 < total_pages:
                        hits.append(i + 1)  # tables often continue on next page

            pages = sorted(set(base_pages + hits))
            if not pages:
                pages = list(range(min(total_pages, max_ocr_pages)))
            if len(pages) > max_ocr_pages:
                pages = pages[:max_ocr_pages]

        pages_to_ocr = len(pages)
        print(f"[OCR] {fname}: {total_pages}页，命中{pages_to_ocr}页待识别（首次，结果将缓存）...")

        ocr_texts = []
        for idx, i in enumerate(pages):
            page = doc[i]
            # 150 DPI — faster and usually readable
            mat = fitz.Matrix(150 / 72, 150 / 72)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            img_np = np.array(img)

            result, _ = ocr_engine(img_np)
            if result:
                page_text = "\n".join([line[1] for line in result])
                ocr_texts.append(f"--- Page {i + 1} ---\n{page_text}")

            # Progress every 3 pages
            if (idx + 1) % 3 == 0:
                print(f"[OCR] {fname}: {idx + 1}/{pages_to_ocr}页完成")

        doc.close()
        if ocr_texts:
            header = f"[OCR识别, {pages_to_ocr}/{total_pages}页]"
            if pages_to_ocr < total_pages:
                header += f" (仅识别{pages_to_ocr}页，按关键词筛选)"
            result_text = header + "\n" + "\n".join(ocr_texts)

            # ★ v7.11: Save cache so next run is instant
            try:
                with open(cache_path, "w", encoding="utf-8") as f:
                    f.write(result_text)
                print(f"[OCR] {fname}: 结果已缓存到 {os.path.basename(cache_path)}")
            except Exception:
                pass  # cache write failure is non-fatal

            return result_text

        return "[PDF无法提取文字，OCR识别也无结果]"

    except ImportError:
        return "[扫描件PDF需要OCR，请安装: pip install pymupdf rapidocr-onnxruntime]"
    except Exception as e:
        return f"[PDF OCR处理失败: {e}]"


def _smart_truncate(text: str, max_chars: int = 15000) -> str:
    if len(text) <= max_chars:
        return text
    keep_start = int(max_chars * 0.7)
    keep_end = max_chars - keep_start - 50
    return (
        text[:keep_start]
        + f"\n\n... （省略约{len(text) - max_chars}字） ...\n\n"
        + text[-keep_end:]
    )


def _expand_paths(file_paths: list[str]) -> list[str]:
    """Expand directories to individual files recursively.

    Supports nested folders (e.g. '4、银行流水/中国银行/*.pdf').
    Skips hidden files, temp files (~$), and non-supported extensions.
    """
    SUPPORTED = {'.txt', '.docx', '.doc', '.pdf', '.xlsx', '.xls', '.csv'}
    expanded = []
    seen = set()
    for fp in file_paths:
        if os.path.isdir(fp):
            for root, dirs, files in os.walk(fp):
                # Skip bank statement folders (user requested to exclude "银行流水")
                dirs[:] = [
                    d for d in dirs
                    if ("银行流水" not in d and "bank" not in d.lower())
                ]
                for fname in sorted(files):
                    if fname.startswith(('~$', '.')):
                        continue
                    ext = os.path.splitext(fname)[1].lower()
                    if ext in SUPPORTED:
                        full = os.path.join(root, fname)
                        if full not in seen:
                            seen.add(full)
                            expanded.append(full)
        else:
            if fp not in seen:
                seen.add(fp)
                expanded.append(fp)
    return expanded


def tool_read_files(file_paths: list[str]) -> str:
    # ★ v7.10: Expand directories recursively
    all_files = _expand_paths(file_paths)
    results = []
    for fp in all_files:
        fname = os.path.basename(fp)
        # Include relative subfolder name for clarity
        parent = os.path.basename(os.path.dirname(fp))
        display_name = f"{parent}/{fname}" if parent else fname
        content = _read_single_file(fp)
        results.append(f"===== {display_name} ({len(content)}字) =====\n{content}")
    return "\n\n".join(results)


def _read_single_file(file_path: str) -> str:
    if not file_path or not os.path.exists(file_path):
        return f"[文件不存在: {file_path}]"
    ext = file_path.lower().rsplit(".", 1)[-1] if "." in file_path else ""
    try:
        if ext == "docx":
            from docx import Document
            doc = Document(file_path)
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
            # ★ v6+: 提取表格内容
            for t_idx, table in enumerate(doc.tables):
                rows_data = []
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    deduped = []
                    for j, c in enumerate(cells):
                        if j == 0 or c != cells[j - 1]:
                            deduped.append(c)
                    rows_data.append("\t".join(deduped))
                if rows_data:
                    parts.append(f"\n[表格{t_idx + 1}]")
                    parts.extend(rows_data)
            return "\n".join(parts)
        elif ext == "doc":
            return _read_doc_binary(file_path)
        elif ext == "pdf":
            return _read_pdf_with_ocr_fallback(file_path)
        elif ext in ("xlsx", "xls"):
            import pandas as pd
            engine = "openpyxl" if ext == "xlsx" else "xlrd"
            sheets = pd.read_excel(file_path, sheet_name=None, engine=engine)
            parts = []
            for name, df in sheets.items():
                parts.append(f"[{name}]")
                parts.append(df.to_string(index=False))
            return "\n\n".join(parts)
        else:
            for enc in ("utf-8-sig", "utf-8", "gbk", "gb2312"):
                try:
                    with open(file_path, encoding=enc) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
    except Exception as e:
        return f"[文件读取错误: {e}]"
    return ""


# ★ v7: 网络搜索实现
def tool_search_web(query: str, max_results: int = 5) -> str:
    """
    使用 DuckDuckGo 搜索（无需 API 密钥）。
    如果 duckduckgo_search 未安装，自动回退到 HTTP 请求方式。
    """
    results = []

    # 方式1: duckduckgo_search 库（推荐）
    try:
        from duckduckgo_search import DDGS
        with DDGS() as ddgs:
            for r in ddgs.text(query, region="cn-zh", max_results=max_results):
                results.append({
                    "title": r.get("title", ""),
                    "snippet": r.get("body", ""),
                    "url": r.get("href", ""),
                })
        if results:
            return _format_search_results(query, results)
    except ImportError:
        pass
    except Exception as e:
        # 库存在但搜索失败
        pass

    # 方式2: 直接用 requests 抓取 DuckDuckGo HTML（后备方案）
    try:
        import requests
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        }
        resp = requests.get(
            f"https://html.duckduckgo.com/html/?q={query}",
            headers=headers,
            timeout=10,
        )
        if resp.status_code == 200:
            # 简单解析搜索结果
            from html.parser import HTMLParser
            import html

            class DDGParser(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.results = []
                    self._in_result = False
                    self._current = {}
                    self._capture = None

                def handle_starttag(self, tag, attrs):
                    attrs_d = dict(attrs)
                    cls = attrs_d.get("class", "")
                    if "result__a" in cls:
                        self._in_result = True
                        self._current["url"] = attrs_d.get("href", "")
                        self._capture = "title"
                    elif "result__snippet" in cls:
                        self._capture = "snippet"

                def handle_data(self, data):
                    if self._capture and data.strip():
                        key = self._capture
                        self._current[key] = self._current.get(key, "") + data.strip()

                def handle_endtag(self, tag):
                    if self._capture == "snippet" and self._current:
                        self.results.append(self._current)
                        self._current = {}
                        self._capture = None
                    elif tag == "a" and self._capture == "title":
                        self._capture = None

            parser = DDGParser()
            parser.feed(resp.text)
            for r in parser.results[:max_results]:
                results.append(r)

            if results:
                return _format_search_results(query, results)
    except Exception:
        pass

    return f"[搜索失败] 未能获取到「{query}」的搜索结果。可能原因：网络不可用或搜索服务暂时不可用。\n建议：可以安装 duckduckgo-search 库（pip install duckduckgo-search）以获得更好的搜索体验。"


def _format_search_results(query: str, results: list[dict]) -> str:
    """格式化搜索结果"""
    lines = [f"搜索「{query}」共找到 {len(results)} 条结果：\n"]
    for i, r in enumerate(results, 1):
        title = r.get("title", "无标题")
        snippet = r.get("snippet", "")
        url = r.get("url", "")
        lines.append(f"[{i}] {title}")
        if snippet:
            lines.append(f"    摘要：{snippet[:200]}")
        if url:
            lines.append(f"    来源：{url}")
        lines.append("")
    return "\n".join(lines)


PURE_TOOL_FUNCTIONS = {
    "read_files": tool_read_files,
    "search_web": tool_search_web,
}


def get_tool_schemas() -> list[dict]:
    return TOOL_SCHEMAS






