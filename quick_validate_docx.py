# -*- coding: utf-8 -*-
"""Structured validation for generated .docx outputs."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass, asdict
from typing import Any

from docx import Document

W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


@dataclass
class ValidationIssue:
    code: str
    severity: str
    message: str
    context: str = ""
    table_index: int | None = None


def _get_xml(z: zipfile.ZipFile, name: str) -> str:
    try:
        with z.open(name) as f:
            return f.read().decode("utf-8", errors="replace")
    except KeyError:
        return ""


def _count_comments(docx_path: str) -> int:
    with zipfile.ZipFile(docx_path, "r") as z:
        cxml = _get_xml(z, "word/comments.xml")
    if not cxml.strip():
        return 0
    try:
        root = ET.fromstring(cxml)
    except Exception:
        return 0
    return len(root.findall(".//w:comment", W_NS))


def _iter_tables(table):
    yield table
    for row in table.rows:
        for cell in row.cells:
            for nt in cell.tables:
                yield from _iter_tables(nt)


def _collect_tables(doc: Document) -> list[list[list[str]]]:
    tables = []
    for table in doc.tables:
        for t in _iter_tables(table):
            rows = []
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                tables.append(rows)
    return tables


def _collect_text(doc: Document) -> str:
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in _collect_tables(doc):
        for row in table:
            for cell in row:
                if cell:
                    parts.append(cell)
    return "\n".join(parts)


def _find_placeholders(text: str) -> list[ValidationIssue]:
    patterns = [
        ("placeholder_xxx", r"XXX", "?????? XXX ???"),
        ("placeholder_pending", r"(?:\u5f85\u8865\u5145|\?\?\?)", "\u5b58\u5728\u5f85\u8865\u5145\u5360\u4f4d"),
        ("placeholder_underline", r"_{3,}", "????????????"),
    ]
    issues = []
    for code, pat, msg in patterns:
        m = re.search(pat, text)
        if m:
            ctx = text[max(0, m.start() - 20):m.end() + 20]
            issues.append(ValidationIssue(code=code, severity="warning", message=msg, context=ctx))
    return issues


def _find_money_decimals(text: str) -> list[ValidationIssue]:
    issues = []
    for m in re.finditer(r"(?<!\d)\d+\.\d+\s*\u4e07\u5143", text):
        ctx = text[max(0, m.start() - 20):min(len(text), m.end() + 20)]
        issues.append(ValidationIssue(
            code="money_decimal",
            severity="blocker",
            message="\u91d1\u989d\u5b57\u6bb5\u51fa\u73b0\u5c0f\u6570\uff0c\u672a\u6309\u4e07\u5143\u6574\u6570\u8f93\u51fa",
            context=ctx,
        ))
    return issues


def _compact_cells(row: list[str]) -> list[str]:
    return [re.sub(r"\s+", "", c or "") for c in row]


def _find_col_index(header: list[str], needles: list[str]) -> int | None:
    compact = _compact_cells(header)
    normalized_needles = [re.sub(r"\s+", "", x or "") for x in needles if x]

    for needle in normalized_needles:
        for i, cell in enumerate(compact):
            if cell == needle:
                return i

    for needle in normalized_needles:
        for i, cell in enumerate(compact):
            if cell and (needle in cell or cell in needle):
                return i
    return None


def _scan_2025_9_tables(tables: list[list[list[str]]]) -> list[ValidationIssue]:
    issues = []
    for ti, table in enumerate(tables):
        if not table or not table[0]:
            continue
        header = _compact_cells(table[0])
        col_20259 = None
        for ci, h in enumerate(header):
            if any(token in h for token in ["2025.9", "2025.09", "2025\u5e749"]):
                col_20259 = ci
                break
        if col_20259 is None:
            continue

        has_year_end = any(h in {"2025\u5e74\u672b(\u65b0\u589e)", "2025\u5e74\u672b\uff08\u65b0\u589e\uff09", "2025\u5e74\u672b\u65b0\u589e"} for h in header)
        if not has_year_end:
            issues.append(ValidationIssue(
                code="missing_2025_year_end_column",
                severity="blocker",
                message="\u5b58\u57282025.9\u5217\u4f46\u7f3a\u5c112025\u5e74\u672b(\u65b0\u589e)\u5217",
                table_index=ti,
                context=" | ".join(table[0]),
            ))

        numeric_samples = []
        for row in table[1:]:
            if col_20259 >= len(row):
                continue
            value = (row[col_20259] or "").strip()
            if re.search(r"\d", value):
                numeric_samples.append(value)
        if numeric_samples:
            issues.append(ValidationIssue(
                code="unexpected_2025_9_values",
                severity="blocker",
                message="2025.9????????????",
                table_index=ti,
                context="; ".join(numeric_samples[:5]),
            ))
    return issues


def _find_supply_chain_tables(tables: list[list[list[str]]]) -> list[ValidationIssue]:
    issues = []
    targets = [
        ("upstream", "主要上游供应商", ["上年采购额", "今年1-9月采购额"]),
        ("downstream", "主要下游销售客户", ["上年销售额", "今年1-9月销售额"]),
    ]
    for code, marker, amount_cols in targets:
        found = None
        for ti, table in enumerate(tables):
            header_text = "\n".join("\t".join(r) for r in table[:2])
            if marker in header_text:
                found = (ti, table)
                break
        if found is None:
            issues.append(ValidationIssue(
                code=f"{code}_table_missing",
                severity="blocker",
                message=f"未找到{marker}表格",
            ))
            continue

        ti, table = found
        header = table[0]
        amount_indices = [ci for ci in (_find_col_index(header, [c]) for c in amount_cols) if ci is not None]
        filled_rows = 0
        for row in table[1:]:
            if not row:
                continue
            first = (row[0] or "").strip()
            if not first or first.startswith("注") or "已核实" in first:
                continue
            if re.search(r"[\u4e00-\u9fffA-Za-z0-9]", first):
                filled_rows += 1
            for ci in amount_indices:
                if ci < len(row):
                    val = (row[ci] or "").strip()
                    if "%" in val:
                        issues.append(ValidationIssue(
                            code=f"{code}_amount_percent",
                            severity="blocker",
                            message=f"{marker}表金额列被填成百分比",
                            table_index=ti,
                            context=" | ".join(row),
                        ))
                        break
        if filled_rows < 5:
            issues.append(ValidationIssue(
                code=f"{code}_rows_incomplete",
                severity="blocker",
                message=f"{marker}表仅识别到{filled_rows}行有效数据，少于5行",
                table_index=ti,
                context=" | ".join(table[0]),
            ))
    return issues


def _find_r_and_d_table_issues(tables: list[list[list[str]]]) -> list[ValidationIssue]:
    issues = []
    required_headers = {
        "年份",
        "研发费用",
        "研发费用占比（研发费用/营收）",
    }
    for ti, table in enumerate(tables):
        if not table or not table[0]:
            continue
        header = table[0]
        compact_header = set(_compact_cells(header))
        if not required_headers.issubset(compact_header):
            continue
        fee_ci = _find_col_index(header, ["研发费用"])
        if fee_ci is None:
            continue
        for row in table[1:]:
            if fee_ci >= len(row):
                continue
            val = (row[fee_ci] or "").strip()
            if "%" in val:
                issues.append(ValidationIssue(
                    code="r_and_d_fee_percent",
                    severity="blocker",
                    message="研发表中研发费用金额列被填成百分比",
                    table_index=ti,
                    context=" | ".join(row),
                ))
                break
    return issues


def _find_financial_metric_misfills(tables: list[list[list[str]]]) -> list[ValidationIssue]:
    issues = []
    ratio_labels = {
        "净利润率",
        "毛利润率",
        "毛利率",
        "资产负债率",
    }
    for ti, table in enumerate(tables):
        if not table or not table[0]:
            continue
        header = _compact_cells(table[0])
        if header and header[0] == "科目":
            for row in table[1:]:
                if not row:
                    continue
                label = (row[0] or "").strip()
                values = [v.strip() for v in row[1:] if v and v.strip()]
                if not values:
                    continue
                if label in ratio_labels:
                    bad = [v for v in values if "%" not in v]
                    if bad:
                        issues.append(ValidationIssue(
                            code="financial_ratio_missing_percent",
                            severity="blocker",
                            message=f"{label}行存在非百分比值，疑似指标错位",
                            table_index=ti,
                            context=" | ".join(row),
                        ))
                        break
                if "周转天数" in label:
                    bad = []
                    for v in values:
                        m = re.fullmatch(r"-?\d+(?:\.\d+)?", v)
                        if m and float(v) > 1000:
                            bad.append(v)
                    if bad:
                        issues.append(ValidationIssue(
                            code="financial_turnover_days_outlier",
                            severity="blocker",
                            message=f"{label}行出现明显异常大值，疑似填入了余额而非天数",
                            table_index=ti,
                            context=" | ".join(row),
                        ))
                        break
        if header and header[0] == "年份" and any(h in header for h in ["总资产", "净资产", "销售收入", "主营业务收入"]):
            for row in table[1:]:
                if not row:
                    continue
                label = (row[0] or "").strip()
                if re.fullmatch(r"20\d{2}\.(?:4|9|X|x)", label):
                    values = [v.strip() for v in row[1:] if v and v.strip()]
                    if values:
                        issues.append(ValidationIssue(
                            code="financial_period_mislabel_filled",
                            severity="blocker",
                            message="行式财务表的占位期间行被填入了实际数值",
                            table_index=ti,
                            context=" | ".join(row),
                        ))
                        break
    return issues


def validate_docx(docx_path: str) -> dict[str, Any]:
    doc = Document(docx_path)
    tables = _collect_tables(doc)
    text = _collect_text(doc)
    issues: list[ValidationIssue] = []
    issues.extend(_find_placeholders(text))
    issues.extend(_find_money_decimals(text))
    issues.extend(_scan_2025_9_tables(tables))
    issues.extend(_find_supply_chain_tables(tables))
    issues.extend(_find_r_and_d_table_issues(tables))
    issues.extend(_find_financial_metric_misfills(tables))

    blockers = [asdict(i) for i in issues if i.severity == "blocker"]
    warnings = [asdict(i) for i in issues if i.severity != "blocker"]
    return {
        "docx": docx_path,
        "comments": _count_comments(docx_path),
        "table_count": len(tables),
        "blockers": blockers,
        "warnings": warnings,
    }


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", help="Path to generated .docx")
    ap.add_argument("--json", action="store_true", help="Print JSON result")
    args = ap.parse_args()

    result = validate_docx(args.docx)
    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(f"DOCX= {result['docx']}")
        print(f"COMMENTS= {result['comments']}")
        print(f"TABLES= {result['table_count']}")
        print(f"BLOCKERS= {len(result['blockers'])}")
        for item in result["blockers"]:
            print(f"[BLOCKER] {item['code']}: {item['message']} :: {item.get('context','')}")
        print(f"WARNINGS= {len(result['warnings'])}")
        for item in result["warnings"]:
            print(f"[WARNING] {item['code']}: {item['message']} :: {item.get('context','')}")
    return int(bool(result["blockers"]))


if __name__ == "__main__":
    raise SystemExit(0 if main() == 0 else 1)
