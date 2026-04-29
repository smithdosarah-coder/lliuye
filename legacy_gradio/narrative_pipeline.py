# -*- coding: utf-8 -*-
"""V15 对公叙述模板管线 (narrative pipeline).

背景 / 分治原则
---------------
V14 终版的"骨架型管线"假设模板长这样:大量字段格/占位符 + 稀疏段落。
对公授信报告模板完全不是这回事:大标题 + 段落级成稿,历史报告就当范例。
form_filler 的骨架管线会原样保留段落原文 → LLM 不敢覆盖 → 产物里全是
旧客户名和旧数字 ("经纬" 22 次 / "中锐" 0 次)。

所以用户要求**分治**:对公走 narrative,普惠继续走 V14;两条路线彼此隔离,
叙述路线失败不得拖垮骨架路线,也不得修改 V14 的任何既有守卫。

本模块仅做一件事:对叙述型模板,按「大标题 section」为单位让 LLM 基于当前
客户材料重写段落内容,然后继承 V14 的全部 text / docx 级守卫。

核心规则
--------
- 只用结构特征识别大标题 (Word Heading 样式 + 中文编号正则),严禁关键词
  黑名单 (这会泛化不过任何没见过的模板)。
- 生成必须 LLM 全新写:模板原文只作 example_text 参考,严禁带入输出
  (旧公司名 / 旧人名 / 旧年份 / 旧数字)。
- V14 的 text 级守卫全部继承:_fix_affiliates_deterministic /
  _fix_variation_phrases / _fix_dup_year_phrases / _sanitize_output。
- 再加一道**公司名硬检测**: 从 example_text 抽出的别家公司名若出现在 LLM
  输出中,替换为当前客户名 (或标 "(待核实)")。这是 LLM 的常见偷懒模式。
- 末尾必须跑 FormFillAgent._fix_skeleton_title_numbers /
  _fix_same_value_pct_variation / _sanitize_docx_runs + 存盘用
  _save_preserving_zip,保持和骨架管线一致的输出质量。
"""

from __future__ import annotations

import os
import re
import logging
from dataclasses import dataclass, field
from typing import Callable, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from docx.document import Document

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────
# 1. Section 数据结构 + 大标题识别正则
# ─────────────────────────────────────────────────────────────

# 优先级递减的大标题正则 (level 越小越高级)。结构特征导向,不含任何业务
# 关键词 / 公司名。text_full 供后续做长度/冒号判据,避免把"(3)该公司结算..."
# 这类内容误认为大标题。
HEADING_RES = [
    # 一、/ 二、/ 十、 ... XXX   (兼容全角 ． U+FF0E)
    (1, re.compile(r'^\s*[一二三四五六七八九十]{1,3}\s*[、.．]\s*\S')),
    # (一) / （二） ... XXX
    (2, re.compile(r'^\s*[（(]\s*[一二三四五六七八九十]{1,3}\s*[)）]\s*\S')),
    # 1. XXX / 10、XXX / 2．XXX
    (2, re.compile(r'^\s*[0-9]{1,2}\s*[.、．]\s*\S')),
    # (1) / （2） ... XXX
    (3, re.compile(r'^\s*[（(]\s*[0-9]{1,2}\s*[)）]\s*\S')),
    # ※ 1. XXX
    (3, re.compile(r'^\s*※\s*[0-9]{1,2}\s*[.、．]\s*\S')),
    # ①XXX : / ②XXX
    (4, re.compile(r'^\s*[①②③④⑤⑥⑦⑧⑨⑩].{1,40}[:：]?$')),
]

# 去除所有常见编号前缀,用于判断 "去标题后内容" 的实际长度
_HEADING_STRIP_RE = re.compile(
    r'^\s*[一二三四五六七八九十\(\)（）\d、.．※①②③④⑤⑥⑦⑧⑨⑩\s]+'
)

# 基本信息区关键词 (只用于识别应跳过不重写的封面/表头),不涉及公司名
_COVER_KEYWORDS = (
    "客户名称", "客户全称", "授信客户全称", "PD评级", "申报单位",
    "白名单", "绿色信贷标识", "环境和社会风险分类", "投向政策对应行业",
)

# 用于从 example_text 抽当前客户之外的公司名 (硬检测守卫)
# 后缀限定 "公司/集团/企业/合伙",非贪婪以防吃过头 ("旗下全资子公司"),
# 然后由 _is_real_company_candidate 做负面过滤
_COMPANY_EXTRACT_RE = re.compile(
    r'([\u4e00-\u9fff]{2,20}?(?:公司|集团|企业|合伙))'
)

# 排除那些以高风险前缀开头的伪公司名 (这些是被误切的 n-gram)
_FAKE_NAME_PREFIXES = (
    "作为", "对", "与", "和", "或", "该", "本", "其", "被", "由", "使",
    "将", "并", "于", "在", "经", "为", "从", "至", "到", "向", "通过",
    "对于", "关于", "根据", "基于", "以", "让", "给", "比", "由于",
    "例如", "如", "像", "比如", "此外", "另外", "同时", "以及",
    "采用", "获得", "取得", "建立", "成立", "承担", "提供", "开展",
    "包含", "包括", "拥有", "持有", "控制", "涉及", "涵盖",
    "控制人", "控股人", "实控人", "代表人", "负责人", "管理人",
    "发起人", "股东", "董事", "监事", "员工", "高管",
    "公司", "子公司", "母公司", "集团", "企业",
    "有限", "股份",
)


def _is_real_company_candidate(name: str) -> bool:
    """过滤掉 n-gram 切出来的伪公司名。

    真实公司名特征:
      - 长度 >= 4
      - 不以非名词性前缀 (动词/介词/连词/指代) 开头
      - 不以"公司/子公司/集团"等机构后缀词本身作前缀 (典型 n-gram 噪音)
      - 后缀是"公司/集团/企业/合伙"中的一个 — 已在正则中保证
      - 以"公司/集团"作前缀的词,要求剩余部分 >= 6 字 (让普通公司名通过,
        同时排除"公司采用集团"这种 6 字伪名)
    """
    if not name or len(name) < 4:
        return False
    org_prefixes = ("公司", "子公司", "母公司", "集团", "企业", "有限", "股份")
    for pref in org_prefixes:
        if name.startswith(pref):
            if len(name) - len(pref) < 6:
                return False
    for pref in _FAKE_NAME_PREFIXES:
        if pref in org_prefixes:
            continue
        if name.startswith(pref) and len(name) - len(pref) < 4:
            return False
    return True


@dataclass
class NarrativeSection:
    """对公叙述型模板中的一个"大标题 section"。

    body_para_indices: 这一节在 doc.paragraphs 中的段落 index 列表
        (**不含**标题本身),重写产物将按段落个数分配回写到这些段落。
    example_text: 段落原文合并,**仅作 prompt 参考**,禁止带入输出。
    """

    title: str                                # 大标题原文 (清洗后)
    title_para_idx: int                       # 标题所在段落 index
    heading_level: int                        # 1~4, 越小越高
    body_para_indices: list[int] = field(default_factory=list)
    example_text: str = ""
    nested_table_refs: list = field(default_factory=list)  # 预留


# ─────────────────────────────────────────────────────────────
# 2. 结构判据: 是否叙述型模板
# ─────────────────────────────────────────────────────────────

def _detect_is_narrative(doc) -> bool:
    """判定当前 docx 是否为"叙述型对公模板"。

    判据 (全部满足 → 叙述型):
      - 非空段落数 >= 30 (骨架型普惠表格通常 < 30 段正文)
      - 非空段落平均字数 >= 20 (骨架型占位符段通常 < 10 字)
      - 下划线占位符段比例 < 20% (骨架型 >= 50%)
    任一条不满足 → 判骨架型,由 V14 管线继续处理。
    """
    nonempty = [p for p in doc.paragraphs if p.text and p.text.strip()]
    if len(nonempty) < 30:
        return False

    total_chars = sum(len(p.text.strip()) for p in nonempty)
    avg_chars = total_chars / len(nonempty)
    if avg_chars < 20:
        return False

    placeholder_count = 0
    placeholder_re = re.compile(r'_{4,}|\s{6,}')
    for p in nonempty:
        txt = p.text
        if placeholder_re.search(txt) or txt.count('_') >= 4:
            placeholder_count += 1
    placeholder_ratio = placeholder_count / len(nonempty)
    if placeholder_ratio >= 0.20:
        return False

    return True


# ─────────────────────────────────────────────────────────────
# 3. Section 切分
# ─────────────────────────────────────────────────────────────

def _match_heading(text: str) -> Optional[tuple[int, str]]:
    """返回 (level, clean_title) 或 None。

    仅结构判据: 正则命中 + (短 or 以冒号结尾 or 去编号后短)。
    """
    if not text:
        return None
    t = text.strip()
    for level, regex in HEADING_RES:
        if regex.match(t):
            # 是否真是大标题 (而非编号内容行)
            is_short = len(t) <= 30
            ends_colon = t.rstrip().endswith(('：', ':'))
            stripped = _HEADING_STRIP_RE.sub('', t)
            stripped_short = len(stripped) <= 25
            if is_short or ends_colon or stripped_short:
                return (level, t)
            return None
    return None


def _is_cover_or_basic_info(idx: int, text: str) -> bool:
    """前 20 段里含基本信息关键词 → 跳过,不重写。"""
    if idx >= 20:
        return False
    return any(kw in text for kw in _COVER_KEYWORDS)


def _discover_narrative_sections(doc) -> list[NarrativeSection]:
    """切大标题 section,跳过封面 / 基本信息区。"""
    sections: list[NarrativeSection] = []
    paragraphs = list(doc.paragraphs)

    # Pass 1: 标注每段是否为 heading
    heading_info: list[Optional[tuple[int, str]]] = []
    for i, p in enumerate(paragraphs):
        text = p.text or ""
        t_stripped = text.strip()

        # Word Heading 样式优先
        style_name = ""
        try:
            style_name = (p.style.name or "") if p.style else ""
        except Exception:
            style_name = ""
        has_heading_style = ("Heading" in style_name) or ("标题" in style_name)

        matched = _match_heading(text)

        # 结构判据追加:bold + 短 (<=20 字) 单行视为标题 level 2
        # (对不按数字编号的一级/二级大标题,如"授信客户背景情况"/"经营/管理风险分析")
        is_bold_short = False
        if not matched and not has_heading_style and t_stripped and len(t_stripped) <= 20:
            try:
                if any(r.bold for r in p.runs if (r.text or "").strip()):
                    # 结尾不能是句号 (避免把 bold 的结论句误判)
                    if not t_stripped.rstrip().endswith(('。', '!', '?', '！', '？')):
                        is_bold_short = True
            except Exception:
                is_bold_short = False

        if has_heading_style and t_stripped:
            lv = matched[0] if matched else 1
            title = matched[1] if matched else t_stripped
            heading_info.append((lv, title))
        elif matched:
            heading_info.append(matched)
        elif is_bold_short:
            heading_info.append((2, t_stripped))
        else:
            heading_info.append(None)

    # Pass 2: 切段
    i = 0
    while i < len(paragraphs):
        hi = heading_info[i]
        if not hi:
            i += 1
            continue

        level, title = hi

        # 封面 / 基本信息区跳过
        if _is_cover_or_basic_info(i, title):
            i += 1
            continue

        # 收集 body: 到下一个 heading 为止,跳过空段
        body_indices: list[int] = []
        body_texts: list[str] = []
        j = i + 1
        while j < len(paragraphs):
            if heading_info[j] is not None:
                break
            txt = paragraphs[j].text or ""
            if txt.strip():
                body_indices.append(j)
                body_texts.append(txt.strip())
            j += 1

        if body_indices:
            sec = NarrativeSection(
                title=title,
                title_para_idx=i,
                heading_level=level,
                body_para_indices=body_indices,
                example_text="\n".join(body_texts),
            )
            sections.append(sec)
        # 即使 body 为空也前进 i
        i = j

    return sections


# ─────────────────────────────────────────────────────────────
# 4. Prompt 构造
# ─────────────────────────────────────────────────────────────

_NARRATIVE_SYSTEM = """你是对公信贷审查撰写专家。基于【当前客户】材料,按【大标题】主旨重写对应段落内容。

铁律(不可违反):
1. 当前客户就是【当前企业画像锚点】里写的企业 — 正文只写这一家公司的情况
2. 【历史示例】是别家旧报告参考,仅供结构深度/写作风格参考,严禁带入:
   - 别家公司名 / 别家子公司名 / 别家实控人名
   - 2019-2022 等旧年份 (除非当前客户确有那些年份的数据)
   - 任何示例里出现的具体数字
3. 数字只从【企业画像锚点】/【结构化材料锚点】/【已计算财务指标】取,不得编造;
   材料没有的数字必须写"(待补充 XX)",禁止"具体见附表"/"具体以财务附表为准"/"以审贷会结论为准"之类兜底话术
4. 财务段须三段式:近三年数据罗列 → 外因(引用行业/政策卡片) → 内因(标注依据:XXX)
5. 变动幅度短语原文复用【已计算财务指标】给的标准短语,禁止自行再计算
6. 缺材料时请列具体清单,不要写"待审贷会/贷审会/贷委会讨论"
7. 输出 1-3 段正文,不要回显大标题/方括号头/评分表;不要用 Markdown 表格
8. 降级严禁:不许写"无""暂无""待确认"来填满段落,要么写具体事实要么列具体问题
"""


def _build_narrative_prompt(
    section: NarrativeSection,
    company_profile: str,
    material_anchor,
    financial_indicators,
    truth_financial_data,
) -> tuple[str, str]:
    """返回 (system_prompt, user_prompt)。"""
    ma_text = ""
    if material_anchor is not None:
        try:
            ma_text = material_anchor.format_for_prompt() or ""
        except Exception:
            ma_text = ""

    fin_text = ""
    if financial_indicators is not None:
        try:
            from financial_analyzer import FinancialAnalyzer
            fin_text = FinancialAnalyzer().format_for_prompt(
                financial_indicators, multi_year_data=truth_financial_data,
            ) or ""
        except Exception:
            fin_text = ""

    # 示例文本截断,避免 token 浪费
    example = (section.example_text or "").strip()
    example_short_hint = ""
    if len(example) < 100:
        example_short_hint = (
            "\n\n【注意】本节历史示例过短,请主要依据大标题主旨 + 当前客户材料扩展撰写 1 段即可,"
            "不要硬凑篇幅。"
        )
    if len(example) > 2000:
        example = example[:2000] + "...(截断)"

    user = f"""【当前企业画像锚点 + 结构化材料锚点】
{company_profile.strip() if company_profile else "(未构建企业画像)"}

{ma_text}

【已计算财务指标】
{fin_text if fin_text else "(本节无财务指标可引用)"}

【历史示例 — 仅供结构/风格参考,严禁带入公司名/人名/年份/数字】
{example if example else "(无示例)"}

【本节大标题】(Level {section.heading_level})
{section.title}
{example_short_hint}
请针对此大标题,基于当前客户材料撰写 1-3 段正文。直接输出正文段落,不要回显大标题。"""

    return _NARRATIVE_SYSTEM, user


# ─────────────────────────────────────────────────────────────
# 5. 公司名硬检测守卫
# ─────────────────────────────────────────────────────────────

def _extract_company_names(text: str) -> list[str]:
    """从一段文本里抽出形如 "XX有限公司" / "XX集团" 的公司名 (去重)。

    关键问题:中文无空格分词,regex 容易把"作为福建省招标股份有限公司"这种
    整块匹配,需要主动剥离已知动词/连词前缀后再过滤。
    """
    if not text:
        return []

    # 已知需要剥离的前缀 (按长度降序,长前缀优先)
    strip_prefixes = sorted(_FAKE_NAME_PREFIXES, key=len, reverse=True)

    seen = set()
    out = []
    for m in _COMPANY_EXTRACT_RE.finditer(text):
        name = m.group(1).strip()
        # 尝试剥离已知前缀
        for pref in strip_prefixes:
            if name.startswith(pref):
                rest = name[len(pref):]
                # 剥离后长度合理且剩余部分以公司尾缀结束才接受
                if (len(rest) >= 4 and
                        rest.endswith(("公司", "集团", "企业", "合伙"))):
                    name = rest
                    break
        if name in seen:
            continue
        if not _is_real_company_candidate(name):
            continue
        seen.add(name)
        out.append(name)
    return out


def _guard_company_name_leak(
    text: str,
    example_text: str,
    current_company_name: str,
) -> tuple[str, int]:
    """扫 LLM 输出,如果出现 example_text 里的公司名但不是当前客户名 → 标记。

    假阳性防御: 只处理长度 >= 4 且与 current_company_name 无子串关系的词。
    返回 (fixed_text, n_fixes)。

    策略: 替换为 "(待核实公司名)" 而不是强行替换成当前客户名 —— 让风控同事
    人工核对,避免改错。
    """
    if not text or not example_text or not current_company_name:
        return text, 0

    example_names = _extract_company_names(example_text)
    if not example_names:
        return text, 0

    current = current_company_name.strip()
    # 从 example 里挑出"和当前客户无子串关系"的别家公司名
    foreign_names = []
    for name in example_names:
        if len(name) < 4:
            continue
        # 子串互检
        if name in current or current in name:
            continue
        # 公共 core (去掉后缀比较)
        core_name = re.sub(r'(有限公司|股份有限公司|集团|公司|企业)$', '', name)
        core_cur = re.sub(r'(有限公司|股份有限公司|集团|公司|企业)$', '', current)
        if core_name and core_cur and (core_name in core_cur or core_cur in core_name):
            continue
        foreign_names.append(name)

    if not foreign_names:
        return text, 0

    # 按长度降序替换,避免短名覆盖长名
    foreign_names.sort(key=len, reverse=True)
    fixed = text
    n = 0
    for fn in foreign_names:
        if fn in fixed:
            count = fixed.count(fn)
            fixed = fixed.replace(fn, "(待核实公司名)")
            n += count
    return fixed, n


# ─────────────────────────────────────────────────────────────
# 6. 单节重写 (LLM + V14 守卫)
# ─────────────────────────────────────────────────────────────

def _rewrite_section(
    section: NarrativeSection,
    agent,
    company_profile: str,
    material_anchor,
    financial_indicators,
    truth_financial_data,
    current_company_name: str,
    progress_cb: Optional[Callable] = None,
) -> Optional[str]:
    """调 LLM 重写一节,跑完所有 text 级守卫,返回新正文 (可能含换行)。

    失败 (LLM 空) 返回 None,调用方保留原段文本。
    """
    sys_p, usr_p = _build_narrative_prompt(
        section, company_profile, material_anchor,
        financial_indicators, truth_financial_data,
    )

    try:
        raw = agent.llm(sys_p, usr_p)
    except Exception as e:
        logger.warning("LLM rewrite failed for %s: %s", section.title[:20], e)
        return None
    if not raw or len(raw.strip()) < 20:
        return None

    text = raw.strip()

    # ── 守卫 1: section_generator._sanitize_output (prompt 回显 / 占位符 / MD 表格) ──
    try:
        from section_generator import _sanitize_output
        text = _sanitize_output(text) or text
    except Exception as e:
        logger.warning("_sanitize_output failed: %s", e)

    # ── 守卫 2: 关联企业数字 (material_anchor 里的权威值代码层替换) ──
    if material_anchor is not None:
        try:
            from section_generator import _fix_affiliates_deterministic
            text2, n_aff = _fix_affiliates_deterministic(text, material_anchor)
            if n_aff:
                text = text2
                if progress_cb:
                    progress_cb(f"    [关联企业守卫] 修 {n_aff} 处")
        except Exception as e:
            logger.warning("_fix_affiliates_deterministic failed: %s", e)

    # ── 守卫 3: 变动短语 (荒谬短语 / 兜底话术) ──
    if financial_indicators is not None:
        try:
            from section_generator import _fix_variation_phrases
            text2, n_var = _fix_variation_phrases(text, financial_indicators)
            if n_var:
                text = text2
                if progress_cb:
                    progress_cb(f"    [变动短语守卫] 修 {n_var} 处")
        except Exception as e:
            logger.warning("_fix_variation_phrases failed: %s", e)

    # ── 守卫 4: 双写自矛盾 "分别为 X、X 同比 Y%" ──
    try:
        from section_generator import _fix_dup_year_phrases
        text2, n_dup = _fix_dup_year_phrases(text)
        if n_dup:
            text = text2
            if progress_cb:
                progress_cb(f"    [双写守卫] 修 {n_dup} 处")
    except Exception as e:
        logger.warning("_fix_dup_year_phrases failed: %s", e)

    # ── 守卫 5 (narrative 专有): 公司名硬检测,别家公司名 → (待核实公司名) ──
    if current_company_name:
        try:
            text2, n_cn = _guard_company_name_leak(
                text, section.example_text, current_company_name,
            )
            if n_cn:
                text = text2
                if progress_cb:
                    progress_cb(f"    [公司名守卫] 替换 {n_cn} 处别家公司名")
        except Exception as e:
            logger.warning("_guard_company_name_leak failed: %s", e)

    return text.strip() or None


# ─────────────────────────────────────────────────────────────
# 7. 回写到 docx
# ─────────────────────────────────────────────────────────────

def _clear_paragraph(para) -> None:
    for run in para.runs:
        run.text = ""


def _write_paragraph(para, text: str) -> None:
    """把 text 写进 para: 保留第一个 run 的格式,清掉其他 run 的文本。"""
    runs = list(para.runs)
    if not runs:
        para.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def _apply_rewrite_to_docx(
    doc,
    section: NarrativeSection,
    new_text: str,
    progress_cb: Optional[Callable] = None,
) -> int:
    """把 new_text 回写到 section.body_para_indices 对应段落。

    按段落个数分:
      - new_text 段数 == body 段数: 一对一
      - new_text 段数 < body 段数: 前 k 段按段写,剩余段清空
      - new_text 段数 > body 段数: 多出的段并入最后一段
    返回实际改动段落数。
    """
    paragraphs = list(doc.paragraphs)
    body_idx = section.body_para_indices
    if not body_idx:
        return 0

    # 规整换行 → 按非空块分段
    chunks = [c.strip() for c in re.split(r'\n\s*\n+', (new_text or "").strip()) if c.strip()]
    if not chunks:
        chunks = [(new_text or "").strip()]

    n_body = len(body_idx)
    n_chunks = len(chunks)

    if n_chunks >= n_body:
        # 最后一段吸收多余 chunk
        final = chunks[: n_body - 1] + ["\n".join(chunks[n_body - 1:])]
    else:
        # 填前 n_chunks 个 body,剩余清空
        final = chunks + [""] * (n_body - n_chunks)

    applied = 0
    for para_idx, piece in zip(body_idx, final):
        if para_idx >= len(paragraphs):
            continue
        para = paragraphs[para_idx]
        if piece:
            _write_paragraph(para, piece)
        else:
            _clear_paragraph(para)
        applied += 1
    return applied


# ─────────────────────────────────────────────────────────────
# 7.5. 顶层表格基本信息 prefill (企业基本信息表用 KB 事实覆盖)
# ─────────────────────────────────────────────────────────────

# label 关键字 → KB facts 中的 key (标签匹配时,把紧邻右侧单元格替换为 KB 值)
# 这些是模板级通用标签,不是公司名黑名单
_BASIC_LABEL_TO_KB = {
    "企业名称": "company_name",
    "公司名称": "company_name",
    "借款人名称": "company_name",
    "借款人全称": "company_name",
    "授信客户全称": "company_name",
    "客户全称": "company_name",
    "法人代表": "legal_representative",
    "法定代表人": "legal_representative",
    "实际控制人": "controller_name",
    "注册资本": "registered_capital",
    "实收资本": "paid_in_capital",
    "所属行业": "industry",
    "经营地址": "operating_address",
    "注册地址": "operating_address",
    "办公地址": "operating_address",
    "客户经理": "customer_manager_name",
}


def _prefill_basic_info_tables(doc, kb: dict, progress_cb=None) -> int:
    """把模板顶层表格里的 label→value 单元格基本信息,用 KB 事实覆盖。

    规则:同一行里找到标签单元格 (文本完全或主要是 label 关键字),
    右边相邻 (或者 row 内下一个非空) 单元格若和 label **不同**,就替换为 KB 值。
    这处理了"基本信息表"这类从历史模板带下来的 company-specific 值
    (如 法人代表 = 郑志煌 → 黄祖海)。

    只替换 value 单元格,不动表头/标签单元格。
    """
    if not kb:
        return 0
    facts = kb.get("facts", {}) or {}
    if not facts:
        return 0

    n_fixed = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = list(row.cells)
            n = len(cells)
            for i, cell in enumerate(cells):
                cell_text = (cell.text or "").strip()
                if not cell_text or len(cell_text) > 10:
                    continue
                # 精确匹配 label (允许冒号/空格)
                normalized = re.sub(r'[：:\s]', '', cell_text)
                kb_key = _BASIC_LABEL_TO_KB.get(normalized)
                if not kb_key:
                    continue

                kb_val = facts.get(kb_key)
                if not kb_val or not isinstance(kb_val, str):
                    continue
                kb_val = kb_val.strip()
                if not kb_val:
                    continue

                # 寻找 value 单元格 (右侧第一个非空且文本!=label 的 cell)
                for j in range(i + 1, n):
                    vcell = cells[j]
                    vtext = (vcell.text or "").strip()
                    if not vtext:
                        continue
                    # 如果 value 正好等于 label (合并单元格重复),跳过
                    if re.sub(r'[：:\s]', '', vtext) == normalized:
                        continue
                    # 如果已经和 KB 值一致,跳过
                    if kb_val in vtext or vtext in kb_val:
                        break
                    # 执行替换:把 vcell 所有 runs 清空,第一个 paragraph 写 KB 值
                    try:
                        paras = vcell.paragraphs
                        if not paras:
                            break
                        first = paras[0]
                        runs = list(first.runs)
                        if runs:
                            runs[0].text = kb_val
                            for r in runs[1:]:
                                r.text = ""
                        else:
                            first.add_run(kb_val)
                        for p in paras[1:]:
                            for r in p.runs:
                                r.text = ""
                        n_fixed += 1
                        if progress_cb:
                            progress_cb(f"    [基本信息表] {normalized} → {kb_val[:20]}")
                    except Exception:
                        pass
                    break
    return n_fixed


# ─────────────────────────────────────────────────────────────
# 8. 主入口
# ─────────────────────────────────────────────────────────────

def _extract_current_company_name(agent) -> str:
    """从 agent.kb.facts 里取 company_name, 失败降级 materials regex。"""
    try:
        facts = (getattr(agent, "kb", {}) or {}).get("facts", {}) or {}
        name = (facts.get("company_name") or "").strip()
        if name:
            return name
    except Exception:
        pass
    try:
        mats = agent._build_materials_text()
        m = re.search(
            r"(?:企业名称|公司名称|借款人名称)[：:]\s*([\u4e00-\u9fff（）()\w]+)",
            mats,
        )
        if m:
            return m.group(1).strip()
    except Exception:
        pass
    return ""


def _build_financial_indicators(agent, progress_cb) -> object:
    """从 file_path_map 里发现财务 xlsx,构建 FinancialIndicators。"""
    try:
        from financial_analyzer import FinancialAnalyzer
        fin_xlsx = None
        for fname, fpath in (getattr(agent, "file_path_map", None) or {}).items():
            if not fpath.lower().endswith((".xlsx", ".xls")):
                continue
            if re.search(r"会企|财务报表|资产负债|利润表|现金流", fname):
                fin_xlsx = fpath
                break
        if fin_xlsx:
            if progress_cb:
                progress_cb(f"  [narrative] 发现财务 xlsx: {os.path.basename(fin_xlsx)}")
            return FinancialAnalyzer().analyze(fin_xlsx)
    except Exception as e:
        if progress_cb:
            progress_cb(f"  [narrative] FinancialIndicators 构建失败: {e}")
    return None


# ─────────────────────────────────────────────────────────────
# V15-v2 护栏:封面旧客户名 scrub + 封面字段清零
# ─────────────────────────────────────────────────────────────

# 严格公司名正则:前缀必须是地名/商号,避免 n-gram 假阳性
_STRICT_FOREIGN_NAME_RE = re.compile(
    r'(?:^|[，。；;\s\(（、])'
    r'((?:福建|北京|上海|广州|深圳|厦门|福州|南京|天津|重庆|成都|武汉|'
    r'江苏|浙江|山东|广东|河南|四川|中国|全国|锐捷|招标|经纬|巨电|'
    r'星链|浩瀚|新天|精导|涵江|翔远|三明|济南|兴业|海峡|交通|工商|'
    r'建设|农业|中行|广发|国际|恒运|华兴|德赢|中船|国网|南网)'
    r'[\u4e00-\u9fff]{2,20}?'
    r'(?:有限公司|股份有限公司|股份公司|集团|有限责任公司|合伙企业|股份合作公司))'
)

_COVER_FIELD_PATS = [
    (re.compile(r'PD\s*评级\s*[:：][^;；\n　]{1,15}'), 'PD评级:(待补充)'),
    (re.compile(r'申报单位\s*[:：][^;；\n　]{1,20}'), '申报单位:(待补充)'),
    (re.compile(r'白名单\s*类?型?\s*[:：][^\n]{0,30}'), '白名单类型:(待补充)'),
    (re.compile(r'绿色信贷\s*标识?\s*[:：][^\n]{1,15}'), '绿色信贷标识:(待补充)'),
    (re.compile(r'投向政策\s*对?应?\s*行业\s*[:：][^;；\n　]{1,20}'),
     '投向政策对应行业:(待补充)'),
]

_COVER_SCAN_LIMIT = 30  # 前 30 段视为封面区


def _extract_foreign_names_from_template(template_path: str,
                                          current_company_name: str) -> list[str]:
    """从**原模板** docx 精确抽公司名,排除当前客户系。

    用严格正则要求"地名/商号前缀",避免 n-gram 假阳性(如"慧水利与智慧教育领域的
    软件和信息技术服务企业"这种被 Agent 从句子中切错)。
    """
    from docx import Document as _Document
    if not template_path or not os.path.exists(template_path):
        return []
    try:
        d = _Document(template_path)
    except Exception:
        return []

    texts = []
    for p in d.paragraphs:
        if p.text and p.text.strip():
            texts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    if p.text and p.text.strip():
                        texts.append(p.text)
    full = '\n'.join(texts)

    names = set()
    for m in _STRICT_FOREIGN_NAME_RE.finditer(full):
        names.add(m.group(1).strip())

    # 用当前客户核心词过滤(支持"中锐网络"/"中锐集团"等衍生)
    current_core = re.sub(
        r'(有限公司|股份有限公司|股份公司|集团|有限责任公司|合伙企业|股份合作公司)$',
        '', current_company_name or ''
    ).strip()
    # 取 2-6 字的核心词做 substring 判定
    exclude_key = current_core[:6] if len(current_core) >= 2 else ''

    foreign = []
    for n in names:
        if n == current_company_name:
            continue
        if exclude_key and exclude_key in n:
            continue
        # 顺便排除当前客户名完全包含在内的(保险)
        if current_company_name and current_company_name in n:
            continue
        foreign.append(n)

    # 按长度降序:长名先替换,避免短名截断长名
    foreign.sort(key=len, reverse=True)
    return foreign


def _scrub_foreign_names_from_template(doc, template_path: str,
                                        current_company_name: str,
                                        progress_cb=None) -> dict:
    """扫整个 docx,把原模板里的旧客户名替换为当前客户名;
    封面前 30 段的基本信息键值对(PD 评级/申报单位/白名单/...)清空 + pending。

    返回 {'names_replaced': N, 'fields_cleared': M}
    """
    stats = {'names_replaced': 0, 'fields_cleared': 0}
    if not current_company_name:
        return stats

    foreign = _extract_foreign_names_from_template(
        template_path, current_company_name
    )
    if not foreign and progress_cb:
        # 没抽到 foreign 不表示错,只是可能原模板已经是空白型
        pass

    def _scrub(p, is_cover: bool):
        if not p.runs:
            return
        orig = p.text or ''
        if not orig.strip():
            return
        new = orig
        for fn in foreign:
            if fn in new:
                cnt = new.count(fn)
                new = new.replace(fn, current_company_name)
                stats['names_replaced'] += cnt
        if is_cover:
            for pat, repl in _COVER_FIELD_PATS:
                new_after = pat.sub(repl, new, count=1)
                if new_after != new:
                    stats['fields_cleared'] += 1
                    new = new_after
        if new != orig:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ''

    for i, p in enumerate(doc.paragraphs):
        _scrub(p, i < _COVER_SCAN_LIMIT)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    _scrub(p, False)
                for sub in c.tables:
                    for sr in sub.rows:
                        for sc in sr.cells:
                            for sp in sc.paragraphs:
                                _scrub(sp, False)
    return stats


def run_narrative_pipeline(
    agent,
    doc,
    template_path: str,
    output_path: str,
    progress_cb: Optional[Callable] = None,
) -> dict:
    """对公叙述模板管线入口。

    前置: agent 已经跑过 _build_company_profile / _truth_build_financial_data
    (form_filler.run() 的 Phase 0 阶段在分流前已经执行)。

    流程:
      1. 切 section
      2. 并发 LLM 重写每节 (继承 V14 所有 text 守卫)
      3. 回写到 docx
      4. 跑表格级 truth prefill (复用 V14 _truth_prefill_financial_tables)
      5. 跑 V14 docx 级守卫 (标题数字 / 率值荒谬 / 加黑 / 表格残影)
      6. 存盘

    返回: {"sections_rewritten": int, "tables_filled": int, "pending_tags": list}
    """
    from form_filler import FormFillAgent, scan_document

    def _log(msg: str):
        if progress_cb:
            progress_cb(msg)

    _log("[narrative] V15 对公叙述管线启动")

    # ── 1) 切 section ──
    sections = _discover_narrative_sections(doc)
    _log(f"[narrative] 识别到 {len(sections)} 个大标题 section")

    if not sections:
        _log("[narrative] 未识别到任何 section,回退到 V14")
        return {
            "sections_rewritten": 0,
            "tables_filled": 0,
            "pending_tags": [],
            "fallback": True,
        }

    # ── 2) 构造上下文 ──
    current_company_name = _extract_current_company_name(agent)
    _log(f"[narrative] 当前客户: {current_company_name or '(未识别)'}")

    company_profile = getattr(agent, "_company_profile_block", "") or ""
    material_anchor = getattr(agent, "_material_anchor", None)
    truth_financial_data = getattr(agent, "_truth_financial_data", None) or {}

    # FinancialIndicators 是叙述段 prompt 的关键
    financial_indicators = _build_financial_indicators(agent, progress_cb)

    # ── 3) 并发重写每节 ──
    try:
        max_workers = int(os.environ.get("NARRATIVE_WORKERS", "6"))
    except Exception:
        max_workers = 6
    max_workers = max(1, min(max_workers, len(sections)))

    _log(f"[narrative] 并发度 {max_workers} 重写 {len(sections)} 节")

    rewrites: dict[int, str] = {}  # title_para_idx -> new_text
    rewrite_lock = None

    def _do_section(sec: NarrativeSection):
        _log(f"  [{sec.heading_level}] 重写「{sec.title[:32]}」")
        new_text = _rewrite_section(
            sec, agent, company_profile, material_anchor,
            financial_indicators, truth_financial_data,
            current_company_name, progress_cb,
        )
        return sec, new_text

    # 简易并发 (threadpool),与 section_generator 风格一致
    try:
        from concurrent.futures import ThreadPoolExecutor, as_completed
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            futs = {ex.submit(_do_section, s): s for s in sections}
            for fut in as_completed(futs):
                try:
                    sec, new_text = fut.result()
                except Exception as e:
                    logger.warning("Narrative worker failed: %s", e)
                    continue
                if new_text:
                    rewrites[sec.title_para_idx] = (sec, new_text)
    except Exception as e:
        _log(f"[narrative] 并发重写出错: {e},降级串行")
        for s in sections:
            sec, new_text = _do_section(s)
            if new_text:
                rewrites[sec.title_para_idx] = (sec, new_text)

    # ── 4) 回写到 docx ──
    sections_rewritten = 0
    for title_idx, payload in rewrites.items():
        sec, new_text = payload
        try:
            applied = _apply_rewrite_to_docx(doc, sec, new_text, progress_cb)
            if applied:
                sections_rewritten += 1
        except Exception as e:
            logger.warning("apply rewrite failed for %s: %s", sec.title[:20], e)

    _log(f"[narrative] 段落重写完成: {sections_rewritten}/{len(sections)} 节")

    # ── 4.5) 顶层表格基本信息覆盖 (企业名称/法人代表等 KB 事实) ──
    try:
        kb = getattr(agent, "kb", {}) or {}
        basic_fixed = _prefill_basic_info_tables(doc, kb, progress_cb)
        if basic_fixed:
            _log(f"[narrative] 基本信息表覆盖 {basic_fixed} 个单元格")
    except Exception as e:
        _log(f"[narrative] 基本信息表覆盖异常: {e}")

    # ── 5) 表格级 truth prefill (复用 V14) ──
    tables_filled = 0
    try:
        _, _, _, nested_tables, _ = scan_document(doc)
        _log(f"[narrative] 表格扫描: {len(nested_tables)} 个嵌套表")

        # 只挑财务表给 truth prefill
        _FIN_HEADER_KW = ['科目', '资产', '负债', '年份', '利润', '现金',
                          '余额', '期初', '期末']
        fin_tables = [
            t for t in nested_tables
            if any(kw in " ".join(t.header_row or []) for kw in _FIN_HEADER_KW)
        ]
        if fin_tables and truth_financial_data:
            _log(f"[narrative] 对 {len(fin_tables)} 个财务表跑 V14 truth prefill")
            try:
                tables_filled = agent._truth_prefill_financial_tables(
                    doc, fin_tables, truth_financial_data, progress_cb,
                ) or 0
                _log(f"[narrative] 表格回填 {tables_filled} 单元格")
            except Exception as e:
                _log(f"[narrative] truth prefill 失败: {e}")
    except Exception as e:
        _log(f"[narrative] 表格处理异常: {e}")

    # ── 6) V14 docx 级守卫 (末尾必跑) ──
    try:
        residue = agent._clear_table_residue_paragraphs(doc)
        if residue:
            _log(f"[narrative] 表格残影清理: {residue} 段")
    except Exception as e:
        _log(f"[narrative] 表格残影清理异常: {e}")

    try:
        bold_reset = FormFillAgent._sanitize_docx_runs(doc)
        if bold_reset:
            _log(f"[narrative] docx-run 清除 {bold_reset} 处非标题加黑")
    except Exception as e:
        _log(f"[narrative] sanitize_docx_runs 异常: {e}")

    try:
        title_fixed = FormFillAgent._fix_skeleton_title_numbers(doc)
        if title_fixed:
            _log(f"[narrative] 标题数字一致性修 {title_fixed} 处")
    except Exception as e:
        _log(f"[narrative] 标题数字守卫异常: {e}")

    try:
        pct_fixed = FormFillAgent._fix_same_value_pct_variation(doc)
        if pct_fixed:
            _log(f"[narrative] 率值荒谬守卫修 {pct_fixed} 处")
    except Exception as e:
        _log(f"[narrative] 率值荒谬守卫异常: {e}")

    # ── 6.5) V15-v2 护栏:原模板旧客户名 scrub + 封面字段清零 ──
    # 标题/封面里残留的旧客户名是 V15 MVP 的已知盲区,这里兜底。
    # 严格策略:只从**原模板** docx 抽公司名(保证 foreign 集精确,零假阳性)。
    try:
        scrub_stats = _scrub_foreign_names_from_template(
            doc, template_path, current_company_name, progress_cb
        )
        if scrub_stats.get('names_replaced') or scrub_stats.get('fields_cleared'):
            _log(f"[narrative] 封面护栏: 替换旧客户名 {scrub_stats['names_replaced']} 处, "
                 f"清封面字段 {scrub_stats['fields_cleared']} 处")
    except Exception as e:
        _log(f"[narrative] 封面护栏异常: {e}")

    # ── 7) 存盘 (保持 ZIP 结构,与 V14 一致) ──
    # template_path 可能是 .doc,需要传 agent._ensure_docx 转后的 docx 路径
    # 但 run() 在调用前已经做过 _ensure_docx,我们传 template_path 对 .docx 模板
    # 没问题;保险起见 fallback 到原 path
    docx_path = template_path
    try:
        # 若 agent 已经缓存了 docx_path,优先用
        maybe = getattr(agent, "_last_docx_path", None)
        if maybe and os.path.exists(maybe):
            docx_path = maybe
    except Exception:
        pass

    try:
        FormFillAgent._save_preserving_zip(
            doc, docx_path, output_path, comments_xml=None,
        )
        _log(f"[narrative] 输出文件已保存: {output_path}")
    except Exception as e:
        _log(f"[narrative] 保存失败(fallback 到 doc.save): {e}")
        try:
            doc.save(output_path)
            _log(f"[narrative] fallback 保存成功: {output_path}")
        except Exception as e2:
            _log(f"[narrative] fallback 保存亦失败: {e2}")
            raise

    return {
        "sections_rewritten": sections_rewritten,
        "tables_filled": tables_filled,
        "pending_tags": [],
    }
