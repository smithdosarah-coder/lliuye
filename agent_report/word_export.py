# -*- coding: utf-8 -*-
"""agent_report.word_export — Stage C.1 Word 导出渲染器(python-docx).

监管底线:与 ``agent_credit.decision_letter_docx`` / ``agent_channel.export_docx``
一致 — 银行业私有化部署 · **禁止调用海外 API 渲染**。全部本地 BytesIO 完成 ·
由 FastAPI 路由作为 attachment 下载返回。

输入 payload 形态(兼容 v16 done · v13 done · refine 后续场景):
    {
        "report_id"  : str,                # session_id 别名
        "session_id" : str,
        "profile"    : EnterpriseProfile dict | None,
        "sections"   : list[ReportSection],   # 见 agent-report-spec §6.3
        "stats"      : {"total_fields", "auto_filled", "unfilled"} | None,
        "pending_questions": list[dict] | None,
        "qc"         : dict | None,           # v16 QC 摘要
        "client_manager"  : str,              # 默认 "客户经理"
        "business_line"   : str,              # corporate / inclusive / reserved
    }

产出: .docx bytes · 文件名 ``build_filename(payload) → str``
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from quality_scorer import DIMENSION_GATES

_DEFAULT_FONT = "Microsoft YaHei"
NA = "—"
QUALITY_GATE_WATERMARK = "质量闸未过 · 内部草稿 · 不得作为审批依据"

_BUSINESS_LABEL = {
    "corporate": "对公授信",
    "inclusive": "普惠 / 个体",
    "reserved":  "预留板块",
}

# 4 chapter 标题(与 v16 / v13 generator 输出 chapter id 对齐)
_CHAPTER_TITLES = {
    "chapter_1_background": "一、企业背景",
    "chapter_2_operation":  "二、经营情况",
    "chapter_3_finance":    "三、财务分析",
    "chapter_4_conclusion": "四、审批意见",
}


# ============================================================================
# 字体 / 段落 helpers (复用 agent_credit / agent_channel 模式)
# ============================================================================

def _set_font(
    run,
    name: str = _DEFAULT_FONT,
    size: float | None = None,
    bold: bool = False,
    italic: bool = False,
    color: tuple[int, int, int] | None = None,
) -> None:
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bool(bold)
    run.italic = bool(italic)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _add_paragraph(
    doc,
    text: str,
    *,
    size: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    align=None,
    color: tuple[int, int, int] | None = None,
) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic, color=color)


def _add_heading(doc, text: str, level: int = 1) -> None:
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=sizes.get(level, 11), bold=True)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: float = 10) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    _set_font(run, size=size, bold=bold)


def _fmt(v: Any) -> str:
    if v is None:
        return NA
    if isinstance(v, (list, tuple)):
        if not v:
            return NA
        return " · ".join(_fmt(x) for x in v)
    if isinstance(v, dict):
        if not v:
            return NA
        return " · ".join(f"{k}: {_fmt(val)}" for k, val in v.items() if val)
    s = str(v).strip()
    return s if s else NA


def _safe_section_title(sec: dict) -> str:
    t = (sec.get("title") or "").strip()
    if t:
        return t
    sid = sec.get("id") or ""
    return _CHAPTER_TITLES.get(sid, sid or "段落")


def _quality_gate_reasons(qc: dict) -> list[str]:
    reasons: list[str] = []

    def _append(message: Any, severity: Any = None) -> None:
        if message in (None, ""):
            return
        prefix = f"{severity}: " if severity not in (None, "") else ""
        reasons.append(prefix + str(message))

    # quality_scorer.QualityReport.to_json() 的真实结构。
    for reason in qc.get("fatal_reasons") or []:
        _append(reason, "fatal")

    for dimension in qc.get("dimensions") or []:
        if not isinstance(dimension, dict):
            continue
        name = str(dimension.get("name") or "未命名维度")
        threshold = dimension.get("pass_threshold", dimension.get("threshold"))
        if not isinstance(threshold, (int, float)):
            threshold = DIMENSION_GATES.get(name)
        raw_score = dimension.get("raw_score")
        gate_failed = (
            threshold is not None
            and isinstance(raw_score, (int, float))
            and raw_score < threshold
        )
        explicit_failed = (
            dimension.get("passed") is False
            or str(dimension.get("status") or "").lower() in {"fail", "failed", "blocked"}
        )
        failed = gate_failed or explicit_failed
        if not failed:
            continue
        missed = dimension.get("missed_items") or []
        score_detail = (
            f"实际分 {raw_score:g} vs 闸值 {threshold:g}"
            if gate_failed else f"实际分 {raw_score:g}，该维度标记失败"
            if isinstance(raw_score, (int, float)) else "该维度标记失败"
        )
        missed_detail = "；".join(str(item) for item in missed if item)
        detail = f"{score_detail}；{missed_detail}" if missed_detail else score_detail
        _append(f"维度「{name}」：{detail}")

    hallucinations = qc.get("hallucinations") or []
    if isinstance(hallucinations, list):
        for item in hallucinations:
            if not isinstance(item, dict):
                continue
            text = item.get("text") or "疑似幻觉"
            reason = item.get("reason") or "需人工复核"
            location = item.get("location")
            location_text = f"（{location}）" if location else ""
            _append(f"{text}{location_text}：{reason}")

    # 兼容历史/外部 QC 载荷；有真实 severity 才显示前缀。
    for key in ("issues", "blocks", "warnings", "blocking_reasons"):
        raw = qc.get(key) or []
        if isinstance(raw, dict):
            raw = list(raw.values())
        if not isinstance(raw, list):
            raw = [raw]
        for item in raw:
            if isinstance(item, dict):
                severity = item.get("severity") or item.get("level")
                message = item.get("message") or item.get("reason") or item.get("detail")
                _append(message, severity)
            elif item:
                _append(item)
    if not reasons:
        reasons.append("会话未通过质量检查")
    return reasons


# ============================================================================
# 主入口
# ============================================================================

def export(payload: dict, output_path: str | Path | None = None) -> bytes:
    """渲染 Agent6 信贷调查报告 .docx · 返字节流."""
    profile = payload.get("profile") or {}
    sections = list(payload.get("sections") or [])
    stats = payload.get("stats") or {}
    pending = list(payload.get("pending_questions") or [])
    qc = payload.get("qc") or {}
    business_line = payload.get("business_line") or "corporate"
    biz_cn = _BUSINESS_LABEL.get(business_line, business_line)
    rm = payload.get("client_manager") or "客户经理"
    report_id = payload.get("report_id") or payload.get("session_id") or ""

    company_name = (profile.get("company_name") or "未命名客户").strip() or "未命名客户"

    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    gate_blocked = qc.get("passed") is not True
    if gate_blocked:
        for section in doc.sections:
            header = section.header.paragraphs[0]
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header.add_run(QUALITY_GATE_WATERMARK)
            _set_font(run, size=10, bold=True, color=(175, 60, 45))

        _add_paragraph(
            doc,
            QUALITY_GATE_WATERMARK,
            size=12,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color=(175, 60, 45),
        )
        _add_paragraph(
            doc,
            "阻断原因：" + "；".join(_quality_gate_reasons(qc)),
            size=10,
            color=(175, 60, 45),
        )

    # ---------- 标题 ----------
    _add_paragraph(
        doc,
        f"{company_name} 授信调查报告",
        size=18, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_paragraph(
        doc,
        f"客户经理：{rm}    日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}    "
        f"业务线：{biz_cn}",
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=(110, 110, 110),
    )
    if report_id:
        _add_paragraph(
            doc,
            f"会话编号：{report_id}",
            size=8.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color=(160, 160, 160),
        )
    _add_paragraph(doc, "")

    # ---------- 一、企业基本信息(profile) ----------
    if profile:
        _render_profile_section(doc, profile)

    # ---------- 二、QC 概览(若有) ----------
    if qc:
        _render_qc_section(doc, qc, stats)

    # ---------- 三、4 章节正文 ----------
    if sections:
        _add_heading(doc, "报告正文(章节)", level=2)
        for sec in sections:
            _render_section(doc, sec)
    else:
        _add_paragraph(
            doc, "（无章节内容 · 请先在 fill 阶段生成 sections）",
            italic=True, color=(150, 150, 150),
        )

    # ---------- 四、待补字段(若有) ----------
    if pending:
        _render_pending_section(doc, pending)

    # ---------- 免责条款 ----------
    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "——本报告由 Agent6 信贷报告生成助手(v16 主管线 · classifier → "
        "generator → QC gate)自动生成 · 仅作客户经理审批参考 · 第 4 章 "
        "审批意见预留 Agent3 决策回写 · 所有数据本地渲染 · 无数据出境 · "
        "私有化合规。",
        size=8.5,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        color=(120, 120, 120),
    )

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(data)
    return data


# ============================================================================
# Sub-renderers
# ============================================================================

# 企业基本字段(EnterpriseProfile · 见 agent_report/enterprise_profile.py)
_PROFILE_FIELDS: list[tuple[str, str]] = [
    ("company_name", "企业名称"),
    ("unified_credit_code", "统一社会信用代码"),
    ("industry", "行业"),
    ("business_line", "业务线"),
    ("establishment_date", "成立日期"),
    ("registered_capital", "注册资本"),
    ("region", "注册地"),
    ("main_business", "主营业务"),
    ("controller_name", "实际控制人"),
    ("controller_share_pct", "控股比例"),
]


def _render_profile_section(doc, profile: dict) -> None:
    _add_heading(doc, "一、企业基本信息", level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    rendered = False
    for key, label in _PROFILE_FIELDS:
        v = profile.get(key)
        if v in (None, "", []):
            continue
        row = table.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=10)
        _set_cell_text(row[1], _fmt(v), size=10)
        rendered = True

    # financial_anchors 单独一行说明
    fa = profile.get("financial_anchors") or {}
    if isinstance(fa, dict) and fa:
        rev = fa.get("revenue_latest")
        np_ = fa.get("net_profit_latest")
        ta = fa.get("total_assets")
        period = fa.get("period")
        bits = []
        if rev:
            bits.append(f"营业收入 {rev}")
        if np_:
            bits.append(f"净利润 {np_}")
        if ta:
            bits.append(f"资产总计 {ta}")
        if period:
            bits.append(f"({period})")
        if bits:
            row = table.add_row().cells
            _set_cell_text(row[0], "财务锚点", bold=True, size=10)
            _set_cell_text(row[1], " · ".join(bits), size=10)
            rendered = True

    if not rendered:
        _add_paragraph(doc, "（profile 字段全空 · 跳过）",
                       italic=True, color=(150, 150, 150))
    _add_paragraph(doc, "")


def _render_qc_section(doc, qc: dict, stats: dict) -> None:
    _add_heading(doc, "二、生成质量与统计", level=2)
    passed = qc.get("passed")
    score = qc.get("score")
    halluc = qc.get("halluc_count", qc.get("hallucinations", 0))
    fatal = qc.get("fatal_fail", False)

    pass_label = "✓ 通过" if passed else ("✕ 阻断" if fatal else "△ 警告")
    pass_color = (46, 112, 76) if passed else (175, 60, 45)
    _add_paragraph(
        doc,
        f"QC 终审：{pass_label}    总分：{score if score is not None else NA}    "
        f"疑似幻觉：{halluc}    一票否决：{'是' if fatal else '否'}",
        size=10.5, bold=True, color=pass_color,
    )
    if stats:
        _add_paragraph(
            doc,
            f"字段统计：总 {stats.get('total_fields', NA)} · "
            f"自动填充 {stats.get('auto_filled', NA)} · "
            f"未填 {stats.get('unfilled', NA)}",
            size=10,
        )
    _add_paragraph(doc, "")


def _render_section(doc, sec: dict) -> None:
    title = _safe_section_title(sec)
    status = sec.get("status") or "done"
    word_count = sec.get("word_count")

    _add_heading(doc, title, level=3)
    if status and status != "done":
        _add_paragraph(
            doc,
            f"（状态：{status}{' · ' + str(word_count) + ' 字' if word_count else ''}）",
            italic=True, color=(150, 150, 150), size=9,
        )

    content = (sec.get("content") or "").strip()
    if not content:
        _add_paragraph(doc, "（暂无内容）", italic=True, color=(150, 150, 150))
        _add_paragraph(doc, "")
        return

    # 简化处理: 按行分段(markdown 列表 / 纯文本 hybrid)
    for line in content.splitlines():
        line = line.rstrip()
        if not line:
            _add_paragraph(doc, "", size=10)
            continue
        # 处理 markdown 列表前缀
        if line.lstrip().startswith(("- ", "* ", "· ")):
            txt = line.lstrip().lstrip("-*· ").strip()
            _add_paragraph(doc, f"  · {txt}", size=10.5)
        elif re.match(r"^\d+[.、]\s", line.lstrip()):
            _add_paragraph(doc, line.strip(), size=10.5)
        else:
            _add_paragraph(doc, line.strip(), size=10.5)

    _add_paragraph(doc, "")


def _render_pending_section(doc, pending: list[dict]) -> None:
    _add_heading(doc, "附录：待补字段清单", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["#", "字段 / 位置", "缺失原因", "建议动作", "建议来源 / 推荐答案"]):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, size=10)
    for i, q in enumerate(pending[:30], 1):
        row = table.add_row().cells
        _set_cell_text(row[0], str(i), size=9.5)
        label = q.get("label") or q.get("text") or q.get("id")
        location = q.get("location")
        field_location = " / ".join(str(v) for v in (label, location) if v not in (None, ""))
        _set_cell_text(row[1], _fmt(field_location), size=9.5)
        _set_cell_text(row[2], _fmt(q.get("reason")), size=9.5)
        _set_cell_text(row[3], _fmt(q.get("suggested_action")), size=9.5)
        rec = q.get("recommended") or q.get("source_ref") or NA
        _set_cell_text(row[4], _fmt(rec), size=9.5)
    _add_paragraph(doc, "")


# ============================================================================
# Filename
# ============================================================================

def build_filename(payload: dict) -> str:
    """生成下载文件名 · ``agent6_报告_<company>_<report_id>.docx``."""
    profile = payload.get("profile") or {}
    company = profile.get("company_name") or payload.get("query") or "客户"
    company = re.sub(r'[\\/:*?"<>|\s]+', "_", str(company)).strip("_")[:30]
    if not company:
        company = "客户"
    rid = (payload.get("report_id") or payload.get("session_id") or "").strip()
    if not rid:
        rid = datetime.now().strftime("%Y%m%d%H%M%S")
    rid = re.sub(r'[\\/:*?"<>|\s]+', "_", rid)[:32]
    return f"agent6_报告_{company}_{rid}.docx"


__all__ = ["export", "build_filename"]
