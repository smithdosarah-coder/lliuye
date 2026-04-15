"""
模板拆解层 (Template Decomposition Layer)

将模板的每个段落分类为两种角色：
  SKELETON — 章节标题、字段标签、结构性元素 → 原样保留
  CONTENT  — 范例文本、写作指导 → 用真实内容替换

核心原则：宁可多重写，不可漏重写。
任何不确定的段落都归为 CONTENT（会被重写），因为：
- 误重写骨架 → LLM 大概率会复现标题（低风险）
- 漏重写范例 → 模板泄漏（高风险，这正是我们要根治的问题）
"""

import re
from dataclasses import dataclass, field as dc_field
from difflib import SequenceMatcher
from enum import Enum
from typing import Any


# ═══════════════════════════════════════════
# 数据结构
# ═══════════════════════════════════════════

class TemplateRole(Enum):
    SKELETON = "skeleton"   # 标题、标签、结构 → 保留
    CONTENT = "content"     # 范例、指导 → 重写
    # 细粒度角色（向下兼容：PRESERVE/DATA_SLOT 等同于 SKELETON，EXAMPLE/INSTRUCTION 等同于 CONTENT）
    PRESERVE = "preserve"       # 保留原样的结构标题、编号、表头
    EXAMPLE = "example"         # 需要完全替换的示例文本
    DATA_SLOT = "data_slot"     # 仅包含 XX/下划线 的占位行
    INSTRUCTION = "instruction" # 模板中的指导说明文字（请、简述、列明等）


@dataclass
class ParaInfo:
    """一个段落的分类结果"""
    para_idx: int
    text: str
    role: TemplateRole
    is_heading: bool = False
    heading_level: int = 0      # 1=一、 2=（一） 3=1. 4=①/※
    section_id: str = ""
    # 从段落中提取的写作指令（括号内的"请..."等）
    embedded_instructions: list = dc_field(default_factory=list)
    # 该段落是否紧跟一个嵌套表格（composite 结构判定用）
    # 由 annotate_composite_structure() 在 decompose_template 中填充
    followed_by_table: bool = False


@dataclass
class SectionInfo:
    """一个逻辑章节"""
    section_id: str
    title: str
    level: int
    paragraphs: list[ParaInfo] = dc_field(default_factory=list)
    # composite 段落索引集合（para_idx）：
    # 这些 SKELETON 段后面紧跟一个嵌套表格（Word XML 中 w:p 之后直接是 w:tbl），
    # 且段落本身是"XX:______"形式的占位骨架（title/data-slot 标签行）。
    # section_generator 据此让 LLM 对该骨架只写 1-2 句总括，不重复列表格数据。
    composite_paragraph_indices: list = dc_field(default_factory=list)
    # V14-v3 循环表 schema：识别"顶部综述 cell + 下方循环子表(列头+示例row)"结构
    #   {
    #     "overview_para_idx": int,              # 综述段 para_idx
    #     "header_cols": list[str],              # 列头文本
    #     "sample_row_indices": list[int],       # 示例/占位 row 的索引
    #     "data_row_template": list[str],        # 空模板 row 的列单元占位
    #   }
    # 未命中则为 None。section_generator / form_filler 据此:
    #   - 综述 cell 只写 1-2 句说明,禁止塞流水明细(Prompt 铁律 + 代码后置长度检测)
    #   - 示例 row 的占位字面 (如 "如未落实请说明原因"/列头复读) 清除
    loop_table_schema: dict = dc_field(default_factory=dict)

    @property
    def skeleton_lines(self) -> list[str]:
        """本节的结构性文本（标题+标签行）"""
        _skeleton_roles = {TemplateRole.SKELETON, TemplateRole.PRESERVE, TemplateRole.DATA_SLOT}
        return [p.text for p in self.paragraphs
                if p.role in _skeleton_roles and p.text.strip()]

    @property
    def content_lines(self) -> list[str]:
        """本节的范例文本（将被重写）"""
        _content_roles = {TemplateRole.CONTENT, TemplateRole.EXAMPLE, TemplateRole.INSTRUCTION}
        return [p.text for p in self.paragraphs
                if p.role in _content_roles and p.text.strip()]

    @property
    def all_instructions(self) -> list[str]:
        """本节所有写作指令（从骨架行括号中提取 + CONTENT 中的指导句）"""
        result = []
        for p in self.paragraphs:
            result.extend(p.embedded_instructions)
        return result

    @property
    def content_para_indices(self) -> list[int]:
        """需要重写的段落索引"""
        _content_roles = {TemplateRole.CONTENT, TemplateRole.EXAMPLE, TemplateRole.INSTRUCTION}
        return [p.para_idx for p in self.paragraphs
                if p.role in _content_roles and p.text.strip()]

    @property
    def has_content_to_rewrite(self) -> bool:
        return len(self.content_lines) > 0


# ═══════════════════════════════════════════
# 分类规则
# ═══════════════════════════════════════════

_HEADING_PATTERNS = [
    (re.compile(r'^[一二三四五六七八九十]+、'), 1),
    (re.compile(r'^[（(][一二三四五六七八九十]+[)）]'), 2),
    (re.compile(r'^\d+[.．、]'), 3),
    (re.compile(r'^[①②③④⑤⑥⑦⑧⑨⑩]'), 4),
    (re.compile(r'^※\s*\d*[.．、]?'), 4),
]

# 签名行关键词 — 这些行属于 SKELETON（银行内部签名，保留不动）
_SIGNATURE_KEYWORDS = [
    '客户经理：', '共同调查人:', '共同调查人：',
    '管理经理/', '分管行长：', '行长：',
    '审查员：', '审批人：',
    '尽职调查：经营单位对上述材料真实性负责',
]


def _detect_heading(text: str) -> tuple[bool, int]:
    """检测标题。返回 (是否标题, 层级)"""
    stripped = text.strip()
    for pat, level in _HEADING_PATTERNS:
        if pat.match(stripped):
            return True, level
    return False, 0


def _extract_instructions(text: str) -> list[str]:
    """从骨架行中提取括号内的写作指令"""
    instructions = []
    # 匹配包含指导性动词的括号内容
    for m in re.finditer(
        r'[（(]([^)）]{4,80}(?:请|如|须|应|若|包括但不限于|可另附|可附)[^)）]{0,80})[）)]',
        text
    ):
        instructions.append(m.group(1).strip())
    return instructions


def _classify_paragraph(text: str) -> tuple[TemplateRole, bool, int]:
    """
    分类单个段落。

    返回: (角色, 是否标题, 标题层级)

    分类逻辑（按优先级）：
    1. 空行 → SKELETON
    2. 标题行 → SKELETON
    3. 签名行 → SKELETON
    4. 短标签行（<80字，以冒号结尾，无大段文字） → SKELETON
    5. 纯占位符行（主要是下划线/空格 + 标签） → SKELETON
    6. 表格注释行（"单位：万元"） → SKELETON
    7. 复选框行（主要由 □/☑ 选项构成） → SKELETON
    8. 其他一切 → CONTENT（安全默认，宁多重写）
    """
    stripped = text.strip()

    # 1. 空行
    if not stripped:
        return TemplateRole.SKELETON, False, 0

    is_h, level = _detect_heading(stripped)

    # 2. 标题行 — 但如果包含范例内容，归为 CONTENT
    #    判断标准：
    #    a) 太长（>=150字）→ 含范例内容
    #    b) 含 XX 占位符且长度>40 → 需要填写的范例
    #    c) 含冒号后接大段文字（标题：范例内容）且总长>100
    #    d) 含冒号后有范例特征（……/、/等连接多描述）且总长>40
    if is_h:
        has_xx = bool(re.search(r'X{2,}', stripped))
        # 冒号后有多少内容？
        colon_pos = max(stripped.find('：'), stripped.find(':'))
        after_colon_len = len(stripped) - colon_pos - 1 if colon_pos > 0 else 0
        # 冒号后的文本
        after_colon_text = stripped[colon_pos + 1:] if colon_pos > 0 else ""
        # 范例特征：省略号、泛指描述、斜杠分支、"如：" 示例引导等
        has_example_markers = bool(re.search(
            r'[\u2026\u00b7]|\.{2,}|/{1}[\u4e00-\u9fff]|'
            r'\u5982[\uff1a:]|'   # "如："
            r'\d{3,}[\u4e07\u5143\u540d\u4eba]',   # 具体数字+单位 (500万, 240名)
            after_colon_text))

        if len(stripped) >= 150:
            return TemplateRole.CONTENT, True, level
        if has_xx and len(stripped) > 40:
            return TemplateRole.CONTENT, True, level
        if after_colon_len > 80 and len(stripped) > 100:
            # 标题后跟了大段范例文字
            return TemplateRole.CONTENT, True, level
        if after_colon_len > 30 and len(stripped) > 50:
            # 冒号后有大量内容（纯标签行不会这么长）
            return TemplateRole.CONTENT, True, level
        if has_example_markers and after_colon_len > 20 and len(stripped) > 40:
            # 标题后跟了含范例特征的描述文字
            return TemplateRole.CONTENT, True, level
        return TemplateRole.SKELETON, True, level

    # 3. 签名/日期行
    for kw in _SIGNATURE_KEYWORDS:
        if kw in stripped:
            return TemplateRole.SKELETON, False, 0
    if re.match(r'^日期[：:]\s', stripped):
        return TemplateRole.SKELETON, False, 0

    # 4. 短标签行（以冒号结尾，且没有大段描述文字）
    #    比如 "关联企业关系：可附图"  "最近三年现金流量表分析："
    if len(stripped) < 80 and re.search(r'[：:]\s*$', stripped):
        return TemplateRole.SKELETON, False, 0

    # 5. 纯占位符/标签行
    #    比如 "上年度末社保人数     人；注册资本：         万元"
    #    特征：短，有大片空白或下划线
    blank_ratio = (stripped.count(' ') + stripped.count('　') +
                   stripped.count('_') + stripped.count('\t'))
    if len(stripped) < 120 and blank_ratio > len(stripped) * 0.3:
        return TemplateRole.SKELETON, False, 0

    # 6. 表格注释
    if re.match(r'^单位[：:]|^备注[：:]', stripped) and len(stripped) < 60:
        return TemplateRole.SKELETON, False, 0

    # 7. 复选框密集行（>= 2 个复选框，且行较短）
    cb_count = len(re.findall(r'[□☐☑✓]', stripped))
    if cb_count >= 2 and len(stripped) < 200:
        return TemplateRole.SKELETON, False, 0
    # 单个复选框但行很短
    if cb_count >= 1 and len(stripped) < 60:
        return TemplateRole.SKELETON, False, 0

    # 8. 兜底 → CONTENT
    return TemplateRole.CONTENT, False, 0


# ═══════════════════════════════════════════
# 细粒度段落角色分类
# ═══════════════════════════════════════════

# 模板示例行业关键词（宁多勿漏）
_TEMPLATE_SAMPLE_KEYWORDS = [
    "注塑", "模具", "塑胶", "某某", "XX公司", "XX有限", "XX集团",
    "XX银行", "XX事务所", "张XX", "李XX", "王XX", "陈XX", "刘XX",
]

# 指导性动词短语（必须是明确的指导结构，避免单字误匹配）
_INSTRUCTION_PHRASES = [
    "简述", "列明", "填写", "阐述", "标注",
]

# XX 占位符模式（宽松匹配）
_XX_PLACEHOLDER_RE = re.compile(r'X{2,}')

# 纯占位符行模式：仅包含 XX、下划线、空格、冒号、标签文字（<30字非占位内容）
_PURE_SLOT_RE = re.compile(r'^[\s_X：:．.、\u3000]*$')


def classify_paragraph_role(
    text: str,
    context_before: str = "",
    context_after: str = "",
) -> TemplateRole:
    """
    细粒度判断段落角色。

    分类优先级：
    1. PRESERVE — 编号格式开头的短标题；表头行；固定标题；签名行
    2. DATA_SLOT — 仅包含 XX 或下划线，无其他叙述内容
    3. INSTRUCTION — 包含指导性动词
    4. EXAMPLE — 包含 XX 占位符 + 叙述性文本；包含模板样例关键词
    5. 兜底 → EXAMPLE（宁可多标记也不要漏标记）

    Args:
        text: 段落文本
        context_before: 前一个段落的文本（可选，用于上下文辅助判断）
        context_after: 后一个段落的文本（可选）

    Returns:
        TemplateRole 枚举值
    """
    stripped = text.strip()

    # 空行 → PRESERVE
    if not stripped:
        return TemplateRole.PRESERVE

    is_h, level = _detect_heading(stripped)

    # ---- PRESERVE 判断 ----
    # 签名行
    for kw in _SIGNATURE_KEYWORDS:
        if kw in stripped:
            return TemplateRole.PRESERVE

    # 日期行
    if re.match(r'^日期[：:]\s', stripped):
        return TemplateRole.PRESERVE

    # 编号格式开头的短标题（无 XX，较短）
    if is_h and len(stripped) < 80:
        has_xx = bool(_XX_PLACEHOLDER_RE.search(stripped))
        if not has_xx:
            return TemplateRole.PRESERVE

    # 表头行
    if re.match(r'^单位[：:]|^备注[：:]', stripped) and len(stripped) < 60:
        return TemplateRole.PRESERVE

    # 复选框密集行
    cb_count = len(re.findall(r'[□☐☑✓]', stripped))
    if cb_count >= 2 and len(stripped) < 200:
        return TemplateRole.PRESERVE
    if cb_count >= 1 and len(stripped) < 60:
        return TemplateRole.PRESERVE

    # 短标签行（以冒号结尾，且没有大段描述文字）
    if len(stripped) < 80 and re.search(r'[：:]\s*$', stripped):
        return TemplateRole.PRESERVE

    # 审查意见等固定标题
    _fixed_titles = ["审查意见", "审批意见", "尽职调查", "意见表"]
    if any(kw in stripped for kw in _fixed_titles) and len(stripped) < 60:
        return TemplateRole.PRESERVE

    # ---- DATA_SLOT 判断 ----
    # 纯占位符行：仅包含 XX 或下划线
    has_xx = bool(_XX_PLACEHOLDER_RE.search(stripped))
    # 去掉 XX 和下划线后剩余的纯文字
    cleaned = re.sub(r'X{2,}|_{2,}|[\s：:．.、\u3000]', '', stripped)
    if has_xx and len(cleaned) < 15:
        return TemplateRole.DATA_SLOT

    # 纯占位符/标签行（大片空白或下划线）
    blank_ratio = (stripped.count(' ') + stripped.count('\u3000') +
                   stripped.count('_') + stripped.count('\t'))
    if len(stripped) < 120 and blank_ratio > len(stripped) * 0.3:
        return TemplateRole.DATA_SLOT

    # ---- INSTRUCTION 判断 ----
    # 包含指导性短语且较短（指导说明通常不长）
    if len(stripped) < 200:
        # "请+动词" 组合（明确的指导结构）
        if re.search(r'请(?:简述|描述|列明|填写|说明|分析|阐述|注明|标注|勿)', stripped):
            return TemplateRole.INSTRUCTION
        # 独立的指导性短语（不容易误匹配的多字词）
        for phrase in _INSTRUCTION_PHRASES:
            if phrase in stripped:
                return TemplateRole.INSTRUCTION

    # ---- EXAMPLE 判断（宽松，宁可多标记也不要漏标记） ----
    # 包含模板样例关键词
    for kw in _TEMPLATE_SAMPLE_KEYWORDS:
        if kw in stripped:
            return TemplateRole.EXAMPLE

    # 包含 XX 占位符 + 叙述性文本（不管数量多少）
    if has_xx and len(stripped) > 20:
        return TemplateRole.EXAMPLE

    # 包含常见模板示例模式
    if re.search(r'XX万元|XX年|XX%|XX人|XX公司|XX银行', stripped):
        return TemplateRole.EXAMPLE

    # 含有范例特征的较长段落
    if len(stripped) > 60:
        has_example_markers = bool(re.search(
            r'[\u2026\u00b7]|\.{2,}|/{1}[\u4e00-\u9fff]|'
            r'\u5982[\uff1a:]|'
            r'\d{3,}[\u4e07\u5143\u540d\u4eba]',
            stripped))
        if has_example_markers:
            return TemplateRole.EXAMPLE

    # 兜底：长段落都当作 EXAMPLE（宁多勿漏）
    if len(stripped) > 60:
        return TemplateRole.EXAMPLE

    # 短且不含任何标记的段落 → PRESERVE
    return TemplateRole.PRESERVE


def extract_template_fingerprints(doc_path: str) -> dict[str, str]:
    """
    提取模板中所有 EXAMPLE 段落的文本指纹。

    遍历模板文档中的所有表格单元格段落，对标记为 EXAMPLE 的段落：
    去掉 XX 标记后作为"指纹"文本，用于后续相似度比对。

    Args:
        doc_path: 模板 .docx 文件路径

    Returns:
        dict[paragraph_location → cleaned_text]
        其中 paragraph_location 格式为 "t{table_idx}_r{row_idx}_c{cell_idx}_p{para_idx}"
    """
    from docx import Document as _Document

    doc = _Document(doc_path)
    fingerprints: dict[str, str] = {}

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                paras = cell.paragraphs
                for p_idx, para in enumerate(paras):
                    text = (para.text or "").strip()
                    if not text or len(text) < 15:
                        continue

                    # 获取上下文
                    ctx_before = paras[p_idx - 1].text.strip() if p_idx > 0 else ""
                    ctx_after = paras[p_idx + 1].text.strip() if p_idx < len(paras) - 1 else ""

                    role = classify_paragraph_role(text, ctx_before, ctx_after)
                    if role == TemplateRole.EXAMPLE:
                        # 去掉 XX 标记，生成指纹
                        cleaned = re.sub(r'X{2,}', '', text)
                        cleaned = re.sub(r'_{2,}', '', cleaned)
                        cleaned = re.sub(r'\s+', '', cleaned)
                        if len(cleaned) >= 10:
                            loc = f"t{t_idx}_r{r_idx}_c{c_idx}_p{p_idx}"
                            fingerprints[loc] = cleaned

    return fingerprints


# ═══════════════════════════════════════════
# 主入口
# ═══════════════════════════════════════════

def classify_cell_paragraphs(cell) -> list[ParaInfo]:
    """对一个单元格中的所有直属段落进行分类"""
    results = []
    for i, para in enumerate(cell.paragraphs):
        text = para.text
        role, is_h, level = _classify_paragraph(text)
        instructions = _extract_instructions(text) if role == TemplateRole.SKELETON else []

        results.append(ParaInfo(
            para_idx=i,
            text=text,
            role=role,
            is_heading=is_h,
            heading_level=level,
            embedded_instructions=instructions,
        ))
    return results


def group_into_sections(paragraphs: list[ParaInfo]) -> list[SectionInfo]:
    """
    将分类后的段落按标题层级分组为章节。

    分割规则：遇到 level <= 3 的标题就开新节。
    level 4 (①/※) 不开新节，视为子项。
    """
    if not paragraphs:
        return []

    sections: list[SectionInfo] = []
    current: SectionInfo | None = None
    counter = 0

    for para in paragraphs:
        if para.is_heading and para.heading_level <= 3:
            counter += 1
            sid = f"sec_{counter:02d}"
            para.section_id = sid
            current = SectionInfo(
                section_id=sid,
                title=para.text.strip(),
                level=para.heading_level,
                paragraphs=[para],
            )
            sections.append(current)
        elif current is not None:
            para.section_id = current.section_id
            current.paragraphs.append(para)
        else:
            # 标题前的段落 → 前言节
            if not sections or sections[0].section_id != "sec_00":
                current = SectionInfo(
                    section_id="sec_00", title="（前言）", level=0
                )
                sections.insert(0, current)
            para.section_id = "sec_00"
            current.paragraphs.append(para)

    return sections


def _scan_para_followed_by_table(body_cell, body_paras: list[ParaInfo]) -> None:
    """
    遍历 body_cell 的 XML children(w:p / w:tbl),按序标记
    "该段落是否紧跟一个嵌套表格"。

    实现方式(结构特征,不依赖任何关键词):
      - body_cell._element 的 child elements 中 w:p 与 w:tbl 以正文顺序交替出现
      - 对每个 w:p,若其 next sibling 是 w:tbl(跨过空行可能的 w:p 会被视作 skeleton
        段的一部分,但通常 skeleton 段紧贴表格,无空行),则该段 followed_by_table=True
      - 段落与 body_paras 的对应:cell.paragraphs 与 body_cell._element 中 w:p
        顺序严格一致(python-docx 契约),按枚举同步即可

    向后兼容:没有嵌套表格的模板,全部 followed_by_table 保持 False,不影响主流程。
    """
    try:
        from docx.oxml.ns import qn
    except Exception:
        return

    p_tag = qn('w:p')
    tbl_tag = qn('w:tbl')

    # 收集 body_cell XML 顺序的 (kind, element) 列表
    ordered = []
    # cell._element 对应 w:tc; 其中直接 children 含 w:p / w:tbl / w:tcPr
    tc = body_cell._element
    for child in list(tc):
        if child.tag == p_tag:
            ordered.append(('p', child))
        elif child.tag == tbl_tag:
            ordered.append(('tbl', child))

    # 把 XML 顺序的 'p' 依次映射到 body_paras（同序）
    p_cursor = 0
    for i, (kind, elem) in enumerate(ordered):
        if kind != 'p':
            continue
        if p_cursor >= len(body_paras):
            break
        # 查找下一个非空 sibling:跨过同级空段、pPr 等不算
        is_followed = False
        for j in range(i + 1, len(ordered)):
            nxt_kind, _ = ordered[j]
            if nxt_kind == 'tbl':
                is_followed = True
                break
            if nxt_kind == 'p':
                # 下一个是段落且非空 → 不构成 composite
                # 若下一个段落是空(text 全空),允许再往后看一层
                pass
                break
        body_paras[p_cursor].followed_by_table = is_followed
        p_cursor += 1


# 匹配"title:______"/"title: 空白"骨架行——composite 判定辅助
#   特征:含冒号,冒号后主要是下划线/全角空格/单位标注("单位:万元"),
#   长度 <= 120 且非 CONTENT。不依赖任何业务关键词。
_COMPOSITE_SKELETON_RE = re.compile(
    r'[：:][ \u3000_]{2,}'                # 冒号后大段空白/下划线
    r'|^[①②③④⑤⑥⑦⑧⑨⑩].{0,40}[：:]'    # ①...:  level-4 带冒号
    r'|单位[：:]\s*(?:万元|元|%)'         # "单位:万元" 强特征
)


def _annotate_composite_sections(sections: list[SectionInfo]) -> int:
    """
    对每个 SectionInfo,识别其内部的 composite 段落:
      段落 followed_by_table + (占位骨架特征 or 长段落兜底)

    V14-v3 扩展:
      任意段落(无论长短、无论 role),只要其"在 cell XML 中紧邻下一个元素为表格"
      (followed_by_table=True),即视为 composite 绑定——下方表格会独立填充,
      段落只写总括/趋势/异常。这样可以覆盖:
        - 短骨架 "①主要上游:单位:万元" + 前五大采购表
        - 长财务段落 + 现金流量表
        - 上游/下游段 + 明细表
        - 关联企业段 + 关联企业表
        - 资产段 + 资产清单
        - 融资段 + 融资清单

    不限制段落 role:分类规则可能把"①主要上游(前五大):"这种后跟表格的 heading
    误分为 CONTENT,也可能把长段落分为 CONTENT。composite 标记与 role 正交,
    role 决定是否重写, composite 决定该段落对应的真实数据由表格填充路径处理。

    返回: 命中的 composite 段落总数
    """
    total = 0
    for sec in sections:
        hits = []
        for p in sec.paragraphs:
            if not p.followed_by_table:
                continue
            txt = (p.text or "").strip()
            if not txt:
                continue
            # V14-v3 放宽判定:只要 followed_by_table,统一绑 composite,
            # 原有短骨架特征保留(用于判定优先级/日志),长段落兜底命中。
            # 上限放宽到 2000 字(避免把整节 body 错绑)。
            if len(txt) > 2000:
                continue
            hits.append(p.para_idx)
        if hits:
            sec.composite_paragraph_indices = hits
            total += len(hits)
    return total


def _detect_loop_table_schemas(
    body_cell,
    body_paras: list[ParaInfo],
    sections: list[SectionInfo],
) -> int:
    """
    V14-v3: 识别"综述段 + 循环子表(列头+示例 row + 空模板 row)"结构。

    识别特征(纯结构,不依赖关键词黑名单):
      - 表格 >= 3 row
      - 第 1 row 是明确列头: 所有 cell 都有文本,每 cell 长度 <= 30 字
        且列头数 >= 2(典型: 序号 / XX / XX)
      - 第 2+ row 至少有 1 row 满足"示例 row"特征:
          a) 与列头存在文本重复(列头复读)
          b) 或含 placeholder 句式(短句、含"如...,请...说明"等通用模板提示,
             单行 <= 50 字)
          c) 或半数以上 cell 空白但不是全空
      - 可用 cell 数量长度弹性(非全填满),说明可循环

    命中的 section 会挂:
        sec.loop_table_schema = {
            "overview_para_idx": <前导综述段 para_idx>,
            "header_cols": [...],
            "sample_row_indices": [...],
            "data_row_template": [...],  # 空模板 row
            "table_cell_key": (row_idx, col_idx, nt_idx),  # 定位嵌套表格
        }
    返回: 命中数
    """
    try:
        from docx.oxml.ns import qn
    except Exception:
        return 0

    # body_cell.tables 仅给出直系嵌套表,按出现顺序与段落在 XML 中交错
    tbl_tag = qn('w:tbl')
    p_tag = qn('w:p')

    # 构建 body_cell 顺序: [(kind, idx-in-its-list)]
    # 同时需要知道每个 tbl 紧邻的"前面最近非空段落"
    # (作为综述段候选)
    order = []
    p_cursor = 0
    tbl_cursor = 0
    for child in list(body_cell._element):
        if child.tag == p_tag:
            order.append(('p', p_cursor))
            p_cursor += 1
        elif child.tag == tbl_tag:
            order.append(('tbl', tbl_cursor))
            tbl_cursor += 1

    # para_idx -> section
    pidx_to_section = {}
    for sec in sections:
        for p in sec.paragraphs:
            pidx_to_section[p.para_idx] = sec

    nested_tables = list(body_cell.tables)
    hits = 0

    for i, (kind, idx) in enumerate(order):
        if kind != 'tbl':
            continue
        if idx >= len(nested_tables):
            continue
        tbl = nested_tables[idx]
        rows = tbl.rows
        if len(rows) < 3:
            continue

        header_cells = [c.text.strip() for c in rows[0].cells]
        if len(header_cells) < 2:
            continue
        if not all(h and len(h) <= 30 for h in header_cells):
            continue

        # 扫描 row 2+: 找示例 row / 模板 row
        sample_indices = []
        template_row = None
        header_set = set(h for h in header_cells if h)
        for r_idx in range(1, len(rows)):
            row_texts = [c.text.strip() for c in rows[r_idx].cells]
            non_empty = [t for t in row_texts if t]
            if not non_empty:
                # 全空 row: 候选 data_row_template
                if template_row is None:
                    template_row = row_texts
                continue

            # a) 列头复读
            overlap = sum(1 for t in row_texts if t in header_set)
            is_header_repeat = overlap >= max(1, len(header_cells) // 2)

            # b) placeholder 句式: 短句 + 含 "如...请...说明"/"需要时"
            placeholder_text = any(
                bool(
                    t and (
                        re.search(r"如[^,。]{0,20}(?:请|需|应)[^,。]{0,30}说明", t)
                        or re.search(r"(?:如未|如有|若.{0,6})(?:请|需|应)", t)
                    )
                )
                for t in row_texts
            )

            # c) 半数 cell 空白但非全空
            empty_ratio = sum(1 for t in row_texts if not t) / max(1, len(row_texts))
            is_half_empty = 0 < empty_ratio < 1 and empty_ratio >= 0.4

            if is_header_repeat or placeholder_text or is_half_empty:
                sample_indices.append(r_idx)

        if not sample_indices:
            continue

        # 找综述段: i 之前最近的 p-kind 段落 (在同 section)
        overview_pidx = None
        for j in range(i - 1, -1, -1):
            k2, idx2 = order[j]
            if k2 == 'p' and idx2 < len(body_paras):
                p = body_paras[idx2]
                if p.text and p.text.strip():
                    overview_pidx = p.para_idx
                    break
        if overview_pidx is None:
            continue

        sec = pidx_to_section.get(overview_pidx)
        if sec is None:
            continue

        sec.loop_table_schema = {
            "overview_para_idx": overview_pidx,
            "header_cols": header_cells,
            "sample_row_indices": sample_indices,
            "data_row_template": template_row or [""] * len(header_cells),
            "nested_table_index": idx,
        }
        hits += 1

    return hits


def decompose_template(doc) -> dict[str, Any]:
    """
    主入口：拆解整个模板文档。

    返回:
        {
            "body_sections": list[SectionInfo],  # R3 主体的章节列表
            "body_paragraphs": list[ParaInfo],    # R3 所有段落分类
            "header_paragraphs": list[ParaInfo],  # R1 头部段落分类
            "stats": {
                "total_paras": int,
                "skeleton_count": int,
                "content_count": int,
                "section_count": int,
                "composite_paragraphs": int,  # 命中"段落+紧邻表格"的骨架段数
            }
        }
    """
    main_table = doc.tables[0]

    # R3 = 主体
    body_cell = main_table.rows[3].cells[0]
    body_paras = classify_cell_paragraphs(body_cell)

    # 标记"段落后紧跟嵌套表格"——composite 结构的底层事实
    _scan_para_followed_by_table(body_cell, body_paras)

    body_sections = group_into_sections(body_paras)

    # 基于上述事实,把 composite 段落挂到所属 section
    composite_count = _annotate_composite_sections(body_sections)

    # V14-v3: 识别循环表 schema (综述段 + 循环子表)
    try:
        loop_table_count = _detect_loop_table_schemas(
            body_cell, body_paras, body_sections
        )
    except Exception:
        loop_table_count = 0

    # R1 = 头部（包含客户名称、授信方案等标签字段和复选框）
    header_paras = []
    if len(main_table.rows) > 1:
        header_cell = main_table.rows[1].cells[0]
        header_paras = classify_cell_paragraphs(header_cell)

    skel_count = sum(1 for p in body_paras if p.role == TemplateRole.SKELETON)
    cont_count = sum(1 for p in body_paras if p.role == TemplateRole.CONTENT)

    return {
        "body_sections": body_sections,
        "body_paragraphs": body_paras,
        "header_paragraphs": header_paras,
        "stats": {
            "total_paras": len(body_paras),
            "skeleton_count": skel_count,
            "content_count": cont_count,
            "section_count": len(body_sections),
            "composite_paragraphs": composite_count,
            "loop_table_sections": loop_table_count,
        },
    }


# ═══════════════════════════════════════════
# 模板泄漏检测（通用性方法）
# ═══════════════════════════════════════════

def compute_similarity(text1: str, text2: str) -> float:
    """计算两段文本的相似度 (0~1)"""
    if not text1 or not text2:
        return 0.0
    # 去空白后比较，避免格式差异干扰
    a = re.sub(r'\s+', '', text1)
    b = re.sub(r'\s+', '', text2)
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()


def detect_leakage(
    generated_paragraphs: dict[int, str],
    template_paragraphs: list[ParaInfo],
    threshold: float = 0.5,
) -> list[dict]:
    """
    检测生成内容是否泄漏了模板原文。

    Args:
        generated_paragraphs: {para_idx: generated_text}
        template_paragraphs: 模板原始段落分类列表
        threshold: 相似度阈值，超过则判定为泄漏

    Returns:
        [{para_idx, similarity, template_text, generated_text}, ...]
    """
    leaks = []
    tpl_map = {p.para_idx: p for p in template_paragraphs}

    for idx, gen_text in generated_paragraphs.items():
        tpl_para = tpl_map.get(idx)
        if not tpl_para:
            continue
        # 检查所有内容角色（CONTENT/EXAMPLE/INSTRUCTION）
        _content_roles = {TemplateRole.CONTENT, TemplateRole.EXAMPLE, TemplateRole.INSTRUCTION}
        if tpl_para.role not in _content_roles:
            continue

        tpl_text = tpl_para.text
        # 短文本不检测
        if len(re.sub(r'\s+', '', tpl_text)) < 20:
            continue
        if len(re.sub(r'\s+', '', gen_text)) < 20:
            continue

        sim = compute_similarity(gen_text, tpl_text)
        if sim > threshold:
            leaks.append({
                "para_idx": idx,
                "similarity": round(sim, 3),
                "template_snippet": tpl_text[:100],
                "generated_snippet": gen_text[:100],
            })

    return leaks


# ═══════════════════════════════════════════
# 章节→KB维度映射
# ═══════════════════════════════════════════

# 通过章节标题关键词推断需要的 KB 维度
_SECTION_DIMENSION_MAP = {
    # 关键词 → 维度列表
    "借款人简介": ["basic_info", "controller", "shareholders", "credit_history", "risk"],
    "借款人概况": ["basic_info", "business"],
    "面访": ["controller"],
    "走访": ["business"],
    "股权融资": ["financing", "shareholders"],
    "历史授信": ["credit_history"],
    "授后": ["credit_history"],
    "风险信号": ["risk", "credit_history"],
    "负面信息": ["risk"],
    "PD评级": ["credit_history", "risk"],
    "法定代表人": ["controller", "basic_info"],
    "实控人": ["controller"],
    "资产情况": ["assets"],
    "关联企业": ["affiliates"],
    "经营": ["business", "supply_chain", "orders"],
    "财务": ["basic_info", "financing"],
    "产品": ["business"],
    "上下游": ["supply_chain"],
    "订单": ["orders"],
    "融资": ["financing", "credit_history"],
    "对外担保": ["financing", "risk"],
    "征信": ["credit_history", "risk"],
    "科技": ["r_and_d", "patents"],
    "研发": ["r_and_d"],
    "专利": ["patents"],
    "综合效益": ["business", "customer_manager", "bank_flows"],
    "借款原因": ["financing", "business", "orders"],
    "还款": ["financing", "business"],
    "担保分析": ["assets", "financing", "controller"],
    "抵押": ["assets"],
    "保证": ["financing"],
    "授信申报结论": ["basic_info", "business", "financing", "risk",
                      "assets", "credit_history", "bank_flows"],
    "六要素": ["financing", "business", "bank_flows", "assets"],
}


def infer_section_dimensions(section: SectionInfo) -> list[str]:
    """根据章节标题和骨架内容推断需要的 KB 维度"""
    text = section.title + " " + " ".join(section.skeleton_lines)
    dims = set()

    for keyword, dim_list in _SECTION_DIMENSION_MAP.items():
        if keyword in text:
            dims.update(dim_list)

    # 兜底：至少给 basic_info
    if not dims:
        dims.add("basic_info")

    return sorted(dims)
