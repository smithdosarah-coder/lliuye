# -*- coding: utf-8 -*-
"""V16 Phase 4.5 治本 hot-fix: 兴业资管 docx · lint byte-level diff 暴露的 18 HIGH 残留清理.

背景:
  lint script `liuye_docx_residue_check.py` (byte-level diff, 不依赖 keyword 黑名单)
  跑出兴业 18 HIGH residue:
    - P179: 长资本叙述段含 '10.5亿元' / '12.12亿元' (canonical 5亿元 substring match) +
            德勤华永会计师事务所 + 195,000.00万元 等真实审计/资本数字
    - P18: 兴业信托经营范围长段 (declared 但未 placeholder 化, 也含 686.86亿元 等)
    - P241: '不良资产管理行业属于资金密集型行业...' (canonical 不良资产管理)
    - P247/P249: 商务服务业 raw 多次出现 (canonical industry)
    - P251/P275: 5亿元授信额度 raw (declared CREDIT_AMOUNT 但未替换)
    - P256/P257: 历史财务大段含 51.52亿/83.25亿等 (canonical 5亿元 substring)
    - P77/P82: 行业百科段含 '不良资产管理' 多次 (canonical industry_tag_long)
    - T0R1C0P0/C1P0/C1P3/C2P2: 额度表 cell · 5亿元 raw
    - T2R5C3P0/T2R6C3P0: 持股比例 cell · 100% raw (declared CLIENT_SHARE_PCT_PRIMARY)
    - T34R0C0P0: '我行商务服务业准入要求' 表头

策略 (deterministic 治本 · 与 jingwei phase4 一致):
  - 段落级整段重写 (PARAGRAPH_OVERRIDES) · 用 placeholder + 通用叙述消除 specific 字面
  - 表格 cell paragraph 替换 (TABLE_CELL_PARA_REPLACEMENTS)
  - 长 narrative 段落 (P77/P82/P179/P241/P256/P257/P18) 改为 schema 已有 narrative key
    或重新通用化文案 (去除任何 5亿/195000万/不良资产管理 等真实数据)

不动:
  - 财务表行/列具体数字 (走 FILL handler) · 仅碰到的 cell 报警才动
  - 涉诉案件表 T43 (走 REWRITE handler)

idempotent: 段落已含目标 placeholder 则 skip
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

DOCX_PATH = Path(__file__).resolve().parent.parent / "samples" / "兴业资管_对公成稿B.docx"


# ────────────────────────────────────────────────────────────────────────────
# 段落级整段替换 (paragraph_index → new_text)
# ────────────────────────────────────────────────────────────────────────────
PARAGRAPH_REPLACEMENTS: dict[int, str] = {
    # P18: 兴业信托经营范围长段 (含 686.86亿元 / 9.35亿元 等真实数字 + '兴业信托' 多次)
    # 原: '兴业信托经营范围包括资金信托、动产信托...截至2021年末资产总额686.86亿元...实现利润总额11.92亿元'
    # 替换策略: 集团叙事 narrative
    18: "{{CLIENT_GROUP_FULL_NAME}}经营范围与战略目标具备综合化、规模化特征，财务表现稳健，{{CLIENT_HISTORY_NARRATIVE}}",

    # P77: 不良资产管理行业百科段 (5 次 '不良资产管理' raw 字面)
    # 原: '不良资产管理公司是特指专业承接、处置不良债务的资产管理公司...'
    # 替换策略: 行业 narrative
    77: "{{CLIENT_INDUSTRY_NARRATIVE}}",

    # P82: 行业格局段 ('不良资产管理' 多次)
    # 原: '至此，我国资产管理行业形成了国有四家金融资产管理公司...不良资产管理的经验丰富...'
    # 替换策略: 行业地位 narrative
    82: "{{CLIENT_INDUSTRY_POSITION}}",

    # P179: 资本与审计长段 (含 195,000.00万元 / 300,000.00万元 / 105,000.00万元 / 10.5亿元 /
    #                       12.12亿元 / 德勤华永会计师事务所 等真实数据)
    # 原: '公司最近三年又一期末的实收资本分别为195,000.00万元...资本金已经德勤华永会计师事务所...'
    # 替换策略: 长 narrative + placeholder
    179: "公司近期实收资本结构稳定，注册及实收资本规模与股权架构与公司战略目标相匹配。"
         "{{CLIENT_PARENT_SHORT_NAME}}作为主要发起人持有相应股份。"
         "资本金已由各发起人以现金缴足并经具备执业资格的会计师事务所审验，验资报告齐备，资本来源合规。",

    # P241: 不良资产管理行业资金密集型段 (canonical 不良资产管理 + 资产管理 多次)
    # 原: '不良资产管理行业属于资金密集型行业，自有资金无法满足{{CLIENT_LONG_CORE_NAME}}业务发展...'
    # 替换策略: 用 INDUSTRY_TAG placeholder + 财务叙事
    241: "{{CLIENT_INDUSTRY_TAG}}行业属于资金密集型行业，自有资金无法满足{{CLIENT_LONG_CORE_NAME}}业务发展所需的资金，"
         "需要通过持续的外部融资来满足资金需求。最近三年及一期末，{{CLIENT_LONG_CORE_NAME}}的合并口径资产负债率维持在较高水平，"
         "随着{{CLIENT_LONG_CORE_NAME}}业务规模的扩大，资产负债率有所上升。未来随着{{CLIENT_LONG_CORE_NAME}}业务规模的扩张，"
         "资产负债率可能进一步提高，进而增加{{CLIENT_LONG_CORE_NAME}}的长期偿债风险。",

    # P247: 商务服务业 raw 多次 (4 次)
    # 原: '{{CLIENT_FULL_NAME}}属于我行授信指引中的商务服务业。该行业整体原则为：商务服务业作为...'
    # 替换策略: 全部用 INDUSTRY_CATEGORY
    247: "{{CLIENT_FULL_NAME}}属于我行授信指引中的{{CLIENT_INDUSTRY_CATEGORY}}。该行业整体原则为：{{CLIENT_INDUSTRY_NARRATIVE}}",

    # P249: '整体来看，{{CLIENT_LONG_CORE_NAME}}的符合我行我2022年商务服务业行业准入条件'
    # 替换策略: INDUSTRY_CATEGORY
    249: "整体来看，{{CLIENT_LONG_CORE_NAME}}符合我行{{CLIENT_INDUSTRY_CATEGORY}}行业准入条件，具体如下：",

    # P251: 申请额度段 (5亿元 raw)
    # 原: '本次申请给予{{CLIENT_FULL_NAME}}总额5亿元综合授信额度（敞口5亿元）。敞口信用。具体品种如下：'
    251: "本次申请给予{{CLIENT_FULL_NAME}}总额{{CREDIT_AMOUNT}}综合授信额度（敞口{{CREDIT_EXPOSURE}}）。敞口信用。具体品种如下：",

    # P256: 历史购买/处置资产包大段 (含 51.52亿 / 83.25亿 / 63.86亿 / 47.31亿 等多 5亿元 substring)
    # 替换策略: 通用化叙述
    256: "从借款人对外公开披露的经营情况来看，借款人近三年通过收购及处置资产包形成的现金流动呈现出资金占用与回收并行的结构，"
         "经营性净现金流受业务节奏影响波动较大，整体表现与{{CLIENT_INDUSTRY_TAG}}行业的资金占用特征一致。",

    # P257: '从借款人最新提供的数据来看，2022年6月末实际投资资产包金额为25.5亿元...本次我行申请5亿元中短期流贷...'
    257: "根据借款人最新提供的经营数据测算，当年新增资产包收购预计带来一定规模的营运资金需求，需通过外部融资予以补充。"
         "本次我行申请的{{CREDIT_AMOUNT}}中短期流贷主要用于补充借款人购买资产包所需的流动资金及置换到期银行借款。",

    # P275: 申请说明段 (5亿元 raw 2 次)
    # 原: '{{CLIENT_FULL_NAME}}是经福建省人民政府批准...本次申请给予{{CLIENT_FULL_NAME}}总额5亿元综合授信额度（敞口5亿元）...'
    275: "{{CLIENT_FULL_NAME}}是经福建省人民政府批准、中国银保监会备案具有金融企业不良资产批量收购处置业务资质的省级资产管理公司。"
         "公司外部评级为AAA，且评级展望为稳定，展现了较好的市场认可和抗风险能力，属于我行积极支持的战略型客户。"
         "本次申请给予{{CLIENT_FULL_NAME}}总额{{CREDIT_AMOUNT}}综合授信额度（敞口{{CREDIT_EXPOSURE}}）。敞口信用。具体品种如下：",
}


# ────────────────────────────────────────────────────────────────────────────
# 表格 cell paragraph 级替换 (table_idx, row_idx, col_idx, para_idx) → new_text
# ────────────────────────────────────────────────────────────────────────────
TABLE_CELL_PARA_REPLACEMENTS: dict[tuple[int, int, int, int], str] = {
    # T0R1C0P0: 现有方案总额段 (5亿元 raw 2 次)
    (0, 1, 0, 0): "申请给予{{CLIENT_FULL_NAME}}总额{{CREDIT_AMOUNT}}综合授信额度（敞口{{CREDIT_EXPOSURE}}）。敞口信用。具体品种如下：",
    # T0R1C1P0: 本次申请方案总额段 (同上)
    (0, 1, 1, 0): "申请给予{{CLIENT_FULL_NAME}}总额{{CREDIT_AMOUNT}}综合授信额度（敞口{{CREDIT_EXPOSURE}}）。敞口信用。具体品种如下：",
    # T0R1C1P3: 债券投资额度段 (5亿元 raw)
    # 原: '（3）债券投资额度5亿元，提用期限1年，业务期限不超过5年，可循环使用。'
    (0, 1, 1, 3): "（3）债券投资额度按业务方案配置，提用期限1年，业务期限不超过5年，可循环使用。",
    # T0R1C2P2: 变化说明 (新增债券投资5亿元 raw)
    # 原: '（3）新增债券投资额度5亿元，提用期限1年，业务期限不超过5年，可循环使用'
    (0, 1, 2, 2): "（3）新增债券投资额度按业务方案配置，提用期限1年，业务期限不超过5年，可循环使用",

    # T2R5C3P0: 持股比例 cell · declared CLIENT_SHARE_PCT_PRIMARY 但未替换
    (2, 5, 3, 0): "{{CLIENT_SHARE_PCT_PRIMARY}}",
    # T2R6C3P0: 合计持股比例 (同上)
    (2, 6, 3, 0): "{{CLIENT_SHARE_PCT_PRIMARY}}",

    # T34R0C0P0: 表头 '我行商务服务业准入要求' (canonical 商务服务业)
    (34, 0, 0, 0): "我行{{CLIENT_INDUSTRY_CATEGORY}}准入要求",
}


def _replace_paragraph_text(para, new_text: str) -> None:
    """整段 paragraph.text 替换 · 用第一 run 装新文本 + 清空其余 run · 不动样式."""
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""


def main() -> int:
    if not DOCX_PATH.is_file():
        print(f"[ERROR] docx 不存在: {DOCX_PATH}")
        return 1

    doc = Document(DOCX_PATH)

    # ── 1. 段落替换 ──
    replaced_paras = 0
    skipped_paras = 0
    for idx, new_text in PARAGRAPH_REPLACEMENTS.items():
        if idx >= len(doc.paragraphs):
            print(f"[WARN] P{idx} 越界 (total {len(doc.paragraphs)}) · skip")
            continue
        para = doc.paragraphs[idx]
        old_text = para.text
        if old_text.strip() == new_text.strip():
            skipped_paras += 1
            print(f"[SKIP] P{idx} 已 placeholder 化")
            continue
        _replace_paragraph_text(para, new_text)
        replaced_paras += 1
        print(f"[OK] P{idx} replaced · old_len={len(old_text)} new_len={len(new_text)}")

    # ── 2. 表格 cell paragraph 级替换 ──
    replaced_cell_paras = 0
    skipped_cell_paras = 0
    for (ti, ri, ci, pi), new_text in TABLE_CELL_PARA_REPLACEMENTS.items():
        if ti >= len(doc.tables):
            print(f"[WARN] T{ti} 越界 (total {len(doc.tables)}) · skip")
            continue
        table = doc.tables[ti]
        if ri >= len(table.rows):
            print(f"[WARN] T{ti}R{ri} 越界 · skip")
            continue
        row = table.rows[ri]
        if ci >= len(row.cells):
            print(f"[WARN] T{ti}R{ri}C{ci} 越界 · skip")
            continue
        cell = row.cells[ci]
        if pi >= len(cell.paragraphs):
            print(f"[WARN] T{ti}R{ri}C{ci}P{pi} 越界 (cell paras={len(cell.paragraphs)}) · skip")
            continue
        para = cell.paragraphs[pi]
        old_text = para.text
        if old_text.strip() == new_text.strip():
            skipped_cell_paras += 1
            print(f"[SKIP] T{ti}R{ri}C{ci}P{pi} 已 placeholder 化")
            continue
        _replace_paragraph_text(para, new_text)
        replaced_cell_paras += 1
        print(f"[OK] T{ti}R{ri}C{ci}P{pi} replaced · old_len={len(old_text)} new_len={len(new_text)}")

    # ── 3. 保存 ──
    doc.save(DOCX_PATH)
    print()
    print(f"=== Phase 4.5 done · {replaced_paras} para replaced ({skipped_paras} skipped) · "
          f"{replaced_cell_paras} cell para replaced ({skipped_cell_paras} skipped) ===")
    return 0


if __name__ == "__main__":
    sys.exit(main())
