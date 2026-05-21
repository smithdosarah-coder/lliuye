# -*- coding: utf-8 -*-
"""V16 Phase 3 worker 1: 把 samples/兴业资管_对公成稿B.docx 改成 placeholder 模板.

设计 (跟 Phase 2 经纬 pilot 同模式 · scripts/v16_placeholderize_jingwei.py):
  - 整段 (paragraph.text) 替换 · 用第一个 run 写整段新文本 + 清空其余 run · 避开 cross-run quirks
  - 完整对照 metadata.json 的 placeholder_locations 字段
  - 替换顺序: 长字符串先 (避免 '兴业资产管理有限公司' 内含 '兴业' 时把核心词替换掉)
  - **不动样式 + 段落格式 + 列表序号** · 只改文本
  - **不动财务表数字** (195,000.00 / 138,542.32 等) · 由 REWRITE/FILL handler 处理
  - **不动 T43 涉诉案件表** · 大量 '兴业资产管理有限公司与...' 案件名 · 走 REWRITE 路径

兴业资管模板特殊 vs 经纬测绘:
  - 三级股权结构: 兴业资管 / 兴业国信资产管理有限公司 / 兴业银行股份有限公司
  - "兴业" 单字 ambiguous (命中 4 个实体) · 不替换 bare '兴业' · 只替换完整词
  - 大量 'XX资产管理有限公司' / 'XX资管' 内嵌结构 · 顺序严格 长→短 保证不破坏

用法:
  py scripts/v16_placeholderize_xingye.py
  - 默认改 samples/兴业资管_对公成稿B.docx (in-place)
  - 改完后用 v16_step1_extract.py 扫一遍验证 placeholder 都标 PLACEHOLDER label

run 一次即完成 · idempotent (再跑一次不会重复替换 · 因为 {{KEY}} 不在 substitution map 内)
"""
from __future__ import annotations

import io
import sys
from pathlib import Path
from docx import Document

# 使 stdout UTF-8 (Windows console 中文)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

DOCX_PATH = Path(__file__).resolve().parent.parent / "samples" / "兴业资管_对公成稿B.docx"

# 替换规则: (硬编码字符串, placeholder) · 顺序 = 长→短 (避免子串被先吃掉)
# 选只在客户档案中明确的硬编码 · 财务数据 / 业务叙事段保留原文 (REWRITE handler 处理)
SUBSTITUTIONS: list[tuple[str, str]] = [
    # ─── 公司名 (最长先, 避免 '兴业资管' 在 '兴业资产管理有限公司' 内被先替换) ───
    # 注意: '兴业资产管理有限公司' 子串在 '兴业财富资产管理有限公司' 内也出现 (T43R23C1P0)
    # 但 T43 是涉诉案件表 · 不属于此 paragraph-level 替换路径覆盖范围 (那些是 cell 内)
    # 走 _walk_paragraphs 时表格 cell 也覆盖 · 但 T43 整段 PARAGRAPH_OVERRIDES 不命中
    # 因此 SUBSTITUTIONS 会把 T43 cell 内的 '兴业资产管理有限公司' 也替换 → 这是 desired
    # (案件标题里出现客户全称属于客户身份信息 · 替换无副作用)
    ("兴业资产管理有限公司", "{{CLIENT_FULL_NAME}}"),

    # ─── 母公司全称 (长→短) ───
    # '兴业国信资产管理有限公司' 含 '资产管理有限公司' · 先替换避免后续误命中
    ("兴业国信资产管理有限公司", "{{CLIENT_PARENT_FULL_NAME}}"),

    # ─── 集团全称 (兴业银行股份有限公司) ───
    ("兴业银行股份有限公司", "{{CLIENT_GROUP_FULL_NAME}}"),

    # ─── 注册资本 (长→短 · 带 '人民币' 前缀的先替换) ───
    ("人民币195,000 万元", "{{CLIENT_REGISTERED_CAPITAL}}"),
    # 195000万元 (T2R5C1P0 / T2R6C1P0 不含逗号 · 替换为同 placeholder)
    # 注意: 不替换 195,000.00 (财务表数字 · 走 FILL handler)
    ("195000万元", "{{CLIENT_REGISTERED_CAPITAL}}"),

    # ─── 注册地址 (T2R0C3P0 only · 完整带括号) ───
    ("福建省福州市马尾区快安路8 号6A 号（自贸试验区内）", "{{CLIENT_REGISTERED_ADDRESS}}"),

    # ─── 成立时间 (T2R0C1P0 only) ───
    ("2017 年2 月20 日", "{{CLIENT_ESTABLISHMENT_DATE}}"),

    # ─── 法人 ───
    ("谢斌", "{{CLIENT_LEGAL_REP}}"),

    # ─── 公司简称 (兴业资管 · 短于全名 · 必须在全名之后替换) ───
    # 关键约束: 此时全名已替换为 {{CLIENT_FULL_NAME}} · 不会再命中 '兴业资管'
    ("兴业资管", "{{CLIENT_LONG_CORE_NAME}}"),

    # ─── 母公司简称 (兴业国信 · 短于母公司全称) ───
    ("兴业国信", "{{CLIENT_PARENT_SHORT_NAME}}"),

    # ─── 集团简称 (兴业银行 · 短于集团全称) ───
    # 注意: 不会命中 '兴业银行股份有限公司' 因前面已替换为 {{CLIENT_GROUP_FULL_NAME}}
    ("兴业银行", "{{CLIENT_GROUP_SHORT_NAME}}"),

    # ─── 主要股东 (T2R5C0P0 only · 与母公司全称同名 → 上面已替换 · 此条 redundant 但保留可读性) ───
    # '兴业国信资产管理有限公司' 已被替换为 {{CLIENT_PARENT_FULL_NAME}}
    # 在 metadata.json 中 CLIENT_SHAREHOLDER_PRIMARY 与 CLIENT_PARENT_FULL_NAME 同值
    # T2R5C0P0 替换后会是 {{CLIENT_PARENT_FULL_NAME}} 而非 {{CLIENT_SHAREHOLDER_PRIMARY}}
    # generator 走 placeholder_replace 时 · alias_map 会双向解析 · 不影响最终填值
    # (这跟经纬模板 T2R5C0P0/T2R5C1P0 同处理 · CLIENT_SHAREHOLDER_PRIMARY 跟 CLIENT_PARENT_FULL_NAME 同值)
]

# 整段重写覆盖 (优先级高于 SUBSTITUTIONS · 用于 1 段含多个 placeholder 的复杂段)
# location → 新 paragraph text
PARAGRAPH_OVERRIDES: dict[str, str] = {
    # 标题: 公司全名 + 授信额度 + 敞口
    "P0": "关于{{CLIENT_FULL_NAME}}{{CREDIT_AMOUNT}}综合授信额度（敞口{{CREDIT_EXPOSURE}}）授信报告",
    # 客户全称行
    "P2": "授信客户全称：{{CLIENT_FULL_NAME}}（简称“{{CLIENT_LONG_CORE_NAME}}”）",
    # 行业政策行
    "P3": "我行投向政策对应行业：{{CLIENT_INDUSTRY_CATEGORY}}",
}


def _replace_paragraph_text(para, new_text: str) -> bool:
    """在保留 run 格式的前提下整段重写 · 与 v16_generator._set_paragraph_text 一致策略.

    返回是否改动.
    """
    if (para.text or "") == new_text:
        return False
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return True
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return True


def _walk_paragraphs(doc):
    """yield (location, paragraph) — 同 v16_step1_extract.extract_elements 逻辑."""
    for pi, p in enumerate(doc.paragraphs):
        yield f"P{pi}", p

    def _walk_table(tbl, loc_prefix: str):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for ppi, p in enumerate(cell.paragraphs):
                    yield f"{loc_prefix}R{ri}C{ci}P{ppi}", p
                for sub in cell.tables:
                    yield from _walk_table(sub, f"{loc_prefix}R{ri}C{ci}NT")

    for ti, tbl in enumerate(doc.tables):
        yield from _walk_table(tbl, f"T{ti}")


def main():
    if not DOCX_PATH.is_file():
        print(f"[fatal] docx 不存在: {DOCX_PATH}", file=sys.stderr)
        sys.exit(1)
    doc = Document(str(DOCX_PATH))

    stats = {"overrides": 0, "substitutions": 0, "checked_paragraphs": 0}
    hit_locs: list[str] = []

    for loc, para in _walk_paragraphs(doc):
        stats["checked_paragraphs"] += 1
        original_text = para.text or ""

        # 1) PARAGRAPH_OVERRIDES 优先 (整段重写)
        if loc in PARAGRAPH_OVERRIDES:
            new_text = PARAGRAPH_OVERRIDES[loc]
            if _replace_paragraph_text(para, new_text):
                stats["overrides"] += 1
                hit_locs.append(loc)
                print(f"  [override] {loc}: {original_text[:60]} → {new_text[:60]}")
            continue

        # 2) SUBSTITUTIONS · 顺序应用 · 不破坏已替换的 placeholder
        new_text = original_text
        for old, ph in SUBSTITUTIONS:
            if old in new_text and ph not in new_text:
                # ph not in new_text 保证 idempotent (再跑不重复)
                new_text = new_text.replace(old, ph)
            elif old in new_text:
                # 已含 placeholder 仍含原文 · 走单次替换 (常见: cross-substr 同时命中)
                new_text = new_text.replace(old, ph)
        if new_text != original_text:
            if _replace_paragraph_text(para, new_text):
                stats["substitutions"] += 1
                hit_locs.append(loc)
                print(f"  [substit] {loc}: {original_text[:60]} → {new_text[:60]}")

    doc.save(str(DOCX_PATH))
    print()
    print(f"[done] checked {stats['checked_paragraphs']} paragraphs / "
          f"{stats['overrides']} overrides / {stats['substitutions']} substitutions")
    print(f"  改动 location 数 (unique): {len(set(hit_locs))}")
    print(f"  改后 docx: {DOCX_PATH}")


if __name__ == "__main__":
    main()
