# -*- coding: utf-8 -*-
"""V16 Phase 3 worker 2: 把 samples/普惠申报书_骨架型.docx 改成 placeholder 模板.

设计 (Branch A · skeleton 路径):
  - 普惠申报书是纯骨架模板 · 无真实客户字面 · 全文是 form fields + 例句指引文
  - 不需要 SUBSTITUTIONS (no client name to replace)
  - 仅 PARAGRAPH_OVERRIDES 在 form 头部 (T0R1/T0R2) 字段标签后注入 {{KEY}}
  - 让 framework REPLACE handler 用 runtime client_metadata 填值

vs Phase 2 经纬 pilot 差异:
  - 经纬: 真客户成稿 · 1023 element · 29 location · 56 occurrence · 10 unique key
  - 普惠: skeleton · 495 element · 6 location · 6 occurrence · 6 unique key
  - 普惠不需 SUBSTITUTIONS (无 '经纬'/'郑志煌' 等真客户字面)

用法:
  py scripts/v16_placeholderize_puhui.py
  - 默认改 samples/普惠申报书_骨架型.docx (in-place)
  - 改完后用 v16_step1_extract.py 扫一遍验证 placeholder 都标 PLACEHOLDER label

run 一次即完成 · idempotent (再跑一次不会重复替换 · 已含 {{KEY}} 跳过)
"""
from __future__ import annotations

import io
import sys
from pathlib import Path
from docx import Document

# 使 stdout UTF-8 (Windows console 中文)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

DOCX_PATH = Path(__file__).resolve().parent.parent / "samples" / "普惠申报书_骨架型.docx"

# Branch A skeleton: 仅整段重写 form 头部字段标签
# location → 新 paragraph text · 仅注入 {{KEY}} placeholder · 保留字段标签 + 单位 + 分隔符
PARAGRAPH_OVERRIDES: dict[str, str] = {
    # 客户名称 field (申报表行 1 · 顶部字段)
    "T0R1C0P0": "客户名称：{{CLIENT_FULL_NAME}}",
    # 集团名称 field (如涉及)
    "T0R1C0P1": "集团名称（如涉及）：{{CLIENT_GROUP_FULL_NAME}}",
    # 国标行业 field (门类-大类-中类-小类)
    "T0R1C0P5": "国标行业（门类-大类-中类-小类）：{{CLIENT_INDUSTRY_FULL}}",
    # 申报额度 + 申报敞口 (上一期保留示例值 · 历史数据归 facts pipeline)
    "T0R2C0P0": "申报额度：{{CREDIT_AMOUNT}}；申报敞口 {{CREDIT_EXPOSURE}}；【上一期授信敞口（若有）：  500 万元】",
    # 业务期限保留示例 (12 个月是普惠通用值) · 仅 placeholder 化 授信期限
    "T0R2C0P1": "业务品种：一般短期流动资金贷款；业务期限： 12个月；授信期限：{{CREDIT_PERIOD}}",
    # PD 评级 · 上一期保留示例值
    "T0R2C0P3": "PD评级：{{PD_RATING}}级【上一期授信评级： 7 级】",
}


def _replace_paragraph_text(para, new_text: str) -> bool:
    """在保留 run 格式的前提下整段重写 · 与 v16_generator._set_paragraph_text 一致策略.

    返回是否改动.
    """
    if (para.text or "") == new_text:
        return False
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return True
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return True


def _walk_paragraphs(doc):
    """yield (location, paragraph) — 同 v16_step1_extract.extract_elements 逻辑."""
    for pi, p in enumerate(doc.paragraphs):
        yield f"P{pi}", p

    def _walk_table(tbl, loc_prefix: str):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for ppi, p in enumerate(cell.paragraphs):
                    yield f"{loc_prefix}R{ri}C{ci}P{ppi}", p
                for sub in cell.tables:
                    yield from _walk_table(sub, f"{loc_prefix}R{ri}C{ci}NT")

    for ti, tbl in enumerate(doc.tables):
        yield from _walk_table(tbl, f"T{ti}")


def main():
    if not DOCX_PATH.is_file():
        print(f"[fatal] docx 不存在: {DOCX_PATH}", file=sys.stderr)
        sys.exit(1)
    doc = Document(str(DOCX_PATH))

    stats = {"overrides": 0, "checked_paragraphs": 0, "skipped_idempotent": 0}
    hit_locs: list[str] = []
    unique_keys: set[str] = set()
    target_locs = set(PARAGRAPH_OVERRIDES.keys())
    seen_target_locs: set[str] = set()

    for loc, para in _walk_paragraphs(doc):
        stats["checked_paragraphs"] += 1
        original_text = para.text or ""

        if loc in PARAGRAPH_OVERRIDES:
            seen_target_locs.add(loc)
            new_text = PARAGRAPH_OVERRIDES[loc]
            # Idempotency check: 如果原文已经包含目标 placeholder set · 跳过
            if original_text == new_text:
                stats["skipped_idempotent"] += 1
                print(f"  [skip-idempotent] {loc}: 已是 placeholder 形式")
                continue
            if _replace_paragraph_text(para, new_text):
                stats["overrides"] += 1
                hit_locs.append(loc)
                # 抓 {{KEY}} 用于统计
                import re
                for key in re.findall(r"\{\{([A-Z_][A-Z0-9_]+)\}\}", new_text):
                    unique_keys.add(key)
                print(f"  [override] {loc}: {original_text[:60]} → {new_text[:80]}")

    # 验证所有 target location 都被处理 (没有 location 丢失 · 例如 cell 嵌套结构变了)
    missing_targets = target_locs - seen_target_locs
    if missing_targets:
        print(f"[warn] 以下 target location 未在 docx walk 时遇到 (可能 docx 被改): {missing_targets}")

    doc.save(str(DOCX_PATH))
    print()
    print(f"[done] checked {stats['checked_paragraphs']} paragraphs / "
          f"{stats['overrides']} overrides / {stats['skipped_idempotent']} skipped-idempotent")
    print(f"  改动 location 数 (unique): {len(set(hit_locs))}")
    print(f"  unique placeholder key 数: {len(unique_keys)} · {sorted(unique_keys)}")
    print(f"  改后 docx: {DOCX_PATH}")


if __name__ == "__main__":
    main()
