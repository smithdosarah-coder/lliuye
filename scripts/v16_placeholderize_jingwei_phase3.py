# -*- coding: utf-8 -*-
"""V16 Phase 3 治本 hot-fix: 把经纬测绘 docx 大段业务描述也 placeholder 化.

背景:
  - DP002 dogfood 报告残留 '测绘地理信息高新技术服务企业' / '1988 年省经贸委直属事业单位' /
    '工程咨询服务' 等业务描述硬编码 (见 docx_scan: P11/P16/P104/P176/P200 多处)
  - Phase 2 只 placeholder 化公司名/法人/资本/地址等 metadata 级字段
  - 长段业务描述原计划走 REWRITE 路径但 LLM 漏改 → 残留泄露

Phase 3 策略 (deterministic 优先 · 治本不治标):
  - 整段替换为 placeholder-only 模板 · 段落内容由 client_metadata 业务描述大段字段填充
  - 选 6 关键段落: P11 (主描述) / P16 (集团沿革) / P12 (资质) / P64+P99+P171 (公司沿革) /
                  P59-P61+P67 (行业上下文) / P104 (环境影响) / P176+P200 (授信结论)
  - 表格 cell T2R3C1/C2/C3 (经营范围 · 3 重复) / T2R3C6 (资质证书表)

不动:
  - 财务分析段 P51-54 / P115-144 (太 sample-specific · DP00X 财务不同)
  - 员工/股东具体名字段 P24/P31-34 (LLM REWRITE 处理)
  - 上下游/供应链段 P72-92 (sample-specific)

idempotent: 段落已含 {{KEY}} 则 skip
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

DOCX_PATH = Path(__file__).resolve().parent.parent / "samples" / "经纬测绘_对公成稿A.docx"


# ────────────────────────────────────────────────────────────────────────────
# 段落级整段替换 (paragraph_index → new_text)
# 新模板必须 placeholder-only 不留任何业务术语硬编码
# ────────────────────────────────────────────────────────────────────────────
PARAGRAPH_REPLACEMENTS: dict[int, str] = {
    # P11: 公司主描述 (原: '{{FULL}}, 是 {{GROUP}} 权属子公司, 是一家集地理信息数据采集...
    # 综合性甲级测绘地理信息高新技术服务企业, 属福建省国资体系中唯一一家测绘地理信息企业')
    11: "{{CLIENT_FULL_NAME}}，是{{CLIENT_GROUP_FULL_NAME}}权属子公司，{{CLIENT_BUSINESS_DESC_LONG}}",

    # P12: 资质证书叙述 (原: '公司目前已经获得福建省高新技术企业证书、福建省科技型企业证书...')
    12: "{{CLIENT_QUALIFICATION_DETAIL}}",

    # P16: 集团沿革 (原: '{{GROUP}} 前身为成立于 1988 年的原省经贸委直属事业单位, 2010 年 7 月省政府...
    # 是全国名列前茅、福建规模最大、实力最强的工程咨询服务和项目管理专业机构')
    16: "{{CLIENT_GROUP_FULL_NAME}}{{CLIENT_HISTORY_NARRATIVE}}",

    # P59: 行业定义 (原: '测绘是以计算机技术、光电技术... 测绘是指对自然地理要素或者地表人工设施...')
    # 这段是纯行业百科 · 整段替换为行业叙事 placeholder
    59: "{{CLIENT_INDUSTRY_NARRATIVE}}",

    # P60-61: 行业产值数据 (硬编码 2022 年总产值 10000 亿 · 2017 年起 20%) · 删除整段 → 由 P59 覆盖
    60: "",
    61: "",

    # P63: 行业风险/政策方向 (太特定 · 但提及"勘测/计量/测绘" · 整段替换)
    63: "{{CLIENT_INDUSTRY_POSITION}}",

    # P64: 公司沿革 (原: '公司成立于 2005 年, 至今已有 15 年历史。目前公司致力于测绘信息的采集及系统研发...')
    64: "{{CLIENT_HISTORY_NARRATIVE}}",

    # P65: '行业成熟度: 测绘行业属于知识密集型企业' · 行业一句话定位
    65: "3、行业成熟度：{{CLIENT_INDUSTRY_POSITION}}",

    # P67: 行业政策 (硬编码 '新《测绘法》/《时空政务地理信息应用服务接口技术规范》/《江苏省测绘地理信息条例》')
    67: "5、行业相关政策：{{CLIENT_INDUSTRY_NARRATIVE}}",

    # P99: 公司沿革重复段 (原: '公司成立于 2005 年, 前身权属于福建省测绘地理信息局...')
    99: "{{CLIENT_HISTORY_NARRATIVE}}",

    # P100: 公司愿景 (原: '根据 {{CLIENT_LONG_CORE_NAME}} 对外公布的企业愿景目标为: 公司以打造全省第一、
    # 全国一流的测绘地理信息产业为己任...')
    100: "根据{{CLIENT_LONG_CORE_NAME}}对外公布的企业愿景目标，公司聚焦{{CLIENT_INDUSTRY_TAG}}领域，{{CLIENT_INDUSTRY_POSITION}}",

    # P104: 环境影响 (原: '{{FULL}} 属于商务服务业主营业务为提供测绘服务及其他专业技术服务不涉及无生产产品制造...')
    104: "{{CLIENT_FULL_NAME}}属于{{CLIENT_INDUSTRY_CATEGORY}}主营业务为{{CLIENT_BUSINESS_DESC}}，不涉及生产制造，对环境不会造成直接影响。",

    # P171: 长段公司情况 + 资质重复 (原: '公司成立于 2005 年, 至今已有 15 年历史。... 公司作为一家集地理信息数据采集...
    # 已获得福建省高新技术企业证书... 60 多项软件著作权...')
    171: "{{CLIENT_HISTORY_NARRATIVE}}{{CLIENT_QUALIFICATION_DETAIL}}{{CLIENT_LONG_CORE_NAME}}作为{{CLIENT_INDUSTRY_TAG}}领域企业，{{CLIENT_INDUSTRY_POSITION}}",

    # P176: 授信结论开头 (原: '{{FULL}} 是福建省内实力较强、专业性较高的, 集高端采集、处理、分析、应用于一体的
    # 综合性甲级测绘地理信息技术服务企业。随着近年来互联网经济、大数据应用...')
    176: "{{CLIENT_FULL_NAME}}{{CLIENT_INDUSTRY_POSITION}}{{CLIENT_BUSINESS_DESC_LONG}}公司依托{{CLIENT_GROUP_FULL_NAME}}的整体实力，在人力、物力、财力方面将得到有效的支持。鉴于该客户资信情况较好，且资金用途真实，",

    # P200: 授信总结 (原: '{{FULL}} 是福建省内实力较强、专业性较高的, 集高端采集... 综合性甲级测绘地理信息技术服务企业')
    200: "   {{CLIENT_FULL_NAME}}{{CLIENT_INDUSTRY_POSITION}}{{CLIENT_BUSINESS_DESC_LONG}}{{CLIENT_PARENT_FULL_NAME}}为{{CLIENT_GROUP_FULL_NAME}}旗下企业，{{CLIENT_GROUP_FULL_NAME}}作为省属一级集团整体抗风险能力强。综上，我行拟给与该客户综合授信额度{{CREDIT_AMOUNT}}（敞口{{CREDIT_EXPOSURE}}），额度循环使用。授信期限{{CREDIT_PERIOD}}，敞口部分由{{CLIENT_FULL_NAME}}自身信用担保。具体品种分布如下：",

    # ── Phase 3 二次清理 · 长段财务分析 / 上下游 / 担保结论 (彻底清残留) ──

    # P21: 子公司列表 (硬编码 '福建省机电设备招标有限公司' 等 9 个子公司名) · 改为通用集团子公司描述
    21: "公司由{{CLIENT_GROUP_FULL_NAME}}发起组建，{{CLIENT_FULL_NAME}}为集团权属子公司之一。集团已建立较为完善的子公司治理体系，{{CLIENT_FULL_NAME}}在集团内承担{{CLIENT_INDUSTRY_TAG}}主业。",

    # P24: 职能机构 + 员工结构 (原 '工程部、移动部、地信部...' + '注册测绘师资格') · 通用化
    24: "公司本部设有综合管理部、业务部、财务部、办公室等综合职能机构（具体情况如下表所示）。{{CLIENT_BUSINESS_STRATEGY_DESC}}",

    # P31: 人才结构 (原 '注册测绘师') · 通用化
    31: "公司目前人才结构较为合理，{{CLIENT_BUSINESS_QUALIFICATION_DESC}}主要骨干人员有：",

    # P32: 法人简介 (含 '高级工程师、测绘系统先进、龙岩市优秀挂村书记...') · 极特定 · 简化为基础信息
    32: "{{CLIENT_LEGAL_REP}}：董事长兼总经理。",

    # P33: 总工程师 (硬编码人名 '王军德') · 简化
    33: "公司另设总工程师、监事长等关键岗位，由具备行业资深背景的专业人员担任。",

    # P34: 监事长 (硬编码 '王锦洋') · 删除
    34: "",

    # P51: 主营业务收入构成 (硬编码 82.79% 主营测绘) · 改为通用业务构成描述
    51: "    该公司属于{{CLIENT_INDUSTRY_CATEGORY}}企业，主要收入来源于其主营业务{{CLIENT_BUSINESS_DESC}}，业务结构与公司战略定位匹配。",

    # P52: 经营情况分析 (硬编码 '测绘服务受疫情') · 通用化
    52: "结合最新一期经营情况来看，公司主营业务发展态势整体稳定，与过往同期相比未发生重大不利变化。",

    # P53: 毛利率分析 (硬编码 '高新技术人才、地理信息技术模块') · 通用化
    53: "    从经营毛利水平来看，公司主营业务毛利率符合所属{{CLIENT_INDUSTRY_CATEGORY}}行业一般水平，盈利质量稳定，未出现重大异常波动。",

    # P54: 销售变化分析 (硬编码 '自然地理水文测绘需求') · 通用化
    54: "目前来看，公司经营周转未出现内部问题，外部行业景气度对营收的影响在可控区间，预计后续利润水平保持稳定。",

    # P57: 客户主营业务总结 (含完整硬编码业务描述) · 简化
    57: "    从客户的主营业务情况来看，{{CLIENT_FULL_NAME}}主要从事{{CLIENT_BUSINESS_DESC}}，我行行业投向指引该客户定位行业为：{{CLIENT_INDUSTRY_FULL}}。",

    # P72: 供给情况 (硬编码 '测绘公司及数据工程公司') · 通用化
    72: "从企业近年来的供给情况来看，供给方主要来源于行业上游服务提供商，企业根据自身经营需要采购相关原材料或服务，形成稳定的上游供应链。",

    # P74: 供应方稳定性 (硬编码 '测绘数据采购') · 通用化
    74: "从企业近年来合作的供应方来看，整体合作年限较长，供应方稳定。供应方各占比均不大，最大供应方占比仅约 20%，趋于分散，避免了对单一供应商的依赖。采购来源稳定，对主要业务板块，市场运行良好供求关系稳定。",

    # P78: 采购方向 (硬编码 '测绘服务') · 通用化
    78: "该公司的主要采购为业务经营所需的原材料、服务及对应数据。",

    # P79: 采购来源稳定 (硬编码 '测绘数据') · 通用化
    79: "（2）采购来源稳定，主要采购方向为业务相关的原材料及服务，对主要业务板块，市场运行良好供求关系稳定。",

    # P80: 结算模式 (硬编码 '测绘进度') · 通用化
    80: "（3）该公司结算模式目前均为现金，合同结算周期平均为一年，根据合同内容规定按照工程或服务进度要求分段付费，起到了风险管控的作用。",

    # P91: 产品价格 (硬编码 '地理测绘') · 通用化
    91: "（1）主要产品及近年价格波动情况：该公司主要产品和服务为{{CLIENT_BUSINESS_DESC}}，价格方面受市场供需、原材料成本、客户结构等多方面情况影响。",

    # P92: 销售渠道 (硬编码 '测绘服务主要为政府...地籍图编绘') · 通用化
    92: "（2）该公司销售渠道稳定性。下游客户主要为公司目标行业客户群体，供求稳定市场运行良好。对于主要业务板块，该公司销售渠道稳定，长期合作关系良好。",

    # P97: 风险应对 (硬编码 '提供测绘及其他服务') · 通用化
    97: "d．风险应对措施：公司提供{{CLIENT_BUSINESS_DESC}}采用分段收费方式，下单客户需先提供一定比率保证金（比率视对手方自身信用决定）之后按照工程进度分段付费。",

    # P115: 负债结构 (含 '测绘及相关专业技术服务') · 简化
    115: "负债结构分析：公司负债主要分布在短期借款、应付账款、其他流动负债及应付职工薪酬。其中其他流动负债为向母公司的借款。负债构成以流动负债为主，主要原因为企业属于轻资产运营企业本身固定资产偏少。负债结构整体较为稳定。",

    # P122: 存货 (硬编码 '测绘项目中收入未确认前') · 通用化
    122: "存货：公司存货整体呈一定波动趋势，企业存货主要为发出商品和劳务成本两类，劳务成本为业务项目中收入未确认前产生的相关成本，在项目收入确认之后会计入企业的主营业务成本中。",

    # P128: 主营业务收入 (硬编码 92.2% 测绘服务) · 通用化
    128: "主营业务收入：公司主营业务为{{CLIENT_BUSINESS_DESC}}，主要收入来源稳定。受合同进度结算特性影响，季度间收入存在波动，全年收入预计保持稳定增长。",

    # P142: 应收款周转 (硬编码 '测绘服务...实地勘察测绘') · 通用化
    142: "应收款周转天数：公司应收款周转天数处于行业相对合理区间。受下游客户结构特性（政府/企事业单位为主）影响，周转速度略高于一般行业水平，整体回款风险可控。",

    # P143: 存货周转 (硬编码 '测绘工程进度') · 通用化
    143: "存货周转天数：公司存货周转水平整体符合企业业务模式特征。受合同进度结算影响，部分时段存货周转天数存在波动，但与主营收入增长趋势一致，整体可控。",

    # P144: 应付账款周转 (硬编码 '测绘项目') · 通用化
    144: "应付账款周转天数：公司应付账款周转水平整体处于行业合理区间。公司在产业链中具备较强的议价能力，能够保持稳健的对外付款节奏。",

    # P190: 国内证及买方押汇 (硬编码 '测绘设备/无人机/测绘车') · 通用化
    190: "（4）国内证及买方押汇相关额度：企业不定期会更新经营所需的设备及原材料采购。为满足企业经营多样性的需求，特按相对于额度配比国内证及买方押汇。",

    # P198: 担保结论 (硬编码 '现金流充足下游企业多为政府及企事业单位' + '招标股份') · 通用化
    198: "本次授信由{{CLIENT_FULL_NAME}}自身信用担保，{{CLIENT_LONG_CORE_NAME}}作为{{CLIENT_PARENT_FULL_NAME}}的全资子公司，近年母公司持续为其注资，体现对公司未来前景的看好。此外{{CLIENT_FULL_NAME}}{{CLIENT_INDUSTRY_POSITION}}从行业长期发展、公司整体运营、集团整体实力来看，公司未来发展前景良好。综上，我行建议采用借款人自身信用担保。",

    # ── Phase 3 三次清理 · 集团描述 + 个别母公司硬编码 ──

    # P17: 集团二级权属企业 (硬编码 '招标股份、巨电新能源、环保设计院 20 多家' + '4000 人/博士 13 人' 详细统计)
    17: "目前，{{CLIENT_GROUP_FULL_NAME}}拥有多家二级权属企业，集团人才结构合理，国家注册各类执业中高级人才占员工总数一半以上，集团本部在我行 PD 评级良好，与我行联系紧密合作融洽，是我行的优质客户。",

    # P127: 借款用途 (硬编码 '向招标股份的借款...购置设备') · 通用化
    127: "其他应付款及其他流动负债：公司其他应付款及其他流动负债主要为关联方往来款和日常经营负债，其中关联方往来款主要用于支付外协费用及购置设备等业务经营开支。其他应付款为企业除与项目相关产生的应付账款之外的其他费用，包括日常经营生产中产生的各项支出。",

    # P169: 上市/集团重组陈述 (硬编码 '福建招标采购股份有限公司计划明年上市' + '招标采购集团旗下被划至招标采购有限公司') · 通用化
    169: "  c．公司目前主营业务收入波动的主要原因来自下游回款周期，{{CLIENT_FULL_NAME}}所属集团对其自身造血能力和经营前景较为看好。同时随着近年宏观环境改善，企业已恢复正常营收并恢复正常盈利。本次的授信方案中包含了快捷保理、商票保贴以及非融资性保函（免保证金）供企业选择应对不同业务场景。",

    # P196: 第二还款 (硬编码 '招标股份已通过募集资金账户向企业注资 1.5 亿元' + '天空地一体化遥感') · 通用化
    196: "从第二还款来看，公司属于{{CLIENT_GROUP_SHORT_NAME}}下属子公司，{{CLIENT_GROUP_SHORT_NAME}}作为大型集团整体抗风险能力较强，{{CLIENT_LONG_CORE_NAME}}作为集团优质资产目前在集团内承担核心业务。{{CLIENT_PARENT_FULL_NAME}}已通过募集资金账户向企业提供必要的运营资金支持。从股东实力方面来看企业第二还款来源充裕。",

    # P142: '政府/企事业单位为主' 残留 '事业单位' 短词 (来自我自己 Phase 3 第一次的通用化 phrasing) · 改通用客户群
    142: "应收款周转天数：公司应收款周转天数处于行业相对合理区间。受下游客户结构特性影响，周转速度处于行业一般水平，整体回款风险可控。",
}


# ────────────────────────────────────────────────────────────────────────────
# 表格 cell 级整段替换 (table_idx, row_idx, col_idx) → new_text
# ────────────────────────────────────────────────────────────────────────────
TABLE_CELL_REPLACEMENTS: dict[tuple[int, int, int], str] = {
    # T2R3C1/C2/C3: 经营范围 (3 个重复 cell · 原 '测绘服务;信息科学与系统科学研究服务;...')
    (2, 3, 1): "{{CLIENT_BUSINESS_SCOPE}}",
    (2, 3, 2): "{{CLIENT_BUSINESS_SCOPE}}",
    (2, 3, 3): "{{CLIENT_BUSINESS_SCOPE}}",
    # T2R3C6: 资质证书时间线 (原: '2008 年公司取得乙级测绘资质证书... 2014 年公司整体划拨至 {{GROUP}}, ...
    # 2016 年公司取得国家测绘地理信息局颁发的甲级测绘资质证书...')
    (2, 3, 6): "{{CLIENT_QUALIFICATION_DETAIL}}",

    # ── T7 主营业务收入构成表 · 业务板块列 (col 0) ──
    (7, 2, 0): "主营业务一",
    (7, 3, 0): "主营业务二",
    (7, 4, 0): "其他业务",

    # ── T8 采购供应商表 · 业务板块列 (col 1) + 供应商名 (col 4) ──
    (8, 1, 1): "主营业务",
    (8, 1, 4): "供应商A",
    (8, 2, 1): "主营业务",
    (8, 2, 4): "供应商B",
    (8, 3, 1): "主营业务",
    (8, 3, 4): "供应商C",
    (8, 4, 1): "主营业务",
    (8, 4, 4): "供应商D",
    (8, 5, 1): "主营业务",
    (8, 5, 4): "供应商E",
    (8, 6, 1): "主营业务",
    (8, 6, 4): "供应商F",
    (8, 7, 1): "主营业务",
    (8, 7, 4): "供应商G",

    # ── T9 销售客户表 · 业务板块列 (col 0) + 客户名 (col 3) ──
    (9, 1, 0): "主营业务",
    (9, 1, 3): "客户A",
    (9, 2, 0): "主营业务",
    (9, 2, 3): "客户B",
    (9, 3, 0): "主营业务",
    (9, 3, 3): "客户C",
    (9, 4, 0): "主营业务",
    (9, 4, 3): "客户D",
    (9, 5, 0): "主营业务",
    (9, 5, 3): "客户E",
}


def _replace_paragraph_text(para, new_text: str) -> None:
    """整段 paragraph.text 替换 · 用第一 run 装新文本 + 清空其余 run · 不动样式.

    new_text 空字符串 ('') 会让段落变空但 paragraph 仍保留 (避免改 paragraph count 影响其他索引)。
    """
    if not para.runs:
        # 空段 · 直接 add_run
        para.add_run(new_text)
        return
    # 第一 run 装全部新文本
    para.runs[0].text = new_text
    # 其余 run 清空
    for r in para.runs[1:]:
        r.text = ""


def _replace_cell_text(cell, new_text: str) -> None:
    """整 cell 替换 · 类似 _replace_paragraph_text 但适用于 cell.paragraphs."""
    if not cell.paragraphs:
        cell.add_paragraph(new_text)
        return
    first_para = cell.paragraphs[0]
    _replace_paragraph_text(first_para, new_text)
    # 清空其他 paragraph
    for p in cell.paragraphs[1:]:
        _replace_paragraph_text(p, "")


def main() -> int:
    if not DOCX_PATH.is_file():
        print(f"[ERROR] docx 不存在: {DOCX_PATH}")
        return 1

    doc = Document(DOCX_PATH)

    # ── 1. 段落替换 ──
    replaced_paras = 0
    skipped_paras = 0
    for idx, new_text in PARAGRAPH_REPLACEMENTS.items():
        if idx >= len(doc.paragraphs):
            print(f"[WARN] P{idx} 越界 (total {len(doc.paragraphs)}) · skip")
            continue
        para = doc.paragraphs[idx]
        old_text = para.text
        # idempotent: 已是纯 placeholder 段则 skip
        if old_text.strip() == new_text.strip():
            skipped_paras += 1
            print(f"[SKIP] P{idx} 已 placeholder 化")
            continue
        _replace_paragraph_text(para, new_text)
        replaced_paras += 1
        print(f"[OK] P{idx} replaced · old_len={len(old_text)} new_len={len(new_text)}")

    # ── 2. 表格 cell 替换 ──
    replaced_cells = 0
    skipped_cells = 0
    for (ti, ri, ci), new_text in TABLE_CELL_REPLACEMENTS.items():
        if ti >= len(doc.tables):
            print(f"[WARN] T{ti} 越界 (total {len(doc.tables)}) · skip")
            continue
        table = doc.tables[ti]
        if ri >= len(table.rows):
            print(f"[WARN] T{ti}R{ri} 越界 · skip")
            continue
        row = table.rows[ri]
        if ci >= len(row.cells):
            print(f"[WARN] T{ti}R{ri}C{ci} 越界 · skip")
            continue
        cell = row.cells[ci]
        old_text = cell.text
        if old_text.strip() == new_text.strip():
            skipped_cells += 1
            print(f"[SKIP] T{ti}R{ri}C{ci} 已 placeholder 化")
            continue
        _replace_cell_text(cell, new_text)
        replaced_cells += 1
        print(f"[OK] T{ti}R{ri}C{ci} replaced · old_len={len(old_text)} new_len={len(new_text)}")

    # ── 3. 保存 ──
    doc.save(DOCX_PATH)
    print()
    print(f"=== Phase 3 done · {replaced_paras} para replaced ({skipped_paras} skipped) · {replaced_cells} cell replaced ({skipped_cells} skipped) ===")
    return 0


if __name__ == "__main__":
    sys.exit(main())
