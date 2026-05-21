# -*- coding: utf-8 -*-
"""V16 Phase 4 治本 hot-fix: 经纬测绘 docx 第 2 次 dogfooding 残留清理.

背景:
  - Phase 3 治本 (commit 19d07ac) 把业务描述长段 placeholder 化, 但用户第 2 次 dogfooding 蓝汀
    家电报告仍看到经纬残留, 这次是财务分析段 + 表格 cell 残留:
      * P150 (筹资活动现金流叙述 · 含真实金额 10500 万 / 1713 万)
      * P168 (货币资金评估叙述 · 含真实金额 3375 万 + 经纬特定情况)
      * P177 (授信结论 · 含 3375 万 / 3000 万)
      * T0R1C0P0 (额度表 · 3375 万 / 2300 万)
      * T0R1C1P0 (额度表 · 3375 万 / 3000 万)
      * T15R1C1P1 (合并报表 · 福建省六一八信息技术 + 三明新基建关联方)

Phase 4 策略 (deterministic 治本):
  - 数字硬编码 → CREDIT_AMOUNT / CREDIT_EXPOSURE / CREDIT_PERIOD placeholder (schema 已有)
  - 关联方公司名 → RELATED_PARTY_1_FULL_NAME / RELATED_PARTY_2_FULL_NAME (schema v1.1.2 新增)
  - 长财务叙述 → FINANCING_ACTIVITY_NARRATIVE / CASHFLOW_ASSESSMENT_NARRATIVE (schema v1.1.2 新增)

idempotent: 段落已含目标 {{KEY}} 则 skip
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

DOCX_PATH = Path(__file__).resolve().parent.parent / "samples" / "经纬测绘_对公成稿A.docx"


# ────────────────────────────────────────────────────────────────────────────
# 段落级整段替换 (paragraph_index → new_text)
# ────────────────────────────────────────────────────────────────────────────
PARAGRAPH_REPLACEMENTS: dict[int, str] = {
    # ── Phase 4.5 治本 hot-fix (lint --json byte-level diff 暴露的 7 HIGH 残留) ──
    # P173: 评级+行业段 · 商务服务业 仍是 raw canonical → 改 placeholder
    # 原: 'a．{{CLIENT_FULL_NAME}}为我行PD评级5级客户，所属行业为商务服务业，目前我行行业投向指引...'
    173: "a．{{CLIENT_FULL_NAME}}为我行PD评级{{PD_RATING}}级客户，所属行业为{{CLIENT_INDUSTRY_CATEGORY}}，目前我行行业投向指引为：总体适度进入，符合投向政策要求。",

    # P150: 筹资活动现金流叙述
    # 原: '筹资活动产生的现金流量流入主要为吸收投资收到的现金10500万元、取得借款收到的现金1713万元...'
    150: "{{FINANCING_ACTIVITY_NARRATIVE}}",

    # P168: 货币资金评估叙述
    # 原: '目前企业的货币资金较大部分需要用以支付应付职工薪酬...该问题主要源头是今年疫情...3375万元综合授信...'
    168: "{{CASHFLOW_ASSESSMENT_NARRATIVE}}",

    # P177: 授信结论 (3375 万 / 3000 万 / 1 年)
    # 原: '拟给与该客户综合授信额度3375万元（敞口3000万元），额度循环使用。授信期限一年，敞口部分由{{CLIENT_FULL_NAME}}自身信用担保。具体品种分布如下：'
    177: "拟给与该客户综合授信额度{{CREDIT_AMOUNT}}（敞口{{CREDIT_EXPOSURE}}），额度循环使用。授信期限{{CREDIT_PERIOD}}，敞口部分由{{CLIENT_FULL_NAME}}自身信用担保。具体品种分布如下：",

    # P167: 短期流贷叙述 (含 3000 万具体金额)
    # 原: '...因此在本次的授信方案中计划给予企业3000万元的短期流贷用以缓解周转压力。'
    167: "  a．由于企业的下游买家行业自身资信状况整体良好、应收账款质量稳健，因此在本次的授信方案中按业务实际周转需求配置短期流贷用以缓解经营周转压力。",

    # P178: 一年期流贷品种分布 (含 3000 万)
    # 原: '（1）一年期流贷3000万元（敞口3000万元），用于公司本部经营周转、代发工资等。'
    178: "（1）一年期流贷：用于公司本部经营周转、代发工资等。",

    # P179-183: 商票保贴/电子银承/国内证/非融资性保函/快捷保理 (含 2300 万 / 2875 万 / 2500 万)
    # 原 (P179): '（2）商票保贴2300万元（敞口2300万元）。'
    179: "（2）商票保贴：用于配合客户结算场景。",
    # 原 (P180): '（3）电子银行承兑汇票2875万元（敞口2300万元），保证金20%，用于公司采购原料及服务等。'
    180: "（3）电子银行承兑汇票：保证金20%，用于公司采购原料及服务等。",
    # 原 (P181): '（4）国内证及买方押汇2500万元（敞口2300万元），保证金10%，用于公司采购原料及服务等。'
    181: "（4）国内证及买方押汇：保证金10%，用于公司采购原料及服务等。",
    # 原 (P182): '（5）非融资性保函2300万元（敞口2300万元），免保证金，用于履约、投标、质量维修保函，业务期限不超过3年。'
    182: "（5）非融资性保函：免保证金，用于履约、投标、质量维修保函，业务期限不超过3年。",
    # 原 (P183): '（6）快捷保理2300万元（敞口2300万元），业务期限1年。'
    183: "（6）快捷保理：业务期限1年。",

    # P187: 短期流贷理由 (含 3000 万 重复 2 次)
    # 原: '（1）短期流贷3000万元：由于企业营业成本中的人工成本占比约69%...因此配比3000万元短期流贷供公司经营周转使用。'
    187: "（1）短期流贷：由于企业经营成本中的人工成本占比较高，对日常经营现金流需求较大，配比相应额度的短期流贷供公司经营周转使用，缓解短期资金压力。",

    # P188: 商票保贴理由 (含 2300 万)
    # 原: '（2）商票保贴2300万元：由于企业的客户多为大型企业和各级政府单位自身信用较好可能会使用商票形式进行付款...'
    188: "（2）商票保贴：由于企业的下游客户多为大型企业及政府/事业单位自身信用较好可能会使用商票形式进行付款，为便利企业在汇票到期前取得资金，因而配比相对于额度的商票保贴供企业灵活使用。",

    # P191: 非融资性保函理由 (含 2300 万)
    # 原: '（5）非融资性保函（免保证金）2300万元：企业行业投向属于商务服务业常有贸易或工程投标等非融资性经营活动...'
    191: "（5）非融资性保函（免保证金）：企业经营涉及一定的非融资性经营活动需求，为不过多占用企业自身资金满足客户需求，特给予对应额度的非融资性保函。",

    # P192: 快捷保理理由 (含 2300 万)
    # 原: '（6）快捷保理2300万元：截至2022年6月份，企业应付账款占比企业营业成本的77.57%...给予企业对应额度的快捷保理。'
    192: "（6）快捷保理：企业上下游账期结构存在一定错配，为满足企业按期交付下游款项的需求，给予企业对应额度的快捷保理工具。",

    # ── Phase 4 二次清理 · 财务真实数字 + 客户真实机构残留 ──

    # P119: 应收账款大段 (含真实客户应收账款金额 + 5 个福建政府机构名 + 真实主营业务收入)
    # 原: '应收账款：截至2021年12月末，公司应收账款5781万元...龙海市自然资源局836万元...武警福建省总队保障部204万元...'
    119: "应收账款：公司应收账款主要集中在下游主要客户群体，前五大客户应收账款占比约 40% 左右，账龄结构以 1 年以内为主，整体账龄结构合理，回款情况与公司业务结算周期吻合。受合同分段收费模式影响，仅完成交付但尚未收款的部分计入应收账款。",

    # P149: 投资活动现金流 (含真实金额 10500 万 / 10539 万)
    # 原: '投资活动现金流入为0万元，投资活动现金流出合计10539万元，具体为构建固定资产、无形资产39万元和投资支付的现金10500万元...'
    149: "投资活动产生的现金流量受公司当期资本性支出节奏与对外财务投资安排影响，整体方向与公司战略相匹配。资金来源以股东支持及自有资金为主，投资活动现金流量净额变动属于正常财务安排范围。",

    # P184: 完全现金保证额度 (含 375 万)
    # 原: '（7）完全现金保证额度375万元 '
    184: "（7）完全现金保证额度",

    # P189: 电子银承理由 (含 2875 万)
    # 原: '（3）电子银行承兑汇票2875万元:...{{CLIENT_FULL_NAME}}本身作为{{CLIENT_PARENT_FULL_NAME}}下属全额子公司...'
    189: "（3）电子银行承兑汇票：{{CLIENT_FULL_NAME}}本身作为{{CLIENT_PARENT_FULL_NAME}}下属全额子公司，自身信用有担保。对下游企业使用商票的概率较大，因而配比相对于额度的电子商票。",

    # P193: 现金保证理由 (含 375 万)
    # 原: '（7）375万元完全现金保证额度：满足企业未来可能签订的各项完全现金业务需求。'
    193: "（7）完全现金保证额度：满足企业未来可能签订的各项完全现金业务需求。",

    # P221: 银行分行机构名 (经纬测绘 specific · 交通银行福建省分行大客户一部)
    # 原: '交通银行福建省分行大客户一部'
    221: "经办行客户部",

    # P108: 审计机构 (含 '福建华兴会计师事务所')
    # 原: '公司已向我行提供2019年、2020年、2021年审计年报及2022年9月报，其中2019-2021年报由福建华兴会计师事务所审计。...'
    108: "公司已向我行提供最近三个年度审计年报及最新一期月报，相关年报由具备执业资格的会计师事务所审计。根据会计师事务所的资质和企业实际运营情况，我行认为以上报表是可信的。具体报表数据如下：",
}


# ────────────────────────────────────────────────────────────────────────────
# 表格 cell 级 paragraph 替换 (table_idx, row_idx, col_idx, para_idx) → new_text
# 仅替换 cell 内特定 paragraph, 不动其余 (避免破坏额度品种分布列表)
# ────────────────────────────────────────────────────────────────────────────
TABLE_CELL_PARA_REPLACEMENTS: dict[tuple[int, int, int, int], str] = {
    # ── Phase 4.5: T2 注册资本/实收资本 cell · 6 处 5000万元 raw → placeholder ──
    # T2R1 row 是 "注册资本 | 5000万元 | 5000万元 | 5000万元 | 实收资本 | 实收资本 | 5000万元"
    # C1-C3 跨度均为注册资本 col · C6 是实收资本 col
    (2, 1, 1, 0): "{{CLIENT_REGISTERED_CAPITAL}}",
    (2, 1, 2, 0): "{{CLIENT_REGISTERED_CAPITAL}}",
    (2, 1, 3, 0): "{{CLIENT_REGISTERED_CAPITAL}}",
    (2, 1, 6, 0): "{{CLIENT_PAID_IN_CAPITAL}}",
    # T2R5 row 是 "{母公司} | {母公司} | 5000万元 | 货币资金 | ..." · 主要股东出资额 col
    (2, 5, 2, 0): "{{CLIENT_REGISTERED_CAPITAL}}",
    # T2R6 row 是 "合计 | 合计 | 5000万元 | 货币资金 | ..." · 总出资合计 col
    (2, 6, 2, 0): "{{CLIENT_REGISTERED_CAPITAL}}",

    # T0R1C0P0: 额度表第 1 段 (3375 万 / 2300 万)
    # 原: '综合授信额度3375万元（敞口2300万元），额度循环使用，授信期限1年，敞口部分以企业自身信用担保。具体分类额度如下：'
    (0, 1, 0, 0): "综合授信额度{{CREDIT_AMOUNT}}（敞口{{CREDIT_EXPOSURE}}），额度循环使用，授信期限{{CREDIT_PERIOD}}，敞口部分以企业自身信用担保。具体分类额度如下：",

    # T0R1C1P0: 额度表第 1 段 (3375 万 / 3000 万)
    # 原: '    申请给予借款人综合授信额度3375万元（敞口3000万元），额度循环使用，授信期限1年，敞口部分以企业自身信用担保。具体分类额度如下：'
    (0, 1, 1, 0): "    申请给予借款人综合授信额度{{CREDIT_AMOUNT}}（敞口{{CREDIT_EXPOSURE}}），额度循环使用，授信期限{{CREDIT_PERIOD}}，敞口部分以企业自身信用担保。具体分类额度如下：",

    # T15R1C1P1: 合并报表 (关联方公司名硬编码)
    # 原: '合并报表含公司本部及福建省六一八信息技术有限公司、三明新基建产业发展有限公司'
    (15, 1, 1, 1): "合并报表含公司本部及{{RELATED_PARTY_1_FULL_NAME}}、{{RELATED_PARTY_2_FULL_NAME}}",

    # T0R1C0P1-6: 额度品种明细 (含 2300 万 / 3000 万 等多次硬编码)
    # 原 P1: '1、短期流贷2000万元（敞口2000万元），用于经营周转'
    (0, 1, 0, 1): "1、短期流贷：用于经营周转",
    # 原 P2: '2、商票保贴2300万元（敞口2300万元），用于经营周转'
    (0, 1, 0, 2): "2、商票保贴：用于经营周转",
    # 原 P3: '3、电子银行承兑汇票2875万元（敞口2300万元），保证金20%，用于经营周转'
    (0, 1, 0, 3): "3、电子银行承兑汇票：保证金20%，用于经营周转",
    # 原 P4: '4、国内证及买方押汇2500万元（敞口2300万元）保证金10%，用于经营周转'
    (0, 1, 0, 4): "4、国内证及买方押汇：保证金10%，用于经营周转",
    # 原 P5: '5、快捷保理2300万元（敞口2300万元），用于经营周转'
    (0, 1, 0, 5): "5、快捷保理：用于经营周转",
    # 原 P6: '6、非融资性保函2300万元（敞口2300万元）'
    (0, 1, 0, 6): "6、非融资性保函",
    # 原 P7: '7、500万元完全现金保证额度'
    (0, 1, 0, 7): "7、完全现金保证额度",

    # T0R1C1P1-7: 同 C0 但 col 1
    # 原 P1: '1、短期流贷3000万元（敞口3000万元），用于经营周转'
    (0, 1, 1, 1): "1、短期流贷：用于经营周转",
    # 原 P2-P6 同 C0
    (0, 1, 1, 2): "2、商票保贴：用于经营周转",
    (0, 1, 1, 3): "3、电子银行承兑汇票：保证金20%，用于经营周转",
    (0, 1, 1, 4): "4、国内证及买方押汇：保证金10%，用于经营周转",
    (0, 1, 1, 5): "5、快捷保理：用于经营周转",
    (0, 1, 1, 6): "6、非融资性保函",
    # 原 P7: '7、375万元完全现金保证额度 '
    (0, 1, 1, 7): "7、完全现金保证额度",

    # T0R1C2P0/P1: 变化说明列 (含 2000 万→3000 万 / 500 万→375 万)
    # 原 P0: '短期流贷由2000万元增加至3000万元。'
    (0, 1, 2, 0): "短期流贷额度按业务需求调整。",
    # 原 P1: '完全现金保证额度从500万元调减为375万元。'
    (0, 1, 2, 1): "完全现金保证额度按业务需求调整。",

    # T6R2: 我行综合收益 (RAROC / 日均存款 / 中收 等 8 列具体数字 · 经纬 specific)
    # 原 C0: '32.99%' (RAROC)
    (6, 2, 0, 0): "—",
    # 原 C1: '50%' (我行结算比例)
    (6, 2, 1, 0): "—",
    # 原 C2: '1000万元' (日均存款)
    (6, 2, 2, 0): "—",
    # 原 C3: '1万元' (中收)
    (6, 2, 3, 0): "—",
    # 原 C4: '2000万元' (日均贷款)
    (6, 2, 4, 0): "—",
    # 原 C5: '3.9%' (放款利率)
    (6, 2, 5, 0): "—",
    # 原 C6: '50' (经营利润)
    (6, 2, 6, 0): "—",
    # 原 C7: '36' (经济利润)
    (6, 2, 7, 0): "—",
}


def _replace_paragraph_text(para, new_text: str) -> None:
    """整段 paragraph.text 替换 · 用第一 run 装新文本 + 清空其余 run · 不动样式."""
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""


def main() -> int:
    if not DOCX_PATH.is_file():
        print(f"[ERROR] docx 不存在: {DOCX_PATH}")
        return 1

    doc = Document(DOCX_PATH)

    # ── 1. 段落替换 ──
    replaced_paras = 0
    skipped_paras = 0
    for idx, new_text in PARAGRAPH_REPLACEMENTS.items():
        if idx >= len(doc.paragraphs):
            print(f"[WARN] P{idx} 越界 (total {len(doc.paragraphs)}) · skip")
            continue
        para = doc.paragraphs[idx]
        old_text = para.text
        if old_text.strip() == new_text.strip():
            skipped_paras += 1
            print(f"[SKIP] P{idx} 已 placeholder 化")
            continue
        _replace_paragraph_text(para, new_text)
        replaced_paras += 1
        print(f"[OK] P{idx} replaced · old_len={len(old_text)} new_len={len(new_text)}")

    # ── 2. 表格 cell paragraph 级替换 ──
    replaced_cell_paras = 0
    skipped_cell_paras = 0
    for (ti, ri, ci, pi), new_text in TABLE_CELL_PARA_REPLACEMENTS.items():
        if ti >= len(doc.tables):
            print(f"[WARN] T{ti} 越界 (total {len(doc.tables)}) · skip")
            continue
        table = doc.tables[ti]
        if ri >= len(table.rows):
            print(f"[WARN] T{ti}R{ri} 越界 · skip")
            continue
        row = table.rows[ri]
        if ci >= len(row.cells):
            print(f"[WARN] T{ti}R{ri}C{ci} 越界 · skip")
            continue
        cell = row.cells[ci]
        if pi >= len(cell.paragraphs):
            print(f"[WARN] T{ti}R{ri}C{ci}P{pi} 越界 (cell paras={len(cell.paragraphs)}) · skip")
            continue
        para = cell.paragraphs[pi]
        old_text = para.text
        if old_text.strip() == new_text.strip():
            skipped_cell_paras += 1
            print(f"[SKIP] T{ti}R{ri}C{ci}P{pi} 已 placeholder 化")
            continue
        _replace_paragraph_text(para, new_text)
        replaced_cell_paras += 1
        print(f"[OK] T{ti}R{ri}C{ci}P{pi} replaced · old_len={len(old_text)} new_len={len(new_text)}")

    # ── 3. 保存 ──
    doc.save(DOCX_PATH)
    print()
    print(f"=== Phase 4 done · {replaced_paras} para replaced ({skipped_paras} skipped) · {replaced_cell_paras} cell para replaced ({skipped_cell_paras} skipped) ===")
    return 0


if __name__ == "__main__":
    sys.exit(main())
