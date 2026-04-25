# -*- coding: utf-8 -*-
"""一次性脚本:生成 Phase 2 Task C 所需 2 份 desensitized 模板 docx。

DoD: L3-12(模板覆盖度,≥2 银行模板)
对应 onboarding: docs/onboarding/agent6-phase-2.md Task C

样本来源:基于公开监管口径 + 行业典型条目人工脱敏构造,不含任何真实客户/银行数据。
产物:
  samples/科创贷申报书_模板.docx
  samples/小微对私授信申报书_模板.docx

命令:
  py scripts/build_phase2_samples.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
SAMPLES = PROJECT_ROOT / "samples"


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def _add_h(doc: Document, text: str, size: int = 13) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _add_p(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(11)


def build_keji_template() -> Path:
    """科创贷申报书模板(脱敏构造,基于科创贷监管口径)。"""
    doc = Document()
    _add_title(doc, "科技创新企业授信申报书(模板)")
    _add_p(doc, "(本模板为脱敏教学版,所有字段均待业务方真实数据填入)")

    _add_h(doc, "一、企业基本情况")
    _add_p(doc, "企业名称:____")
    _add_p(doc, "统一社会信用代码:____")
    _add_p(doc, "成立日期:____   注册资本:____ 万元   实缴资本:____ 万元")
    _add_p(doc, "所属行业:____   高新认定:[ ] 是  [ ] 否   科技型中小企业入库:[ ] 是  [ ] 否")

    _add_h(doc, "二、研发情况")
    _add_p(doc, "近三年研发投入(万元):2022 ____ / 2023 ____ / 2024 ____")
    _add_p(doc, "研发投入强度(研发费用 / 营业收入):2024 ____ %(科创贷认定口径 ≥5%)")
    _add_p(doc, "研发人员占比:____ %    发明专利数:____ 项")

    _add_h(doc, "三、财务情况")
    _add_p(doc, "营业收入(万元):2022 ____ / 2023 ____ / 2024 ____")
    _add_p(doc, "净利润(万元):2022 ____ / 2023 ____ / 2024 ____")
    _add_p(doc, "资产负债率:____ %    经营性现金流净额:____ 万元")

    _add_h(doc, "四、授信申请方案")
    _add_p(doc, "申请综合授信:____ 万元   敞口:____ 万元   期限:____ 个月")
    _add_p(doc, "用途:____    还款来源:____")
    _add_p(doc, "担保方式:[ ] 信用 [ ] 知识产权质押 [ ] 第三方担保 [ ] 抵押")

    _add_h(doc, "五、风险评估及结论")
    _add_p(doc, "主要风险点:____   缓释措施:____")
    _add_p(doc, "建议授信结论:____   前提条件:____")

    out = SAMPLES / "科创贷申报书_模板.docx"
    doc.save(out)
    return out


def build_xiaowei_duisi_template() -> Path:
    """小微对私授信申报书模板(脱敏构造,基于普惠小微+对私授信口径)。"""
    doc = Document()
    _add_title(doc, "小微对私授信申报书(模板)")
    _add_p(doc, "(本模板为脱敏教学版,所有字段均待业务方真实数据填入)")

    _add_h(doc, "一、申请人基本情况")
    _add_p(doc, "申请人姓名:____   身份证号:____  ")
    _add_p(doc, "工作单位/经营实体:____   职务/经营年限:____")
    _add_p(doc, "婚姻状况:[ ] 已婚 [ ] 未婚    家庭年收入(万元):____")
    _add_p(doc, "家庭住址:____   联系方式:____")

    _add_h(doc, "二、征信情况")
    _add_p(doc, "本人查询次数(近 2 年):____ 次")
    _add_p(doc, "贷款审批查询(近 6 个月):____ 次")
    _add_p(doc, "信用卡审批查询(近 6 个月):____ 次")
    _add_p(doc, "在贷余额(万元):____   逾期记录:____")

    _add_h(doc, "三、收入与负债情况")
    _add_p(doc, "月均收入:____ 元    月均负债支出:____ 元    DSR:____ %")
    _add_p(doc, "可还款能力评估:____")

    _add_h(doc, "四、授信申请方案")
    _add_p(doc, "申请额度(万元):____   期限:____   还款方式:[ ] 等额本息 [ ] 等额本金 [ ] 一次还本付息")
    _add_p(doc, "用途:____   担保方式:[ ] 信用 [ ] 抵押 [ ] 第三方保证")

    _add_h(doc, "五、风险评估及结论")
    _add_p(doc, "主要风险点:____   缓释措施:____")
    _add_p(doc, "建议授信结论:____   前提条件:____")

    out = SAMPLES / "小微对私授信申报书_模板.docx"
    doc.save(out)
    return out


def main() -> int:
    SAMPLES.mkdir(exist_ok=True)
    p1 = build_keji_template()
    p2 = build_xiaowei_duisi_template()
    print(f"[ok] {p1.name} · {p1.stat().st_size} bytes")
    print(f"[ok] {p2.name} · {p2.stat().st_size} bytes")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
