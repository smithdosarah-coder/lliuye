# -*- coding: utf-8 -*-
"""V16 Phase 3 worker 4: 把 samples/小微对私授信申报书_模板.docx 改成 placeholder 模板.

设计 (Branch A · skeleton 路径 · 对私 retail personal 场景):
  - docx 全文 21 element / 473 chars · 0 真客户字面 · 全是 字段:____ 空槽 + [ ] checkbox + 章节号
  - basename 已显式标 "(模板)" · P1 自带 "本模板为脱敏教学版" 声明
  - 不需 SUBSTITUTIONS (no client name to replace)
  - 仅 PARAGRAPH_OVERRIDES 把字段标签后的 `:____` 空槽替换为 `:{{KEY}}`
  - 保留 [ ] checkbox 字面 (runtime checkbox_tick handler 走 client_metadata
    `CLIENT_MARITAL_STATUS` / `REPAYMENT_METHOD` / `GUARANTEE_METHOD` 等枚举值勾选)
  - 让 framework REPLACE handler 用 runtime client_metadata 填值

对私场景 placeholder schema 提议 (schema v1.1):
  - 23 个新 key (对私专属) + 3 个复用 v1.0 (CLIENT_FULL_NAME / CREDIT_AMOUNT / CREDIT_PERIOD)
  - 不修改 templates/placeholder-schema.json (避免污染对公 worker)
  - 完整 key 列表见 samples/小微对私授信申报书_模板.metadata.json schema_v1_1_proposed_keys

vs Phase 2 经纬 pilot / Phase 3 worker 1 兴业 / Phase 3 worker 2 普惠 差异:
  - 经纬 (worker 0): 1023 element · 29 location · 10 key · 对公真稿 (双路径)
  - 兴业 (worker 1): 1500+ element · 141 location · 13 key · 对公真稿 (双路径 · 三级股权)
  - 普惠 (worker 2): 495 element · 6 location · 6 key · 对公骨架 (仅 PARAGRAPH_OVERRIDES)
  - 小微 (worker 4): 21 element · 14 location · 26 key · 对私骨架 (仅 PARAGRAPH_OVERRIDES · schema v1.1 提议)

用法:
  py scripts/v16_placeholderize_xiaowei.py
  - 默认改 samples/小微对私授信申报书_模板.docx (in-place)
  - 改完后用 v16_step1_extract.py 扫一遍验证 placeholder 都标 PLACEHOLDER label

run 一次即完成 · idempotent (再跑一次不会重复替换 · 已含 {{KEY}} 跳过)
"""
from __future__ import annotations

import io
import re
import sys
from pathlib import Path
from docx import Document

# 使 stdout UTF-8 (Windows console 中文)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

DOCX_PATH = Path(__file__).resolve().parent.parent / "samples" / "小微对私授信申报书_模板.docx"

# Branch A skeleton: 仅整段重写 form 字段段落
# location → 新 paragraph text · 注入 {{KEY}} placeholder · 保留字段标签 + checkbox 字面 + 单位
PARAGRAPH_OVERRIDES: dict[str, str] = {
    # 一、申请人基本情况
    "P3": "申请人姓名:{{CLIENT_FULL_NAME}}   身份证号:{{CLIENT_ID_NUMBER}}",
    "P4": "工作单位/经营实体:{{CLIENT_EMPLOYER_OR_BUSINESS}}   职务/经营年限:{{CLIENT_POSITION_OR_YEARS}}",
    # P5 含 [ ] 已婚 [ ] 未婚 checkbox · runtime checkbox_tick handler 走 CLIENT_MARITAL_STATUS 值勾选
    "P5": "婚姻状况:{{CLIENT_MARITAL_STATUS}} [ ] 已婚 [ ] 未婚    家庭年收入(万元):{{CLIENT_HOUSEHOLD_INCOME}}",
    "P6": "家庭住址:{{CLIENT_HOME_ADDRESS}}   联系方式:{{CLIENT_PHONE}}",
    # 二、征信情况
    "P8": "本人查询次数(近 2 年):{{CREDIT_QUERY_COUNT_2Y}} 次",
    "P9": "贷款审批查询(近 6 个月):{{CREDIT_LOAN_QUERY_6M}} 次",
    "P10": "信用卡审批查询(近 6 个月):{{CREDIT_CARD_QUERY_6M}} 次",
    "P11": "在贷余额(万元):{{CREDIT_OUTSTANDING_BALANCE}}   逾期记录:{{CREDIT_OVERDUE_RECORD}}",
    # 三、收入与负债情况
    "P13": "月均收入:{{CLIENT_MONTHLY_INCOME}} 元    月均负债支出:{{CLIENT_MONTHLY_DEBT_PAYMENT}} 元    DSR:{{CLIENT_DSR}} %",
    "P14": "可还款能力评估:{{REPAYMENT_ABILITY_ASSESSMENT}}",
    # 四、授信申请方案
    # P16 含 [ ] 等额本息 [ ] 等额本金 [ ] 一次还本付息 checkbox · runtime checkbox_tick 走 REPAYMENT_METHOD 值勾选
    "P16": "申请额度(万元):{{CREDIT_AMOUNT}}   期限:{{CREDIT_PERIOD}}   还款方式:{{REPAYMENT_METHOD}} [ ] 等额本息 [ ] 等额本金 [ ] 一次还本付息",
    # P17 含 [ ] 信用 [ ] 抵押 [ ] 第三方保证 checkbox · runtime checkbox_tick 走 GUARANTEE_METHOD 值勾选
    "P17": "用途:{{CREDIT_PURPOSE}}   担保方式:{{GUARANTEE_METHOD}} [ ] 信用 [ ] 抵押 [ ] 第三方保证",
    # 五、风险评估及结论
    "P19": "主要风险点:{{RISK_KEY_POINTS}}   缓释措施:{{RISK_MITIGATION}}",
    "P20": "建议授信结论:{{CREDIT_CONCLUSION}}   前提条件:{{CREDIT_PRECONDITIONS}}",
}


def _replace_paragraph_text(para, new_text: str) -> bool:
    """在保留 run 格式的前提下整段重写 · 与 v16_generator._set_paragraph_text 一致策略.

    返回是否改动.
    """
    if (para.text or "") == new_text:
        return False
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return True
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return True


def _walk_paragraphs(doc):
    """yield (location, paragraph) — 同 v16_step1_extract.extract_elements 逻辑."""
    for pi, p in enumerate(doc.paragraphs):
        yield f"P{pi}", p

    def _walk_table(tbl, loc_prefix: str):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for ppi, p in enumerate(cell.paragraphs):
                    yield f"{loc_prefix}R{ri}C{ci}P{ppi}", p
                for sub in cell.tables:
                    yield from _walk_table(sub, f"{loc_prefix}R{ri}C{ci}NT")

    for ti, tbl in enumerate(doc.tables):
        yield from _walk_table(tbl, f"T{ti}")


def main():
    if not DOCX_PATH.is_file():
        print(f"[fatal] docx 不存在: {DOCX_PATH}", file=sys.stderr)
        sys.exit(1)
    doc = Document(str(DOCX_PATH))

    stats = {"overrides": 0, "checked_paragraphs": 0, "skipped_idempotent": 0}
    hit_locs: list[str] = []
    unique_keys: set[str] = set()
    target_locs = set(PARAGRAPH_OVERRIDES.keys())
    seen_target_locs: set[str] = set()

    for loc, para in _walk_paragraphs(doc):
        stats["checked_paragraphs"] += 1
        original_text = para.text or ""

        if loc in PARAGRAPH_OVERRIDES:
            seen_target_locs.add(loc)
            new_text = PARAGRAPH_OVERRIDES[loc]
            # Idempotency check: 如果原文已经是目标 placeholder 形式 · 跳过
            if original_text == new_text:
                stats["skipped_idempotent"] += 1
                print(f"  [skip-idempotent] {loc}: 已是 placeholder 形式")
                continue
            if _replace_paragraph_text(para, new_text):
                stats["overrides"] += 1
                hit_locs.append(loc)
                # 抓 {{KEY}} 用于统计
                for key in re.findall(r"\{\{([A-Z_][A-Z0-9_]+)\}\}", new_text):
                    unique_keys.add(key)
                print(f"  [override] {loc}: {original_text[:60]} -> {new_text[:90]}")

    # 验证所有 target location 都被处理 (没有 location 丢失 · 例如 docx 被改)
    missing_targets = target_locs - seen_target_locs
    if missing_targets:
        print(f"[warn] 以下 target location 未在 docx walk 时遇到: {missing_targets}")

    doc.save(str(DOCX_PATH))
    print()
    print(f"[done] checked {stats['checked_paragraphs']} paragraphs / "
          f"{stats['overrides']} overrides / {stats['skipped_idempotent']} skipped-idempotent")
    print(f"  改动 location 数 (unique): {len(set(hit_locs))}")
    print(f"  unique placeholder key 数: {len(unique_keys)}")
    print(f"  placeholder keys: {sorted(unique_keys)}")
    print(f"  改后 docx: {DOCX_PATH}")


if __name__ == "__main__":
    main()
