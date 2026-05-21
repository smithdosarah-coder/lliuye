# -*- coding: utf-8 -*-
"""V16 Phase 3 worker 3: 把 samples/科创贷申报书_模板.docx 改成 placeholder 模板.

设计 (Branch A · skeleton 路径 · 跟普惠 worker 2 同模式):
  - 科创贷申报书是纯骨架模板 · 副标题 P1 显式声明 "本模板为脱敏教学版,所有字段均待业务方真实数据填入"
  - 全文 22 段 paragraph (无 table cell) 全是 form 字段 + ____ 空槽 + checkbox · 0 真客户字面
  - 不需要 SUBSTITUTIONS (no client name to replace)
  - 仅 PARAGRAPH_OVERRIDES 在 form 主干字段 (P3/P5/P6/P16) 注入 {{KEY}}
  - 让 framework REPLACE handler 用 runtime client_metadata 填值

vs Phase 2/3 其他 docx 差异:
  - 经纬 (worker pilot): 1023 element / 真客户成稿 / 56+ occurrence (Branch B)
  - 兴业资管 (worker 1): 141 location / 真客户成稿 / 三级股权 (Branch B)
  - 普惠 (worker 2): 495 element / skeleton / 6 placeholder / 表格 form (Branch A)
  - 科创贷 (worker 3 本工件): 22 element / skeleton / 8 placeholder / 纯 paragraph form (Branch A)
    · 数量级最小 · location 全是 P0..P21 (无 T*R*C*P* prefix)
    · canonical schema 覆盖 8 key (FULL_NAME / ESTABLISHMENT_DATE / REGISTERED_CAPITAL / PAID_IN_CAPITAL
      / INDUSTRY_CATEGORY / CREDIT_AMOUNT / CREDIT_EXPOSURE / CREDIT_PERIOD)
    · 研发情况 (P8-P10) / 财务情况 (P12-P14) / 担保/风险/结论 (P17-P21) 留 ____ 走 SLOT/FILL/REWRITE handler

用法:
  py scripts/v16_placeholderize_kechuang.py
  - 默认改 samples/科创贷申报书_模板.docx (in-place)
  - 改完后用 v16_step1_extract.py 扫一遍验证 placeholder 都标 PLACEHOLDER label

run 一次即完成 · idempotent (再跑一次不会重复替换 · 已含 {{KEY}} 跳过)
"""
from __future__ import annotations

import io
import sys
from pathlib import Path
from docx import Document

# 使 stdout UTF-8 (Windows console 中文)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

DOCX_PATH = Path(__file__).resolve().parent.parent / "samples" / "科创贷申报书_模板.docx"

# Branch A skeleton: 仅整段重写 form 主干字段
# location → 新 paragraph text · 仅注入 {{KEY}} placeholder · 保留字段标签 + checkbox + 非 canonical 槽位 ____
# 注意: 单位 (万元 / 个月) 不在模板中保留 · 单位归并到 placeholder value
#       (eg client_metadata.CREDIT_AMOUNT="5亿元" / CLIENT_REGISTERED_CAPITAL="195000万元" 单位 baked-in)
PARAGRAPH_OVERRIDES: dict[str, str] = {
    # 企业名称 field (一、企业基本情况)
    "P3": "企业名称:{{CLIENT_FULL_NAME}}",
    # 成立日期 + 注册资本 + 实缴资本 (同一段三字段)
    # 原: "成立日期:____   注册资本:____ 万元   实缴资本:____ 万元"
    # 单位 万元 归并到 placeholder value
    "P5": "成立日期:{{CLIENT_ESTABLISHMENT_DATE}}   注册资本:{{CLIENT_REGISTERED_CAPITAL}}   实缴资本:{{CLIENT_PAID_IN_CAPITAL}}",
    # 所属行业 + 高新认定 checkbox + 科技型中小企业入库 checkbox (同一段)
    # 仅 placeholder 化 所属行业 · checkbox 组保留原样 走 checkbox_tick handler
    "P6": "所属行业:{{CLIENT_INDUSTRY_CATEGORY}}   高新认定:[ ] 是  [ ] 否   科技型中小企业入库:[ ] 是  [ ] 否",
    # 申请综合授信 + 敞口 + 期限 (四、授信申请方案首段)
    # 原: "申请综合授信:____ 万元   敞口:____ 万元   期限:____ 个月"
    # 单位 万元 / 个月 归并到 placeholder value (兴业 CREDIT_AMOUNT='5亿元' / CREDIT_PERIOD='一年')
    "P16": "申请综合授信:{{CREDIT_AMOUNT}}   敞口:{{CREDIT_EXPOSURE}}   期限:{{CREDIT_PERIOD}}",
}


def _replace_paragraph_text(para, new_text: str) -> bool:
    """在保留 run 格式的前提下整段重写 · 与 v16_generator._set_paragraph_text 一致策略.

    返回是否改动.
    """
    if (para.text or "") == new_text:
        return False
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return True
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return True


def _walk_paragraphs(doc):
    """yield (location, paragraph) — 同 v16_step1_extract.extract_elements 逻辑."""
    for pi, p in enumerate(doc.paragraphs):
        yield f"P{pi}", p

    def _walk_table(tbl, loc_prefix: str):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for ppi, p in enumerate(cell.paragraphs):
                    yield f"{loc_prefix}R{ri}C{ci}P{ppi}", p
                for sub in cell.tables:
                    yield from _walk_table(sub, f"{loc_prefix}R{ri}C{ci}NT")

    for ti, tbl in enumerate(doc.tables):
        yield from _walk_table(tbl, f"T{ti}")


def main():
    if not DOCX_PATH.is_file():
        print(f"[fatal] docx 不存在: {DOCX_PATH}", file=sys.stderr)
        sys.exit(1)
    doc = Document(str(DOCX_PATH))

    stats = {"overrides": 0, "checked_paragraphs": 0, "skipped_idempotent": 0}
    hit_locs: list[str] = []
    unique_keys: set[str] = set()
    target_locs = set(PARAGRAPH_OVERRIDES.keys())
    seen_target_locs: set[str] = set()

    for loc, para in _walk_paragraphs(doc):
        stats["checked_paragraphs"] += 1
        original_text = para.text or ""

        if loc in PARAGRAPH_OVERRIDES:
            seen_target_locs.add(loc)
            new_text = PARAGRAPH_OVERRIDES[loc]
            # Idempotency check: 如果原文已经包含目标 placeholder set · 跳过
            if original_text == new_text:
                stats["skipped_idempotent"] += 1
                print(f"  [skip-idempotent] {loc}: 已是 placeholder 形式")
                continue
            if _replace_paragraph_text(para, new_text):
                stats["overrides"] += 1
                hit_locs.append(loc)
                # 抓 {{KEY}} 用于统计
                import re
                for key in re.findall(r"\{\{([A-Z_][A-Z0-9_]+)\}\}", new_text):
                    unique_keys.add(key)
                print(f"  [override] {loc}: {original_text[:60]} → {new_text[:80]}")

    # 验证所有 target location 都被处理 (没有 location 丢失 · 例如 docx 被改)
    missing_targets = target_locs - seen_target_locs
    if missing_targets:
        print(f"[warn] 以下 target location 未在 docx walk 时遇到 (可能 docx 被改): {missing_targets}")

    doc.save(str(DOCX_PATH))
    print()
    print(f"[done] checked {stats['checked_paragraphs']} paragraphs / "
          f"{stats['overrides']} overrides / {stats['skipped_idempotent']} skipped-idempotent")
    print(f"  改动 location 数 (unique): {len(set(hit_locs))}")
    print(f"  unique placeholder key 数: {len(unique_keys)} · {sorted(unique_keys)}")
    print(f"  改后 docx: {DOCX_PATH}")


if __name__ == "__main__":
    main()
