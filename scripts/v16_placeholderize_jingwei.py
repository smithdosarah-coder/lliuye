# -*- coding: utf-8 -*-
"""V16 Phase 2 pilot: 把 samples/经纬测绘_对公成稿A.docx 改成 placeholder 模板.

设计:
  - 整段 (paragraph.text) 替换 · 用第一个 run 写整段新文本 + 清空其余 run · 避开 cross-run quirks
  - 完整对照 metadata.json 的 placeholder_locations 字段
  - 替换顺序: 长字符串先 (避免 '福建经纬数字科技信息有限公司' 内含 '经纬' 时把核心词替换掉)
  - **不动样式 + 段落格式 + 列表序号** · 只改文本

用法:
  py scripts/v16_placeholderize_jingwei.py
  - 默认改 samples/经纬测绘_对公成稿A.docx (in-place)
  - 改完后用 v16_step1_extract.py 扫一遍验证 placeholder 都标 PLACEHOLDER label

run 一次即完成 · idempotent (再跑一次不会重复替换 · 因为 {{KEY}} 不在 substitution map 内)
"""
from __future__ import annotations

import io
import sys
from pathlib import Path
from docx import Document

# 使 stdout UTF-8 (Windows console 中文)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

DOCX_PATH = Path(__file__).resolve().parent.parent / "samples" / "经纬测绘_对公成稿A.docx"

# 替换规则: (硬编码字符串, placeholder) · 顺序 = 长→短 (避免子串被先吃掉)
# 选只在客户档案中明确的硬编码 · 财务数据 / 业务叙事段保留原文 (REWRITE handler 处理)
SUBSTITUTIONS: list[tuple[str, str]] = [
    # 公司名 (最长先, 避免 '经纬' 在 '福建经纬...' 内被先替换)
    ("福建经纬数字科技信息有限公司", "{{CLIENT_FULL_NAME}}"),
    ("经纬数字科技", "{{CLIENT_LONG_CORE_NAME}}"),
    # '经纬数字' 在被 ↑ 替换后仅剩独立出现 (如 P196 '经纬数字作为集团优质资产')
    ("经纬数字", "{{CLIENT_LONG_CORE_NAME}}"),
    # 单字 '经纬' (作公司简称) · 注意此时长字符串已替换 不会重复命中
    # 但仍可能误命中"经纬地理"等非客户用法 · 经检查 docx 内"经纬"仅作公司简称 OK

    # 母公司 (长→短)
    ("福建省招标采购集团有限公司", "{{CLIENT_GROUP_FULL_NAME}}"),
    ("福建省招标股份有限公司", "{{CLIENT_PARENT_FULL_NAME}}"),
    # 短形式 (仅长形式不在的位置才有 · P196 出现"福建省招标采购集团")
    ("福建省招标采购集团", "{{CLIENT_GROUP_SHORT_NAME}}"),
    # 短"招标股份" / "招标采购集团" 简称 · P196/P198 出现
    # 不替换 · 在长 REWRITE 段保留 · 不破坏 LLM 生成质量

    # 法人 · 注册地 · 注册资本
    ("郑志煌", "{{CLIENT_LEGAL_REP}}"),
    ("福建省福州市鼓楼区洪山园路68号", "{{CLIENT_REGISTERED_ADDRESS}}"),
    ("2005年8月4日", "{{CLIENT_ESTABLISHMENT_DATE}}"),

    # 授信额度主标题 (仅 P0 用 · 避免和正文 3375 段串干扰)
    # 标题 P0 只出现 1 次 · 用整段重写 (见 PARAGRAPH_OVERRIDES) 不依赖此规则
]

# 整段重写覆盖 (优先级高于 SUBSTITUTIONS · 用于 1 段含多个 placeholder 的复杂段)
# location → 新 paragraph text
PARAGRAPH_OVERRIDES: dict[str, str] = {
    # 标题: 公司全名 + 授信额度
    "P0": "{{CLIENT_FULL_NAME}}维持{{CREDIT_AMOUNT}}综合授信额度的授信报告",
    # 标题字段行: 全角空格 8 个 → 字段分隔
    "P2": "授信客户全称：{{CLIENT_FULL_NAME}}       我行投向政策对应行业：{{CLIENT_INDUSTRY_CATEGORY}}",
}


def _replace_paragraph_text(para, new_text: str) -> bool:
    """在保留 run 格式的前提下整段重写 · 与 v16_generator._set_paragraph_text 一致策略.

    返回是否改动.
    """
    if (para.text or "") == new_text:
        return False
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return True
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return True


def _walk_paragraphs(doc):
    """yield (location, paragraph) — 同 v16_step1_extract.extract_elements 逻辑."""
    for pi, p in enumerate(doc.paragraphs):
        yield f"P{pi}", p

    def _walk_table(tbl, loc_prefix: str):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for ppi, p in enumerate(cell.paragraphs):
                    yield f"{loc_prefix}R{ri}C{ci}P{ppi}", p
                for sub in cell.tables:
                    yield from _walk_table(sub, f"{loc_prefix}R{ri}C{ci}NT")

    for ti, tbl in enumerate(doc.tables):
        yield from _walk_table(tbl, f"T{ti}")


def main():
    if not DOCX_PATH.is_file():
        print(f"[fatal] docx 不存在: {DOCX_PATH}", file=sys.stderr)
        sys.exit(1)
    doc = Document(str(DOCX_PATH))

    stats = {"overrides": 0, "substitutions": 0, "checked_paragraphs": 0}
    hit_locs: list[str] = []

    for loc, para in _walk_paragraphs(doc):
        stats["checked_paragraphs"] += 1
        original_text = para.text or ""

        # 1) PARAGRAPH_OVERRIDES 优先 (整段重写)
        if loc in PARAGRAPH_OVERRIDES:
            new_text = PARAGRAPH_OVERRIDES[loc]
            if _replace_paragraph_text(para, new_text):
                stats["overrides"] += 1
                hit_locs.append(loc)
                print(f"  [override] {loc}: {original_text[:60]} → {new_text[:60]}")
            continue

        # 2) SUBSTITUTIONS · 顺序应用 · 不破坏已替换的 placeholder
        new_text = original_text
        for old, ph in SUBSTITUTIONS:
            if old in new_text and ph not in new_text:
                # ph not in new_text 保证 idempotent (再跑不重复)
                new_text = new_text.replace(old, ph)
            elif old in new_text:
                # 已含 placeholder 仍含原文 · 走单次替换 (常见: cross-substr 同时命中)
                new_text = new_text.replace(old, ph)
        if new_text != original_text:
            if _replace_paragraph_text(para, new_text):
                stats["substitutions"] += 1
                hit_locs.append(loc)
                print(f"  [substit] {loc}: {original_text[:60]} → {new_text[:60]}")

    doc.save(str(DOCX_PATH))
    print()
    print(f"[done] checked {stats['checked_paragraphs']} paragraphs / "
          f"{stats['overrides']} overrides / {stats['substitutions']} substitutions")
    print(f"  改动 location 数 (unique): {len(set(hit_locs))}")
    print(f"  改后 docx: {DOCX_PATH}")


if __name__ == "__main__":
    main()
