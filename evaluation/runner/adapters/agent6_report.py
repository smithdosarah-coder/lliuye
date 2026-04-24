# -*- coding: utf-8 -*-
"""
evaluation.runner.adapters.agent6_report — Agent6 报告助手评估 adapter

HOLDING H-A · v16 summary parse 接入 (2026-04-24):
  artifact 侧自动查找 v16_pipeline_summary.json 及同目录 QC md, 把下列三项
  从 method=manual 升级为 deterministic:
    - quality_score_total    ← qc.score (0-100, 直接)
    - hallucination_rate     ← qc.hallucinations / paragraph_count (段落级率)
    - evidence_rate          ← QC md 9 维度中"无 miss 行"的维度占比
  artifact 无 summary 时这三项 fallback 回 manual (保持向后兼容)。

已确定性指标:
  [domain]
    - template_leakage_rate   — SequenceMatcher vs samples/ 基线
    - unfilled_marker_accuracy — 占位符残留 + "未能自动填写" 标记一致性
  [common]
    - task_completion_rate    — docx 是否正常打开、段落数 > 0

Phase 2 pending (留 Batch 2):
  - financial_ratio_consistency (EV-12) — 需 financial_analyzer runtime 对照
  - section_length_calibration — 需真人范文库
"""

from __future__ import annotations

import json
import re
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

from docx import Document

from ..base_evaluator import REPO_ROOT, BaseEvaluator
from ..registry import register_evaluator
from ..schemas import EvalRun, MetricOutcome


SAMPLES_DIR = REPO_ROOT / "samples"
SIMILARITY_HIT = 0.85           # 段落视作"模板泄漏"的 SequenceMatcher 阈值
PLACEHOLDER_PAT = re.compile(r"\{\{[^}]+\}\}|XXX|xxx|【[^】]*】")
UNFILLED_MARK = "未能自动填写"

V16_SUMMARY_FILENAME = "v16_pipeline_summary.json"

# QC md 解析: 9 维度行形如 "- <name>: <raw_score> (权重 <w>, 加权 <a>)"
_QC_DIM_LINE_RE = re.compile(r"^-\s+(.+?):\s*([0-9.]+)\s*\(权重")
# miss 行: "    - miss: <...>"
_QC_MISS_LINE_RE = re.compile(r"^\s+-\s+miss:")


def _read_paragraphs(path: Path) -> list[str]:
    doc = Document(str(path))
    out = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            out.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    out.append(t)
    return out


def _baseline_for(artifact: Path) -> Path | None:
    """按文件名骨架找对应的 samples/ 基线."""
    stem = artifact.stem
    # strip 常见 v16/_v15/_test 后缀
    core = re.sub(r"_v\d+$|_test.*$", "", stem)
    for candidate in SAMPLES_DIR.glob("*.docx"):
        if candidate.stem in core or core in candidate.stem:
            return candidate
    return None


def _locate_v16_summary(artifact: Path) -> Path | None:
    """在 artifact 附近查找 v16_pipeline_summary.json.

    查找顺序:
      1. artifact 本身是 v16_pipeline_summary.json → 直接用
      2. artifact 是目录 → <artifact>/v16_pipeline_summary.json
      3. artifact 是文件 → <artifact.parent>/v16_pipeline_summary.json
      4. 否则 None
    """
    if artifact.is_file() and artifact.name == V16_SUMMARY_FILENAME:
        return artifact
    if artifact.is_dir():
        p = artifact / V16_SUMMARY_FILENAME
        return p if p.exists() else None
    if artifact.is_file():
        p = artifact.parent / V16_SUMMARY_FILENAME
        return p if p.exists() else None
    return None


def _pick_summary_entry(summary_doc: Any, artifact: Path) -> dict | None:
    """summary JSON 是 list[dict], 每 dict 对应一次 source_docx 跑分.

    artifact = .docx 时按 source_docx basename / output_docx basename 匹配;
    artifact 非 .docx 时返回第一条 entry.
    """
    if isinstance(summary_doc, dict):
        # 单条 pipeline 也可能直接是 dict
        return summary_doc
    if not isinstance(summary_doc, list) or not summary_doc:
        return None
    art_name = artifact.name.lower() if artifact and artifact.suffix == ".docx" else None
    if art_name:
        for entry in summary_doc:
            src = Path(entry.get("source_docx", "")).name.lower()
            out = Path(entry.get("output_docx", "")).name.lower()
            if art_name in (src, out):
                return entry
    return summary_doc[0]


def _parse_qc_md(md_path: Path) -> dict[str, Any]:
    """解析 v16 QC md 的 9 维度行, 返回:
      {
        'dim_count': int,
        'dims_without_miss': int,
        'dim_scores': {name: raw_score},
      }
    未能解析时返回 dim_count=0.
    """
    try:
        text = md_path.read_text(encoding="utf-8")
    except OSError:
        return {"dim_count": 0, "dims_without_miss": 0, "dim_scores": {}}

    lines = text.splitlines()
    in_dim_section = False
    dim_scores: dict[str, float] = {}
    dim_names_ordered: list[str] = []
    # miss 归属最近一个 dim
    current_dim: str | None = None
    dims_with_miss: set[str] = set()

    for line in lines:
        if line.startswith("## "):
            in_dim_section = line.startswith("## 9 维度评分")
            current_dim = None
            continue
        if not in_dim_section:
            continue
        m = _QC_DIM_LINE_RE.match(line)
        if m:
            name = m.group(1).strip()
            try:
                score = float(m.group(2))
            except ValueError:
                continue
            dim_scores[name] = score
            dim_names_ordered.append(name)
            current_dim = name
            continue
        if current_dim and _QC_MISS_LINE_RE.match(line):
            dims_with_miss.add(current_dim)

    dim_count = len(dim_names_ordered)
    dims_without_miss = sum(1 for d in dim_names_ordered if d not in dims_with_miss)
    return {
        "dim_count": dim_count,
        "dims_without_miss": dims_without_miss,
        "dim_scores": dim_scores,
    }


@register_evaluator("report")
class Agent6ReportEvaluator(BaseEvaluator):
    agent_id = "report"
    config_name = "agent6_report.yaml"

    def load_artifacts(self, run: EvalRun) -> dict[str, Any]:
        # 空 artifact 时所有依赖 artifact 的指标走 manual fallback
        if not run.artifacts:
            return {
                "paragraphs": [],
                "artifact_path": None,
                "baseline_path": None,
                "v16_summary_entry": None,
                "v16_summary_path": None,
                "qc_report": None,
                "qc_md_path": None,
            }
        art_path = Path(run.artifacts[0])
        if not art_path.is_absolute():
            art_path = REPO_ROOT / art_path
        if not art_path.exists():
            return {
                "paragraphs": [],
                "artifact_path": str(art_path),
                "baseline_path": None,
                "v16_summary_entry": None,
                "v16_summary_path": None,
                "qc_report": None,
                "qc_md_path": None,
                "error": f"artifact missing: {art_path}",
            }

        # --- docx 侧 (仅当 artifact 是 docx 文件时才读段落, 目录/json 跳过) ---
        paragraphs: list[str] = []
        baseline_path: Path | None = None
        baseline_paragraphs: list[str] = []
        if art_path.is_file() and art_path.suffix.lower() == ".docx":
            paragraphs = _read_paragraphs(art_path)
            baseline_path = _baseline_for(art_path)
            if baseline_path:
                baseline_paragraphs = _read_paragraphs(baseline_path)

        # --- v16 summary 查找 + 解析 ---
        summary_path = _locate_v16_summary(art_path)
        summary_entry: dict | None = None
        qc_report: dict[str, Any] | None = None
        qc_md_path: Path | None = None
        if summary_path and summary_path.exists():
            try:
                summary_doc = json.loads(summary_path.read_text(encoding="utf-8"))
                summary_entry = _pick_summary_entry(summary_doc, art_path)
            except (json.JSONDecodeError, OSError):
                summary_entry = None

            # 若 summary 指向的 output_docx 存在, 从它读段落做 hallucination_rate 分母
            if summary_entry and not paragraphs:
                out_docx_ref = summary_entry.get("output_docx")
                if out_docx_ref:
                    out_docx = Path(out_docx_ref)
                    if not out_docx.is_absolute():
                        out_docx = REPO_ROOT / out_docx
                    if out_docx.exists():
                        paragraphs = _read_paragraphs(out_docx)
                        if not baseline_path:
                            baseline_path = _baseline_for(out_docx)
                            if baseline_path:
                                baseline_paragraphs = _read_paragraphs(baseline_path)

            # QC md 解析 (9 维度 miss 情况)
            if summary_entry:
                qc = summary_entry.get("qc") or {}
                md_ref = qc.get("md")
                if md_ref:
                    md_p = Path(md_ref)
                    if not md_p.is_absolute():
                        md_p = REPO_ROOT / md_p
                    if md_p.exists():
                        qc_md_path = md_p
                        qc_report = _parse_qc_md(md_p)

        return {
            "paragraphs": paragraphs,
            "artifact_path": str(art_path),
            "baseline_path": str(baseline_path) if baseline_path else None,
            "baseline_paragraphs": baseline_paragraphs,
            "v16_summary_entry": summary_entry,
            "v16_summary_path": str(summary_path) if summary_path else None,
            "qc_report": qc_report,
            "qc_md_path": str(qc_md_path) if qc_md_path else None,
        }

    def compute_common_metrics(self, artifacts: dict[str, Any]) -> list[MetricOutcome]:
        paragraphs = artifacts.get("paragraphs") or []
        art_path = artifacts.get("artifact_path")
        summary_entry: dict | None = artifacts.get("v16_summary_entry")
        qc_report: dict | None = artifacts.get("qc_report")
        summary_path = artifacts.get("v16_summary_path")
        qc_md_path = artifacts.get("qc_md_path")

        # task_completion_rate: 能打开 + 段落数非零 ⇒ 1.0, 否则 0.0
        task_completion = 1.0 if paragraphs else 0.0
        out: list[MetricOutcome] = [
            self.mark(
                "task_completion_rate",
                task_completion,
                method="deterministic",
                evidence=[art_path] if art_path else [],
                note=f"{len(paragraphs)} paragraphs",
                kind="common",
            )
        ]

        # --- field_completeness: 仍走 manual stub (v16 summary 无此字段) ---
        out.append(
            MetricOutcome(
                name="field_completeness",
                value=None,
                target=self._lookup_target("field_completeness", "common") or "n/a",
                passed=None,
                method="manual",
                note="Phase B stub — v16 summary 未暴露 filled/expected 字段统计",
            )
        )

        # --- evidence_rate: v16 summary + QC md → 9 维度无 miss 的占比 ---
        if qc_report and qc_report.get("dim_count"):
            er = qc_report["dims_without_miss"] / qc_report["dim_count"]
            out.append(
                self.mark(
                    "evidence_rate",
                    er,
                    method="deterministic",
                    evidence=[qc_md_path] if qc_md_path else [],
                    note=(
                        f"{qc_report['dims_without_miss']}/{qc_report['dim_count']} "
                        "QC 9 维度无 miss 行 (QC md parse)"
                    ),
                    kind="common",
                )
            )
        else:
            out.append(
                MetricOutcome(
                    name="evidence_rate",
                    value=None,
                    target=self._lookup_target("evidence_rate", "common") or "n/a",
                    passed=None,
                    method="manual",
                    note="fallback: 无 v16 summary 或 QC md 未找到",
                )
            )

        # --- hallucination_rate: qc.hallucinations / paragraph_count ---
        if summary_entry and isinstance(summary_entry.get("qc"), dict):
            qc = summary_entry["qc"]
            halluc = qc.get("hallucinations")
            if isinstance(halluc, (int, float)) and paragraphs:
                hr = halluc / max(len(paragraphs), 1)
                out.append(
                    self.mark(
                        "hallucination_rate",
                        hr,
                        method="deterministic",
                        evidence=[summary_path] if summary_path else [],
                        note=(
                            f"{int(halluc)} halluc findings / {len(paragraphs)} paragraphs "
                            "(v16 summary qc.hallucinations ÷ 段落数)"
                        ),
                        kind="common",
                    )
                )
            elif isinstance(halluc, (int, float)):
                out.append(
                    MetricOutcome(
                        name="hallucination_rate",
                        value=None,
                        target=self._lookup_target("hallucination_rate", "common") or "n/a",
                        passed=None,
                        method="heuristic",
                        note=f"summary qc.hallucinations={halluc} 但段落分母为 0",
                    )
                )
            else:
                out.append(
                    MetricOutcome(
                        name="hallucination_rate",
                        value=None,
                        target=self._lookup_target("hallucination_rate", "common") or "n/a",
                        passed=None,
                        method="manual",
                        note="fallback: v16 summary qc.hallucinations 字段缺失",
                    )
                )
        else:
            out.append(
                MetricOutcome(
                    name="hallucination_rate",
                    value=None,
                    target=self._lookup_target("hallucination_rate", "common") or "n/a",
                    passed=None,
                    method="manual",
                    note="fallback: 无 v16 summary",
                )
            )

        # --- tool_success_rate: 仍 manual (v16 summary 无 tool 元数据) ---
        out.append(
            MetricOutcome(
                name="tool_success_rate",
                value=None,
                target=self._lookup_target("tool_success_rate", "common") or "n/a",
                passed=None,
                method="manual",
                note="Phase B stub — v16 summary 未暴露 tool_calls 元数据",
            )
        )
        return out

    def compute_domain_metrics(self, artifacts: dict[str, Any]) -> list[MetricOutcome]:
        paragraphs: list[str] = artifacts.get("paragraphs") or []
        baseline_paragraphs: list[str] = artifacts.get("baseline_paragraphs") or []
        art_path = artifacts.get("artifact_path")
        base_path = artifacts.get("baseline_path")

        out: list[MetricOutcome] = []

        # --- template_leakage_rate ---
        if paragraphs and baseline_paragraphs:
            leaks = 0
            for p in paragraphs:
                best = 0.0
                for bp in baseline_paragraphs:
                    r = SequenceMatcher(None, p, bp).ratio()
                    if r > best:
                        best = r
                        if best >= SIMILARITY_HIT:
                            break
                if best >= SIMILARITY_HIT:
                    leaks += 1
            leakage = leaks / max(len(paragraphs), 1)
            out.append(
                self.mark(
                    "template_leakage_rate",
                    leakage,
                    method="deterministic",
                    evidence=[base_path] if base_path else [],
                    note=f"{leaks}/{len(paragraphs)} paragraphs >= {SIMILARITY_HIT} vs baseline",
                )
            )
        else:
            out.append(
                MetricOutcome(
                    name="template_leakage_rate",
                    value=None,
                    target=self._lookup_target("template_leakage_rate", "domain") or "n/a",
                    passed=None,
                    method="heuristic",
                    note="缺 artifact 或 baseline, 跳过",
                )
            )

        # --- unfilled_marker_accuracy ---
        placeholder_hits = sum(bool(PLACEHOLDER_PAT.search(p)) for p in paragraphs)
        unfilled_marks = sum(UNFILLED_MARK in p for p in paragraphs)
        total = len(paragraphs) or 1
        # 简化公式: 1 - 异常率 (占位符残留视为漏标, 无占位符但有 unfilled_mark 视为正常)
        abnormal = placeholder_hits
        accuracy = 1.0 - abnormal / total
        out.append(
            self.mark(
                "unfilled_marker_accuracy",
                max(0.0, min(1.0, accuracy)),
                method="heuristic",
                evidence=[art_path] if art_path else [],
                note=f"placeholder_residue={placeholder_hits}, unfilled_marks={unfilled_marks}",
            )
        )

        # --- quality_score_total: v16 summary qc.score (0-100 scale, 直接) ---
        summary_entry = artifacts.get("v16_summary_entry")
        summary_path = artifacts.get("v16_summary_path")
        if summary_entry and isinstance(summary_entry.get("qc"), dict):
            score = summary_entry["qc"].get("score")
            if isinstance(score, (int, float)):
                out.append(
                    self.mark(
                        "quality_score_total",
                        float(score),
                        method="deterministic",
                        evidence=[summary_path] if summary_path else [],
                        note=(
                            f"v16 summary qc.score={score:.2f}/100 "
                            f"(passed={summary_entry['qc'].get('passed')}, "
                            f"fatal_fail={summary_entry['qc'].get('fatal_fail')})"
                        ),
                    )
                )
            else:
                out.append(
                    MetricOutcome(
                        name="quality_score_total",
                        value=None,
                        target=self._lookup_target("quality_score_total", "domain") or "n/a",
                        passed=None,
                        method="manual",
                        note="fallback: v16 summary qc.score 字段缺失",
                    )
                )
        else:
            out.append(
                MetricOutcome(
                    name="quality_score_total",
                    value=None,
                    target=self._lookup_target("quality_score_total", "domain") or "n/a",
                    passed=None,
                    method="manual",
                    note="fallback: 无 v16 summary · 需 quality_scorer runtime",
                )
            )

        # --- financial_ratio_consistency — EV-12 Batch 2 交由
        # evaluation.runner.cross_agent.ratio_consistency 主编排计算 · 本 adapter
        # 不 inline 算 (保持 agent6 adapter 单 artifact 单 DP 语义), 此处保留 manual
        # stub, 由 baseline JSON 顶层 ev_12_ratio_calc_consistency 承载.
        out.append(
            MetricOutcome(
                name="financial_ratio_consistency",
                value=None,
                target=self._lookup_target("financial_ratio_consistency", "domain") or "n/a",
                passed=None,
                method="manual",
                note=(
                    "由 evaluation.runner.cross_agent.ratio_consistency 跨 5 DP 聚合算 "
                    "(baseline JSON 顶层 ev_12_ratio_calc_consistency 字段, method 升级为 "
                    "deterministic_cross_agent)"
                ),
            )
        )
        # --- section_length_calibration 仍 pending (真人范文库依赖) ---
        out.append(
            MetricOutcome(
                name="section_length_calibration",
                value=None,
                target=self._lookup_target("section_length_calibration", "domain") or "n/a",
                passed=None,
                method="manual",
                note="Phase B stub — 真人范文库依赖 (Batch 3)",
            )
        )

        return out

    def _extract_financial_ratios(self, enterprise_id: str) -> dict[str, float | None]:
        """EV-12 入口 (instance method; 委派给 module-level extract_financial_ratios)."""
        return extract_financial_ratios(enterprise_id)


# ═══════════════════════════════════════════════════════════════════
# Batch 2 Task B · EV-12 跨 Agent 一致率用
#   agent6_report 侧独立实现 (与 agent3_credit 并行不 import 对方),
#   都读 data/mock/deep-pillar/<eid>_*/ 下 xlsx → FinancialAnalyzer.analyze.
#   若有朝一日某侧改走 LLM / 硬编数字, EV-12 的 consistency 会下降至 < 99%
#   触发 blocker_threshold 阻断发布.
# ═══════════════════════════════════════════════════════════════════

DEEP_PILLAR_ROOT = REPO_ROOT / "data" / "mock" / "deep-pillar"


def _resolve_financial_xlsx_agent6(enterprise_id: str) -> Path | None:
    """Agent6 侧: 按报告管线 material_kb.parse_client_folder 的"财报优先级".

    v16 pipeline 的 material_kb 按文件名特征排: 财务报表 > 年度 > 审计.
    """
    candidates = sorted(DEEP_PILLAR_ROOT.glob(f"{enterprise_id}_*"))
    if not candidates:
        return None
    client_dir = candidates[0]
    patterns = ["*财务报表2025*.xlsx", "*财务报表2024*.xlsx", "*财务报表*.xlsx", "*.xlsx"]
    for pat in patterns:
        xlsx = sorted(client_dir.glob(pat), reverse=True)
        for p in xlsx:
            if p.is_file():
                return p
    return None


def extract_financial_ratios(enterprise_id: str) -> dict[str, float | None]:
    """Agent6 独立调用 financial_analyzer 的入口 (EV-12 双侧独立之二).

    红线: 只读消费 financial_analyzer.FinancialAnalyzer, 不改 v16_* 业务代码.
    """
    xlsx = _resolve_financial_xlsx_agent6(enterprise_id)
    empty = {n: None for n in ("current_ratio", "debt_ratio", "roe", "gross_margin")}
    if xlsx is None or not xlsx.exists():
        return empty

    import sys as _sys
    root_str = str(REPO_ROOT)
    if root_str not in _sys.path:
        _sys.path.insert(0, root_str)
    try:
        from financial_analyzer import FinancialAnalyzer  # noqa: E402
    except ImportError:
        return empty

    try:
        ind = FinancialAnalyzer().analyze(str(xlsx))
    except Exception:
        return empty

    current_ratio = ind.current_ratio.current if ind.current_ratio else None
    debt_ratio = ind.debt_to_asset_ratio.current if ind.debt_to_asset_ratio else None
    gross_margin = ind.gross_margin.current if ind.gross_margin else None
    # ROE = net_profit / total_equity (Agent6 侧就地算, 与 Agent3 对齐)
    roe: float | None = None
    try:
        np_ = ind.net_profit.current if ind.net_profit else None
        eq_ = ind.total_equity.current if ind.total_equity else None
        if np_ is not None and eq_ not in (None, 0):
            roe = float(np_) / float(eq_) * 100.0
    except (TypeError, ValueError, ZeroDivisionError):
        roe = None

    return {
        "current_ratio": current_ratio,
        "debt_ratio": debt_ratio,
        "roe": roe,
        "gross_margin": gross_margin,
    }
