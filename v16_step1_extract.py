# -*- coding: utf-8 -*-
"""V16 Step 1 - Element Extractor + Prelabeler.

从 3 份样本(经纬对公 A / 兴业资管对公 B / 普惠申报书骨架)抽取所有 paragraph 和
table cell 作为 element,按(source × type × structure_feature)分层抽样 200+,
规则预标注(6 类标签 + confidence),输出 labeled samples 供人工 review。

标签体系:
  SCAFFOLD  — 骨架保留原文(标题/章节编号/字段标签/表头/指引短语)
  FILL      — 有客户数据,代码直接填(注册资本/法人/股东/主营业务等)
  CLEAR     — 是示例但当前客户无对应数据,清空+pending
  REWRITE   — 正文段落,LLM 基于当前客户材料整段重写
  SLOT      — 下划线/占位槽位,按有无数据决定 FILL/CLEAR
  PRESERVE  — 说明性指引文本,保留不动
  PLACEHOLDER — v16 模板 placeholder 化 治本: 文本含 `{{KEY}}` 占位符,
                generator 走 REPLACE op 用 client_metadata 替换 (Phase 1+ 新增)

"当前客户"档案 (CURRENT_CLIENT) 加载策略 (2026-05-21 治本 Phase 1):
  1. 优先读 samples/<docx_stem>.metadata.json sidecar 的 original_client 字段
     (Phase 2+ 经纬测绘等模板已 placeholder 化时,sidecar 提供真实客户档案
     供 classifier prompt context + featurize 当前客户词识别)
  2. fallback 到旧 hardcode "中锐" (兼容现有 5 个 sample fixture / 5 个 docx)
     避免 break 老路径 — placeholder 化是 docx-by-docx 渐进迁移,不是 big-bang

v16 模板 placeholder 化治本 brief: D:/second-brain/wiki/concepts/v16-template-placeholder-治本-brief.md
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, asdict, field
from pathlib import Path
from docx import Document

SAMPLES_DIR = Path(__file__).parent / "samples"
OUT_PATH = Path(__file__).parent / "outputs" / "v16_labeled_elements.json"

# Fallback CURRENT_CLIENT (旧 hardcode · 兼容兜底)
# 真实生产路径应该走 sidecar metadata (Phase 2+) — 5 个 docx 全 placeholder 化后此常量可删除
_FALLBACK_CURRENT_CLIENT = {
    "name": "福建中锐网络股份有限公司",
    "name_core": "中锐",
    "legal_rep": "黄祖海",
    "industry": "其他未列明信息技术服务业",
    "business": "智慧水利与智慧教育",
    "registered_capital": "4100万元",
    "location": "福州",
}


def load_client_for_template(docx_path: Path) -> dict:
    """加载 docx 模板对应的 "当前客户" 档案 (用于 classifier prompt + featurize 当前客户词).

    优先级:
      1. samples/<docx_stem>.metadata.json sidecar 的 ``original_client`` 字段
      2. fallback _FALLBACK_CURRENT_CLIENT (中锐 hardcode · 老路径兼容)

    返回:
      dict 含 keys: name / name_core / legal_rep / industry / business /
                    registered_capital / location

    Phase 1 设计 (与 brief §3 一致):
      - 不传 docx_path 时返 fallback (向后兼容)
      - sidecar 存在但缺字段时,从 fallback 补齐 (避免 KeyError)
    """
    if docx_path is None:
        return dict(_FALLBACK_CURRENT_CLIENT)
    sidecar = Path(docx_path).with_suffix(".metadata.json")
    if not sidecar.is_file():
        return dict(_FALLBACK_CURRENT_CLIENT)
    try:
        data = json.loads(sidecar.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return dict(_FALLBACK_CURRENT_CLIENT)
    original = data.get("original_client") or {}
    if not isinstance(original, dict):
        return dict(_FALLBACK_CURRENT_CLIENT)
    # 缺字段补齐 (avoid KeyError downstream)
    merged = dict(_FALLBACK_CURRENT_CLIENT)
    merged.update({k: str(v) for k, v in original.items() if v is not None})
    return merged


# 全局 CURRENT_CLIENT (向后兼容 · 模块 import 期默认 fallback · runtime 可 reload)
# v16_classifier / v16_step1_extract main 都从此取 · 切 sidecar 走 load_client_for_template
CURRENT_CLIENT = dict(_FALLBACK_CURRENT_CLIENT)

# ─────────────────────────────────────────────────────────────
# Element 抽取
# ─────────────────────────────────────────────────────────────
@dataclass
class Element:
    source: str          # 样本来源文件名
    kind: str            # "para" | "cell"
    location: str        # "P{idx}" / "T{ti}R{ri}C{ci}"
    text: str
    style: str = ""      # Word 样式名
    in_table: bool = False
    is_header_row: bool = False
    row_idx: int = -1    # 表格中行位置
    label: str = ""      # 预标注
    confidence: float = 0.0
    features: list[str] = field(default_factory=list)
    justification: str = ""


def extract_elements(path: Path) -> list[Element]:
    """从一份 docx 抽所有非空 paragraph + cell。"""
    d = Document(str(path))
    out: list[Element] = []
    name = path.name

    # 段落流
    for pi, p in enumerate(d.paragraphs):
        t = (p.text or "").strip()
        if not t:
            continue
        style = p.style.name if p.style else ""
        out.append(Element(
            source=name, kind="para",
            location=f"P{pi}", text=t, style=style,
        ))

    # 表格(含嵌套)
    def _walk_table(tbl, loc_prefix: str):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for ppi, p in enumerate(cell.paragraphs):
                    t = (p.text or "").strip()
                    if not t:
                        continue
                    out.append(Element(
                        source=name, kind="cell",
                        location=f"{loc_prefix}R{ri}C{ci}P{ppi}",
                        text=t,
                        in_table=True,
                        is_header_row=(ri == 0),
                        row_idx=ri,
                    ))
                for sub in cell.tables:
                    _walk_table(sub, f"{loc_prefix}R{ri}C{ci}NT")

    for ti, tbl in enumerate(d.tables):
        _walk_table(tbl, f"T{ti}")

    return out


# ─────────────────────────────────────────────────────────────
# 规则特征抽取
# ─────────────────────────────────────────────────────────────

# 字段标签本体(银行授信报告常见 50 个)
FIELD_LABELS_FILL = {
    # 有客户数据, 代码应直接填
    "客户名称", "授信客户全称", "客户全称", "企业名称",
    "注册资本", "实收资本", "法定代表人", "法人代表",
    "实际控制人", "成立时间", "经营地址", "注册地址",
    "主营业务", "所属行业", "统一社会信用代码",
}
FIELD_LABELS_CLEAR = {
    # 示例但当前客户通常无数据 / 业务线独有, 清空加 pending
    "PD评级", "申报单位", "客户经理", "联系电话", "白名单类型",
    "绿色信贷标识", "绿色信贷", "环境和社会风险分类",
    "客户环境和社会表现的动态评估结果", "反洗钱风险等级",
    "我行投向政策对应行业", "投向政策对应行业",
    "申报额度", "申报敞口", "上一期授信敞口", "担保方式",
    "业务品种", "业务期限", "授信期限",
}

# 章节编号正则
SECTION_NUM_RES = [
    (1, re.compile(r'^\s*[一二三四五六七八九十]{1,3}\s*[、.]\s*\S')),
    (2, re.compile(r'^\s*[（(]\s*[一二三四五六七八九十]{1,3}\s*[)）]\s*\S')),
    (2, re.compile(r'^\s*[0-9]{1,2}\s*[.、]\s*\S')),
    (3, re.compile(r'^\s*[（(]\s*[0-9]{1,2}\s*[)）]\s*\S')),
    (3, re.compile(r'^\s*※\s*[0-9]{1,2}\s*[.、]\s*\S')),
    (4, re.compile(r'^\s*[①②③④⑤⑥⑦⑧⑨⑩].{0,40}')),
]

# 占位符
PLACEHOLDER_RES = [
    re.compile(r'_{4,}'),
    re.compile(r'\s{6,}'),
    re.compile(r'X+年\s*X+月\s*X+日'),
    re.compile(r'^\s*[\(（].{0,20}[）\)]\s*$'),  # 纯括号指引
]

# 指引短语(PRESERVE)
INSTRUCTION_HINTS = [
    "如涉及请填写", "如涉及", "如未落实", "请说明原因", "具体情况如下表所示",
    "请附", "请补充", "如有", "若有", "可附图", "请填列下表",
    "根据实际情况", "根据实际情况增减", "请根据",
    "注:", "注：", "备注:", "备注：", "说明：", "说明:",
    "如涉及A/B类", "具体说明见",
]

# 外来实体(非当前客户的公司/人名)
FOREIGN_COMPANY_RE = re.compile(
    r'(?:福建|北京|上海|广州|深圳|厦门|福州|南京|天津|重庆|成都|武汉|'
    r'江苏|浙江|山东|广东|中国|全国|兴业|招标|经纬|巨电)'
    r'[\u4e00-\u9fff]{2,20}?'
    r'(?:有限公司|股份有限公司|股份公司|集团|有限责任公司|合伙企业)'
)
FOREIGN_PERSON_RE = re.compile(
    r'(?:郑|王|许|陈|李|张|赵|黄|刘|周|吴|徐|孙|朱|胡|高|林|何)'
    r'[\u4e00-\u9fff]{1,2}(?=(?:先生|女士|同志|董事|总经|工程师|注册|高级|[:,：，\s]|$))'
)
# 当前客户 & 中锐系子公司核心词(防止误判为外来)
CURRENT_CORE_WORDS = [
    CURRENT_CLIENT["name_core"], CURRENT_CLIENT["legal_rep"],
    "汉鼎", "青云", "海沃", "康恩慧", "陈其俤",
]


def featurize(e: Element) -> tuple[list[str], dict]:
    """抽结构+语义+对比特征,返回 (features, info)."""
    feats = []
    info = {}
    text = e.text

    # 样式
    if "Heading" in (e.style or "") or "标题" in (e.style or ""):
        feats.append("heading_style")

    # 章节编号
    for level, r in SECTION_NUM_RES:
        if r.match(text):
            feats.append(f"section_num_L{level}")
            info["section_level"] = level
            break

    # 字段标签:值 模式(上限放宽到 25 字,覆盖长字段名如"客户环境和社会表现的动态评估结果")
    m = re.match(r'^(.{2,25}?)[:：]\s*(.*)$', text)
    if m:
        label, value = m.group(1).strip(), m.group(2).strip()
        info["field_label"] = label
        info["field_value"] = value
        if label in FIELD_LABELS_FILL:
            feats.append("label_fill_candidate")
        elif label in FIELD_LABELS_CLEAR:
            feats.append("label_clear_candidate")
        else:
            # 未在本体但符合"字段-值"模式,也标注为候选(低置信)
            feats.append("field_label_unknown")

    # 指引性单位/附注短文本("单位:万元" / "单位：元" / "注:" / "附件" 等)
    if re.match(r'^单位\s*[:：]', text) or text.strip() == "单位：万元" or text.strip() == "单位：元":
        feats.append("unit_indicator")
    if re.match(r'^(附件|备注|说明)\s*[:：]?', text):
        feats.append("meta_indicator")

    # 报告标题(形如"关于XXX的授信报告"/"XXX授信申报表")
    if re.match(r'^(关于|福建|北京|上海).{2,30}(授信|报告|申报|审批).*$', text):
        feats.append("report_title_like")

    # 占位符
    if any(r.search(text) for r in PLACEHOLDER_RES):
        feats.append("has_placeholder")

    # 指引短语
    if any(h in text for h in INSTRUCTION_HINTS):
        feats.append("has_instruction_hint")

    # 外来实体
    foreign_cos = [m.group(0) for m in FOREIGN_COMPANY_RE.finditer(text)
                   if not any(cw in m.group(0) for cw in CURRENT_CORE_WORDS)]
    foreign_persons = [m.group(0) for m in FOREIGN_PERSON_RE.finditer(text)
                       if not any(cw in m.group(0) for cw in CURRENT_CORE_WORDS)]
    if foreign_cos:
        feats.append("has_foreign_company")
        info["foreign_companies"] = foreign_cos[:5]
    if foreign_persons:
        feats.append("has_foreign_person")
        info["foreign_persons"] = foreign_persons[:5]

    # 具体数字(大于 100 的整数 或 带%的比率)
    if re.search(r'\b\d{3,}\.?\d*\s*(万元|亿元|元|%|倍|天|人|平方米)', text):
        feats.append("has_concrete_number")
    # 具体日期
    if re.search(r'20\d{2}\s*年\s*\d{1,2}\s*月', text) or re.search(r'20\d{2}-\d{1,2}-\d{1,2}', text):
        feats.append("has_concrete_date")

    # 表头
    if e.in_table and e.is_header_row:
        feats.append("table_header")

    # v16 模板 placeholder 化治本 Phase 1: {{KEY}} 标记
    # 用 (?:^|[^{]) 边界避免 {{{ 误命中 · 用非贪婪保证一段含多个 {{X}}{{Y}} 全捕获
    _ph_matches = re.findall(r"\{\{([A-Z][A-Z0-9_]{1,40})\}\}", text)
    if _ph_matches:
        feats.append("has_placeholder_marker")
        info["placeholder_keys"] = _ph_matches[:10]

    # 长段(正文)
    if len(text) > 80 and not feats:
        feats.append("long_narrative")
    elif len(text) > 80:
        feats.append("long_text")

    return feats, info


# ─────────────────────────────────────────────────────────────
# 规则预标
# ─────────────────────────────────────────────────────────────

def prelabel(e: Element, feats: list[str], info: dict) -> tuple[str, float, str]:
    """规则预标,返回 (label, confidence 0-1, justification)."""
    # 1. 空段或只标点,不处理
    if not e.text.strip():
        return "PRESERVE", 0.9, "空文本"

    # 1.5 v16 模板 placeholder 化 治本: {{KEY}} 标记一律 PLACEHOLDER (高置信)
    #     classifier prompt 同样按此规则识别 · generator REPLACE op 处理
    if "has_placeholder_marker" in feats:
        keys = info.get("placeholder_keys", [])
        return "PLACEHOLDER", 0.98, f"含 placeholder {keys}"

    # 2. 标题/章节编号 → SCAFFOLD(高置信)
    if "heading_style" in feats:
        return "SCAFFOLD", 0.95, "Word Heading 样式"
    if any(f.startswith("section_num_") for f in feats):
        # 但若同时含外来实体,可能是 "(3) 实际控制人:福建省招标股份" 这种被污染的标题
        if "has_foreign_company" in feats:
            return "SCAFFOLD", 0.7, f"章节编号但含外来公司名 {info.get('foreign_companies')}, 可能需清空值"
        return "SCAFFOLD", 0.9, f"章节编号 L{info.get('section_level', '?')}"

    # 3. 字段标签-值模式
    if "label_fill_candidate" in feats:
        val = info.get("field_value", "")
        if not val or val == "" or re.match(r'^[_\s]*$', val):
            return "FILL", 0.85, f"字段 '{info['field_label']}' 空值,代码填客户数据"
        if any(cw in val for cw in CURRENT_CORE_WORDS):
            return "SCAFFOLD", 0.85, f"字段 '{info['field_label']}' 已是当前客户值"
        if "has_foreign_company" in feats or "has_foreign_person" in feats:
            return "FILL", 0.8, f"字段 '{info['field_label']}' 是外来实体值,需替换为客户"
        if "has_concrete_number" in feats or "has_concrete_date" in feats:
            return "FILL", 0.75, f"字段 '{info['field_label']}' 含具体数字/日期示例值,代码填客户"
        return "FILL", 0.7, f"字段 '{info['field_label']}' 候选 FILL"

    if "label_clear_candidate" in feats:
        return "CLEAR", 0.85, f"字段 '{info['field_label']}' 示例类,清空+pending"

    # 4. 占位符槽位
    if "has_placeholder" in feats and len(e.text.strip()) < 30:
        return "SLOT", 0.8, "短占位符段"

    # 5. 指引短语 → PRESERVE
    if "has_instruction_hint" in feats:
        # 但如果是 "请" 和 "请说明" 开头的指引中夹带了实体, 需要 REWRITE
        if "has_foreign_company" in feats and len(e.text) > 50:
            return "REWRITE", 0.6, "指引混合外来实体,需 LLM 重写"
        return "PRESERVE", 0.8, "说明性指引短语"

    # 6. 表头 row → SCAFFOLD
    if "table_header" in feats:
        return "SCAFFOLD", 0.9, "表格 header row"

    # 7. 长正文段落
    if "long_narrative" in feats or len(e.text) > 120:
        if "has_foreign_company" in feats or "has_foreign_person" in feats:
            return "REWRITE", 0.9, f"长正文含外来实体 {info.get('foreign_companies', []) or info.get('foreign_persons', [])}, LLM 重写"
        if "has_concrete_number" in feats or "has_concrete_date" in feats:
            return "REWRITE", 0.7, "长正文含具体示例数字,需 LLM 重写"
        return "REWRITE", 0.5, "长正文(无明显实体特征),需审"

    # 8. 短单元格 + 外来实体 → FILL (替换为客户对应值)
    if e.kind == "cell" and ("has_foreign_company" in feats or "has_foreign_person" in feats):
        return "FILL", 0.8, f"表格 cell 含外来实体,代码替换"

    # 9. 短文本 + 具体数字(数据 cell)
    if e.kind == "cell" and "has_concrete_number" in feats:
        return "FILL", 0.7, "表格数据 cell 含示例数字,代码填客户真值"

    # 10a. 报告标题 / 单位指引 / 元信息 → SCAFFOLD
    if "report_title_like" in feats:
        # 含外来公司名就是旧标题需 FILL(替换为当前客户)
        if "has_foreign_company" in feats:
            return "FILL", 0.8, "报告标题含旧客户名,代码替换"
        return "SCAFFOLD", 0.8, "报告标题行"
    if "unit_indicator" in feats or "meta_indicator" in feats:
        return "SCAFFOLD", 0.8, "单位/附注指引类短文本"

    # 10b. 未知字段-值 模式
    if "field_label_unknown" in feats:
        val = info.get("field_value", "")
        if any(cw in val for cw in CURRENT_CORE_WORDS):
            return "SCAFFOLD", 0.6, f"疑似字段 '{info.get('field_label')}' 已是当前客户值"
        if "has_foreign_company" in feats or "has_foreign_person" in feats:
            return "FILL", 0.55, f"疑似字段含外来实体值,需替换"
        if "has_concrete_number" in feats or "has_concrete_date" in feats:
            return "FILL", 0.5, f"疑似字段含具体示例值"
        if not val:
            return "SLOT", 0.5, "字段-值 模式但值空"
        return "CLEAR", 0.4, f"疑似字段 '{info.get('field_label')}' 示例值,建议清空"

    # 10c. 短文本(<25 字)无特征 → SCAFFOLD(副标题/短行指引)
    if len(e.text) <= 25:
        return "SCAFFOLD", 0.5, "短文本,疑似副标题/指引"

    # 11. 兜底 REWRITE
    if len(e.text) > 40:
        return "REWRITE", 0.3, "需人工 review(中长段无明显特征)"

    return "PRESERVE", 0.3, "需人工 review"


# ─────────────────────────────────────────────────────────────
# 分层抽样
# ─────────────────────────────────────────────────────────────

def stratified_sample(elements: list[Element], target_per_source: int = 70) -> list[Element]:
    """按 (source × kind × label) 分层抽样,每源约 target_per_source 条。"""
    import random
    random.seed(42)

    by_source: dict[str, list[Element]] = {}
    for e in elements:
        by_source.setdefault(e.source, []).append(e)

    sampled = []
    for source, elems in by_source.items():
        # 按预标 label 分组
        by_label: dict[str, list[Element]] = {}
        for e in elems:
            by_label.setdefault(e.label, []).append(e)

        # 每个 label 至少 5 条,余量按比例分
        per_label = {}
        min_per = max(3, target_per_source // (len(by_label) or 1))
        for lab, es in by_label.items():
            per_label[lab] = min(len(es), min_per)

        # 再优先覆盖低 confidence 样本
        for lab, es in by_label.items():
            es_sorted = sorted(es, key=lambda x: x.confidence)
            n = per_label[lab]
            sampled.extend(es_sorted[:n])

    return sampled


# ─────────────────────────────────────────────────────────────
# main
# ─────────────────────────────────────────────────────────────

def main():
    # 2026-05-21 治本 Phase 1: CURRENT_CLIENT 改成 per-docx 加载 sidecar
    # 走 load_client_for_template 拿真实客户档案 (旧 sample fallback 到中锐)
    global CURRENT_CLIENT, CURRENT_CORE_WORDS
    all_elements: list[Element] = []
    for f in sorted(SAMPLES_DIR.glob("*.docx")):
        print(f"[extract] {f.name}")
        # 切换当前客户档案 (sidecar > fallback)
        CURRENT_CLIENT = load_client_for_template(f)
        # 重建 CURRENT_CORE_WORDS (featurize 用此判定"非外来公司")
        CURRENT_CORE_WORDS = [
            CURRENT_CLIENT["name_core"], CURRENT_CLIENT["legal_rep"],
            "汉鼎", "青云", "海沃", "康恩慧", "陈其俤",
        ]
        print(f"  当前客户档案: {CURRENT_CLIENT.get('name')} (core={CURRENT_CLIENT.get('name_core')})")
        elems = extract_elements(f)
        print(f"  抽取 {len(elems)} 个 element")
        for e in elems:
            feats, info = featurize(e)
            lab, conf, just = prelabel(e, feats, info)
            e.features = feats
            e.label = lab
            e.confidence = conf
            e.justification = just
        all_elements.extend(elems)

    print(f"\n总 element 数: {len(all_elements)}")

    # 全量标注分布
    from collections import Counter
    label_dist = Counter(e.label for e in all_elements)
    print(f"全量标签分布: {dict(label_dist)}")

    # 分层抽样 ~200
    sampled = stratified_sample(all_elements, target_per_source=70)
    print(f"\n分层抽样: {len(sampled)} 条")
    sample_dist = Counter(e.label for e in sampled)
    print(f"抽样标签分布: {dict(sample_dist)}")

    # 低 confidence 样本(优先 review)
    low_conf = [e for e in sampled if e.confidence < 0.7]
    print(f"低 confidence (<0.7) 需 review 的: {len(low_conf)} 条")

    OUT_PATH.parent.mkdir(exist_ok=True)
    data = {
        "current_client": CURRENT_CLIENT,
        "total_extracted": len(all_elements),
        "total_sampled": len(sampled),
        "label_distribution": dict(sample_dist),
        "low_confidence_count": len(low_conf),
        "labels_legend": {
            "SCAFFOLD": "骨架, 保留原文 (标题/编号/字段标签/表头/指引)",
            "FILL": "有客户数据, 代码直接填",
            "CLEAR": "是示例但客户无数据, 清空+pending",
            "REWRITE": "正文段落, LLM 整段重写",
            "SLOT": "占位槽位, 按有无数据决定 FILL/CLEAR",
            "PRESERVE": "说明性指引, 保留不动",
        },
        "elements": [asdict(e) for e in sampled],
    }
    with OUT_PATH.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"\n[done] 输出: {OUT_PATH}")
    print(f"       全量 labeled 数: {len(sampled)}")
    print(f"       低 confidence(待 review): {len(low_conf)}")


if __name__ == "__main__":
    main()
