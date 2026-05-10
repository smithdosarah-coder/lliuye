# -*- coding: utf-8 -*-
"""agent_compliance.demo_loader — Phase B.2 (PM 2026-05-10 真意 reframe).

`/api/compliance/demo/run` 改真后端跑 · 输入来自 `data/mock/compliance-kb/manifest.json`
描述的 sample 政策原文 + 内部制度库 docx · 反 5 原则 §3.5 内部稳态 mock + 外部 (Tavily) 改真接.

mock 只能 mock 输入 · 不能 mock 结果 · 真后端 pipeline (LLM/Tavily/算法) 跑出真违规.
"""
from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any

PROJECT_ROOT = Path(__file__).resolve().parent.parent
MANIFEST_PATH = PROJECT_ROOT / "data" / "mock" / "compliance-kb" / "manifest.json"


class DemoBatchError(Exception):
    """manifest / sample 文件加载失败 · /demo/run 用 typed code 暴露 · 不 silent fallback."""

    def __init__(self, message: str, *, code: str):
        super().__init__(message)
        self.code = code


@dataclass
class DemoBatch:
    scenario_id: str
    label: str
    policy_doc: str
    policy_meta: dict[str, Any]
    business_docs: list[str]
    business_doc_sources: list[str]


def _load_manifest() -> dict:
    if not MANIFEST_PATH.exists():
        raise DemoBatchError(
            f"manifest.json 缺失: {MANIFEST_PATH}",
            code="DEMO_MANIFEST_MISSING",
        )
    try:
        return json.loads(MANIFEST_PATH.read_text("utf-8"))
    except json.JSONDecodeError as e:
        raise DemoBatchError(
            f"manifest.json 解析失败: {e}",
            code="DEMO_MANIFEST_PARSE",
        ) from e


def list_scenarios() -> list[dict[str, Any]]:
    manifest = _load_manifest()
    out: list[dict[str, Any]] = []
    for sid, body in (manifest.get("scenarios") or {}).items():
        out.append({
            "scenario_id": sid,
            "label": body.get("label", sid),
            "policy_title": (body.get("policy_meta") or {}).get("title", ""),
            "doc_count": len(body.get("business_doc_paths") or []),
        })
    return out


def _read_text(path: Path) -> str:
    return path.read_text("utf-8")


def _read_docx(path: Path) -> str:
    """python-docx 读 docx · paragraph + table cell 全平铺."""
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as e:
        raise DemoBatchError(
            "python-docx 未安装 · 无法解析制度库 docx",
            code="DEMO_DOCX_DEP_MISSING",
        ) from e

    try:
        doc = docx.Document(str(path))
    except (OSError, ValueError) as e:
        raise DemoBatchError(
            f"docx 打开失败: {path.name}: {type(e).__name__}: {e}",
            code="DEMO_DOCX_READ_FAIL",
        ) from e

    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)


def load_scenario(scenario_id: str) -> DemoBatch:
    """读 manifest 返完整 batch (policy_doc + business_docs 真 text)."""
    manifest = _load_manifest()
    scenarios = manifest.get("scenarios") or {}
    if scenario_id not in scenarios:
        allowed = "/".join(sorted(scenarios.keys())) or "<empty>"
        raise DemoBatchError(
            f"unknown scenario_id={scenario_id} (allowed: {allowed})",
            code="DEMO_SCENARIO_INVALID",
        )

    body = scenarios[scenario_id]
    policy_path = PROJECT_ROOT / body.get("policy_path", "")
    if not policy_path.is_file():
        raise DemoBatchError(
            f"sample 政策文件缺失: {policy_path.relative_to(PROJECT_ROOT)}",
            code="DEMO_POLICY_MISSING",
        )

    try:
        policy_doc = _read_text(policy_path)
    except OSError as e:
        raise DemoBatchError(
            f"sample 政策读失败: {policy_path.name}: {e}",
            code="DEMO_POLICY_READ_FAIL",
        ) from e

    if not policy_doc.strip():
        raise DemoBatchError(
            f"sample 政策内容为空: {policy_path.name}",
            code="DEMO_POLICY_EMPTY",
        )

    business_doc_paths = body.get("business_doc_paths") or []
    if not business_doc_paths:
        raise DemoBatchError(
            f"scenario={scenario_id} 未配置 business_doc_paths",
            code="DEMO_BUSINESS_DOCS_EMPTY",
        )

    business_docs: list[str] = []
    business_doc_sources: list[str] = []
    for rel in business_doc_paths:
        path = PROJECT_ROOT / rel
        if not path.is_file():
            raise DemoBatchError(
                f"业务制度库文件缺失: {rel}",
                code="DEMO_BUSINESS_DOC_MISSING",
            )
        suffix = path.suffix.lower()
        if suffix == ".docx":
            text = _read_docx(path)
        elif suffix in (".txt", ".md"):
            text = _read_text(path)
        else:
            raise DemoBatchError(
                f"不支持的业务制度库格式: {path.name} (仅 .docx/.txt/.md)",
                code="DEMO_BUSINESS_DOC_FORMAT",
            )
        if text.strip():
            business_docs.append(text)
            business_doc_sources.append(rel)

    if not business_docs:
        raise DemoBatchError(
            f"scenario={scenario_id} 业务制度库读出全空",
            code="DEMO_BUSINESS_DOCS_BLANK",
        )

    policy_meta = dict(body.get("policy_meta") or {})

    return DemoBatch(
        scenario_id=scenario_id,
        label=body.get("label", scenario_id),
        policy_doc=policy_doc,
        policy_meta=policy_meta,
        business_docs=business_docs,
        business_doc_sources=business_doc_sources,
    )
