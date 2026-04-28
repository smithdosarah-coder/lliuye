# -*- coding: utf-8 -*-
"""Agent5 Compliance · 修订意见书 Word 导出 (Stage C.4 · onboarding W-C3-A3 compli-backend-complete).

新建 · 与既有 ledger_exporter.py 并存 (那个是整改报告 · 本文件是修订意见书 · 用途不同):
- ledger_exporter.export_remediation_word: 整改计划表 (责任分派 + 期限) · 给运营
- 本模块 build_revision_docx: 修订意见书 · 改 / 补 / 强 三类 RevisionPanel 格式 · 给合规

Schema (修订意见书结构):
  封面 → 摘要 (违规统计 + scan 元信息) →
  改 类 RevisionPanel (modify) →
  补 类 RevisionPanel (supplement) →
  强 类 RevisionPanel (strengthen) →
  附录 (规则原文 + 事件原始记录 · 可追溯)

返 docx 字节 · 调用方写文件 / Response。
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from .scan_engine import REVISION_CATEGORIES


SEVERITY_LABEL = {
    "critical": "🔴 严重",
    "major": "🟡 一般",
    "minor": "🟢 观察",
    "none": "—",
}

CATEGORY_LABEL = {
    "改": "改 · 修改业务条款",
    "补": "补 · 补充缺失条款",
    "强": "强 · 强化既有要求",
}


def build_revision_docx(scan_payload: dict, *, title: str = "合规修订意见书") -> bytes:
    """生成修订意见书 docx · 返 bytes.

    Args:
        scan_payload: load_scan_result 返回的 payload (含 violations + revisions)
        title: 封面标题

    Returns:
        docx 字节流
    """
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Pt, RGBColor
    except ImportError as e:
        raise RuntimeError(f"python-docx 未安装: {e}") from e

    doc = Document()

    # 默认字体大小
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)

    # ---- 封面 ----
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f"扫描 ID: {scan_payload.get('scan_id', '-')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(
        f"生成时间: {scan_payload.get('generated_at', datetime.now().isoformat(timespec='seconds'))}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ---- 摘要 ----
    doc.add_heading("一、扫描摘要", level=1)
    stats = scan_payload.get("stats", {})
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = "Light Grid Accent 1"
    headers = summary_table.rows[0].cells
    headers[0].text = "规则数"
    headers[1].text = "事件数"
    headers[2].text = "矩阵单元"
    headers[3].text = "命中违规"
    values = summary_table.rows[1].cells
    values[0].text = str(stats.get("rule_count", "-"))
    values[1].text = str(stats.get("event_count", "-"))
    values[2].text = str(stats.get("cell_count", "-"))
    values[3].text = str(stats.get("violation_count", "-"))

    doc.add_paragraph()
    sev_table = doc.add_table(rows=2, cols=3)
    sev_table.style = "Light Grid Accent 1"
    sev_headers = sev_table.rows[0].cells
    sev_headers[0].text = SEVERITY_LABEL["critical"]
    sev_headers[1].text = SEVERITY_LABEL["major"]
    sev_headers[2].text = SEVERITY_LABEL["minor"]
    sev_values = sev_table.rows[1].cells
    sev_values[0].text = str(stats.get("severe_count", 0))
    sev_values[1].text = str(stats.get("normal_count", 0))
    sev_values[2].text = str(stats.get("observation_count", 0))

    # ---- 三类 RevisionPanel ----
    violations = scan_payload.get("violations", []) or []
    panels: dict[str, list[dict]] = {cat: [] for cat in REVISION_CATEGORIES}
    for v in violations:
        for rev in v.get("revisions", []) or []:
            cat = rev.get("category", "")
            if cat in panels:
                panels[cat].append({"violation": v, "revision": rev})

    section_counter = 2
    for cat in REVISION_CATEGORIES:
        items = panels.get(cat, [])
        doc.add_heading(f"{_chinese_section_no(section_counter)}、{CATEGORY_LABEL[cat]}", level=1)
        section_counter += 1
        if not items:
            doc.add_paragraph("（本类无修订建议）")
            continue
        for idx, entry in enumerate(items, 1):
            v = entry["violation"]
            rev = entry["revision"]
            doc.add_heading(f"{cat}-{idx} · {rev.get('title') or '修订建议'}", level=2)

            # 命中信息
            tbl = doc.add_table(rows=4, cols=2)
            tbl.style = "Light Grid Accent 1"
            tbl.rows[0].cells[0].text = "违规 ID"
            tbl.rows[0].cells[1].text = v.get("violation_id", "-")
            tbl.rows[1].cells[0].text = "规则"
            tbl.rows[1].cells[1].text = (
                f"{v.get('rule_id', '-')} {v.get('rule_article', '')} · "
                f"{v.get('rule_condition', '')[:120]}"
            )
            tbl.rows[2].cells[0].text = "事件"
            tbl.rows[2].cells[1].text = (
                f"{v.get('event_id', '-')} ({v.get('event_type', '-')}) · "
                f"{v.get('match_reason', '')[:120]}"
            )
            tbl.rows[3].cells[0].text = "严重度"
            tbl.rows[3].cells[1].text = SEVERITY_LABEL.get(v.get("severity", "none"), "-")

            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("修订建议：")
            run.bold = True
            doc.add_paragraph(rev.get("text") or "")

            doc.add_paragraph()  # 间距

    # ---- 附录 (rules + events 摘录 · 仅前 20 条防爆) ----
    doc.add_heading(f"{_chinese_section_no(section_counter)}、附录 · 数据可追溯", level=1)
    doc.add_heading("A. 政策规则集", level=2)
    rules = scan_payload.get("rules", []) or []
    if rules:
        rule_table = doc.add_table(rows=min(len(rules), 20) + 1, cols=4)
        rule_table.style = "Light Grid Accent 1"
        h = rule_table.rows[0].cells
        h[0].text = "rule_id"
        h[1].text = "条款"
        h[2].text = "类别"
        h[3].text = "条件"
        for i, r in enumerate(rules[:20], 1):
            row = rule_table.rows[i].cells
            row[0].text = r.get("rule_id", "-")
            row[1].text = r.get("article", "-")
            row[2].text = r.get("category", "-")
            row[3].text = (r.get("condition", "") or "")[:80]
    else:
        doc.add_paragraph("（规则集为空）")

    doc.add_paragraph()
    doc.add_heading("B. 业务事件集", level=2)
    events = scan_payload.get("events", []) or []
    if events:
        ev_table = doc.add_table(rows=min(len(events), 20) + 1, cols=3)
        ev_table.style = "Light Grid Accent 1"
        h = ev_table.rows[0].cells
        h[0].text = "event_id"
        h[1].text = "类型"
        h[2].text = "字段摘要"
        for i, e in enumerate(events[:20], 1):
            row = ev_table.rows[i].cells
            row[0].text = e.get("event_id", "-")
            row[1].text = e.get("event_type", "-")
            row[2].text = _fields_brief(e.get("fields") or {})
    else:
        doc.add_paragraph("（事件集为空）")

    # 保存到内存
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_revision_filename(scan_payload: dict) -> str:
    sid = scan_payload.get("scan_id", "scan")
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"compliance_revision_{sid}_{ts}.docx"


def _chinese_section_no(n: int) -> str:
    """1 → 一, 2 → 二, 3 → 三 ... 11 → 十一"""
    table = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if 1 <= n <= 10:
        return table[n]
    if 11 <= n <= 19:
        return f"十{table[n - 10]}"
    return str(n)


def _fields_brief(fields: dict, max_chars: int = 80) -> str:
    if not isinstance(fields, dict):
        return str(fields)[:max_chars]
    parts = []
    for k, v in fields.items():
        s = f"{k}={v}"
        parts.append(s)
        if sum(len(p) for p in parts) > max_chars:
            break
    return "; ".join(parts)[:max_chars]
