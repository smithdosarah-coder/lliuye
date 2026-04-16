# -*- coding: utf-8 -*-
"""LedgerExporter: ComplianceLedger → 违规榜单 Excel + 整改报告 Word。

导出物：
  1. Excel：5 sheet（严重 / 一般 / 观察 / 规则清单 / 事件清单）
  2. Word：封面 + 摘要 + 分级明细 + 整改计划表 + 责任分派 + 附录
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from .matrix_matcher import ComplianceLedger, ViolationRecord


SEVERITY_LABEL = {"critical": "🔴 严重", "major": "🟡 一般", "minor": "🟢 观察"}


# ======================================================================
# Excel 导出
# ======================================================================

def export_ledger_excel(ledger: ComplianceLedger,
                        rules: list,
                        events: list[dict],
                        output_dir: str | Path = "outputs",
                        scenario_title: str = "合规巡检") -> str:
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = out_dir / f"compliance_ledger_{ts}.xlsx"

    from openpyxl import Workbook
    from openpyxl.styles import PatternFill, Font, Alignment
    wb = Workbook()

    # 封面 sheet
    cover = wb.active
    cover.title = "概览"
    cover.append([f"{scenario_title} 违规榜单"])
    cover["A1"].font = Font(bold=True, size=16)
    cover.append([])
    cover.append(["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    cover.append(["规则总数", ledger.rule_count])
    cover.append(["事件总数", ledger.event_count])
    cover.append(["矩阵单元", ledger.cell_count])
    cover.append(["🔴 严重违规", len(ledger.severe)])
    cover.append(["🟡 一般缺陷", len(ledger.normal)])
    cover.append(["🟢 观察项", len(ledger.observation)])
    cover.append(["扫描用时(秒)", round(ledger.duration_seconds, 2)])
    cover.column_dimensions["A"].width = 18
    cover.column_dimensions["B"].width = 40

    _write_violations_sheet(wb, "严重违规",
                            ledger.severe, PatternFill("solid", start_color="FFE6E6"))
    _write_violations_sheet(wb, "一般缺陷",
                            ledger.normal, PatternFill("solid", start_color="FFF9E6"))
    _write_violations_sheet(wb, "观察项",
                            ledger.observation, PatternFill("solid", start_color="E6F9E6"))

    # 规则清单
    ws = wb.create_sheet("规则清单")
    ws.append(["规则ID", "来源", "条号", "类别", "标题", "触发条件", "严重度"])
    for r in rules:
        ws.append([r.rule_id, r.source_doc, r.source_page,
                   r.category, r.title, r.trigger_condition, r.severity.value])
    for i, w in enumerate([15, 25, 10, 14, 28, 36, 10]):
        ws.column_dimensions[chr(ord("A") + i)].width = w

    # 事件清单
    ws = wb.create_sheet("事件清单")
    ws.append(["事件ID", "类型", "日期", "来源文件", "关键字段"])
    for e in events:
        keyinfo = ", ".join(f"{k}={v}" for k, v in list(e.get("fields", {}).items())[:6])
        ws.append([e["event_id"], e["event_type"], e.get("event_date", ""),
                   e.get("source_file", ""), keyinfo])
    for i, w in enumerate([18, 14, 14, 24, 70]):
        ws.column_dimensions[chr(ord("A") + i)].width = w

    wb.save(str(path))
    return str(path)


def _write_violations_sheet(wb, title: str,
                            violations: list[ViolationRecord],
                            fill) -> None:
    ws = wb.create_sheet(title)
    ws.append(["编号", "严重度", "规则ID", "来源", "条号", "标题",
               "涉及业务单号", "事件数", "原因", "整改建议", "责任部门", "期限"])
    for v in violations:
        event_ids = "; ".join(e["event_id"] for e in v.events)
        ws.append([v.violation_id,
                   SEVERITY_LABEL.get(v.severity, v.severity),
                   v.rule.rule_id, v.rule.source_doc, v.rule.source_page,
                   v.rule.title, event_ids, len(v.events),
                   v.match_reason, v.recommendation,
                   " / ".join(v.responsible_depts), v.deadline])
        for c in ws[ws.max_row]:
            c.fill = fill
    widths = [10, 10, 16, 24, 10, 22, 32, 8, 40, 48, 20, 16]
    for i, w in enumerate(widths):
        ws.column_dimensions[chr(ord("A") + i)].width = w


# ======================================================================
# Word 整改报告
# ======================================================================

def export_remediation_word(ledger: ComplianceLedger,
                            output_dir: str | Path = "outputs",
                            scenario_title: str = "合规巡检",
                            kb_summary: str = "") -> str:
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = out_dir / f"compliance_remediation_{ts}.docx"

    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 封面
    t = doc.add_heading(f"{scenario_title}·整改报告", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(datetime.now().strftime("生成时间：%Y 年 %m 月 %d 日 %H:%M"))
    doc.add_paragraph()
    if kb_summary:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"知识库：{kb_summary}")

    doc.add_page_break()

    # 摘要
    doc.add_heading("一、总体摘要", level=1)
    doc.add_paragraph(
        f"本次合规专项巡检扫描规则 {ledger.rule_count} 条，业务事件 {ledger.event_count} 条，"
        f"比对单元 {ledger.cell_count}，累计用时 {ledger.duration_seconds:.1f} 秒。"
    )
    t = doc.add_table(rows=2, cols=4)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "🔴 严重"
    t.rows[0].cells[1].text = "🟡 一般"
    t.rows[0].cells[2].text = "🟢 观察"
    t.rows[0].cells[3].text = "扫描用时"
    t.rows[1].cells[0].text = str(len(ledger.severe))
    t.rows[1].cells[1].text = str(len(ledger.normal))
    t.rows[1].cells[2].text = str(len(ledger.observation))
    t.rows[1].cells[3].text = f"{ledger.duration_seconds:.1f} 秒"

    # 分级明细
    for level_title, bucket in (
        ("二、🔴 严重违规明细", ledger.severe),
        ("三、🟡 一般缺陷明细", ledger.normal),
        ("四、🟢 观察项明细", ledger.observation),
    ):
        doc.add_heading(level_title, level=1)
        if not bucket:
            doc.add_paragraph("（无）")
            continue
        for v in bucket:
            doc.add_heading(f"{v.violation_id}  {v.rule.title}", level=2)
            doc.add_paragraph(f"触发规则：{v.rule.rule_id}（{v.rule.source_doc} · {v.rule.source_page}）")
            doc.add_paragraph(f"条款原文：{v.rule.content}")
            eids = "；".join(e["event_id"] for e in v.events)
            doc.add_paragraph(f"涉及业务单号（{len(v.events)}条）：{eids}")
            doc.add_paragraph(f"命中原因：{v.match_reason}")
            # 证据链
            if v.evidence_chain:
                doc.add_paragraph("证据：")
                for ev in v.evidence_chain:
                    doc.add_paragraph(
                        f"· {ev.get('event_id')}：{ev.get('evidence')}",
                        style="List Bullet",
                    )
            doc.add_paragraph(f"责任部门：{' / '.join(v.responsible_depts)}")
            doc.add_paragraph(f"整改期限：{v.deadline}")
            doc.add_paragraph("整改建议：")
            for line in v.recommendation.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip(), style="List Bullet")

    # 整改计划表
    doc.add_heading("五、整改计划汇总表", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "编号"
    hdr[1].text = "严重度"
    hdr[2].text = "议题"
    hdr[3].text = "责任部门"
    hdr[4].text = "期限"
    for v in ledger.severe + ledger.normal + ledger.observation:
        row = t.add_row().cells
        row[0].text = v.violation_id
        row[1].text = SEVERITY_LABEL.get(v.severity, v.severity)
        row[2].text = v.rule.title
        row[3].text = " / ".join(v.responsible_depts)
        row[4].text = v.deadline

    doc.save(str(path))
    return str(path)
