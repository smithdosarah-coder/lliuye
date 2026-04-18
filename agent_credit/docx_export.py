"""agent_credit.docx_export — 决策意见书本地 docx 导出（python-docx）。

监管底线：按银行业私有化部署要求，禁止调用海外 API 渲染文档。全部渲染在本地
完成，产物落 `data/exports/agent3_credit/` 目录，由 FastAPI 作为文件下载返回。

字段契约：消费 DecisionAdvice dict；数值口径与前端一致（approved_amount 以万元为
单位，interest_rate 为小数 0-1）。
"""
from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SEVERITY_LABEL = {"red": "红灯", "yellow": "黄灯", "green": "绿灯"}
_DEFAULT_FONT = "Microsoft YaHei"


def _set_font(run, name: str = _DEFAULT_FONT, size: float | None = None,
              bold: bool = False, color: tuple[int, int, int] | None = None) -> None:
    run.font.name = name
    # 中英字体都绑定
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bool(bold)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _add_paragraph(doc: Document, text: str, *, size: float = 10.5,
                   bold: bool = False, align=None,
                   color: tuple[int, int, int] | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=sizes.get(level, 11), bold=True)


def _format_amount_wan(v: Any) -> str:
    try:
        f = float(v)
    except (TypeError, ValueError):
        return "—"
    if f <= 0:
        return "—"
    if f == int(f):
        return f"{int(f):,} 万元"
    return f"{f:,.1f} 万元"


def _format_rate(v: Any) -> str:
    try:
        f = float(v)
    except (TypeError, ValueError):
        return "—"
    if f <= 0:
        return "—"
    return f"{f * 100:.2f}%"


def render_decision_letter(advice: dict, output_path: str | Path | None = None) -> bytes:
    """渲染决策意见书。

    Args:
        advice: DecisionAdvice.to_dict() 结果
        output_path: 落盘路径；None 表示仅返回 bytes（不落盘）

    Returns:
        docx 文件字节流
    """
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    subject = advice.get("subject_name", "—")
    segment_cn = {"corporate": "对公", "retail": "对私/零售"}.get(
        advice.get("segment", ""), advice.get("segment", ""))

    # 标题
    _add_paragraph(
        doc,
        f"{subject} 授信决策意见书",
        size=18, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_paragraph(
        doc,
        f"决策编号：{advice.get('advice_id', '—')}    "
        f"业务线：{segment_cn}    "
        f"生成时间：{advice.get('decision_time', datetime.now().isoformat(timespec='seconds'))}",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=(120, 120, 120),
    )
    _add_paragraph(doc, "")

    # 一、核心结论
    _add_heading(doc, "一、核心结论", level=2)
    decision = advice.get("decision", "—")
    color = {
        "批准": (46, 112, 76),
        "有条件批准": (167, 120, 47),
        "拒绝": (175, 60, 45),
    }.get(decision, (0, 0, 0))
    _add_paragraph(doc, f"审批结论：{decision}", size=12, bold=True, color=color)
    _add_paragraph(
        doc,
        f"建议额度：{_format_amount_wan(advice.get('approved_amount'))}    "
        f"期限：{advice.get('approved_term_months') or '—'} 个月    "
        f"利率：{_format_rate(advice.get('interest_rate'))}（{advice.get('rate_benchmark', '—')}）",
        size=10.5,
    )
    _add_paragraph(
        doc,
        f"综合评分：{advice.get('composite_score', 0)}    风险等级：{advice.get('risk_grade', '—')}",
        size=10.5,
    )

    reason = advice.get("decision_reason") or ""
    if reason:
        _add_paragraph(doc, "决策说明：", size=10.5, bold=True)
        for line in str(reason).splitlines():
            if line.strip():
                _add_paragraph(doc, line, size=10.5)

    # 二、批准条件 / 附加要求
    conditions = advice.get("conditions") or []
    if conditions:
        _add_heading(doc, "二、批准条件 / 附加要求", level=2)
        for i, c in enumerate(conditions, 1):
            _add_paragraph(doc, f"{i}. {c}", size=10.5)

    # 三、Top-5 决策理由码
    reasons = advice.get("top_reason_codes") or []
    if reasons:
        _add_heading(doc, "三、Top-5 决策理由码", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["#", "严重度", "理由码", "名称", "证据/阈值"]):
            run = hdr[i].paragraphs[0].add_run(h)
            _set_font(run, size=10, bold=True)
        for i, r in enumerate(reasons, 1):
            row = table.add_row().cells
            ev_parts = []
            if r.get("evidence_path"):
                ev_parts.append(f"path={r['evidence_path']}")
            if r.get("evidence_value") not in (None, ""):
                ev_parts.append(f"实际={r['evidence_value']}")
            if r.get("threshold") not in (None, ""):
                ev_parts.append(f"阈值={r['threshold']}")
            values = [
                str(i),
                SEVERITY_LABEL.get(r.get("severity", ""), r.get("severity", "")),
                r.get("code", ""),
                r.get("title", ""),
                "  ".join(ev_parts),
            ]
            for j, v in enumerate(values):
                run = row[j].paragraphs[0].add_run(v)
                _set_font(run, size=9.5)

    # 四、红线命中明细
    red_lines = advice.get("red_line_hits") or []
    if red_lines:
        _add_heading(doc, "四、红线命中明细", level=2)
        for h in red_lines:
            sev = h.get("severity", "")
            _add_paragraph(
                doc,
                f"[{SEVERITY_LABEL.get(sev, sev)}] {h.get('rule_name', h.get('rule_id', '—'))} "
                f"（实际 {h.get('actual_value', '—')}，阈值 {h.get('threshold', '—')}）",
                size=10.5, bold=(sev == "red"),
            )
            if h.get("description"):
                _add_paragraph(doc, f"    {h.get('description')}", size=9.5,
                               color=(100, 100, 100))

    # 五、额度三估算
    amounts = advice.get("amount_methods") or {}
    if amounts:
        _add_heading(doc, "五、额度估算方法", level=2)
        label_map = {
            "revenue_method": "营收法",
            "netasset_method": "净资产法",
            "cashflow_method": "现金流法",
            "collateral_method": "抵押法",
            "income_method": "收入法",
            "dscr_method": "DSCR 法",
        }
        for k, v in amounts.items():
            _add_paragraph(
                doc,
                f"· {label_map.get(k, k)}：{_format_amount_wan(v)}",
                size=10.5,
            )

    # 六、同业案例摘要
    cases_summary = advice.get("similar_cases_summary")
    if cases_summary:
        _add_heading(doc, "六、同业参考案例", level=2)
        _add_paragraph(doc, str(cases_summary), size=10.5)

    # 免责条款
    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "——本意见书由 Agent3 授信决策辅助系统自动生成，仅作审批参考，"
        "最终决定以行内审批流程为准。所有数据与判断均可追溯证据链，"
        "本地渲染，无数据出境。",
        size=8.5,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        color=(120, 120, 120),
    )

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(data)
    return data


def build_filename(advice: dict) -> str:
    """产出文件名：{subject}_授信决策意见书_{advice_id}.docx"""
    subject = advice.get("subject_name") or "无名客户"
    # 去掉非法字符
    subject = "".join(c for c in subject if c not in r'\/:*?"<>|').strip() or "客户"
    advice_id = advice.get("advice_id") or datetime.now().strftime("%Y%m%d%H%M%S")
    return f"{subject}_授信决策意见书_{advice_id}.docx"
