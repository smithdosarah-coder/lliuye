# -*- coding: utf-8 -*-
"""agent_channel.kb_upload — Channel KB upload + parse + persist.

Stage B.6 backend · master plan §B.6 (file 上传 KB · 3 类: 客户名录 / 政策 / 行业指引).

Endpoint shape (per onboarding W-B-A2-channel-kb-upload.md):
  POST /api/channel/upload_kb
  multipart/form-data:
    - kb_type: "customer_list" | "policy" | "industry_guide"
    - file:    xlsx | pdf | docx (single file per request · ≤ 50MB)
  return JSON:
    { kb_id: "<uuid4>", kb_type, n_rows | n_pages, summary_text }

持久化: data/channel_kb/{kb_id}.json (struct: type/source_filename/parsed payload).

错误:
  400 — kb_type 非法 / file extension 不支持 / file 内容损坏
  413 — file > 50MB (FastAPI request body size 由 reverse proxy 控 · 这里仅做 content-len 兜底)

解析:
  xlsx → pandas.read_excel(BytesIO, sheet_name=None) → 多 sheet 合并 row dicts
  pdf  → PyMuPDF (fitz) page.get_text() · 不走 OCR (KB 通常 text-based · OCR 路径在 tools.py 备用)
  docx → python-docx · paragraphs + tables 拼文本 (复用 tools._read_single_file docx 模式)

设计:
- 解析逻辑全在 buffer 内 · 不写临时文件 · 利于单元测试
- 持久化为 JSON (xlsx 行展开为 list[dict] · pdf/docx 文本为 str)
- summary_text = LLM-grounded 抽取的"该 KB 内容摘要" (≤ 200 字 · 取 head/tail truncate · 真 LLM 摘要在 Stage C C2 IdealProfile 阶段做)
"""
from __future__ import annotations

import io
import json
import uuid
from pathlib import Path
from typing import Any, Literal

from fastapi import HTTPException, UploadFile

# ----------------------------------------------------------------------------
# Constants
# ----------------------------------------------------------------------------

KB_TYPES = ("customer_list", "policy", "industry_guide")
KbType = Literal["customer_list", "policy", "industry_guide"]

EXT_XLSX = {"xlsx", "xls"}
EXT_PDF = {"pdf"}
EXT_DOCX = {"docx"}
ALL_EXT = EXT_XLSX | EXT_PDF | EXT_DOCX

MAX_BYTES = 50 * 1024 * 1024  # 50MB hard cap (per onboarding)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
KB_DIR = PROJECT_ROOT / "data" / "channel_kb"

SUMMARY_HEAD_CHARS = 120
SUMMARY_TAIL_CHARS = 40


# ----------------------------------------------------------------------------
# Parse handlers (buffer-based · no temp file)
# ----------------------------------------------------------------------------


def _ext_of(filename: str) -> str:
    if not filename or "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def parse_xlsx(content: bytes) -> dict[str, Any]:
    """Parse xlsx/xls bytes → multi-sheet rows.

    Returns:
      {
        "sheets":   { sheet_name: [ {col: val, ...}, ... ] },
        "n_rows":   int (total across sheets),
        "n_sheets": int,
        "columns":  { sheet_name: [col1, col2, ...] }
      }
    """
    import pandas as pd

    try:
        sheets = pd.read_excel(io.BytesIO(content), sheet_name=None, engine="openpyxl")
    except Exception as e:
        raise HTTPException(400, f"xlsx 解析失败: {type(e).__name__}: {e}") from e

    out_sheets: dict[str, list[dict[str, Any]]] = {}
    out_columns: dict[str, list[str]] = {}
    n_rows_total = 0
    for name, df in sheets.items():
        df = df.fillna("")
        rows = df.to_dict(orient="records")
        out_sheets[name] = [{str(k): _coerce(v) for k, v in r.items()} for r in rows]
        out_columns[name] = [str(c) for c in df.columns]
        n_rows_total += len(rows)
    return {
        "sheets": out_sheets,
        "n_rows": n_rows_total,
        "n_sheets": len(sheets),
        "columns": out_columns,
    }


def parse_pdf(content: bytes) -> dict[str, Any]:
    """Parse pdf bytes → page-wise text.

    Returns:
      { "pages": ["text page 1", ...], "n_pages": int, "text": "joined" }
    """
    import fitz  # PyMuPDF

    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as e:
        raise HTTPException(400, f"pdf 解析失败: {type(e).__name__}: {e}") from e

    pages: list[str] = []
    try:
        for page in doc:
            pages.append(page.get_text("text") or "")
    finally:
        doc.close()
    return {
        "pages": pages,
        "n_pages": len(pages),
        "text": "\n\n".join(pages).strip(),
    }


def parse_docx(content: bytes) -> dict[str, Any]:
    """Parse docx bytes → paragraphs + tables → joined text.

    Returns:
      { "text": "...", "n_paragraphs": int, "n_tables": int }
    """
    from docx import Document

    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(400, f"docx 解析失败: {type(e).__name__}: {e}") from e

    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for t_idx, table in enumerate(doc.tables, start=1):
        rows_data: list[str] = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            deduped: list[str] = []
            for j, c in enumerate(cells):
                if j == 0 or c != cells[j - 1]:
                    deduped.append(c)
            rows_data.append("\t".join(deduped))
        if rows_data:
            parts.append(f"[表格{t_idx}]")
            parts.extend(rows_data)
    return {
        "text": "\n".join(parts),
        "n_paragraphs": len(doc.paragraphs),
        "n_tables": len(doc.tables),
    }


def _coerce(v: Any) -> Any:
    """Make pandas cell JSON-serialisable."""
    import pandas as pd

    if v is None:
        return ""
    if isinstance(v, (int, float, str, bool)):
        try:
            if isinstance(v, float) and (pd.isna(v) or v != v):  # noqa: PLR0124
                return ""
        except (TypeError, ValueError):
            pass
        return v
    return str(v)


# ----------------------------------------------------------------------------
# Summary (deterministic · no LLM)
# ----------------------------------------------------------------------------


def make_summary(kb_type: str, parsed: dict[str, Any], filename: str) -> str:
    """Deterministic 1-line summary · LLM-grade summary 留 Stage C IdealProfile 阶段."""
    if "n_rows" in parsed:
        cols_per_sheet = parsed.get("columns") or {}
        first_sheet_cols = next(iter(cols_per_sheet.values()), [])
        cols_preview = ", ".join(first_sheet_cols[:6])
        return (
            f"[{kb_type}] {filename} · {parsed['n_sheets']} sheet · "
            f"{parsed['n_rows']} rows · cols: {cols_preview}"
        )
    if "n_pages" in parsed:
        head = (parsed.get("text") or "")[:SUMMARY_HEAD_CHARS].replace("\n", " ").strip()
        return f"[{kb_type}] {filename} · {parsed['n_pages']} pages · 首段: {head}…"
    if "n_paragraphs" in parsed:
        head = (parsed.get("text") or "")[:SUMMARY_HEAD_CHARS].replace("\n", " ").strip()
        return (
            f"[{kb_type}] {filename} · {parsed['n_paragraphs']} para · "
            f"{parsed['n_tables']} table · 首段: {head}…"
        )
    return f"[{kb_type}] {filename} · 已解析"


# ----------------------------------------------------------------------------
# Persist
# ----------------------------------------------------------------------------


def persist(kb_id: str, kb_type: str, filename: str, parsed: dict[str, Any]) -> Path:
    KB_DIR.mkdir(parents=True, exist_ok=True)
    path = KB_DIR / f"{kb_id}.json"
    payload = {
        "kb_id": kb_id,
        "kb_type": kb_type,
        "source_filename": filename,
        "parsed": parsed,
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


# ----------------------------------------------------------------------------
# Public handler (FastAPI endpoint glue)
# ----------------------------------------------------------------------------


async def handle_upload(kb_type: str, file: UploadFile) -> dict[str, Any]:
    """Validate + parse + persist + return summary JSON.

    Raises HTTPException 400 (bad input) / 413 (too big).
    """
    if kb_type not in KB_TYPES:
        raise HTTPException(
            400,
            f"kb_type 必须 ∈ {KB_TYPES} · 收到 {kb_type!r}",
        )

    filename = file.filename or "unnamed"
    ext = _ext_of(filename)
    if ext not in ALL_EXT:
        raise HTTPException(
            400,
            f"file extension {ext!r} 不支持 · 仅接受 xlsx/xls/pdf/docx",
        )

    content = await file.read()
    if len(content) > MAX_BYTES:
        raise HTTPException(
            413,
            f"file 过大 · {len(content)} bytes > {MAX_BYTES} bytes (50MB)",
        )
    if not content:
        raise HTTPException(400, "file 内容为空")

    if ext in EXT_XLSX:
        parsed = parse_xlsx(content)
    elif ext in EXT_PDF:
        parsed = parse_pdf(content)
    elif ext in EXT_DOCX:
        parsed = parse_docx(content)
    else:  # 不可达 (上面已校验) · 留给 type-checker
        raise HTTPException(400, f"unreachable: ext={ext}")

    kb_id = str(uuid.uuid4())
    persist(kb_id, kb_type, filename, parsed)
    summary_text = make_summary(kb_type, parsed, filename)

    response: dict[str, Any] = {
        "kb_id": kb_id,
        "kb_type": kb_type,
        "source_filename": filename,
        "summary_text": summary_text,
    }
    if "n_rows" in parsed:
        response["n_rows"] = parsed["n_rows"]
        response["n_sheets"] = parsed["n_sheets"]
    if "n_pages" in parsed:
        response["n_pages"] = parsed["n_pages"]
    if "n_paragraphs" in parsed:
        response["n_paragraphs"] = parsed["n_paragraphs"]
        response["n_tables"] = parsed["n_tables"]
    return response
