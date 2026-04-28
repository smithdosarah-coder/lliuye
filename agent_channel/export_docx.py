"""agent_channel.export_docx — Agent1 全渠道获客 · 候选线索 Word 报告本地导出 (python-docx).

监管底线:与 ``agent_credit.decision_letter_docx`` 同 — 银行业私有化部署 ·
**禁止调用海外 API 渲染**。全部渲染本地 BytesIO 完成 · 由 FastAPI 路由作为
attachment 下载返回。

输入 payload 形态(兼容 mock + live):
    {
        "session_id": str,                # 文件名后缀 · 空时按 timestamp 兜底
        "ideal_profile": dict | None,     # IdealProfile 12 维(B.6b 输出·可空)
        "candidates": list[dict],         # B.5 输出 candidates · Top10
        "business_line": str,             # "corporate" | "small" | "retail"
        "client_manager": str,            # 客户经理姓名(默认"客户经理")
        "query": str,                     # 用户原 query · 副标题用
    }

字段契约:
  - candidate dict 字段消费 ``agent_channel.realtime_stream._build_final_output``
    的输出格式 (legacy camelCase + B.5 snake_case 双键并存) · 优先取 snake_case
    新键(score / geo / scale / similarity / radar_8axis / match_dimensions /
    product_recommendations / pitch_scripts) · 老 camelCase 作 fallback。
  - ``ideal_profile`` 字段按 channel-spec.md §C2 12 维 IdealProfile 形态消费 ·
    缺字段优雅降级显 "—" · 不抛异常。

入口:
  - ``export(payload, output_path=None) -> bytes``
  - ``build_filename(payload) -> str``

Author: Worker A1 (Stage B 第 2 批 · §B.7) · 2026-04-28
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_DEFAULT_FONT = "Microsoft YaHei"
NA = "—"

_BUSINESS_LABEL = {
    "corporate": "对公",
    "small": "对公小微 / 普惠",
    "retail": "对私 / 零售",
}


# ============================================================================
# 字体 / 段落 helpers (复用 agent_credit/decision_letter_docx 模式)
# ============================================================================

def _set_font(
    run,
    name: str = _DEFAULT_FONT,
    size: float | None = None,
    bold: bool = False,
    italic: bool = False,
    color: tuple[int, int, int] | None = None,
) -> None:
    """中英字体绑定 + 字号 / 粗体 / 斜体 / 颜色一次写入。"""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bool(bold)
    run.italic = bool(italic)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _add_paragraph(
    doc,
    text: str,
    *,
    size: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    align=None,
    color: tuple[int, int, int] | None = None,
) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic, color=color)


def _add_heading(doc, text: str, level: int = 1) -> None:
    sizes = {1: 17, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=sizes.get(level, 11), bold=True)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: float = 10) -> None:
    """覆盖 cell 默认空段 · 写文字 + 字体绑定。"""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    _set_font(run, size=size, bold=bold)


# ============================================================================
# 字段格式化
# ============================================================================

def _fmt_str(v: Any) -> str:
    """str/None/list/dict 统一转显示文本。"""
    if v is None:
        return NA
    if isinstance(v, list):
        if not v:
            return NA
        return " · ".join(_fmt_str(x) for x in v)
    if isinstance(v, dict):
        if not v:
            return NA
        return " · ".join(f"{k}={_fmt_str(val)}" for k, val in v.items())
    s = str(v).strip()
    return s if s else NA


def _fmt_pct(v: Any) -> str:
    """0~1 float → "%xx.x" · 0/None → "—"。"""
    try:
        f = float(v)
    except (TypeError, ValueError):
        return NA
    if f == 0:
        return NA
    return f"{f * 100:.1f}%"


def _candidate_field(c: dict, snake: str, *camel_alts: str, default=None) -> Any:
    """优先取 snake_case (B.5 新键) · 兜底 camelCase (legacy)。"""
    if snake in c and c[snake] not in (None, "", []):
        return c[snake]
    for alt in camel_alts:
        if alt in c and c[alt] not in (None, "", []):
            return c[alt]
    return default


# ============================================================================
# 主入口
# ============================================================================

def export(payload: dict, output_path: str | Path | None = None) -> bytes:
    """渲染 Agent1 候选线索 Word 报告。

    Args:
        payload: 输入 dict · 形态见模块 docstring
        output_path: 落盘路径(可空 · 仅返 bytes 时传 None)

    Returns:
        .docx 文件字节流
    """
    candidates = list(payload.get("candidates") or [])
    ideal_profile = payload.get("ideal_profile") or {}
    business_line = payload.get("business_line") or "corporate"
    client_manager = payload.get("client_manager") or "客户经理"
    query = payload.get("query") or ""
    session_id = payload.get("session_id") or ""

    biz_cn = _BUSINESS_LABEL.get(business_line, business_line)

    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---------- 标题 ----------
    _add_paragraph(
        doc,
        "Agent1 全渠道获客 · 候选企业线索报告",
        size=18, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    sub_meta = (
        f"客户经理：{client_manager}    "
        f"日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}    "
        f"业务线：{biz_cn}    "
        f"候选数：{len(candidates)}"
    )
    _add_paragraph(
        doc, sub_meta,
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=(110, 110, 110),
    )
    if query:
        _add_paragraph(
            doc, f"原始诉求：{query}",
            size=9.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color=(110, 110, 110), italic=True,
        )
    if session_id:
        _add_paragraph(
            doc, f"会话编号：{session_id}",
            size=8.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color=(160, 160, 160),
        )
    _add_paragraph(doc, "")

    # ---------- 一、IdealProfile (理想客户画像 12 维) ----------
    if ideal_profile:
        _render_ideal_profile(doc, ideal_profile)

    # ---------- 二、Top10 候选概览 ----------
    _render_overview_table(doc, candidates, header_index=2 if ideal_profile else 1)

    # ---------- 三、候选企业明细 ----------
    detail_idx = 3 if ideal_profile else 2
    _add_heading(doc, f"{_chinese_num(detail_idx)}、候选企业明细", level=2)
    if not candidates:
        _add_paragraph(doc, "（无候选企业 · 检查 query 与知识库覆盖度）", italic=True)
    for i, cand in enumerate(candidates, 1):
        _render_candidate_detail(doc, cand, idx=i)

    # ---------- 免责条款 ----------
    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "——本报告由 Agent1 全渠道获客系统 (look-alike + 信号驱动) 自动生成 · "
        "仅作客户经理 outreach 参考 · 候选线索需人工尽调验证后方可进入 Agent3 "
        "授信流程 · 所有数据本地渲染 · 无数据出境 · 私有化合规。",
        size=8.5,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        color=(120, 120, 120),
    )

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(data)
    return data


# ============================================================================
# Sub-renderers
# ============================================================================

# IdealProfile 12 维 key → 显示标签 (channel-spec.md §C2)
_IDEAL_PROFILE_FIELDS: list[tuple[str, str]] = [
    ("benchmark", "标杆客户"),
    ("target_industries", "目标行业"),
    ("target_regions", "目标区域"),
    ("scale_range", "规模区间"),
    ("revenue_range", "营收区间"),
    ("must_have_tags", "必备特征"),
    ("nice_to_have_tags", "加分特征"),
    ("exclude_tags", "排除特征"),
    ("policy_context", "政策依据"),
    ("qualifications", "资质要求"),
    ("growth_stage", "成长阶段"),
    ("key_signals", "关键信号"),
]


def _render_ideal_profile(doc, ip: dict) -> None:
    _add_heading(doc, "一、理想客户画像 (IdealProfile · 12 维)", level=2)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"

    # 适当列宽
    for col_idx, w_cm in [(0, 4.0), (1, 12.0)]:
        for cell in table.columns[col_idx].cells:
            cell.width = Cm(w_cm)

    rendered_any = False
    for key, label in _IDEAL_PROFILE_FIELDS:
        if key not in ip:
            continue
        val = _fmt_str(ip.get(key))
        if val == NA:
            continue
        row = table.add_row().cells
        _set_cell_text(row[0], label, bold=True, size=10)
        _set_cell_text(row[1], val, size=10)
        rendered_any = True

    if not rendered_any:
        _add_paragraph(doc, "（IdealProfile 字段全空 · 跳过）", italic=True, color=(150, 150, 150))

    # reasoning 段落 (LLM 抽取的解释 · channel-spec C2)
    reasoning = ip.get("reasoning") or ""
    if reasoning and reasoning.strip():
        _add_paragraph(doc, "")
        _add_paragraph(doc, "画像推导：", bold=True, size=10.5)
        for line in str(reasoning).splitlines():
            line = line.strip()
            if line:
                _add_paragraph(doc, line, size=10)
    _add_paragraph(doc, "")


def _render_overview_table(doc, candidates: list[dict], header_index: int) -> None:
    _add_heading(
        doc,
        f"{_chinese_num(header_index)}、Top{len(candidates)} 候选企业概览",
        level=2,
    )
    if not candidates:
        _add_paragraph(doc, "（候选清单为空）", italic=True)
        return

    headers = ["#", "企业名", "总分", "行业", "地域", "规模", "相似度", "信号数"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, size=10)

    for i, c in enumerate(candidates, 1):
        score = _candidate_field(c, "score", "signalScore", default=0)
        industry = _candidate_field(c, "industry", default=NA)
        geo = _candidate_field(c, "geo", "region", default=NA)
        scale = _candidate_field(c, "scale", default=NA)
        similarity = _candidate_field(c, "similarity", default=0)
        signal_count = _candidate_field(c, "signalCount", default=0)
        if not signal_count:
            sigs = c.get("signals") or []
            signal_count = len(sigs)

        row = table.add_row().cells
        values = [
            str(i),
            _fmt_str(c.get("name")),
            str(score),
            _fmt_str(industry),
            _fmt_str(geo),
            _fmt_str(scale),
            _fmt_pct(similarity),
            str(signal_count),
        ]
        for j, v in enumerate(values):
            _set_cell_text(row[j], v, size=9.5)

    _add_paragraph(doc, "")


def _render_candidate_detail(doc, cand: dict, idx: int) -> None:
    name = _fmt_str(cand.get("name"))
    score = _candidate_field(cand, "score", "signalScore", default=0)
    industry = _candidate_field(cand, "industry", default=NA)
    geo = _candidate_field(cand, "geo", "region", default=NA)
    scale = _candidate_field(cand, "scale", default=NA)
    similarity = _candidate_field(cand, "similarity", default=0)

    _add_heading(doc, f"{idx}. {name}（总分 {score}）", level=3)

    # 基础卡片
    base_line = (
        f"行业：{_fmt_str(industry)}    "
        f"地域：{_fmt_str(geo)}    "
        f"规模：{_fmt_str(scale)}    "
        f"相似度：{_fmt_pct(similarity)}"
    )
    _add_paragraph(doc, base_line, size=10.5)

    # 工商基础信息 (legacy qcc 字段 · 选择性显)
    qcc_pairs = [
        ("统一社会信用代码", _fmt_str(cand.get("uscc"))),
        ("注册资本", _fmt_str(cand.get("registeredCapital"))),
        ("法定代表人", _fmt_str(cand.get("legalRep"))),
        ("成立日期", _fmt_str(cand.get("founded"))),
        ("员工数", str(cand.get("employees") or 0) if cand.get("employees") else NA),
    ]
    qcc_filtered = [(k, v) for k, v in qcc_pairs if v != NA]
    if qcc_filtered:
        _add_paragraph(
            doc,
            "工商信息：" + " · ".join(f"{k} {v}" for k, v in qcc_filtered),
            size=9.5, color=(100, 100, 100),
        )

    # ---------- 8 维 radar 表 ----------
    radar = _candidate_field(cand, "radar_8axis", default=None) or {}
    if radar:
        _add_paragraph(doc, "")
        _add_paragraph(doc, "8 维评分：", bold=True, size=10.5)
        radar_items = list(radar.items())
        # 4 col 表 (axis | score · 两组并排显得紧凑)
        rows_count = (len(radar_items) + 1) // 2
        radar_table = doc.add_table(rows=rows_count, cols=4)
        radar_table.style = "Light Grid Accent 1"
        for r in range(rows_count):
            left = radar_items[r * 2] if r * 2 < len(radar_items) else ("", "")
            right = radar_items[r * 2 + 1] if r * 2 + 1 < len(radar_items) else ("", "")
            cells = radar_table.rows[r].cells
            _set_cell_text(cells[0], left[0], bold=True, size=9.5)
            _set_cell_text(cells[1], str(left[1]) if left[0] else "", size=9.5)
            _set_cell_text(cells[2], right[0], bold=True, size=9.5)
            _set_cell_text(cells[3], str(right[1]) if right[0] else "", size=9.5)

    # ---------- 信号 timeline ----------
    signals = cand.get("signals") or []
    if signals:
        _add_paragraph(doc, "")
        _add_paragraph(doc, "信号事件：", bold=True, size=10.5)
        for s in signals[:6]:
            stype = _fmt_str(s.get("type"))
            title = _fmt_str(s.get("title"))
            date = _fmt_str(s.get("date"))
            source = _fmt_str(s.get("source"))
            _add_paragraph(
                doc,
                f"  · [{stype}] {title}    （{date} · {source}）",
                size=9.5,
            )
            detail = (s.get("detail") or "").strip()
            if detail:
                _add_paragraph(
                    doc, f"      {detail[:120]}",
                    size=9, color=(120, 120, 120),
                )

    # ---------- 匹配维度 (PRD v2 "为什么像") ----------
    match_dims = _candidate_field(cand, "match_dimensions", default=None) or []
    if match_dims:
        _add_paragraph(doc, "")
        _add_paragraph(doc, "匹配维度（为什么像）：", bold=True, size=10.5)
        for md in match_dims:
            dim = _fmt_str(md.get("dim_name"))
            ev = _fmt_str(md.get("hit_evidence"))
            sc = md.get("score", 0)
            _add_paragraph(
                doc,
                f"  · {dim}（分 {sc}）：{ev}",
                size=9.5,
            )
    else:
        # legacy matchTags 兜底
        match_tags = cand.get("matchTags") or []
        if match_tags:
            _add_paragraph(doc, "")
            _add_paragraph(doc, "匹配标签：", bold=True, size=10.5)
            for t in match_tags:
                lbl = _fmt_str(t.get("label"))
                detail = _fmt_str(t.get("detail"))
                _add_paragraph(doc, f"  · {lbl}：{detail}", size=9.5)

    # ---------- Top3 产品推荐 (structured) ----------
    products = _candidate_field(cand, "product_recommendations", default=None) or []
    if products:
        _add_paragraph(doc, "")
        _add_paragraph(doc, "推荐产品（Top3）：", bold=True, size=10.5)
        for p in products[:3]:
            pname = _fmt_str(p.get("product_name"))
            fit = p.get("fit_score", 0)
            cat = _fmt_str(p.get("category"))
            intro = _fmt_str(p.get("intro"))
            _add_paragraph(
                doc,
                f"  · {pname}（适配 {fit}・{cat}）",
                size=10, bold=True,
            )
            _add_paragraph(
                doc, f"    {intro}",
                size=9.5, color=(80, 80, 80),
            )
    else:
        # legacy recommendedProducts (str list) 兜底
        legacy_products = cand.get("recommendedProducts") or []
        if legacy_products:
            _add_paragraph(doc, "")
            _add_paragraph(doc, "推荐产品：", bold=True, size=10.5)
            for pname in legacy_products[:3]:
                _add_paragraph(doc, f"  · {pname}", size=10)

    # ---------- 切入话术 ----------
    pitch_scripts = _candidate_field(cand, "pitch_scripts", default=None) or []
    if pitch_scripts:
        _add_paragraph(doc, "")
        _add_paragraph(doc, "切入话术：", bold=True, size=10.5)
        for ps in pitch_scripts:
            text = _fmt_str(ps.get("script_text"))
            placeholder = ps.get("customer_name_placeholder") or ""
            note = f"（占位符 {placeholder} 实际使用替换为「{name}」）" if placeholder else ""
            _add_paragraph(doc, f"  {text}", size=10, italic=True)
            if note:
                _add_paragraph(
                    doc, f"  {note}",
                    size=8.5, color=(150, 150, 150),
                )
    elif cand.get("pitch"):
        # legacy pitch (str) 兜底
        _add_paragraph(doc, "")
        _add_paragraph(doc, "切入话术：", bold=True, size=10.5)
        _add_paragraph(doc, f"  {_fmt_str(cand.get('pitch'))}", size=10, italic=True)

    # 来源
    sources = cand.get("dataSources") or []
    if sources:
        _add_paragraph(doc, "")
        labels = " · ".join(_fmt_str(s.get("label")) for s in sources if s.get("label"))
        if labels:
            _add_paragraph(doc, f"数据来源：{labels}", size=9, color=(120, 120, 120))

    _add_paragraph(doc, "")  # 候选间空行隔


# ============================================================================
# 辅助
# ============================================================================

_CN_NUMS = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]


def _chinese_num(n: int) -> str:
    """1 → 一 · 2 → 二 · ... · 11+ → 数字."""
    if 0 <= n <= 10:
        return _CN_NUMS[n]
    return str(n)


def build_filename(payload: dict) -> str:
    """生成下载文件名: ``agent1_候选线索_<benchmark>_<session_id>.docx``."""
    ip = payload.get("ideal_profile") or {}
    benchmark = (
        ip.get("benchmark")
        or payload.get("query", "")
        or "look-alike"
    )
    benchmark = re.sub(r'[\\/:*?"<>|\s]+', "_", str(benchmark)).strip("_")[:30]
    if not benchmark:
        benchmark = "客户"
    sid = (payload.get("session_id") or "").strip()
    if not sid:
        sid = datetime.now().strftime("%Y%m%d%H%M%S")
    sid = re.sub(r'[\\/:*?"<>|\s]+', "_", sid)[:32]
    return f"agent1_候选线索_{benchmark}_{sid}.docx"


__all__ = ["export", "build_filename"]
