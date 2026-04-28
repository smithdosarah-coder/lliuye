# -*- coding: utf-8 -*-
"""Tests for POST /api/channel/upload_kb (Stage B.6 backend KB upload).

Coverage:
  - xlsx case · 客户名录解析 · n_rows / kb_id / 落盘 verify
  - pdf  case · 政策文件解析 · n_pages / 文本提取
  - docx case · 行业指引解析 · n_paragraphs + n_tables
  - 400 — kb_type 非法 / extension 不支持 / 空文件
  - 413 — file > 50MB

不需 LLM key · 不依赖 .env · pure CPU 解析 + FastAPI TestClient.
"""
from __future__ import annotations

import io
import json
from pathlib import Path

import pytest
from fastapi.testclient import TestClient


@pytest.fixture(scope="module")
def client():
    """Channel app · 与 portal 装载方式无关 · 直接挂 sub-app 测."""
    from agent_channel.api import app
    return TestClient(app)


@pytest.fixture(scope="module")
def kb_dir() -> Path:
    from agent_channel.kb_upload import KB_DIR
    KB_DIR.mkdir(parents=True, exist_ok=True)
    return KB_DIR


# ----------------------------------------------------------------------------
# Sample bytes builders (in-memory · no fixture file)
# ----------------------------------------------------------------------------


def _make_xlsx_bytes() -> bytes:
    """Build a 3-row customer list xlsx (in-memory · no fs)."""
    import pandas as pd
    df = pd.DataFrame([
        {"company_name": "杭州精工智造", "industry": "机械制造", "revenue_yuan": 50_000_000},
        {"company_name": "宁波东方智能", "industry": "电子",     "revenue_yuan": 30_000_000},
        {"company_name": "苏州优材科技", "industry": "新材料",   "revenue_yuan": 80_000_000},
    ])
    buf = io.BytesIO()
    df.to_excel(buf, index=False, engine="openpyxl")
    buf.seek(0)
    return buf.read()


def _make_pdf_bytes() -> bytes:
    """Build a 2-page PDF text doc (PyMuPDF native build · no template).

    Use ASCII text · PyMuPDF default Helvetica font 不 embed CJK · 真政策文件
    走真 PDF + OCR fallback · 这里只验 endpoint pipeline (parse/persist/return).
    """
    import fitz
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((50, 80), "Manufacturing SME Loan Guide v2.0")
    page1.insert_text((50, 110), "Annual revenue under 50M CNY")
    page2 = doc.new_page()
    page2.insert_text((50, 80), "Application Conditions and Amounts")
    page2.insert_text((50, 110), "Credit line 1M - 5M CNY")
    out = io.BytesIO()
    doc.save(out)
    doc.close()
    out.seek(0)
    return out.read()


def _make_docx_bytes() -> bytes:
    """Build a 3-paragraph + 1-table docx."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("行业指引 · 浙江省制造业转型升级 2025")
    doc.add_paragraph("重点扶持方向: 智能制造 / 新材料 / 工业软件")
    doc.add_paragraph("授信偏好: 营收 5000 万 -3 亿 · B 轮后 · 华东地区")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "维度"
    table.rows[0].cells[1].text = "权重"
    table.rows[0].cells[2].text = "说明"
    table.rows[1].cells[0].text = "营收"
    table.rows[1].cells[1].text = "0.30"
    table.rows[1].cells[2].text = "1-3 亿优先"
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


# ----------------------------------------------------------------------------
# Happy-path: 3 file format
# ----------------------------------------------------------------------------


def test_upload_xlsx_customer_list(client, kb_dir):
    """xlsx · 3 行客户名录 · 落 data/channel_kb/{kb_id}.json."""
    content = _make_xlsx_bytes()
    resp = client.post(
        "/api/channel/upload_kb",
        data={"kb_type": "customer_list"},
        files={"file": ("customers.xlsx", content,
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["kb_type"] == "customer_list"
    assert data["source_filename"] == "customers.xlsx"
    assert data["n_rows"] == 3
    assert data["n_sheets"] == 1
    assert "kb_id" in data and len(data["kb_id"]) == 36  # uuid4
    assert "customer_list" in data["summary_text"]

    # Persist verify · file 落地 + 内容 reload-able
    p = kb_dir / f"{data['kb_id']}.json"
    assert p.exists()
    payload = json.loads(p.read_text(encoding="utf-8"))
    assert payload["kb_type"] == "customer_list"
    assert payload["parsed"]["n_rows"] == 3
    sheet0 = next(iter(payload["parsed"]["sheets"].values()))
    assert sheet0[0]["company_name"] == "杭州精工智造"


def test_upload_pdf_policy(client, kb_dir):
    """pdf · 2 页政策文件."""
    content = _make_pdf_bytes()
    resp = client.post(
        "/api/channel/upload_kb",
        data={"kb_type": "policy"},
        files={"file": ("policy.pdf", content, "application/pdf")},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["kb_type"] == "policy"
    assert data["n_pages"] == 2
    assert "kb_id" in data
    p = kb_dir / f"{data['kb_id']}.json"
    payload = json.loads(p.read_text(encoding="utf-8"))
    text = payload["parsed"]["text"]
    assert "Manufacturing SME Loan Guide" in text
    assert "Credit line" in text


def test_upload_docx_industry_guide(client, kb_dir):
    """docx · 3 段 + 1 表的行业指引."""
    content = _make_docx_bytes()
    resp = client.post(
        "/api/channel/upload_kb",
        data={"kb_type": "industry_guide"},
        files={"file": ("industry.docx", content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["kb_type"] == "industry_guide"
    assert data["n_paragraphs"] == 3
    assert data["n_tables"] == 1
    p = kb_dir / f"{data['kb_id']}.json"
    payload = json.loads(p.read_text(encoding="utf-8"))
    txt = payload["parsed"]["text"]
    assert "智能制造" in txt
    assert "[表格1]" in txt


# ----------------------------------------------------------------------------
# Error path
# ----------------------------------------------------------------------------


def test_invalid_kb_type(client):
    """kb_type 不在白名单 → 400."""
    resp = client.post(
        "/api/channel/upload_kb",
        data={"kb_type": "garbage"},
        files={"file": ("a.xlsx", _make_xlsx_bytes(),
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert resp.status_code == 400
    assert "kb_type" in resp.json()["detail"]


def test_unsupported_extension(client):
    """ext .txt → 400."""
    resp = client.post(
        "/api/channel/upload_kb",
        data={"kb_type": "policy"},
        files={"file": ("a.txt", b"hello", "text/plain")},
    )
    assert resp.status_code == 400
    assert "extension" in resp.json()["detail"]


def test_empty_file(client):
    """空 body → 400."""
    resp = client.post(
        "/api/channel/upload_kb",
        data={"kb_type": "policy"},
        files={"file": ("a.pdf", b"", "application/pdf")},
    )
    assert resp.status_code == 400


def test_oversize_file(client, monkeypatch):
    """>50MB 触发 413 · 用 monkeypatch 把上限调到极小 · 1KB 即可触发."""
    import agent_channel.kb_upload as kb
    monkeypatch.setattr(kb, "MAX_BYTES", 1024)
    payload = b"x" * 2048  # 2KB > 1KB cap
    resp = client.post(
        "/api/channel/upload_kb",
        data={"kb_type": "policy"},
        files={"file": ("big.pdf", payload, "application/pdf")},
    )
    assert resp.status_code == 413
