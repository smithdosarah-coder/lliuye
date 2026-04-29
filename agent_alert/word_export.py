"""agent_alert.word_export — Agent4 贷中预警 · 命中清单 Word 报告本地导出 (python-docx).

监管底线 (复用 agent_channel.export_docx / agent_credit.decision_letter_docx 模式):
  - 银行业私有化部署 · **禁止调用海外 API 渲染** · 全部本地 BytesIO 完成
  - 由 FastAPI ``POST /api/alert/export_docx`` 作为 attachment 下载返回

输入 payload 形态 (兼容 frontend AlertWorkspace 命中清单 · API 序列化):
    {
        "session_id": str,                # 文件名后缀 · 空时按 timestamp 兜底
        "summary": str,                   # 扫描总览短句 (e.g. "扫描 100 家 · 红 12 / 黄 38 / 绿 50")
        "cases": list[dict],              # 命中清单 (TopCase-like / HitItem 序列化)
        "scan_range": str,                # 扫描范围标签 (e.g. "在贷 · 全量")
        "client_manager": str,            # 客户经理姓名 (默认 "客户经理")
        "stage": str,                     # 扫描阶段 (如 "已完成") · 元信息
        "totals": dict,                   # {"red": int, "yellow": int, "green": int} · 元信息
    }

字段契约 (case dict):
  - 优先 snake_case (后端 HitItem 序列化键) · 兜底 camelCase (frontend mock TopCase 键)
  - 必备显示字段:
        customer / customer_name / company_name → 客户名
        risk_level / level / tier             → red | yellow | green | RED | YELLOW | GREEN
        triggers / matched_rules / reasons    → 命中规则简述 list
        amount                                → 余额 / 敞口 (str)
        advice / disposition / summary        → 建议处置短句
        last_update / updated / lastUpdate    → 时间戳
  - 缺字段优雅降级 "—" · 不抛异常

入口:
  - ``export(payload, output_path=None) -> bytes``
  - ``build_filename(payload) -> str``

Author: Worker A1 (Stage W-FIX2 · 修 bug #6) · 2026-04-29
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_DEFAULT_FONT = "Microsoft YaHei"
NA = "—"

_TIER_LABEL = {
    "red": "🔴 红色",
    "yellow": "🟡 黄色",
    "green": "🟢 绿色",
}

_TIER_COLOR = {
    "red": (192, 0, 0),
    "yellow": (191, 144, 0),
    "green": (0, 110, 70),
}


# ============================================================================
# 字体 / 段落 helpers (复用 agent_channel/export_docx 模式)
# ============================================================================

def _set_font(
    run,
    name: str = _DEFAULT_FONT,
    size: float | None = None,
    bold: bool = False,
    italic: bool = False,
    color: tuple[int, int, int] | None = None,
) -> None:
    """中英字体绑定 + 字号 / 粗体 / 斜体 / 颜色一次写入。"""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bool(bold)
    run.italic = bool(italic)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _add_paragraph(
    doc,
    text: str,
    *,
    size: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    align=None,
    color: tuple[int, int, int] | None = None,
) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic, color=color)


def _add_heading(doc, text: str, level: int = 1) -> None:
    sizes = {1: 17, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=sizes.get(level, 11), bold=True)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10,
    color: tuple[int, int, int] | None = None,
) -> None:
    """覆盖 cell 默认空段 · 写文字 + 字体绑定。"""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    _set_font(run, size=size, bold=bold, color=color)


# ============================================================================
# 字段格式化
# ============================================================================

def _fmt_str(v: Any) -> str:
    """str/None/list/dict 统一转显示文本。"""
    if v is None:
        return NA
    if isinstance(v, list):
        if not v:
            return NA
        return " · ".join(_fmt_str(x) for x in v)
    if isinstance(v, dict):
        if not v:
            return NA
        return " · ".join(f"{k}={_fmt_str(val)}" for k, val in v.items())
    s = str(v).strip()
    return s if s else NA


def _case_field(c: dict, *keys: str, default: Any = None) -> Any:
    """优先取列表前置 key (snake_case) · 兜底 camelCase / legacy 字段。"""
    for k in keys:
        if k in c and c[k] not in (None, "", []):
            return c[k]
    return default


def _normalize_tier(v: Any) -> str:
    """red / RED / RiskLevel.RED 统一为 lowercase {red, yellow, green, unknown}."""
    if v is None:
        return "unknown"
    s = str(v).strip().lower()
    # 兼容 "RiskLevel.RED" / "tier:red" / "🔴"
    if "red" in s or "🔴" in s:
        return "red"
    if "yellow" in s or "🟡" in s:
        return "yellow"
    if "green" in s or "🟢" in s:
        return "green"
    return "unknown"


# ============================================================================
# 主入口
# ============================================================================

def export(payload: dict, output_path: str | Path | None = None) -> bytes:
    """渲染 Agent4 贷中预警命中清单 Word 报告。

    Args:
        payload: 输入 dict · 形态见模块 docstring
        output_path: 落盘路径 (可空 · 仅返 bytes 时传 None)

    Returns:
        .docx 文件字节流
    """
    cases = list(payload.get("cases") or [])
    summary = (payload.get("summary") or "").strip()
    session_id = payload.get("session_id") or ""
    scan_range = payload.get("scan_range") or "在贷客户 · 全量"
    client_manager = payload.get("client_manager") or "客户经理"
    stage = payload.get("stage") or ""
    totals = payload.get("totals") or {}

    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---------- 标题 ----------
    _add_paragraph(
        doc,
        "Agent4 贷中预警 · 命中清单报告",
        size=18, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    red = totals.get("red")
    yellow = totals.get("yellow")
    green = totals.get("green")
    totals_line = ""
    if red is not None or yellow is not None or green is not None:
        totals_line = (
            f"🔴 {red or 0}    🟡 {yellow or 0}    🟢 {green or 0}    "
            f"命中数 {len(cases)}"
        )

    sub_meta = (
        f"客户经理：{client_manager}    "
        f"日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}    "
        f"扫描范围：{scan_range}"
    )
    _add_paragraph(
        doc, sub_meta,
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=(110, 110, 110),
    )
    if totals_line:
        _add_paragraph(
            doc, totals_line,
            size=10.5, bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    if stage:
        _add_paragraph(
            doc, f"扫描阶段：{stage}",
            size=9.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color=(110, 110, 110), italic=True,
        )
    if session_id:
        _add_paragraph(
            doc, f"会话编号：{session_id}",
            size=8.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color=(160, 160, 160),
        )
    _add_paragraph(doc, "")

    # ---------- 一、扫描总览 ----------
    _add_heading(doc, "一、扫描总览", level=2)
    if summary:
        _add_paragraph(doc, summary, size=10.5)
    else:
        _add_paragraph(doc, "（未提供扫描总览短句）", italic=True, color=(150, 150, 150))
    _add_paragraph(doc, "")

    # ---------- 二、命中清单概览 ----------
    _render_overview_table(doc, cases)

    # ---------- 三、命中明细 ----------
    _add_heading(doc, "三、命中明细", level=2)
    if not cases:
        _add_paragraph(doc, "（无命中客户 · 检查规则覆盖度与扫描范围）", italic=True)
    for i, case in enumerate(cases, 1):
        _render_case_detail(doc, case, idx=i)

    # ---------- 免责条款 ----------
    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "——本报告由 Agent4 贷中预警系统 (知识库驱动 · 双路交叉) 自动生成 · "
        "命中客户需人工核实后方可触发处置流程 · 所有数据本地渲染 · "
        "无数据出境 · 私有化合规。",
        size=8.5,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        color=(120, 120, 120),
    )

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    if output_path is not None:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(data)
    return data


# ============================================================================
# Sub-renderers
# ============================================================================

def _render_overview_table(doc, cases: list[dict]) -> None:
    _add_heading(doc, f"二、命中客户概览 ({len(cases)} 家)", level=2)
    if not cases:
        _add_paragraph(doc, "（命中清单为空）", italic=True)
        return

    headers = ["#", "客户", "档位", "余额 / 敞口", "命中规则", "更新时间"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, size=10)

    for i, c in enumerate(cases, 1):
        customer = _case_field(
            c, "customer", "customer_name", "company_name", "name",
            default=NA,
        )
        tier = _normalize_tier(_case_field(c, "risk_level", "level", "tier"))
        triggers = _case_field(
            c, "triggers", "matched_rules", "reasons", "rules",
            default=[],
        )
        if isinstance(triggers, str):
            triggers = [triggers]
        triggers_text = "、".join(_fmt_str(t) for t in (triggers or []) if t)
        if not triggers_text:
            triggers_text = NA
        amount = _case_field(c, "amount", "balance", "exposure", default=NA)
        last_update = _case_field(
            c, "last_update", "updated", "lastUpdate", "ts",
            default=NA,
        )

        row = table.add_row().cells
        values = [
            str(i),
            _fmt_str(customer),
            _TIER_LABEL.get(tier, NA),
            _fmt_str(amount),
            triggers_text,
            _fmt_str(last_update),
        ]
        for j, v in enumerate(values):
            _set_cell_text(row[j], v, size=9.5)

    _add_paragraph(doc, "")


def _render_case_detail(doc, case: dict, idx: int) -> None:
    customer = _case_field(
        case, "customer", "customer_name", "company_name", "name",
        default=NA,
    )
    tier = _normalize_tier(_case_field(case, "risk_level", "level", "tier"))
    tier_label = _TIER_LABEL.get(tier, NA)
    color = _TIER_COLOR.get(tier)

    _add_heading(doc, f"{idx}. {_fmt_str(customer)}（{tier_label}）", level=3)

    # 基础卡片
    amount = _case_field(case, "amount", "balance", "exposure", default=NA)
    last_update = _case_field(
        case, "last_update", "updated", "lastUpdate", "ts",
        default=NA,
    )
    base_line = (
        f"档位：{tier_label}    "
        f"余额 / 敞口：{_fmt_str(amount)}    "
        f"更新：{_fmt_str(last_update)}"
    )
    _add_paragraph(doc, base_line, size=10.5, color=color)

    # ---------- 命中规则 ----------
    triggers = _case_field(
        case, "triggers", "matched_rules", "reasons", "rules",
        default=[],
    )
    if isinstance(triggers, str):
        triggers = [triggers]
    if triggers:
        _add_paragraph(doc, "")
        _add_paragraph(doc, "命中规则：", bold=True, size=10.5)
        for t in triggers:
            _add_paragraph(doc, f"  · {_fmt_str(t)}", size=10)

    # ---------- 处置建议 ----------
    advice = _case_field(
        case, "advice", "disposition", "summary",
        default="",
    )
    if advice:
        _add_paragraph(doc, "")
        _add_paragraph(doc, "处置建议：", bold=True, size=10.5)
        for line in str(advice).splitlines():
            line = line.strip()
            if line:
                _add_paragraph(doc, f"  {line}", size=10)

    # ---------- 信号 / 原因明细 ----------
    signals = case.get("signals") or []
    if signals and isinstance(signals, list):
        _add_paragraph(doc, "")
        _add_paragraph(doc, "信号事件：", bold=True, size=10.5)
        for s in signals[:6]:
            if isinstance(s, dict):
                stype = _fmt_str(s.get("type"))
                title = _fmt_str(s.get("title") or s.get("desc"))
                date = _fmt_str(s.get("date") or s.get("ts"))
                _add_paragraph(
                    doc, f"  · [{stype}] {title}    （{date}）",
                    size=9.5,
                )
            else:
                _add_paragraph(doc, f"  · {_fmt_str(s)}", size=9.5)

    _add_paragraph(doc, "")  # 案件间空行隔


# ============================================================================
# Filename + 兼容入口
# ============================================================================

def build_filename(payload: dict) -> str:
    """生成下载文件名: ``agent4_命中清单_<session_id>.docx``."""
    sid = (payload.get("session_id") or "").strip()
    if not sid:
        sid = datetime.now().strftime("%Y%m%d%H%M%S")
    sid = re.sub(r'[\\/:*?"<>|\s]+', "_", sid)[:40]
    return f"agent4_命中清单_{sid}.docx"


def export_hitlist_docx(
    session_id: str = "",
    summary: str = "",
    cases: list[dict] | None = None,
    *,
    scan_range: str = "",
    client_manager: str = "",
    stage: str = "",
    totals: dict | None = None,
    output_dir: str | Path | None = None,
) -> str:
    """便捷入口 (per onboarding spec) · 渲染 + 落盘 · 返绝对路径.

    Args:
        session_id: 会话 ID · 出现在文件名 + 内容
        summary: 扫描总览短句
        cases: 命中清单 list (TopCase-like dict)
        scan_range: 扫描范围标签
        client_manager: 客户经理姓名
        stage: 阶段标签
        totals: {"red", "yellow", "green"}
        output_dir: 落盘目录 (默认 <project_root>/data/exports/agent4)

    Returns:
        生成的 .docx 绝对路径 (调用方用 FileResponse 回传)
    """
    payload = {
        "session_id": session_id or "",
        "summary": summary or "",
        "cases": list(cases or []),
        "scan_range": scan_range or "",
        "client_manager": client_manager or "",
        "stage": stage or "",
        "totals": dict(totals or {}),
    }
    filename = build_filename(payload)
    if output_dir is None:
        project_root = Path(__file__).resolve().parent.parent
        output_dir = project_root / "data" / "exports" / "agent4"
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / filename
    export(payload, out_path)
    return str(out_path.resolve())


__all__ = [
    "build_filename",
    "export",
    "export_hitlist_docx",
    "NA",
]
