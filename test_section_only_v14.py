# -*- coding: utf-8 -*-
"""
Short-circuit test: 只跑 section generation,跳过 nested-table 填充与 phase C。
用于快速验证 V14 改动:composite 分支 + 三段式 + 问卷 + 措辞硬规则。
输出 outputs/section_gen_test_v14.docx + dump 到 outputs/_v14_dump.txt。
"""
import sys, io, os, time
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', line_buffering=True)

from openai import OpenAI
from docx import Document
from template_decomposer import decompose_template
from section_generator import (
    generate_all_sections, apply_section_results,
    _extract_missing_aspects, _generate_pending_questions,
)
from material_kb import build_material_kb, build_dimension_text
from material_index import MaterialIndex
from material_anchor import build_material_anchor
from tools import _read_single_file

TEMPLATE = r"templates_cache\福建普惠授信申报及审查审批意见表2025新版.docx"
MATERIAL_DIR = r"D:\刘野\众安\新建文件夹\2026.3.25续贷材料"
OUTPUT_DOCX = r"outputs\section_gen_test_v14v2.docx"
DUMP_TXT = r"outputs\_v14v2_dump.txt"


def load_materials(directory):
    exts = {'.txt', '.docx', '.doc', '.pdf', '.xlsx', '.xls'}
    result = {}
    path_map = {}
    for root, dirs, files in os.walk(directory):
        for fname in files:
            if os.path.splitext(fname)[1].lower() in exts:
                fp = os.path.join(root, fname)
                try:
                    content = _read_single_file(fp)
                    if content:
                        result[os.path.basename(fp)] = content
                        path_map[os.path.basename(fp)] = fp
                except Exception as e:
                    print(f"Skip {fname}: {e}")
    return result, path_map


def make_llm_caller():
    api_key = os.environ["DEEPSEEK_API_KEY"]
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
    n = [0]
    def caller(sys_p, usr_p):
        n[0] += 1
        for attempt in range(3):
            try:
                r = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[{"role": "system", "content": sys_p},
                              {"role": "user", "content": usr_p}],
                    max_tokens=8192, temperature=0.3, timeout=120,
                )
                return r.choices[0].message.content or ""
            except Exception as e:
                if attempt == 2:
                    print(f"    LLM retry fail: {e}")
                    return ""
    return caller


def progress(msg):
    print(f"[{time.strftime('%H:%M:%S')}] {msg}")


def main():
    print("=" * 60)
    print("V14 Section-only Pipeline Test")
    print("=" * 60)

    print("\n[1] Load materials...")
    file_contents, file_path_map = load_materials(MATERIAL_DIR)
    print(f"    {len(file_contents)} files")

    print("\n[2] Build KB + MaterialIndex + MaterialAnchor...")
    kb = build_material_kb(file_contents)
    material_index = MaterialIndex(file_contents)
    material_anchor = build_material_anchor(file_path_map)
    try:
        n_rc = len(getattr(material_anchor, "related_companies", []) or [])
        print(f"    related_companies in anchor: {n_rc}")
    except Exception:
        pass

    print("\n[3] Load template + decompose...")
    doc = Document(TEMPLATE)
    decomp = decompose_template(doc)
    print(f"    {decomp['stats']}")

    print("\n[4] Build FinancialIndicators...")
    financial_indicators = None
    try:
        from financial_analyzer import FinancialAnalyzer
        import re
        fin_xlsx = None
        for fname, fpath in file_path_map.items():
            if not fpath.lower().endswith((".xlsx", ".xls")):
                continue
            if re.search(r"会企|财务报表|资产负债|利润表|现金流", fname):
                fin_xlsx = fpath
                break
        if fin_xlsx:
            print(f"    xlsx: {os.path.basename(fin_xlsx)}")
            financial_indicators = FinancialAnalyzer().analyze(fin_xlsx)
    except Exception as e:
        print(f"    FI failed: {e}")

    print("\n[5] LLM caller + run section generation...")
    llm = make_llm_caller()

    sections_to_rewrite = [s for s in decomp['body_sections'] if s.has_content_to_rewrite]
    print(f"    {len(sections_to_rewrite)} sections to rewrite")

    pending_questions = []
    section_results = generate_all_sections(
        sections=sections_to_rewrite,
        kb=kb,
        company_profile="",
        llm_fn=llm,
        build_dimension_text_fn=build_dimension_text,
        truth_financial_data=None,
        template_paragraphs=decomp['body_paragraphs'],
        progress_cb=progress,
        max_retries=1,
        material_index=material_index,
        use_evidence_protocol=True,
        use_audit=True,
        financial_indicators=financial_indicators,
        material_anchor=material_anchor,
        pending_questions_sink=pending_questions,
    )

    print(f"\n[6] Apply section results: {len(section_results)} paragraphs")
    applied = apply_section_results(doc, section_results)
    print(f"    applied={applied}")

    os.makedirs("outputs", exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"    saved: {OUTPUT_DOCX}")

    # Dump body paragraphs
    with open(DUMP_TXT, "w", encoding="utf-8") as f:
        body_cell = doc.tables[0].rows[3].cells[0]
        for i, p in enumerate(body_cell.paragraphs):
            f.write(f"{i}\t{p.text}\n")
    print(f"    dump: {DUMP_TXT}")

    print(f"\n[7] Pending questions collected: {len(pending_questions)}")
    for q in pending_questions[:10]:
        print(f"    [{q['id']}] {q['section_title'][:30]}")
        print(f"      Q: {q['question']}")
        print(f"      hint: {q['hint']} (input_type={q['input_type']})")

    # Save pending_questions separately
    import json as _json
    with open("outputs/_v14v2_pending_questions.json", "w", encoding="utf-8") as f:
        _json.dump(pending_questions, f, ensure_ascii=False, indent=2)
    print("    saved: outputs/_v14v2_pending_questions.json")

    print("\n" + "=" * 60)
    print("Done")


if __name__ == "__main__":
    main()
