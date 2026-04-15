# -*- coding: utf-8 -*-
"""truth_fill.py

Truth-first deterministic fillers that bypass the LLM.

These functions are intentionally conservative: if KB or parsed statements
lack evidence, the target cells stay blank rather than guessed.
"""

from __future__ import annotations

import re
from copy import deepcopy
from typing import Any, Callable

from period_matcher import normalize_period, period_sort_key


def _norm(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"\s+", "", s)
    s = s.replace("\u3000", "")
    s = s.replace("\uff08", "(").replace("\uff09", ")")
    s = s.replace("\uff1a", ":")
    return s


def _contains(text: str, needles: list[str]) -> bool:
    t = _norm(text)
    return any(_norm(nd) in t for nd in needles if nd)


def _find_col(headers: list[str], needles: list[str]) -> int | None:
    hn = [_norm(h) for h in headers]
    normalized_needles = [_norm(nd) for nd in needles if _norm(nd)]

    for n in normalized_needles:
        for i, h in enumerate(hn):
            if h and h == n:
                return i

    for n in normalized_needles:
        for i, h in enumerate(hn):
            if not h:
                continue
            if n in h:
                return i
    return None


def _write_cell_text(cell, text: str) -> None:
    if cell is None:
        return
    try:
        text = "" if text is None else str(text)
        if cell.paragraphs:
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].text = text
                for r in p.runs[1:]:
                    r.text = ""
            else:
                new_run = p.add_run(text)
                template_rpr = None
                for other_p in cell.paragraphs:
                    for other_run in other_p.runs:
                        rpr = other_run._r.rPr
                        if rpr is not None:
                            template_rpr = deepcopy(rpr)
                            break
                    if template_rpr is not None:
                        break
                if template_rpr is not None:
                    new_run._r.insert(0, template_rpr)
        else:
            cell.text = text
    except Exception:
        try:
            cell.text = str(text)
        except Exception:
            pass


def _infer_data_start_row(nt) -> int:
    data_start = 1
    try:
        if len(nt.rows) >= 2:
            row1_texts = [c.text.strip() for c in nt.rows[1].cells]
            row1_filled = sum(1 for t in row1_texts if t.strip())
            row1_numeric = sum(
                1 for t in row1_texts if t and re.match(r"^[\d,.\-%]+$", t.replace(" ", ""))
            )
            if (
                row1_numeric == 0
                and row1_filled > len(row1_texts) * 0.5
                and all(len(t) < 30 for t in row1_texts if t)
            ):
                data_start = 2
    except Exception:
        data_start = 1
    return data_start


def _first_nonempty(item: dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = item.get(key)
        if value is None:
            continue
        text = str(value).strip()
        if text:
            return text
    return ""


def _value_by_key_contains(item: dict[str, Any], key_needles: list[str]) -> str:
    normalized_needles = [_norm(nd) for nd in key_needles if _norm(nd)]

    for n in normalized_needles:
        for k, v in (item or {}).items():
            if v is None:
                continue
            if _norm(str(k)) == n:
                text = str(v).strip()
                if text:
                    return text

    for k, v in (item or {}).items():
        if v is None:
            continue
        if _contains(str(k), key_needles):
            text = str(v).strip()
            if text:
                return text
    return ""


def _drop_filled_from_empty(ts) -> None:
    try:
        filled_pos = set(ts.existing_data.keys())
        ts.empty_cells = [(r, c) for (r, c) in ts.empty_cells if (r, c) not in filled_pos]
    except Exception:
        pass


def prefill_supply_chain_tables(
    doc,
    nested_tables: list[Any],
    kb: dict[str, Any],
    log: Callable[[str], None] | None = None,
) -> int:
    """Fill upstream/downstream Top5 tables using KB facts."""
    facts = (kb or {}).get("facts", {}) or {}
    upstream = facts.get("upstream_top5")
    downstream = facts.get("downstream_top5")

    if not isinstance(upstream, list) and not isinstance(downstream, list):
        return 0

    filled = 0

    def _extract_header_period_columns(headers: list[str], metric_keyword: str) -> dict[tuple[str, str], int]:
        period_cols: dict[tuple[str, str], int] = {}
        for ci, header in enumerate(headers):
            hs = str(header or "").strip()
            if metric_keyword not in hs:
                continue
            metric = "pct" if "占比" in hs else "amt"
            # Try canonical YYYY / YYYY.M period first
            period = normalize_period(hs)
            if re.fullmatch(r"20\d{2}(?:\.\d{1,2})?", period):
                period_cols[(metric, period)] = ci
                continue
            # Fallback: relative period headers without a 4-digit year
            if re.search(r"上年|去年", hs):
                period_cols[(metric, "annual")] = ci
            elif re.search(r"今年|本年|当年", hs):
                period_cols[(metric, "interim")] = ci
        return period_cols

    def _resolve_relative_periods(periods: list[str]) -> tuple[str | None, str | None]:
        annuals = sorted(
            [p for p in periods if re.fullmatch(r"20\d{2}", p)],
            key=period_sort_key,
        )
        interim = sorted(
            [p for p in periods if re.fullmatch(r"20\d{2}\.\d{1,2}", p)],
            key=period_sort_key,
        )
        annual_period = annuals[-1] if annuals else None
        interim_period = interim[-1] if interim else None
        return annual_period, interim_period

    def _log(msg: str):
        if log:
            try:
                log(msg)
            except Exception:
                pass

    up_markers = ["\u4e3b\u8981\u4e0a\u6e38\u4f9b\u5e94\u5546", "\u4e0a\u6e38\u4f9b\u5e94\u5546"]
    dn_markers = ["\u4e3b\u8981\u4e0b\u6e38\u9500\u552e\u5ba2\u6237", "\u4e0b\u6e38\u9500\u552e\u5ba2\u6237"]

    for ts in nested_tables or []:
        try:
            headers = [h or "" for h in (ts.header_row or [])]
            header_text = " ".join(h for h in headers if h)
        except Exception:
            continue

        is_up = isinstance(upstream, list) and upstream and _contains(header_text, up_markers)
        is_dn = isinstance(downstream, list) and downstream and _contains(header_text, dn_markers)
        if not is_up and not is_dn:
            continue

        try:
            parent_cell = doc.tables[ts.cell_path[0]].rows[ts.cell_path[1]].cells[ts.cell_path[2]]
            nt = parent_cell.tables[ts.nested_table_idx]
        except Exception:
            continue

        data_start = _infer_data_start_row(nt)
        items = upstream if is_up else downstream
        if not isinstance(items, list) or not items:
            continue

        if is_up:
            col_name = _find_col(headers, ["\u4e3b\u8981\u4e0a\u6e38\u4f9b\u5e94\u5546", "\u4e0a\u6e38\u4f9b\u5e94\u5546", "\u4f9b\u5e94\u5546"])
            col_product = _find_col(headers, ["\u91c7\u8d2d\u5546\u54c1", "\u91c7\u8d2d\u5185\u5bb9", "\u5546\u54c1"])
            period_cols = _extract_header_period_columns(headers, "\u91c7\u8d2d")
        else:
            col_name = _find_col(headers, ["\u4e3b\u8981\u4e0b\u6e38\u9500\u552e\u5ba2\u6237", "\u4e0b\u6e38\u9500\u552e\u5ba2\u6237", "\u5ba2\u6237"])
            col_product = _find_col(headers, ["\u9500\u552e\u4ea7\u54c1", "\u9500\u552e\u5546\u54c1", "\u9500\u552e\u5185\u5bb9", "\u5546\u54c1"])
            period_cols = _extract_header_period_columns(headers, "\u9500\u552e")

        all_periods = [period for _, period in period_cols.keys()]
        annual_period, interim_period = _resolve_relative_periods(all_periods)
        # Resolve amt/pct columns: try canonical YYYY period, then fallback "annual"/"interim"
        col_amt = (period_cols.get(("amt", annual_period)) if annual_period else None) or period_cols.get(("amt", "annual"))
        col_pct = (period_cols.get(("pct", annual_period)) if annual_period else None) or period_cols.get(("pct", "annual"))
        col_amt_ytd = (period_cols.get(("amt", interim_period)) if interim_period else None) or period_cols.get(("amt", "interim"))
        col_pct_ytd = (period_cols.get(("pct", interim_period)) if interim_period else None) or period_cols.get(("pct", "interim"))

        col_settle = _find_col(headers, ["\u7ed3\u7b97\u65b9\u5f0f"])
        col_term = _find_col(headers, ["\u8d26\u671f"])
        col_years = _find_col(headers, ["\u5408\u4f5c\u5e74\u9650"])

        max_rows = min(5, len(items), max(0, len(nt.rows) - data_start))
        for i in range(max_rows):
            ri = data_start + i
            row = nt.rows[ri]
            item = items[i]
            if not isinstance(item, dict):
                continue

            values = {
                col_name: _value_by_key_contains(item, ["\u4e0a\u6e38\u4f9b\u5e94\u5546", "\u4e0b\u6e38\u9500\u552e\u5ba2\u6237", "\u5ba2\u6237", "\u4f9b\u5e94\u5546"]),
                col_product: _value_by_key_contains(item, ["\u91c7\u8d2d\u5546\u54c1", "\u9500\u552e\u4ea7\u54c1", "\u9500\u552e\u5546\u54c1", "\u5546\u54c1", "\u5185\u5bb9"]),
                col_amt: _value_by_key_contains(item, ["\u4e0a\u5e74\u91c7\u8d2d\u989d", "\u4e0a\u5e74\u9500\u552e\u989d", "\u4e0a\u5e74\u91d1\u989d"]),
                col_pct: _value_by_key_contains(item, ["\u4e0a\u5e74\u91c7\u8d2d\u989d\u5360\u6bd4", "\u4e0a\u5e74\u9500\u552e\u989d\u5360\u6bd4", "\u5360\u6bd4"]),
                col_amt_ytd: _value_by_key_contains(item, ["\u4eca\u5e741-9\u6708\u91c7\u8d2d\u989d", "\u4eca\u5e741-9\u6708\u9500\u552e\u989d", "1-9\u6708\u91d1\u989d"]),
                col_pct_ytd: _value_by_key_contains(item, ["\u4eca\u5e741-9\u6708\u91c7\u8d2d\u989d\u5360\u6bd4", "\u4eca\u5e741-9\u6708\u9500\u552e\u989d\u5360\u6bd4"]),
                col_settle: _value_by_key_contains(item, ["\u7ed3\u7b97\u65b9\u5f0f"]),
                col_term: _value_by_key_contains(item, ["\u8d26\u671f"]),
                col_years: _value_by_key_contains(item, ["\u5408\u4f5c\u5e74\u9650"]),
            }

            for ci, value in values.items():
                if ci is None or ci >= len(row.cells) or not value:
                    continue
                if row.cells[ci].text.strip():
                    continue
                _write_cell_text(row.cells[ci], value)
                filled += 1
                ts.existing_data[(ri, ci)] = value

        for ri in range(data_start + max_rows, len(nt.rows)):
            row = nt.rows[ri]
            row_text = " ".join((c.text or "").strip() for c in row.cells)
            if any(marker in row_text for marker in ["已核实", "银行账户流水", "账户流水"]):
                for ci in range(len(row.cells)):
                    _write_cell_text(row.cells[ci], "")
                    ts.existing_data.pop((ri, ci), None)

        _drop_filled_from_empty(ts)
        if max_rows:
            _log(f"KB prefill {ts.table_id}: {'upstream' if is_up else 'downstream'} {max_rows} rows")

    return filled


def prefill_kb_structured_tables(
    doc,
    nested_tables: list[Any],
    kb: dict[str, Any],
    log: Callable[[str], None] | None = None,
) -> int:
    """Fill other objective KB-driven tables (affiliates/financing/R&D)."""
    facts = (kb or {}).get("facts", {}) or {}
    filled = 0

    def _log(msg: str):
        if log:
            try:
                log(msg)
            except Exception:
                pass

    for ts in nested_tables or []:
        headers = [h or "" for h in (ts.header_row or [])]
        header_text = " ".join(headers)

        try:
            parent_cell = doc.tables[ts.cell_path[0]].rows[ts.cell_path[1]].cells[ts.cell_path[2]]
            nt = parent_cell.tables[ts.nested_table_idx]
        except Exception:
            continue

        # Affiliates
        if (
            isinstance(facts.get("affiliates"), list)
            and facts.get("affiliates")
            and _contains(header_text, ["\u5173\u8054", "\u6301\u80a1", "\u5f80\u6765", "\u4f01\u4e1a"])
        ):
            items = facts.get("affiliates") or []
            data_start = _infer_data_start_row(nt)
            max_rows = max(0, len(nt.rows) - data_start)
            if nt.rows and nt.rows[-1].cells and _contains(nt.rows[-1].cells[0].text or "", ["\u6ce8", "\u8bf4\u660e"]):
                max_rows -= 1
            max_rows = min(max_rows, len(items))

            col_specs = {
                _find_col(headers, ["\u4f01\u4e1a\u540d\u79f0", "\u5173\u8054\u4f01\u4e1a\u540d\u79f0", "\u540d\u79f0"]): ["\u4f01\u4e1a\u540d\u79f0", "\u5173\u8054\u4f01\u4e1a\u540d\u79f0", "\u540d\u79f0"],
                _find_col(headers, ["\u6210\u7acb\u65f6\u95f4"]): ["\u6210\u7acb\u65f6\u95f4"],
                _find_col(headers, ["\u80a1\u6743\u7ed3\u6784"]): ["\u80a1\u6743\u7ed3\u6784"],
                _find_col(headers, ["\u4e3b\u8425\u4e1a\u52a1\u53ca\u4ea7\u54c1", "\u4e3b\u8425\u4e1a\u52a1"]): ["\u4e3b\u8425\u4e1a\u52a1\u53ca\u4ea7\u54c1", "\u4e3b\u8425\u4e1a\u52a1"],
                _find_col(headers, ["\u6ce8\u518c\u8d44\u672c", "\u5b9e\u6536\u8d44\u672c"]): ["\u6ce8\u518c\u8d44\u672c/\u5b9e\u6536\u8d44\u672c", "\u6ce8\u518c\u8d44\u672c", "\u5b9e\u6536\u8d44\u672c"],
                _find_col(headers, ["\u4e0a\u5e74\u4e3b\u8425\u4e1a\u52a1\u6536\u5165", "\u8425\u4e1a\u6536\u5165"]): ["\u4e0a\u5e74\u4e3b\u8425\u4e1a\u52a1\u6536\u5165", "\u4e0a\u5e74\u8425\u4e1a\u6536\u5165", "\u8425\u4e1a\u6536\u5165"],
                _find_col(headers, ["\u4e0a\u5e74\u51c0\u5229\u6da6", "\u51c0\u5229\u6da6"]): ["\u4e0a\u5e74 \u51c0\u5229\u6da6", "\u4e0a\u5e74\u51c0\u5229\u6da6", "\u51c0\u5229\u6da6"],
                _find_col(headers, ["\u6700\u65b0\u878d\u8d44", "\u878d\u8d44"]): ["\u6700\u65b0\u878d\u8d44\u91d1\u989d", "\u6700\u65b0\u878d\u8d44"],
                _find_col(headers, ["\u5173\u8054\u4ea4\u6613", "\u8d44\u91d1\u5f80\u6765", "\u4e1a\u52a1\u5f80\u6765"]): ["\u4e1a\u52a1\u5f80\u6765", "\u4e1a\u52a1\u5f80\u6765\u60c5\u51b5"],
                _find_col(headers, ["\u6d89\u8bc9", "\u8bc9\u8bbc"]): ["\u6d89\u8bc9\u60c5\u51b5", "\u5907\u6ce8"],
            }

            for i in range(max_rows):
                ri = data_start + i
                row = nt.rows[ri]
                item = items[i]
                if not isinstance(item, dict):
                    continue
                for ci, keys in col_specs.items():
                    if ci is None or ci >= len(row.cells):
                        continue
                    value = _value_by_key_contains(item, keys)
                    if value and not row.cells[ci].text.strip():
                        _write_cell_text(row.cells[ci], value)
                        filled += 1
                        ts.existing_data[(ri, ci)] = value

            _drop_filled_from_empty(ts)
            if max_rows:
                _log(f"KB prefill {ts.table_id}: affiliates {max_rows} rows")
            continue

        # Financing
        if (
            isinstance(facts.get("financing"), list)
            and facts.get("financing")
            and _contains(header_text, ["\u878d\u8d44\u673a\u6784", "\u6388\u4fe1\u91d1\u989d", "\u62c5\u4fdd\u65b9\u5f0f"])
        ):
            items = facts.get("financing") or []
            data_start = _infer_data_start_row(nt)
            max_rows = max(0, len(nt.rows) - data_start)
            if nt.rows and nt.rows[-1].cells and _contains(nt.rows[-1].cells[0].text or "", ["\u6ce8", "\u8bf4\u660e"]):
                max_rows -= 1
            max_rows = min(max_rows, len(items))

            col_specs = {
                _find_col(headers, ["\u878d\u8d44\u673a\u6784"]): ["\u878d\u8d44\u673a\u6784"],
                _find_col(headers, ["\u54c1\u79cd"]): ["\u54c1\u79cd"],
                _find_col(headers, ["\u6388\u4fe1\u91d1\u989d", "\u6388\u4fe1"]): ["\u6388\u4fe1\u91d1\u989d", "\u6388\u4fe1"],
                _find_col(headers, ["\u5df2\u7528\u91d1\u989d", "\u5df2\u7528"]): ["\u5df2\u7528\u91d1\u989d", "\u5df2\u7528"],
                _find_col(headers, ["\u4e1a\u52a1\u8d77\u59cb\u65e5", "\u8d77\u59cb\u65e5"]): ["\u4e1a\u52a1\u8d77\u59cb\u65e5", "\u8d77\u59cb\u65e5"],
                _find_col(headers, ["\u4e1a\u52a1\u5230\u671f\u65e5", "\u5230\u671f\u65e5"]): ["\u4e1a\u52a1\u5230\u671f\u65e5", "\u5230\u671f\u65e5"],
                _find_col(headers, ["\u62c5\u4fdd\u65b9\u5f0f"]): ["\u62c5\u4fdd\u65b9\u5f0f"],
            }

            for i in range(max_rows):
                ri = data_start + i
                row = nt.rows[ri]
                item = items[i]
                if not isinstance(item, dict):
                    continue
                for ci, keys in col_specs.items():
                    if ci is None or ci >= len(row.cells):
                        continue
                    value = _value_by_key_contains(item, keys)
                    if value and not row.cells[ci].text.strip():
                        _write_cell_text(row.cells[ci], value)
                        filled += 1
                        ts.existing_data[(ri, ci)] = value

            _drop_filled_from_empty(ts)
            if max_rows:
                _log(f"KB prefill {ts.table_id}: financing {max_rows} rows")
            continue

        # R&D
        if (
            isinstance(facts.get("r_and_d"), list)
            and facts.get("r_and_d")
            and _contains(header_text, ["\u7814\u53d1\u8d39\u7528"])
            and _contains(header_text, ["\u5e74\u4efd"])
        ):
            by_year = {}
            for item in facts.get("r_and_d") or []:
                if not isinstance(item, dict):
                    continue
                year = _value_by_key_contains(item, ["\u5e74\u4efd"])
                if year:
                    by_year[_norm(year)] = item

            data_start = _infer_data_start_row(nt)
            col_specs = {
                _find_col(headers, ["\u5e74\u4efd"]): ["\u5e74\u4efd"],
                _find_col(headers, ["\u7814\u53d1\u8d39\u7528", "\u7814\u53d1\u6295\u5165"]): ["\u7814\u53d1\u8d39\u7528", "\u7814\u53d1\u6295\u5165"],
                _find_col(headers, ["\u7814\u53d1\u8d39\u7528\u5360\u6bd4", "\u7814\u53d1\u8d39\u7528/\u8425\u6536"]): ["\u7814\u53d1\u8d39\u7528\u5360\u6bd4", "\u7814\u53d1\u8d39\u7528/\u8425\u6536"],
                _find_col(headers, ["\u53d1\u660e\u4e13\u5229"]): ["\u53d1\u660e\u4e13\u5229", "\u53d1\u660e\u4e13\u5229\u9879\u6570"],
                _find_col(headers, ["\u5b9e\u7528\u65b0\u578b\u4e13\u5229"]): ["\u5b9e\u7528\u65b0\u578b\u4e13\u5229", "\u5b9e\u7528\u65b0\u578b\u4e13\u5229\u9879\u6570"],
            }

            for ri in range(data_start, len(nt.rows)):
                row = nt.rows[ri]
                if not row.cells:
                    continue
                year = _norm(row.cells[0].text)
                item = by_year.get(year)
                if not item:
                    continue
                for ci, keys in col_specs.items():
                    if ci is None or ci >= len(row.cells):
                        continue
                    value = _value_by_key_contains(item, keys)
                    if value and not row.cells[ci].text.strip():
                        _write_cell_text(row.cells[ci], value)
                        filled += 1
                        ts.existing_data[(ri, ci)] = value

            _drop_filled_from_empty(ts)
            _log(f"KB prefill {ts.table_id}: r_and_d")
            continue

    return filled


# ---------------------------------------------------------------------------
# New deterministic fillers (改造五)
# ---------------------------------------------------------------------------

def prefill_shareholder_table(
    doc,
    nested_tables: list[Any],
    kb: dict[str, Any],
    log: Callable[[str], None] | None = None,
) -> int:
    """Fill shareholder structure table using KB facts['shareholders'].

    Critically, this fills the *shareholder* table (股东名称/出资比例),
    NOT the affiliate table.  The semantic distinction prevents the previous
    bug where subsidiary names were inserted as shareholders.
    """
    facts = (kb or {}).get("facts", {}) or {}
    shareholders = facts.get("shareholders")
    if not isinstance(shareholders, list) or not shareholders:
        return 0

    filled = 0

    def _log(msg: str):
        if log:
            try:
                log(msg)
            except Exception:
                pass

    for ts in nested_tables or []:
        try:
            headers = [h or "" for h in (ts.header_row or [])]
            header_text = " ".join(h for h in headers if h)
        except Exception:
            continue

        # Must have shareholder-specific columns — NOT affiliate columns
        is_shareholder = (
            _contains(header_text, ["股东名称", "认缴出资", "出资比例"])
            and not _contains(header_text, ["成立时间", "净利润", "融资"])
        )
        if not is_shareholder:
            continue

        try:
            parent_cell = doc.tables[ts.cell_path[0]].rows[ts.cell_path[1]].cells[ts.cell_path[2]]
            nt = parent_cell.tables[ts.nested_table_idx]
        except Exception:
            continue

        data_start = _infer_data_start_row(nt)

        col_name = _find_col(headers, ["股东名称", "股东"])
        col_amount = _find_col(headers, ["认缴出资额", "认缴出资", "出资额", "持股数"])
        col_pct = _find_col(headers, ["出资比例%", "出资比例", "持股比例"])
        col_actual = _find_col(headers, ["实际出资额", "实际出资", "实缴出资"])
        col_form = _find_col(headers, ["出资形式", "出资方式"])

        # Find the "合计" row if it exists — we fill above it
        total_row_idx = None
        for ri in range(data_start, len(nt.rows)):
            row_text = " ".join((c.text or "").strip() for c in nt.rows[ri].cells)
            if "合计" in row_text:
                total_row_idx = ri
                break

        max_data_rows = (total_row_idx or len(nt.rows)) - data_start
        max_rows = min(max_data_rows, len(shareholders))

        for i in range(max_rows):
            ri = data_start + i
            if ri >= len(nt.rows):
                break
            row = nt.rows[ri]
            item = shareholders[i]
            if not isinstance(item, dict):
                continue

            values = {
                col_name: _value_by_key_contains(item, ["股东名称", "股东"]),
                col_amount: _value_by_key_contains(item, ["认缴出资额", "认缴出资", "出资额", "持股数"]),
                col_pct: _value_by_key_contains(item, ["出资比例%", "出资比例", "持股比例"]),
                col_actual: _value_by_key_contains(item, ["实际出资额", "实际出资", "实缴出资"]),
                col_form: _value_by_key_contains(item, ["出资形式", "出资方式"]),
            }

            for ci, value in values.items():
                if ci is None or ci >= len(row.cells) or not value:
                    continue
                if row.cells[ci].text.strip():
                    continue
                _write_cell_text(row.cells[ci], value)
                filled += 1
                ts.existing_data[(ri, ci)] = value

        _drop_filled_from_empty(ts)
        if max_rows:
            _log(f"KB prefill {ts.table_id}: shareholders {max_rows} rows")

    return filled


def prefill_asset_table(
    doc,
    nested_tables: list[Any],
    kb: dict[str, Any],
    log: Callable[[str], None] | None = None,
) -> int:
    """Fill asset tables (借款人资产 / 实控人家庭资产) using KB facts['assets']."""
    facts = (kb or {}).get("facts", {}) or {}
    assets = facts.get("assets")
    if not isinstance(assets, list) or not assets:
        return 0

    filled = 0

    def _log(msg: str):
        if log:
            try:
                log(msg)
            except Exception:
                pass

    for ts in nested_tables or []:
        try:
            headers = [h or "" for h in (ts.header_row or [])]
            header_text = " ".join(h for h in headers if h)
        except Exception:
            continue

        if not _contains(header_text, ["资产名称"]):
            continue
        if not _contains(header_text, ["所有权人", "权属", "资产类型"]):
            continue

        try:
            parent_cell = doc.tables[ts.cell_path[0]].rows[ts.cell_path[1]].cells[ts.cell_path[2]]
            nt = parent_cell.tables[ts.nested_table_idx]
        except Exception:
            continue

        data_start = _infer_data_start_row(nt)

        col_specs = {
            _find_col(headers, ["资产名称", "坐落位置"]): ["资产名称", "坐落位置"],
            _find_col(headers, ["所有权人"]): ["所有权人"],
            _find_col(headers, ["资产类型"]): ["资产类型"],
            _find_col(headers, ["权属", "权属证号"]): ["权属", "权属证号"],
            _find_col(headers, ["资产概况", "面积"]): ["资产概况", "面积"],
            _find_col(headers, ["账面", "购置价"]): ["账面", "购置价"],
            _find_col(headers, ["市场", "评估价"]): ["市场", "评估价"],
            _find_col(headers, ["自用", "出租"]): ["自用", "出租"],
            _find_col(headers, ["抵质押"]): ["抵质押"],
            _find_col(headers, ["融资金额"]): ["融资金额"],
        }

        max_rows = min(max(0, len(nt.rows) - data_start), len(assets))

        for i in range(max_rows):
            ri = data_start + i
            if ri >= len(nt.rows):
                break
            row = nt.rows[ri]
            item = assets[i]
            if not isinstance(item, dict):
                continue
            for ci, keys in col_specs.items():
                if ci is None or ci >= len(row.cells):
                    continue
                value = _value_by_key_contains(item, keys)
                if value and not row.cells[ci].text.strip():
                    _write_cell_text(row.cells[ci], value)
                    filled += 1
                    ts.existing_data[(ri, ci)] = value

        _drop_filled_from_empty(ts)
        if max_rows:
            _log(f"KB prefill {ts.table_id}: assets {max_rows} rows")

    return filled


def prefill_bank_flow_table(
    doc,
    nested_tables: list[Any],
    kb: dict[str, Any],
    log: Callable[[str], None] | None = None,
) -> int:
    """Fill bank flow verification table using KB facts['bank_flows']."""
    facts = (kb or {}).get("facts", {}) or {}
    flows = facts.get("bank_flows")
    if not isinstance(flows, list) or not flows:
        return 0

    filled = 0

    def _log(msg: str):
        if log:
            try:
                log(msg)
            except Exception:
                pass

    for ts in nested_tables or []:
        try:
            headers = [h or "" for h in (ts.header_row or [])]
            header_text = " ".join(h for h in headers if h)
        except Exception:
            continue

        if not _contains(header_text, ["开户行"]):
            continue
        if not _contains(header_text, ["流入量", "流入", "户名"]):
            continue

        try:
            parent_cell = doc.tables[ts.cell_path[0]].rows[ts.cell_path[1]].cells[ts.cell_path[2]]
            nt = parent_cell.tables[ts.nested_table_idx]
        except Exception:
            continue

        data_start = _infer_data_start_row(nt)

        col_specs = {
            _find_col(headers, ["开户行"]): ["开户行"],
            _find_col(headers, ["户名"]): ["户名"],
            _find_col(headers, ["账号"]): ["账号"],
            _find_col(headers, ["关系", "与借款人关系"]): ["关系", "与借款人关系"],
            _find_col(headers, ["时间段", "时间"]): ["时间段", "时间"],
            _find_col(headers, ["流入量", "流入"]): ["流入量", "流入"],
        }

        max_rows = min(max(0, len(nt.rows) - data_start), len(flows))

        for i in range(max_rows):
            ri = data_start + i
            if ri >= len(nt.rows):
                break
            row = nt.rows[ri]
            item = flows[i]
            if not isinstance(item, dict):
                continue
            for ci, keys in col_specs.items():
                if ci is None or ci >= len(row.cells):
                    continue
                value = _value_by_key_contains(item, keys)
                if value and not row.cells[ci].text.strip():
                    _write_cell_text(row.cells[ci], value)
                    filled += 1
                    ts.existing_data[(ri, ci)] = value

        _drop_filled_from_empty(ts)
        if max_rows:
            _log(f"KB prefill {ts.table_id}: bank_flows {max_rows} rows")

    return filled


def prefill_labeled_fields_from_kb(
    labeled_fields: list[Any],
    kb: dict[str, Any],
    financial_data: dict[str, Any] | None = None,
) -> dict[str, str]:
    """Pre-fill labeled fields (label:blank patterns) with KB facts.

    Returns a dict of {lf_id: value} for fields that can be filled
    deterministically from KB, so they don't need LLM processing.

    F4 扩展：覆盖全部可从KB/财务数据确定性填充的字段，
    减少LLM猜测，从根本上消除幻觉和遗漏。
    """
    facts = (kb or {}).get("facts", {}) or {}
    if not facts and not financial_data:
        return {}

    result: dict[str, str] = {}

    # Helper: get latest annual financial data
    def _latest_annual_value(*keys):
        if not financial_data:
            return ""
        annual_periods = sorted(
            [k for k in financial_data.keys() if re.fullmatch(r"20\d{2}", str(k))],
        )
        if not annual_periods:
            return ""
        latest = financial_data.get(annual_periods[-1], {})
        for key in keys:
            val = latest.get(key)
            if val is not None:
                return str(val)
        return ""

    # Comprehensive mapping: label keywords -> KB fact key or callable
    # 覆盖表头区、借款人简介区、科技区、担保区等全部标签字段
    label_to_fact = [
        # ── 表头基本信息 ──
        (["客户名称"], "company_name"),
        (["申报单位"], "reporting_unit"),
        (["客户经理"], "customer_manager_name"),
        (["联系电话"], "customer_manager_phone"),
        (["申报额度"], "applied_credit_amount"),
        (["申报敞口"], "applied_exposure"),
        (["上一期授信敞口"], "prev_credit_exposure"),
        (["业务品种"], "business_product"),
        (["业务期限"], "business_term"),
        (["授信期限"], "credit_term"),
        (["担保方式"], "guarantee_method"),
        (["上一期授信担保方式"], "prev_guarantee_method"),
        (["PD评级"], "pd_rating"),
        (["上一期授信评级"], "prev_pd_rating"),
        (["国标行业"], "industry"),
        (["绿通标识"], "green_channel"),

        # ── 原有额度批复日期 ──
        (["原有额度批复日期"], "prev_approval_date"),

        # ── 企业基本信息 ──
        (["注册资本"],
         lambda: facts.get("registered_capital") or _latest_annual_value(
             "registered_capital", "注册资本") or ""),
        (["实收资本"],
         lambda: facts.get("paid_in_capital") or _fmt_wan(
             _latest_annual_value("paid_in_capital", "实收资本", "实收资本（或股本）")) or ""),
        (["法定代表人", "法人代表"], "legal_representative"),
        (["实际控制人姓名", "实控人姓名"], "controller_name"),
        (["经营地址", "实际经营地址"], "operating_address"),
        (["社保人数", "上年度末社保"], "social_insurance_count"),
        (["员工人数", "上年度末员工"],
         lambda: facts.get("social_insurance_count") or facts.get("employee_count", "")),

        # ── 划型要素（从财务数据+KB获取）──
        (["上年度资产总额"],
         lambda: _latest_annual_value("total_assets", "资产总计", "资产合计")),
        (["上年度销售额"],
         lambda: _latest_annual_value("revenue", "operating_revenue", "营业收入")),

        # ── 实控人信息 ──
        (["实际控制人持股", "股权穿透后"], "controller_share_pct"),
        (["过往两年企业股权变动"], "equity_change"),

        # ── 风险信息 ──
        (["反洗钱风险等级"], "aml_risk_level"),

        # ── 科技信息 ──
        (["科技型企业评分"], "tech_score"),

        # ── 在手订单 ──
        (["在手订单"], lambda: _format_orders(facts)),

        # ── 综合效益 ──
        (["日均存款"], "avg_daily_deposit"),
        (["代发签约"], "payroll_signup"),
    ]

    for lf in labeled_fields:
        lab = (getattr(lf, 'label', '') or '').strip()

        for keywords, fact_key in label_to_fact:
            if not any(kw in lab for kw in keywords):
                continue
            if fact_key is None:
                break
            if callable(fact_key):
                value = fact_key()
            else:
                value = str(facts.get(fact_key, "")).strip()
            if value and value.lower() not in ("none", "false", "true"):
                # Strip unit suffix to avoid duplication (e.g., field already has "万元" after blank)
                unit = getattr(lf, 'unit', '').strip()
                if unit:
                    # unit='万' means template has "万元" after blank; strip "万元"/"万" from value
                    for suffix in [unit + "元", unit]:
                        if value.endswith(suffix):
                            value = value[:-len(suffix)].strip()
                            break
                result[lf.lf_id] = value
            break

    return result


def _fmt_wan(val_str: str) -> str:
    """Convert raw numeric string (possibly in yuan) to '万元' format."""
    if not val_str:
        return ""
    try:
        v = float(str(val_str).replace(",", ""))
        if v >= 10000:
            return f"{v / 10000:.0f}万元"
        elif v > 0:
            return f"{v:.0f}元"
    except (ValueError, TypeError):
        pass
    # Already has unit
    if "万" in str(val_str):
        return str(val_str)
    return str(val_str)


def _format_orders(facts: dict) -> str:
    """Format in-hand orders from KB facts."""
    count = facts.get("orders_count", "")
    total = facts.get("orders_total_amount", "")
    uncollected = facts.get("orders_uncollected", "")
    if not count and not total:
        return ""
    parts = []
    if count:
        parts.append(f"{count}笔")
    if total:
        parts.append(f"合同总金额{total}")
    if uncollected:
        parts.append(f"未回款{uncollected}")
    return "，".join(parts)


def prefill_checkboxes_from_kb(
    checkboxes: list[Any],
    kb: dict[str, Any],
) -> dict[str, bool]:
    """Deterministically set checkboxes based on KB inference rules.

    F4 框架：基于已知事实推断checkbox状态，不依赖LLM猜测。
    返回 {cb_id: should_be_checked} 字典。

    设计原则：
    - 只在有充分证据时才设置，不确定则不动（留给LLM）
    - 推断规则基于业务逻辑，不是关键词匹配
    """
    facts = (kb or {}).get("facts", {}) or {}
    if not facts:
        return {}

    result: dict[str, bool] = {}

    for cb in checkboxes:
        ctx = (getattr(cb, 'group_context', '') or getattr(cb, 'context_line', '') or '').strip()
        label = (getattr(cb, 'option_text', '') or getattr(cb, 'label', '') or '').strip()
        opt_text = label  # option_text only, for rules that must NOT match group context
        key_text = ctx + ' ' + label
        cb_id = getattr(cb, 'cb_id', None) or getattr(cb, 'field_id', None)
        if not cb_id:
            continue

        # ── 存量/新客户判断 ──
        # 使用opt_text避免同组checkbox交叉污染
        if "新客户" in opt_text:
            is_existing = facts.get("is_existing_customer")
            if is_existing is True:
                result[cb_id] = False  # 不勾新客户
            elif is_existing is False:
                result[cb_id] = True
        elif "存量客户" in opt_text:
            is_existing = facts.get("is_existing_customer")
            if is_existing is True:
                result[cb_id] = True  # 勾存量客户
            elif is_existing is False:
                result[cb_id] = False

        # ── 维持/增加/减少 ──
        elif "维持" in opt_text:
            change = facts.get("credit_change_direction")
            if change == "maintain":
                result[cb_id] = True
        elif "增加" in opt_text:
            change = facts.get("credit_change_direction")
            if change == "increase":
                result[cb_id] = True
        elif "减少" in opt_text and "维持" not in opt_text:
            change = facts.get("credit_change_direction")
            if change == "decrease":
                result[cb_id] = True

        # ── 授信方案调整 ──
        elif "授信方案调整" in opt_text:
            if facts.get("guarantee_changed") or facts.get("credit_change_direction"):
                result[cb_id] = True

        # ── 规模划型 ──
        elif "微型" in opt_text and "小型" not in opt_text:
            scale = facts.get("enterprise_scale")
            if scale:
                result[cb_id] = ("微型" in scale)
        elif "小型" in opt_text and "微型" not in opt_text:
            scale = facts.get("enterprise_scale")
            if scale:
                result[cb_id] = ("小型" in scale)
        elif "中型" in opt_text:
            scale = facts.get("enterprise_scale")
            if scale:
                result[cb_id] = ("中型" in scale)

        # ── 科技型企业类型 ──
        # Match against option_text (not group context) to avoid cross-contamination
        elif "科技型中小企业" in opt_text:
            tech_types = facts.get("tech_enterprise_types", "")
            if tech_types:
                result[cb_id] = ("科技型中小企业" in tech_types)
        elif "高新技术企业" in opt_text:
            tech_types = facts.get("tech_enterprise_types", "")
            if tech_types:
                result[cb_id] = ("高新技术企业" in tech_types)
        elif "创新型中小企业" in opt_text:
            tech_types = facts.get("tech_enterprise_types", "")
            if tech_types:
                result[cb_id] = ("创新型中小企业" in tech_types)
        elif "专精特新" in opt_text and "小巨人" not in opt_text:
            tech_types = facts.get("tech_enterprise_types", "")
            if tech_types:
                result[cb_id] = ("专精特新" in tech_types)

        # ── ESG评级 ──
        elif re.search(r"ESG.*[ABC]", opt_text) or re.search(r"[ABC].*ESG", opt_text):
            esg = facts.get("esg_rating")
            if esg:
                for rating in ["A", "B", "C"]:
                    if rating in opt_text:
                        result[cb_id] = (esg == rating)
                        break

        # ── 客群分类 ──
        elif "支持类" in opt_text:
            cg = facts.get("customer_group")
            if cg:
                result[cb_id] = ("支持" in cg)
        elif "审慎类" in opt_text:
            cg = facts.get("customer_group")
            if cg:
                result[cb_id] = ("审慎" in cg)
        elif "禁止类" in opt_text:
            cg = facts.get("customer_group")
            if cg:
                result[cb_id] = ("禁止" in cg)

    return result


# ---------------------------------------------------------------------------
# V14 G1:集团合并报表溯源 + "(若有)"表 gating
# ---------------------------------------------------------------------------

# 需要严格溯源的"若有"类表头 → pending_tag 提示
_CONSOLIDATED_HEADER_PATTERNS = [
    # 关键词清单,任一命中即视为"需合并口径才能填"的表
    ("集团合并财务数据", "需补充近三年合并财务报表"),
    ("合并财务报表", "需补充近三年合并财务报表"),
    ("合并资产负债表", "需补充近三年合并财务报表"),
    ("合并利润表", "需补充近三年合并财务报表"),
    ("合并现金流量表", "需补充近三年合并财务报表"),
]

# 合并报表 sheet/段落的识别词(docx/xlsx 共用)
_CONSOLIDATED_SOURCE_MARKERS = [
    "合并资产负债表", "合并利润表", "合并现金流量表",
    "合并财务报表", "合并报表", "集团合并",
    "consolidated balance", "consolidated income",
    "consolidated cash flow", "consolidated financial",
]


def _scan_xlsx_sheetnames_for_consolidated(xlsx_path: str) -> bool:
    """扫 xlsx 所有 sheet 名,命中任一合并标识返回 True。解析失败返回 False。"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return False
    try:
        wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    except Exception:
        return False
    try:
        for sn in wb.sheetnames:
            s = (sn or "").lower()
            for mk in _CONSOLIDATED_SOURCE_MARKERS:
                if mk.lower() in s:
                    return True
    finally:
        try:
            wb.close()
        except Exception:
            pass
    return False


def _scan_docx_for_consolidated(docx_path: str) -> bool:
    """扫 docx 段落 + 表格首格文字,命中任一合并标识返回 True。"""
    try:
        from docx import Document
    except ImportError:
        return False
    try:
        doc = Document(docx_path)
    except Exception:
        return False
    markers_lower = [m.lower() for m in _CONSOLIDATED_SOURCE_MARKERS]
    try:
        for p in doc.paragraphs:
            t = (p.text or "").lower()
            if not t:
                continue
            for m in markers_lower:
                if m in t:
                    return True
        for tbl in doc.tables:
            for row in tbl.rows[:3]:  # 仅前 3 行检查表头区
                for cell in row.cells:
                    t = (cell.text or "").lower()
                    for m in markers_lower:
                        if m in t:
                            return True
    except Exception:
        return False
    return False


def _has_source_for_consolidated_report(
    file_path_map: dict[str, str] | None,
    material_anchor: Any | None = None,
) -> bool:
    """判断客户材料里是否存在"合并报表"的真实来源。

    扫描规则:
      1. file_path_map 中每个 xlsx → 检查 sheet 名是否含合并标识
      2. file_path_map 中每个 docx → 检查段落 + 表格前 3 行是否含合并标识
      3. material_anchor 若附带合并来源标记(如新字段)也视为 True

    False → 调用方必须把"集团合并财务数据"等 (若有) 表标记为 pending,
            严禁拿借款人单体数据冒充。
    """
    # 先查 material_anchor(快捷通道)
    if material_anchor is not None:
        has_flag = getattr(material_anchor, "has_consolidated_source", None)
        if has_flag is True:
            return True

    if not file_path_map:
        return False

    for fname, fpath in (file_path_map or {}).items():
        if not fpath:
            continue
        low = fpath.lower()
        base = (fname or "").lower()
        # 文件名命中"合并"也算(通常带"合并财务报表.xlsx"等命名)
        if "合并" in fname or "合并" in os.path.basename(fpath or ""):
            return True
        try:
            if low.endswith((".xlsx", ".xls")):
                if _scan_xlsx_sheetnames_for_consolidated(fpath):
                    return True
            elif low.endswith(".docx"):
                if _scan_docx_for_consolidated(fpath):
                    return True
        except Exception:
            continue
    return False


def should_skip_consolidated_financial_table(
    header_text: str,
    file_path_map: dict[str, str] | None,
    material_anchor: Any | None = None,
) -> tuple[bool, str]:
    """针对识别为"合并/集团合并/若有"类的财务表,决定是否跳过填写。

    返回 (should_skip, pending_tag):
      - should_skip=True:调用方**不得**用借款人单体数据冒充合并,应留空并将 pending_tag
        写入 pending_questions / pending_tags,让审贷员人工补材料。
      - should_skip=False:该表可按常规单体数据填写(或不是合并类表)。

    设计:缺失即跳过,严禁编造。
    """
    text = _norm(header_text or "")
    matched_tag: str | None = None
    for kw, tag in _CONSOLIDATED_HEADER_PATTERNS:
        if _norm(kw) in text:
            matched_tag = tag
            break
    if not matched_tag:
        # 兜底:含"(若有)"/"如有"且同时提"合并/集团"字样的表也纳入
        if ("若有" in text or "如有" in text) and ("合并" in text or "集团" in text):
            matched_tag = "需补充近三年合并财务报表"
    if not matched_tag:
        return (False, "")

    has_source = _has_source_for_consolidated_report(file_path_map, material_anchor)
    if has_source:
        return (False, "")
    return (True, matched_tag)


# 确保 os 可用(该文件顶部未 import os,此处 lazy import 避免引起破坏)
import os  # noqa: E402
