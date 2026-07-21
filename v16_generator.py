# -*- coding: utf-8 -*-
"""V16 Generator — element-level docx 产出器.

输入:
  - classifier 产物: `outputs/v16_llm_classified.json`
    (每个 location → {op, label, confidence, justification})
  - 模板 docx: `samples/*.docx`
  - 材料目录: 一组客户资料(docx/xlsx/pdf/txt)
输出:
  - 填充后的 docx
  - pending_tags.json (未能自动填写的字段清单,供前端展示)

架构参见 plans/dapper-cuddling-cupcake.md。Step 1 范围:骨架 + dry-run round-trip。
"""
from __future__ import annotations

import glob
import json
import os
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document

# 复用 v16 classifier 的 Element / section 索引
sys.path.insert(0, str(Path(__file__).parent))
from v16_classifier import build_section_index
from v16_step1_extract import Element, extract_elements

OUTPUT_DIR = Path(__file__).parent / "outputs"
DEFAULT_CLASSIFIED_JSON = OUTPUT_DIR / "v16_llm_classified.json"


# ────────────────────────────────────────────────────────────
# 数据结构
# ────────────────────────────────────────────────────────────

@dataclass
class Classification:
    """单个 element 的分类结果."""
    location: str
    op: str           # PRESERVE / FILL / REWRITE / REPLACE (2026-05-21 治本 Phase 1 新增 REPLACE)
    label: str        # SCAFFOLD / PRESERVE / FILL / CLEAR / SLOT / CHECKBOX / REWRITE / PLACEHOLDER
    confidence: float
    justification: str = ""


@dataclass
class Materials:
    """材料层 facade — 给 handler 统一消费接口.

    Step 1 仅作为占位;KB / financial / anchors 的实际加载在 Step 2+ 接入.
    body_gaps: Step 5 预处理,location → gap_info 映射,用于 SCAFFOLD 注入提示。
    client_metadata: 2026-05-21 治本 Phase 1 新增 — v16 模板 placeholder 化的
        当前客户档案 (key 见 templates/placeholder-schema.json) · REPLACE op handler
        从此读 placeholder 替换值 · None 时 placeholder 整段保留原文 + 写 pending_tag.
    """
    file_contents: dict[str, str] = field(default_factory=dict)
    kb: dict[str, Any] = field(default_factory=dict)
    financial: Any = None
    anchors: dict[str, Any] = field(default_factory=dict)
    index: Any = None
    body_gaps: dict[str, dict[str, Any]] = field(default_factory=dict)
    client_metadata: dict[str, Any] | None = None

    @property
    def facts(self) -> dict[str, Any]:
        return self.kb.get("facts", {}) if self.kb else {}

    @property
    def tables(self) -> dict[str, Any]:
        return self.kb.get("tables", {}) if self.kb else {}


@dataclass
class GenResult:
    """单个 element 的处理结果.

    action:
      - keep  : 保留原文,不改 docx
      - fill  : 用 new_text 替换 element 的文本
      - clear : 清空为 空字符串 / pending 标签
    pending_tag: 若非 None,会写入 pending_tags.json
    """
    location: str
    action: str
    new_text: str | None = None
    pending_tag: dict[str, Any] | list[dict[str, Any]] | None = None
    debug: str = ""


# ────────────────────────────────────────────────────────────
# I/O
# ────────────────────────────────────────────────────────────

def load_classifier_output(json_path: Path) -> dict[str, Classification]:
    """读 classifier 产出的 JSON,返回 location → Classification."""
    data = json.loads(Path(json_path).read_text(encoding="utf-8"))
    out: dict[str, Classification] = {}
    for loc, rec in data.get("classifications_by_location", {}).items():
        out[loc] = Classification(
            location=loc,
            op=rec.get("op", ""),
            label=rec.get("label", ""),
            confidence=float(rec.get("confidence", 0.0) or 0.0),
            justification=rec.get("justification", ""),
        )
    return out


def load_template(docx_path: Path) -> tuple[Any, list[Element], dict[str, list[str]]]:
    """读模板 docx,返回 (doc, elements, section_by_loc).

    - doc:  python-docx Document 对象(可直接修改)
    - elements: 按出现顺序的 Element 列表
    - section_by_loc: location → 最近 3 层 section 标题栈
    """
    doc = Document(str(docx_path))
    elements = extract_elements(Path(docx_path))
    section_by_loc = build_section_index(elements)
    return doc, elements, section_by_loc


def load_materials(material_dir: Path | None) -> Materials:
    """读材料目录,构建 KB.

    Step 2: 加载 file_contents + build_material_kb。
    Step 4+: 延迟接入 financial / anchors / index。
    """
    if material_dir is None or not Path(material_dir).exists():
        return Materials()

    # 延迟 import 避免 Step 1 强依赖 tools 整个模块
    from tools import _read_single_file
    from material_kb import build_material_kb

    exts = {".txt", ".docx", ".doc", ".pdf", ".xlsx", ".xls"}
    file_contents: dict[str, str] = {}
    for fp in glob.glob(os.path.join(str(material_dir), "*")):
        ext = os.path.splitext(fp)[1].lower()
        if ext not in exts:
            continue
        try:
            content = _read_single_file(fp)
            if content:
                file_contents[os.path.basename(fp)] = content
        except Exception as e:
            print(f"  [skip] {os.path.basename(fp)}: {e}")

    kb = build_material_kb(file_contents) if file_contents else {}

    # P2: 财务分析(deterministic) — 解析 xlsx 财务报表,供 REWRITE 做
    # 同比/趋势/三大活动现金流等深度叙事。若无 xlsx 或解析失败,静默跳过。
    financial_block = ""
    try:
        from financial_analyzer import FinancialAnalyzer
        fa = FinancialAnalyzer()
        xlsx_candidates = [
            fp for fp in glob.glob(os.path.join(str(material_dir), "*"))
            if os.path.splitext(fp)[1].lower() in (".xlsx", ".xls")
            and any(kw in os.path.basename(fp) for kw in ("财务", "报表", "资产负债", "利润"))
        ]
        if xlsx_candidates:
            xlsx_candidates.sort(reverse=True)  # 取最新(文件名含日期时倾向较新)
            ind = fa.analyze(xlsx_candidates[0])
            financial_block = fa.format_for_prompt(ind)
            print(f"  financial: parsed {os.path.basename(xlsx_candidates[0])}"
                  f" → {len(financial_block)} chars")
    except Exception as exc:
        print(f"  [skip financial] {type(exc).__name__}: {exc}")

    # P1: 硬字段格式正则抽取 — 补 material_kb 关键词启发式的召回缺口
    if file_contents:
        from hard_fields import extract_hard_fields
        hard = extract_hard_fields(file_contents)
        facts = kb.setdefault("facts", {})
        # hard_fields 对以下字段拥有更强的抽取逻辑(格式正则 / 源优先),
        # 应覆盖 material_kb 可能存在的脏值(如 operating_address 拼接错位)
        _HARD_OVERRIDE = {
            "uscc", "controller_id", "controller_birth",
            "phone", "post_code", "bank_account",
            "operating_address", "registered_address",
            "auditor", "lease_area", "lease_location",
            "controller_home_address",
        }
        for k, v in hard.items():
            if k.startswith("_"):
                continue
            if k in _HARD_OVERRIDE and v:
                facts[k] = v
            elif facts.get(k) in (None, ""):
                facts[k] = v
        kb.setdefault("hard_fields_debug", hard.get("_source_counts", {}))

    mats = Materials(file_contents=file_contents, kb=kb)
    if financial_block:
        mats.financial = {"prompt_block": financial_block}
    return mats


# ────────────────────────────────────────────────────────────
# docx location → element 解析(写回 docx 时用)
# ────────────────────────────────────────────────────────────

def _resolve_docx_elements(doc, location: str) -> list:
    """把 location 字符串解析到 docx Paragraph 列表.

    支持:
      - P{n}           → [doc.paragraphs[n]]
      - T{t}R{r}C{c}P{p}            → [cell 内第 p 个 paragraph]
      - T{t}R{r}C{c}NTR{r2}C{c2}P{p} → 父 cell 下所有子表里对应 (r2,c2,p) 的
        paragraph 列表(extract_elements 对同 cell 多张子表共用相同 NT location,
        resolve 时对所有匹配子表都返回,由 apply 层统一处理)
    """
    # 简单段落
    m = re.match(r"^P(\d+)$", location)
    if m:
        pi = int(m.group(1))
        if 0 <= pi < len(doc.paragraphs):
            return [doc.paragraphs[pi]]
        return []

    m = re.match(r"^T(\d+)(.+)$", location)
    if not m:
        return []
    t_idx = int(m.group(1))
    rest = m.group(2)
    if t_idx >= len(doc.tables):
        return []

    # 候选 cell 集合(支持 NT 层多分支)
    current_cells = [None]  # 占位,第一步用 doc.tables[t_idx]
    current_tables = [doc.tables[t_idx]]

    cursor = 0
    while cursor < len(rest):
        mm = re.match(r"^(NT)?R(\d+)C(\d+)", rest[cursor:])
        if not mm:
            break
        is_nt, r_idx, c_idx = bool(mm.group(1)), int(mm.group(2)), int(mm.group(3))

        next_cells = []
        if is_nt:
            # 展开所有当前 cell 下的子表为新 tables
            next_tables = []
            for cell in current_cells:
                if cell is None:
                    continue
                for sub in cell.tables:
                    next_tables.append(sub)
            current_tables = next_tables
            current_cells = []

        for tbl in current_tables:
            if r_idx >= len(tbl.rows):
                continue
            row = tbl.rows[r_idx]
            if c_idx >= len(row.cells):
                continue
            next_cells.append(row.cells[c_idx])

        current_cells = next_cells
        # 下一轮的 tables 默认继承(若下一段又是 NT,再重新展开)
        current_tables = [] if not current_cells else [c for c in current_cells if c is not None]
        cursor += mm.end()

    mp = re.match(r"^P(\d+)$", rest[cursor:])
    if not mp:
        return []
    p_idx = int(mp.group(1))
    out = []
    for cell in current_cells:
        if cell is None or p_idx >= len(cell.paragraphs):
            continue
        out.append(cell.paragraphs[p_idx])
    return out


def _resolve_docx_element(doc, location: str):
    """兼容 API: 返回第一个匹配(无则 None)."""
    lst = _resolve_docx_elements(doc, location)
    return lst[0] if lst else None


def _set_paragraph_text(para, new_text: str) -> None:
    """在保留 run 格式的前提下,把 paragraph 文本替换为 new_text.

    策略:第一个 run 写全部文本,其余 run 清空。若 para 完全无 runs 则 add_run。
    """
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    # 若 add_run 在父结构存在 field/hyperlink 干扰,保险地用 _p 检查一次
    if not para.text:
        # runs 写入未成功(罕见 case),追加 run 兜底
        para.add_run(new_text)


# ────────────────────────────────────────────────────────────
# 预处理:section 分组 / body-gap 检测(Step 5 实装,此处仅留钩子)
# ────────────────────────────────────────────────────────────

def _index_elements_by_location(elements: list[Element]) -> dict[str, Element]:
    return {e.location: e for e in elements}


def _group_rewrite_by_section(
    elements: list[Element],
    classifications: dict[str, Classification],
    section_by_loc: dict[str, list[str]],
) -> dict[tuple[str, ...], list[Element]]:
    """将 op==REWRITE 的 element 按 section heading 栈分组,准备合批处理."""
    groups: dict[tuple[str, ...], list[Element]] = {}
    for e in elements:
        cls = classifications.get(e.location)
        if not cls or cls.op != "REWRITE":
            continue
        sec_key = tuple(section_by_loc.get(e.location, []))
        groups.setdefault(sec_key, []).append(e)
    return groups


_HEADER_SCAFFOLD_RE = re.compile(r"[：:]\s*$")
_SECTION_NUM_RE = re.compile(
    r"^\s*(?:[一二三四五六七八九十]+、|[（(][一二三四五六七八九十0-9]{1,2}[)）]|[0-9]{1,2}[．\.、](?!\d))"
)


def _is_header_scaffold(text: str) -> bool:
    """判定文本是否为'章节/字段标题'型 SCAFFOLD.

    硬约束:
      - 总长度 < 40 字符(避免把长叙述句误判为 header)
      - 以 : 或 ：结尾,或有小节编号前缀(1. / （1） / 一、)
      - 不含句号、分号(避免多句混杂的描述段)
    """
    if not text:
        return False
    t = text.strip()
    if len(t) >= 40:
        return False
    if "。" in t or "；" in t or ";" in t:
        return False
    if _HEADER_SCAFFOLD_RE.search(t):
        return True
    if _SECTION_NUM_RE.search(t):
        return True
    return False


def _detect_body_gaps(
    elements: list[Element],
    classifications: dict[str, Classification],
) -> dict[str, dict[str, Any]]:
    """识别"只有 header、没有 body"的 SCAFFOLD 元素.

    规则:
      - 对每个 PRESERVE/SCAFFOLD 且 _is_header_scaffold 的元素
      - 取同一 source 内下一个 classified 元素
      - 若下一元素也是 PRESERVE/SCAFFOLD(header 型) → 当前元素无 body → 标 gap
      - 若下一元素是 FILL/REWRITE 或非 header-SCAFFOLD → 认为有后续内容,不标 gap

    Returns: {location: {"reason": ..., "header_text": ..., "next_loc": ...}}
    """
    gaps: dict[str, dict[str, Any]] = {}
    # 按 source 分组保序
    from collections import defaultdict
    by_source: dict[str, list[Element]] = defaultdict(list)
    for e in elements:
        by_source[e.source].append(e)

    # 检测窗口:当前 scaffold 后续 5 个同 source 元素
    WINDOW = 5
    BODY_MIN_CHARS = 30

    for src_elems in by_source.values():
        for i, e in enumerate(src_elems):
            cls = classifications.get(e.location)
            if not cls or cls.op != "PRESERVE" or cls.label != "SCAFFOLD":
                continue
            text = (e.text or "").strip()
            if not _is_header_scaffold(text):
                continue
            # 扫描后续窗口:若有任一元素是"body-like"(非 header、非空、≥阈值)→ 非 gap
            found_body = False
            last_checked_loc = None
            for j in range(i + 1, min(i + 1 + WINDOW, len(src_elems))):
                nxt = src_elems[j]
                last_checked_loc = nxt.location
                nxt_text = (nxt.text or "").strip()
                if not nxt_text:
                    continue
                if _is_header_scaffold(nxt_text):
                    continue
                # 长 body 或 非 SCAFFOLD 分类 → 认为后续有正文
                nxt_cls = classifications.get(nxt.location)
                is_non_scaffold = (
                    nxt_cls is not None
                    and not (nxt_cls.op == "PRESERVE" and nxt_cls.label == "SCAFFOLD")
                )
                if len(nxt_text) >= BODY_MIN_CHARS or is_non_scaffold:
                    found_body = True
                    break
            if not found_body:
                gaps[e.location] = {
                    "reason": "SCAFFOLD 标题后窗口内无 body 内容",
                    "header_text": text,
                    "next_loc": last_checked_loc,
                }
    return gaps


# ────────────────────────────────────────────────────────────
# 核心:用 handler 跑完整流程
# ────────────────────────────────────────────────────────────

def _finalize_doc_hard_facts(doc, facts: dict) -> None:
    """doc-level post-process: 替换 PRESERVE 文本里的占位符 & 补遗漏硬字段.

    REWRITE 路径已在 section_batch_rewrite 里调用 _inject_hard_facts;
    但对 PRESERVE/SCAFFOLD 保留的模板文本(如 `XX事务所`/`XX年报`),
    需要在 doc 级别做一次确定性替换,否则审贷员收到的是未填占位符。
    """
    if not facts:
        return
    auditor = facts.get("auditor")
    bank_account = facts.get("bank_account")

    all_text = []
    for p in doc.paragraphs:
        if p.text.strip():
            all_text.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        all_text.append(p.text)
    joined_text = "\n".join(all_text)

    def _rewrite_para(para) -> None:
        txt = para.text
        if not txt:
            return
        new = txt
        if auditor:
            new = re.sub(r"XX会计师事务所", f"{auditor}会计师事务所", new)
            new = re.sub(r"XX事务所", f"{auditor}会计师事务所", new)
        if bank_account and bank_account not in joined_text:
            # 仅在本 paragraph 含 '基本户'/'开户许可证'/'基本账户' 时追加一次
            if re.search(r"(开户许可证|基本户|基本账户)", new) and bank_account not in new:
                if new.endswith("。"):
                    new = new[:-1] + f"(基本账户:{bank_account})。"
                else:
                    new = new + f"(基本账户:{bank_account})"
        if new != txt:
            _set_paragraph_text(para, new)

    for p in doc.paragraphs:
        _rewrite_para(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _rewrite_para(p)


def apply_to_docx(doc, results: list["GenResult"]) -> dict[str, int]:
    """把 handler 产物写回 doc.

    每个 GenResult 的 location 可能对应多个 paragraph(NT 共用 location 的场景),
    此时对所有匹配 paragraph 都应用同一结果(语义等价的子表统一处理).
    Returns: {"fill": n, "keep": n, "clear": n, "miss": n} 统计
    """
    stats = {"fill": 0, "keep": 0, "clear": 0, "miss": 0}
    for r in results:
        if r.action == "keep":
            stats["keep"] += 1
            continue
        paras = _resolve_docx_elements(doc, r.location)
        if not paras:
            stats["miss"] += 1
            continue
        for para in paras:
            if r.action == "clear":
                _set_paragraph_text(para, "")
                stats["clear"] += 1
            elif r.action == "fill":
                _set_paragraph_text(para, r.new_text or "")
                stats["fill"] += 1
    return stats


def generate(
    classified_json: Path,
    template_docx: Path,
    material_dir: Path | None,
    output_docx: Path,
    client_metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """V16 主入口:按 classifier 路由处理所有 element,写回 docx.

    Step 2 范围: PRESERVE/*、FILL/FILL、FILL/CLEAR、FILL/CHECKBOX 走确定性 handler;
                 FILL/SLOT、REWRITE/REWRITE 暂走 _not_impl(保留原状).

    2026-05-21 治本 Phase 1:
      - client_metadata 可选 dict (key 见 templates/placeholder-schema.json) —
        驱动 REPLACE op handler 替换 `{{KEY}}` placeholder.
      - client_metadata=None → 走旧路径 (CURRENT_CLIENT hardcode + 老 docx 模板)
        不 break 现有 5 个 sample fixture
      - 若模板含 placeholder 但 client_metadata 缺对应 key,该 placeholder 整段保留
        原文 + 写 pending_tag (符合 "幻觉零容忍 · 字段填不了标 未能自动填写")
    """
    # 延迟 import handler(它依赖 v16_generator 的 GenResult)
    from v16_op_handlers import dispatch, section_batch_rewrite

    print(f"[v16_generator.generate] {template_docx.name}")
    classifications = load_classifier_output(classified_json)
    doc, elements, section_by_loc = load_template(template_docx)
    mats = load_materials(material_dir)

    # 2026-05-21 治本 Phase 1: client_metadata 优先级 — 显式参数 > sidecar > None
    if client_metadata is None:
        sidecar = template_docx.with_suffix(".metadata.json")
        if sidecar.is_file():
            try:
                _data = json.loads(sidecar.read_text(encoding="utf-8"))
                # original_client 用作 metadata 兜底值的 source (Phase 2 经纬测绘 sidecar 结构)
                # 真实生产路径 client_metadata 应该从 credit handoff payload 来
                if isinstance(_data, dict):
                    client_metadata = _data.get("client_metadata") or _data.get("original_client")
                    if client_metadata is not None and not isinstance(client_metadata, dict):
                        client_metadata = None
            except (OSError, ValueError):
                client_metadata = None
    mats.client_metadata = client_metadata
    if client_metadata:
        print(f"  client_metadata: {len(client_metadata)} 字段 "
              f"(client={client_metadata.get('CLIENT_FULL_NAME') or client_metadata.get('name', '?')})")
    else:
        print("  client_metadata: None (走旧路径 · placeholder 保留原文)")
    # classifier 是按 sampled elements 产出的,面对新模板会有大批 element
    # 没有分类落位(location 不匹配)。兜底策略:
    #   - 超短(<4 ch) 或纯符号/分隔:PRESERVE/PRESERVE (原样保留)
    #   - 其余含正文的 element: REWRITE/REWRITE (LLM + facts 重写)
    # 目的:避免 stale classified.json 让主要正文 element 被整段跳过。
    # stale classified.json 产物对新模板定位不命中,需要补齐未分类 element。
    # 启发:
    #   - 超短(<4 ch):PRESERVE
    #   - 含大量 colon/占位结构 ("XX:"/"___"/"XX、XX、XX") 且整体较长:REWRITE 重写
    #   - 其他(短标题/空白):PRESERVE,保底不改动原文
    # 目的:保证 R1C0 这类"字段密集的大文本块"进入 REWRITE,同时不把所有
    # 段落都扫进来(避免 REWRITE 批次过大导致 LLM JSON 截断)。
    pre_count = len(classifications)
    from v16_op_handlers import FIELD_TO_KB_KEY
    import re as _re
    # "XX:" 或 "XX：" 结尾的字段标签(单槽或多槽)
    _label_pat = _re.compile(r"^([\u4e00-\u9fff（）()0-9A-Za-z、/ -]{2,40})\s*[:：]\s*$")
    fb_stats = {"preserve": 0, "fill": 0, "slot": 0, "rewrite": 0, "upgrade": 0, "placeholder": 0}

    # ───────────────────────────────────────────────────────────────
    # 治本 · 全量 deterministic placeholder pre-pass (2026-05-21 codex R1 抓的根因 3 闭环)
    # ───────────────────────────────────────────────────────────────
    # 问题: classifier 只 sample 235/1023 element,未采样的含 {{KEY}} element 会被
    #       下面 fallback 路径错路由到 PRESERVE,原文输出字面 "{{CLIENT_FULL_NAME}}".
    # 治本: deterministic regex 扫全量 element,含 {{KEY}} 立即标 REPLACE/PLACEHOLDER
    #       (不依赖 LLM / sampled flag · 强制覆盖已分类即使 classifier 误分 PRESERVE).
    # 2026-05-21 红线 fix · 允许 KEY 中含数字 (如 RELATED_PARTY_1_FULL_NAME) ·
    # 原 [A-Z_]+ regex 漏掉 · 致 LLM REWRITE 漏字面 {{RELATED_PARTY_1_*}} 进 docx.
    # 与 v16_op_handlers._PLACEHOLDER_KEY_RE 对齐.
    _placeholder_re = _re.compile(r"\{\{([A-Z][A-Z0-9_]{1,40})\}\}")
    for e in elements:
        text = (e.text or "")
        if _placeholder_re.search(text):
            classifications[e.location] = Classification(
                location=e.location, op="REPLACE", label="PLACEHOLDER",
                confidence=1.0,
                justification="deterministic placeholder pre-pass · 全量 regex 识别 {{KEY}}"
            )
            fb_stats["placeholder"] += 1

    # ───────────────────────────────────────────────────────────────
    # 治本 · 全量 deterministic financial pre-pass (2026-05-22 agent21 抓的真根因)
    # ───────────────────────────────────────────────────────────────
    # 问题: classifier 把财务分析段 (营业收入/资产负债率/经营活动现金流/应收账款 + 具体数字)
    #       误标 PRESERVE → handler 直接保留原文 → 经纬模板财务数字渗进 docx ·
    #       是 user 第 3+4 次 dogfood 看到"分析逻辑还是经纬度"的真根因.
    # 治本: deterministic regex 扫全量 element · 含【财务关键词 + 具体数字】立即标 REWRITE
    #       走 LLM 改写 (LLM 已有 raw material 剔除数字 + financial_metrics 真值 block).
    # 强制覆盖 classifier 误分的 PRESERVE/SCAFFOLD/FILL · 不让经纬数字漏过.
    _FINANCIAL_KEYWORDS = (
        "营业收入", "总资产", "总负债", "所有者权益", "应收账款", "应付账款",
        "存货", "净利润", "毛利", "资产负债率", "流动比率", "速动比率",
        "经营活动", "投资活动", "筹资活动", "现金流", "毛利率", "净利率",
        "周转", "应收", "应付", "净资产收益率", "ROA", "ROE",
        "财务费用", "营业成本", "管理费用", "销售费用",
    )
    # 2026-05-22 user dogfood 5th: 扩 client-specific 关键词 catch 业务背景/历史/关联方/地名段
    _NARRATIVE_CLIENT_SPECIFIC_KEYWORDS = (
        # 业务/行业描述类
        "测绘", "地理信息", "招标", "投标", "工程咨询", "信托", "资产管理",
        "国信", "兴业", "经纬", "六一八", "三明",
        # 政府/公司类专有 (经纬模板里出现的)
        "省经贸委", "省政府", "省管国有", "拟上市", "上市子公司",
        # 历史时间 specific (经纬模板里的)
        "1988年", "2010年", "2017年", "2019年", "2020年", "2021年", "2022年", "2023年",
        # 人名/具体身份 (经纬模板里的)
        "郑志煌", "谢斌",
        # 地名 specific
        "福州市", "福建省", "马尾区", "仓山区", "厦门",
    )
    _NUM_PATTERN = _re.compile(r"[\d,]+(?:\.\d+)?\s*(?:万元|亿元|百万元|千万元|%|‰|个百分点)")
    _financial_skipped_keys = ("PLACEHOLDER",)  # 已 placeholder 化的不动
    for e in elements:
        text = (e.text or "")
        if not text or len(text) < 30:
            continue
        # 已是 PLACEHOLDER (pre-pass 优先) 跳过 · 走 REPLACE 路径不冲突
        existing = classifications.get(e.location)
        if existing and existing.label in _financial_skipped_keys:
            continue
        # 含财务关键词 + 含具体数字 = 财务分析段 · force REWRITE
        has_fin_kw = any(kw in text for kw in _FINANCIAL_KEYWORDS)
        has_num = bool(_NUM_PATTERN.search(text))
        if has_fin_kw and has_num:
            classifications[e.location] = Classification(
                location=e.location, op="REWRITE", label="FINANCIAL",
                confidence=1.0,
                justification="deterministic financial pre-pass · 强制 REWRITE 防经纬数字 PRESERVE 残留"
            )
            fb_stats["rewrite"] += 1
            continue
        # 2026-05-22 user dogfood 5th: 含 client-specific narrative 关键词 (经纬业务/历史/股权/地名/人名)
        # = 经纬原模板的 client-specific 叙述段 · force REWRITE 防 PRESERVE 漏过
        has_narrative_kw = any(kw in text for kw in _NARRATIVE_CLIENT_SPECIFIC_KEYWORDS)
        if has_narrative_kw and len(text) >= 50:
            classifications[e.location] = Classification(
                location=e.location, op="REWRITE", label="FINANCIAL",
                confidence=1.0,
                justification="deterministic narrative pre-pass · 含经纬模板 client-specific 关键词 · 强制 REWRITE"
            )
            fb_stats["rewrite"] += 1

    # Pre-pass: 现有 PRESERVE/SCAFFOLD 若文本是已知字段标签(如"客户名称：")且 KB
    # 有对应值,上调为 FILL/FILL。分类器会把只写了标签的短段判为 SCAFFOLD,
    # 错失填充机会;这里做事实驱动的补救。
    for e in elements:
        existing = classifications.get(e.location)
        if existing is None:
            continue
        if existing.op != "PRESERVE" or existing.label != "SCAFFOLD":
            continue
        t = (e.text or "").strip()
        m = _label_pat.match(t)
        if not m:
            continue
        label_name = m.group(1).strip()
        stripped_label = _re.sub(r"[（(][^）)]*[）)]", "", label_name).strip()
        if label_name in FIELD_TO_KB_KEY or stripped_label in FIELD_TO_KB_KEY:
            classifications[e.location] = Classification(
                location=e.location, op="FILL", label="FILL",
                confidence=0.0,
                justification=f"upgrade: SCAFFOLD→FILL for known label {label_name}"
            )
            fb_stats["upgrade"] += 1

    for e in elements:
        if e.location in classifications:
            continue
        t = (e.text or "").strip()
        if not t or len(t) < 2:
            classifications[e.location] = Classification(
                location=e.location, op="PRESERVE", label="PRESERVE",
                confidence=0.0, justification="fallback: empty"
            )
            fb_stats["preserve"] += 1
            continue
        # 识别"字段密集型"大文本块:至少 3 个中文冒号/全角占位 + 长度 >= 80 字
        has_dense_fields = (
            (t.count("：") + t.count(":")) >= 3 or
            "□" in t or "☑" in t
        ) and len(t) >= 80
        if has_dense_fields:
            classifications[e.location] = Classification(
                location=e.location, op="REWRITE", label="REWRITE",
                confidence=0.0, justification="fallback: dense-field block"
            )
            fb_stats["rewrite"] += 1
            continue
        # "XX:" / "XX：" 结尾的短标签:优先按 FILL 路由(FIELD_TO_KB_KEY 决定是否命中)
        m = _label_pat.match(t)
        if m:
            label_name = m.group(1).strip()
            stripped_label = _re.sub(r"[（(][^）)]*[）)]", "", label_name).strip()
            colon_count = t.count("：") + t.count(":")
            if colon_count >= 2:
                classifications[e.location] = Classification(
                    location=e.location, op="FILL", label="SLOT",
                    confidence=0.0, justification="fallback: multi-slot label"
                )
                fb_stats["slot"] += 1
            elif label_name in FIELD_TO_KB_KEY or stripped_label in FIELD_TO_KB_KEY:
                classifications[e.location] = Classification(
                    location=e.location, op="FILL", label="FILL",
                    confidence=0.0, justification=f"fallback: FILL label {label_name}"
                )
                fb_stats["fill"] += 1
            else:
                classifications[e.location] = Classification(
                    location=e.location, op="PRESERVE", label="PRESERVE",
                    confidence=0.0, justification=f"fallback: unknown label {label_name}"
                )
                fb_stats["preserve"] += 1
            continue
        # 其他短文本 / 标题 / 散文片段:PRESERVE 兜底
        classifications[e.location] = Classification(
            location=e.location, op="PRESERVE", label="PRESERVE",
            confidence=0.0, justification="fallback: preserve default"
        )
        fb_stats["preserve"] += 1
    if len(classifications) > pre_count:
        print(f"  [fallback] 补齐 {len(classifications) - pre_count} 个未分类 element "
              f"({pre_count} 已分类 → {len(classifications)}): {fb_stats}")

    # ───────────────────────────────────────────────────────────────
    # 治本 · 长段 pre-pass (2026-05-22 user dogfood 5th · agent diag 实测)
    # ───────────────────────────────────────────────────────────────
    # 必须放在 fallback pre-pass 之后 · 否则未分类 element 没 entry · long-para 跳过.
    # 真根因: classifier 把"经纬模板的财务/业务/授信品种叙述段"标 PRESERVE,
    #         这些段不含 financial keyword 也不含 narrative keyword (业务通用描述),
    #         但**就是经纬原模板逐字保留** → 报告主体跟原模板高相似度 (e2e 结构断言 FAIL).
    # 实测样本 (corp-dingsheng 跑通后 dump):
    #   - P127 (121): "其他应付款及其他流动负债：公司其他应付款主要为关联方往来款..."
    #   - P115 (105): "负债结构分析：公司负债主要分布在短期借款、应付账款..."
    #   - P149 (88): "投资活动产生的现金流量受公司当期资本性支出节奏..."
    #   - P122 (83): "存货：公司存货整体呈一定波动趋势..."
    #   - P188 (89): "（2）商票保贴：由于企业的下游客户多为大型企业..."
    #   - P190/P191/P192: "国内证及买方押汇 / 非融资性保函 / 快捷保理" 授信品种叙述
    # 治本: 长段 PRESERVE (≥ 50 char) 一律 force REWRITE · 让 LLM 改写为客户 specific.
    # 排除 _is_obvious_generic_template (table header / checkbox 模板 / 章节号 / 落款)
    # 避免 force LLM 编 nonsense 标题.
    _LONG_PARA_THRESHOLD = 50

    def _is_obvious_generic_template(t: str) -> bool:
        """识别"不该 force REWRITE"的 generic template 段 (header / 章节号 / checkbox / 落款)."""
        if not t:
            return True
        stripped = t.strip()
        if len(stripped) < _LONG_PARA_THRESHOLD:
            return True
        # checkbox 模板 (含 □ 或 ☑ 或 "是  否" 表格 header)
        if "□" in stripped or "☑" in stripped:
            return True
        # 落款行 (含多个 ":" / "：" 但每个后面跟空格 · 没实质内容)
        # e.g. "客户经理:              总经理助理:              部门负责人:"
        if stripped.count("：") + stripped.count(":") >= 3:
            # 且总字数不多 (落款典型 < 80) 且没句号
            if len(stripped) < 80 and "。" not in stripped and "，" not in stripped:
                return True
        # 章节号 / 标题 (短 + 编号前缀 + 无句号)
        if _SECTION_NUM_RE.search(stripped) and "。" not in stripped and len(stripped) < 80:
            return True
        # 全 placeholder ({{KEY}}{{KEY}}...) — 已 PLACEHOLDER 化的 (pre-pass 已经覆盖,这里双保险)
        if _placeholder_re.search(stripped):
            return True
        # 全数字 / 全标点 / 表格数据行
        non_punct = _re.sub(r"[\s\d\.,，。;；:：\-—、()（）%]+", "", stripped)
        if len(non_punct) < 5:
            return True
        # 表格 header 枚举型 (含"系统/网站/类目"逗号枚举 + 短 + 无句号)
        # e.g. "企查查、天眼查、启信宝、裁判文书网、中国执行网、信用中国、各地市环保局网站..."
        comma_count = stripped.count("、")
        if comma_count >= 3 and "。" not in stripped and len(stripped) < 80:
            return True
        return False

    long_para_forced = 0
    for e in elements:
        text = (e.text or "").strip()
        if len(text) < _LONG_PARA_THRESHOLD:
            continue
        existing = classifications.get(e.location)
        if existing is None:
            continue
        if existing.op != "PRESERVE":
            continue
        if _is_obvious_generic_template(text):
            continue
        # force REWRITE · 用 FINANCIAL label 走 section_batch_rewrite 路径
        classifications[e.location] = Classification(
            location=e.location, op="REWRITE", label="FINANCIAL",
            confidence=1.0,
            justification=(
                f"long-para pre-pass · 长段 PRESERVE (len={len(text)}) "
                f"force REWRITE 防经纬模板原文残留"
            ),
        )
        fb_stats["rewrite"] += 1
        long_para_forced += 1
    if long_para_forced:
        print(f"  [long-para pre-pass] 强制 {long_para_forced} 长段 PRESERVE → REWRITE/FINANCIAL")

    print(f"  classifier: {len(classifications)} loc / template: {len(elements)} elems"
          f" / materials: {len(mats.file_contents)} files")
    if mats.kb:
        print(f"  KB facts: {len(mats.facts)}")

    # Step 5: body-gap 预处理,结果挂到 mats 供 handler 读取
    mats.body_gaps = _detect_body_gaps(elements, classifications)
    if mats.body_gaps:
        print(f"  body-gap 检测: {len(mats.body_gaps)} 个 SCAFFOLD 缺 body")

    # Step 4: 先对 REWRITE 按 section 合批,结果登记到 rewrite_results
    rewrite_groups = _group_rewrite_by_section(elements, classifications, section_by_loc)
    rewrite_results: dict[str, GenResult] = {}
    if rewrite_groups:
        print(f"  REWRITE 合批: {len(rewrite_groups)} 个 section,"
              f"{sum(len(v) for v in rewrite_groups.values())} 个段落")
        for sec_key, sec_elems in rewrite_groups.items():
            sec_heading = list(sec_key)
            out = section_batch_rewrite(sec_elems, sec_heading, mats)
            rewrite_results.update(out)

    # 对每个有分类的 element 调 dispatch(REWRITE 走 rewrite_results)
    results: list[GenResult] = []
    pending_tags: list[dict[str, Any]] = []
    handler_stats: dict[str, int] = {}

    for elem in elements:
        cls = classifications.get(elem.location)
        if cls is None:
            continue
        if cls.op == "REWRITE" and elem.location in rewrite_results:
            r = rewrite_results[elem.location]
        else:
            r = dispatch(elem, cls, mats)
        results.append(r)
        if isinstance(r.pending_tag, list):
            pending_tags.extend(r.pending_tag)
        elif r.pending_tag:
            pending_tags.append(r.pending_tag)
        key = f"{cls.op}/{cls.label}"
        handler_stats[key] = handler_stats.get(key, 0) + 1

    apply_stats = apply_to_docx(doc, results)
    _finalize_doc_hard_facts(doc, mats.facts if mats.kb else {})
    doc.save(str(output_docx))

    # 导出 pending tags
    pending_path = output_docx.with_name(output_docx.stem + "_pending.json")
    pending_path.write_text(
        json.dumps({"pending_tags": pending_tags}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print(f"  dispatch 分布: {dict(sorted(handler_stats.items()))}")
    print(f"  apply 统计: {apply_stats}")
    print(f"  pending tags: {len(pending_tags)}")
    print(f"  └ output: {output_docx}")
    print(f"  └ pending: {pending_path}")

    return {
        "handler_stats": handler_stats,
        "apply_stats": apply_stats,
        "pending_count": len(pending_tags),
        "output_docx": str(output_docx),
        "pending_json": str(pending_path),
    }


# ────────────────────────────────────────────────────────────
# Dry-run:验证 classifier 输出与 docx element round-trip
# ────────────────────────────────────────────────────────────

def dry_run(
    classified_json: Path,
    template_docx: Path,
    material_dir: Path | None = None,
    max_print: int = 40,
) -> dict[str, Any]:
    """把 classifier 输出与 docx element 对齐一次,打印前 N 条.

    返回统计:命中率、op/label 分布、section 分组数.
    """
    print(f"[v16_generator.dry_run] classified_json={classified_json.name}")
    print(f"                        template_docx ={template_docx.name}")
    classifications = load_classifier_output(classified_json)
    doc, elements, section_by_loc = load_template(template_docx)
    mats = load_materials(material_dir)
    print(f"  classifications: {len(classifications)}")
    print(f"  template elements: {len(elements)}")
    print(f"  material files loaded: {len(mats.file_contents)}")

    # 命中统计
    el_by_loc = _index_elements_by_location(elements)
    cls_locs = set(classifications.keys())
    elem_locs = set(el_by_loc.keys())
    hit = cls_locs & elem_locs
    only_cls = cls_locs - elem_locs
    only_elem = elem_locs - cls_locs
    print(f"  ├ loc 命中:  {len(hit)} / {len(cls_locs)} = "
          f"{(len(hit) / max(len(cls_locs), 1)) * 100:.1f}%")
    print(f"  ├ 只在 classifier 出现: {len(only_cls)} (跨文档的 element)")
    print(f"  └ 只在 template 出现:   {len(only_elem)} (classifier 未覆盖)")

    # op / label 分布
    op_dist: dict[str, int] = {}
    lab_dist: dict[str, int] = {}
    for loc in hit:
        c = classifications[loc]
        op_dist[c.op] = op_dist.get(c.op, 0) + 1
        lab_dist[c.label] = lab_dist.get(c.label, 0) + 1
    print(f"  op 分布:    {dict(sorted(op_dist.items()))}")
    print(f"  label 分布: {dict(sorted(lab_dist.items()))}")

    # REWRITE 分组
    rw_groups = _group_rewrite_by_section(elements, classifications, section_by_loc)
    print(f"  REWRITE 分组: {len(rw_groups)} 个 section")
    for i, (sec, items) in enumerate(list(rw_groups.items())[:5]):
        head = sec[-1] if sec else "<root>"
        print(f"    └ [{i+1}] {head[:30]}  ({len(items)} elems)")

    # 逐条打印前 N
    print(f"\n  [sample] 前 {max_print} 条 (location, op, label, text):")
    print("  " + "-" * 100)
    count = 0
    for e in elements:
        if e.location not in classifications:
            continue
        c = classifications[e.location]
        text_preview = (e.text or "").replace("\n", " ")[:50]
        print(f"    {e.location:20}  {c.op:8}/{c.label:9}  {text_preview}")
        count += 1
        if count >= max_print:
            break

    return {
        "classifications": len(classifications),
        "elements": len(elements),
        "hit": len(hit),
        "only_cls": len(only_cls),
        "only_elem": len(only_elem),
        "op_dist": op_dist,
        "label_dist": lab_dist,
        "rewrite_groups": len(rw_groups),
    }


# ────────────────────────────────────────────────────────────
# Entry
# ────────────────────────────────────────────────────────────

def _pick_template_for(docx_name: str, samples_dir: Path) -> Path:
    return samples_dir / docx_name


def main():
    """入口:

        py -u v16_generator.py                       # dry-run 3 samples
        py -u v16_generator.py 经纬测绘_对公成稿A      # dry-run 特定 sample
        py -u v16_generator.py --run                  # 生成 3 docx
        py -u v16_generator.py --run 普惠             # 生成特定 sample

    --run 下若提供 --material=DIR,会加载 KB;否则 KB 为空,FILL/FILL 全部 pending。
    """
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

    samples_dir = Path(__file__).parent / "samples"
    if not DEFAULT_CLASSIFIED_JSON.exists():
        print(f"[err] classifier 输出不存在: {DEFAULT_CLASSIFIED_JSON}")
        sys.exit(1)

    do_run = "--run" in sys.argv
    material_dir = None
    for a in sys.argv:
        if a.startswith("--material="):
            material_dir = Path(a.split("=", 1)[1])

    key_args = [a for a in sys.argv[1:] if not a.startswith("--")]
    candidates = sorted(samples_dir.glob("*.docx"))
    if key_args:
        candidates = [p for p in candidates if any(k in p.stem for k in key_args)]

    for docx in candidates:
        print("\n" + "=" * 80)
        print(f"Template: {docx.name}")
        print("=" * 80)
        if do_run:
            out = OUTPUT_DIR / f"{docx.stem}_v16.docx"
            generate(
                classified_json=DEFAULT_CLASSIFIED_JSON,
                template_docx=docx,
                material_dir=material_dir,
                output_docx=out,
            )
        else:
            dry_run(
                classified_json=DEFAULT_CLASSIFIED_JSON,
                template_docx=docx,
                material_dir=material_dir,
                max_print=30,
            )


if __name__ == "__main__":
    main()
