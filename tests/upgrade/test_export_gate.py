# -*- coding: utf-8 -*-
import io
import json
import zipfile
from pathlib import Path

import pytest
from fastapi.testclient import TestClient
from docx import Document
from pypdf import PdfReader

from agent_report.api import app
from agent_report.v16_runner import _load_pending_questions
from agent_report.word_export import _quality_gate_reasons
from agent_report.session_store import store
from auth_service.dependencies import COOKIE_NAME
from auth_service.jwt_util import issue
from v16_generator import generate, load_template


WATERMARK = "质量闸未过 · 内部草稿 · 不得作为审批依据"


@pytest.fixture
def client():
    test_client = TestClient(app)
    test_client.cookies.set(COOKIE_NAME, issue("u_test", "admin"))
    return test_client


_MISSING = object()


def _session(qc=_MISSING) -> str:
    done_payload = {
        "profile": {"company_name": "闸门测试企业"},
        "sections": [{"id": "s1", "title": "正文", "content": "可信会话正文"}],
        "stats": {},
    }
    if qc is not _MISSING:
        done_payload["qc"] = qc
    return store.create({
        "enterprise_profile": {"company_name": "闸门测试企业"},
        "pending_questions": [],
        "done_payload": done_payload,
    })


def _docx_xml(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        names = [name for name in archive.namelist() if name.startswith("word/header")]
        return archive.read("word/document.xml").decode("utf-8") + "".join(
            archive.read(name).decode("utf-8") for name in names
        )


def _pdf_text(data: bytes) -> str:
    return "".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(data)).pages)


def _docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join([*(p.text for p in doc.paragraphs), *cells])


def test_client_forged_qc_cannot_bypass_blocked_session(client):
    sid = _session({
        "passed": False,
        "issues": [{"severity": "block", "message": "关键字段缺少证据"}],
    })
    try:
        response = client.post("/api/report/export_docx", json={
            "session_id": sid,
            "qc": {"passed": True},
            "profile": {"company_name": "伪造企业"},
            "sections": [{"title": "伪造正文", "content": "伪造"}],
        })
        assert response.status_code == 200, response.text
        xml = _docx_xml(response.content)
        assert WATERMARK in xml
        assert "关键字段缺少证据" in xml
        assert "可信会话正文" in xml
        assert "伪造正文" not in xml
    finally:
        store.delete(sid)


def test_passed_session_exports_without_watermark(client):
    sid = _session({"passed": True, "issues": []})
    try:
        response = client.post("/api/report/export_docx", json={"session_id": sid})
        assert response.status_code == 200, response.text
        assert WATERMARK not in _docx_xml(response.content)
    finally:
        store.delete(sid)


def test_pending_questions_dict_is_accepted_and_normalized(client):
    sid = _session({"passed": True, "issues": []})
    try:
        response = client.post("/api/report/export_docx", json={
            "session_id": sid,
            "pending_questions": {
                "missing_tax": {"id": "missing_tax", "label": "补充纳税材料"},
            },
        })
        assert response.status_code == 200, response.text
        assert "补充纳税材料" in _docx_xml(response.content)
    finally:
        store.delete(sid)


def test_pending_tags_real_generator_to_api_export_docx_end_to_end(tmp_path, client):
    template = tmp_path / "minimal.docx"
    doc = Document()
    doc.add_paragraph("员工人数：XXX人；注册资本：XXX万元")
    doc.save(template)
    _, elements, _ = load_template(template)
    target_locations = [element.location for element in elements if "XXX" in element.text]
    assert len(target_locations) == 1

    classified = tmp_path / "classified.json"
    classified.write_text(json.dumps({"classifications_by_location": {
        location: {
            "op": "FILL", "label": "SLOT", "confidence": 1.0,
            "justification": "R3 real offline multi-slot pending chain",
        }
        for location in target_locations
    }}, ensure_ascii=False), encoding="utf-8")
    output = tmp_path / "generated.docx"
    empty_materials = tmp_path / "materials"
    empty_materials.mkdir()
    summary = generate(classified, template, empty_materials, output)

    pending_payload = json.loads(
        Path(summary["pending_json"]).read_text(encoding="utf-8")
    )
    raw_pending = pending_payload["pending_tags"]
    assert isinstance(raw_pending, list)
    assert len(raw_pending) >= 2
    assert all(isinstance(item, dict) for item in raw_pending)
    assert not any(isinstance(item, list) for item in raw_pending)

    pending = _load_pending_questions(summary["pending_json"])
    assert pending
    assert all(isinstance(item, dict) for item in pending)
    for item in pending:
        assert item.get("reason")
        assert item.get("location")
        assert item.get("suggested_action")

    sid = store.create({
        "enterprise_profile": {"company_name": "真实 pending 企业"},
        "pending_questions": pending,
        "done_payload": {
            "profile": {"company_name": "真实 pending 企业"},
            "sections": [{"id": "s1", "title": "正文", "content": "生成链正文"}],
            "qc": {"passed": True},
        },
    })
    try:
        response = client.post("/api/report/export_docx", json={"session_id": sid})
        assert response.status_code == 200, response.text
        exported = _docx_text(response.content)
        for item in pending:
            assert item["reason"] in exported
            assert item["location"] in exported
            assert item["suggested_action"] in exported
        assert "缺失原因\n—" not in exported
        assert "建议动作\n—" not in exported
    finally:
        store.delete(sid)


def test_quality_reasons_use_scorer_gates_and_include_unknown_failed_dimension():
    reasons = _quality_gate_reasons({
        "passed": False,
        "fatal_reasons": ["真实一票否决"],
        "dimensions": [
            {"name": "财务分析深度", "raw_score": 4.0, "missed_items": []},
            {"name": "新增审查维度", "raw_score": 2.0, "passed": False,
             "missed_items": ["缺少关键证据"]},
            {"name": "载荷阈值维度", "raw_score": 3.0, "threshold": 4.0,
             "missed_items": []},
            {"name": "无失败标记维度", "raw_score": 1.0, "missed_items": []},
        ],
        "hallucinations": [{"text": "待复核表述", "reason": "缺少来源", "location": "P2"}],
    })
    assert any("实际分 4" in reason and "闸值 7" in reason for reason in reasons)
    assert any("新增审查维度" in reason and "缺少关键证据" in reason for reason in reasons)
    assert any("载荷阈值维度" in reason and "闸值 4" in reason for reason in reasons)
    assert not any("无失败标记维度" in reason for reason in reasons)
    assert not any(reason.startswith(("block:", "warn:")) for reason in reasons)
    assert any("一票否决" in reason for reason in reasons)


def test_pdf_passed_has_no_watermark(client):
    sid = _session({"passed": True, "dimensions": [], "hallucinations": []})
    try:
        response = client.post("/api/report/export_pdf", json={"session_id": sid})
        assert response.status_code == 200, response.text
        assert WATERMARK not in _pdf_text(response.content)
    finally:
        store.delete(sid)


def test_pdf_unchecked_has_watermark_and_fallback_reason(client):
    sid = _session()
    try:
        response = client.post("/api/report/export_pdf", json={"session_id": sid})
        assert response.status_code == 200, response.text
        text = _pdf_text(response.content)
        assert WATERMARK in text
        assert "会话未通过质量检查" in text
    finally:
        store.delete(sid)


def test_pdf_blocked_has_real_quality_scorer_details(client):
    sid = _session({
        "passed": False,
        "fatal_fail": True,
        "fatal_reasons": ["维度「财务分析深度」低于闸值"],
        "dimensions": [{
            "name": "财务分析深度",
            "raw_score": 4.0,
            "missed_items": ["缺少现金流分析"],
        }],
        "hallucinations": [{
            "text": "营业收入 9000 万",
            "reason": "与已知营业收入差异超过 5%",
            "location": "P3",
        }],
    })
    try:
        response = client.post("/api/report/export_pdf", json={"session_id": sid})
        assert response.status_code == 200, response.text
        text = _pdf_text(response.content)
        assert WATERMARK in text
        assert "财务分析深度" in text
        assert "缺少现金流分析" in text
        assert "营业收入" in text
        assert "与已知营业收入差异超过 5%" in text
        assert "会话未通过质量检查" not in text
    finally:
        store.delete(sid)
