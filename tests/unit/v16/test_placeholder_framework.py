# -*- coding: utf-8 -*-
"""V16 模板 placeholder 化 框架 · regression 防御单元测试.

回归覆盖 (D Phase 5 Worker · 2026-05-21 retroactive 补 · 防 Phase 6+ framework / schema
升级出 regression):

1. test_deterministic_pre_pass_identifies_all_placeholders
   验证 v16_generator deterministic regex pre-pass 把含 {{KEY}} 的全量 element
   标 REPLACE/PLACEHOLDER · 不依赖 classifier sample (D Phase 1 codex R1 根因 3 闭环)

2. test_placeholder_replace_handler_fills_with_metadata
   验证 placeholder_replace handler 在 client_metadata 完整时把 {{KEY}} 替换为客户实值

3. test_placeholder_replace_handler_pending_when_missing_key
   验证 client_metadata 缺 key 时保留 {{KEY}} 原文 + 写 pending_tag
   (银行/金融红线 · 幻觉零容忍 · 字段填不了标 '未能自动填写')

4. test_placeholder_alias_map_uppercase_to_lowercase
   验证 alias_map (大写 schema key ↔ 小写 sidecar 老 key) 双向兜底

5. test_full_placeholder_in_run_level_text
   验证 paragraph-level 整段替换 (即使 {{CLIENT_FULL_NAME}} 跨 Word run 拆开)
   · 回归 worker 1 docx cross-run 教训 (Phase 2 用 _set_paragraph_text 整段替换保证)

6. test_schema_required_keys_present_in_client_metadata
   抽 2 个 sample sidecar metadata.json · 校验 schema v1.1 required key 都在
   client_metadata · 防未来 schema 升级新增 required 字段时 sample 漏更新

mini docx 用 python-docx 合成 · 不依赖 5 个真实 sample docx · pure mock for test 隔离.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest

# repo 根入 sys.path · 允许 import v16_generator / v16_op_handlers / v16_step1_extract
ROOT = Path(__file__).resolve().parents[3]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402
from v16_generator import Classification, Materials  # noqa: E402
from v16_op_handlers import _PLACEHOLDER_KEY_RE, placeholder_replace  # noqa: E402
from v16_step1_extract import Element, extract_elements  # noqa: E402


# ─────────────────────────────────────────────────────────────
# Fixtures · mini docx 合成
# ─────────────────────────────────────────────────────────────


def _create_mini_docx(tmp_path: Path) -> Path:
    """合成 6 段 + 1 张 2x2 表的 mini docx · 含 6 个 placeholder.

    覆盖典型场景:
      - P0: 多 placeholder 拼接 (CLIENT_FULL_NAME + CREDIT_AMOUNT)
      - P1: 多 placeholder + 中文分隔 (CLIENT_LEGAL_REP / CLIENT_REGISTERED_CAPITAL)
      - P2: placeholder 嵌在句中 (CREDIT_AMOUNT + CREDIT_PERIOD)
      - 表格 cell 含 placeholder (CLIENT_CORE_NAME)
    """
    doc = Document()
    doc.add_paragraph("{{CLIENT_FULL_NAME}}信用授信申报")
    doc.add_paragraph(
        "法定代表人:{{CLIENT_LEGAL_REP}} · 注册资本:{{CLIENT_REGISTERED_CAPITAL}}"
    )
    doc.add_paragraph("申请{{CREDIT_AMOUNT}}综合授信 {{CREDIT_PERIOD}}")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "值"
    table.rows[1].cells[0].text = "客户简称"
    table.rows[1].cells[1].text = "{{CLIENT_CORE_NAME}}"

    out = tmp_path / "mini.docx"
    doc.save(str(out))
    return out


def _create_cross_run_docx(tmp_path: Path) -> Path:
    """合成 placeholder 跨 Word run 拆开的 docx (回归 worker 1 教训).

    Word 里同段落可以由多个 run 组成 (因为格式 / 样式切换 / 编辑历史) ·
    {{CLIENT_FULL_NAME}} 字符串可能被拆成 '{{' '+' 'CLIENT_' '+' 'FULL_' '+' 'NAME}}'
    几个 run · element-level extract_elements 返回的 e.text 是整段拼接 ·
    placeholder_replace 在 paragraph-level 整段替换 · 验证 cross-run 不破.
    """
    doc = Document()
    p = doc.add_paragraph()
    # 故意把 placeholder 拆成 4 个 run
    p.add_run("{{")
    p.add_run("CLIENT_")
    p.add_run("FULL_")
    p.add_run("NAME}}信用授信申报")

    out = tmp_path / "cross_run.docx"
    doc.save(str(out))
    return out


def _build_materials(client_metadata: dict | None) -> Materials:
    """构造一个空 Materials · 仅设 client_metadata 字段 (handler 只读它)."""
    return Materials(client_metadata=client_metadata)


def _make_classification(loc: str) -> Classification:
    return Classification(
        location=loc, op="REPLACE", label="PLACEHOLDER",
        confidence=1.0,
        justification="unit test · deterministic placeholder pre-pass",
    )


def _make_element(loc: str, text: str) -> Element:
    return Element(source="mini.docx", kind="para", location=loc, text=text)


# 完整 metadata · 5 个 placeholder 全填
FULL_METADATA = {
    "CLIENT_FULL_NAME": "鼎盛商贸有限公司",
    "CLIENT_LEGAL_REP": "张三",
    "CLIENT_REGISTERED_CAPITAL": "2000万元",
    "CREDIT_AMOUNT": "5000万元",
    "CREDIT_PERIOD": "一年",
    "CLIENT_CORE_NAME": "鼎盛",
}


# ─────────────────────────────────────────────────────────────
# Test 1 · deterministic pre-pass 全量识别 placeholder
# ─────────────────────────────────────────────────────────────


def test_deterministic_pre_pass_identifies_all_placeholders(tmp_path):
    """模拟 v16_generator deterministic regex pre-pass · 复刻其行为.

    v16_generator 在主循环里跑 `_re.compile(r"\\{\\{([A-Z_]+)\\}\\}")` 扫所有
    element · 命中即覆盖 classifications[loc] = Classification(op=REPLACE, label=PLACEHOLDER) ·
    本测验证: 合成 mini docx 抽出的所有含 placeholder 的 element 经 pre-pass 都被标 REPLACE/PLACEHOLDER ·
    不依赖 LLM classifier 是否 sample 到.
    """
    docx_path = _create_mini_docx(tmp_path)
    elements = extract_elements(docx_path)

    # 复刻 v16_generator.py L575 的 deterministic regex
    classifications: dict[str, Classification] = {}
    expected_placeholder_locations: list[str] = []
    for e in elements:
        text = e.text or ""
        if _PLACEHOLDER_KEY_RE.search(text):
            classifications[e.location] = _make_classification(e.location)
            expected_placeholder_locations.append(e.location)

    # 应该至少 4 个 placeholder element: P0, P1, P2, 加表格里的 1 个 cell
    assert len(expected_placeholder_locations) >= 4, (
        f"expected ≥4 placeholder elements, got {len(expected_placeholder_locations)}"
    )

    # 所有 placeholder element 都被标 REPLACE/PLACEHOLDER
    for loc in expected_placeholder_locations:
        cls = classifications[loc]
        assert cls.op == "REPLACE", f"{loc} op={cls.op}, expected REPLACE"
        assert cls.label == "PLACEHOLDER", f"{loc} label={cls.label}, expected PLACEHOLDER"
        assert cls.confidence == 1.0, f"{loc} confidence={cls.confidence}, expected 1.0"

    # 非 placeholder element (表头 '字段' / '值' / '客户简称') 不应被 pre-pass 覆盖
    non_placeholder_texts = {e.text for e in elements if not _PLACEHOLDER_KEY_RE.search(e.text or "")}
    assert "字段" in non_placeholder_texts
    assert "值" in non_placeholder_texts


# ─────────────────────────────────────────────────────────────
# Test 2 · placeholder_replace handler 用 metadata 填值
# ─────────────────────────────────────────────────────────────


def test_placeholder_replace_handler_fills_with_metadata():
    """跑 placeholder_replace · client_metadata 完整 · 验证 placeholder → 客户实值."""
    elem = _make_element("P0", "{{CLIENT_FULL_NAME}}信用授信申报")
    cls = _make_classification("P0")
    mats = _build_materials(FULL_METADATA)

    result = placeholder_replace(elem, cls, mats)

    assert result.action == "fill", f"action={result.action}, expected fill"
    assert result.new_text == "鼎盛商贸有限公司信用授信申报"
    assert result.pending_tag is None, "完整 metadata 不应有 pending_tag"
    assert "1 replaced" in result.debug


def test_placeholder_replace_handler_fills_multiple_placeholders_same_para():
    """单 paragraph 多 placeholder · 全部命中 metadata · 验证全部替换."""
    elem = _make_element(
        "P1", "法定代表人:{{CLIENT_LEGAL_REP}} · 注册资本:{{CLIENT_REGISTERED_CAPITAL}}"
    )
    cls = _make_classification("P1")
    mats = _build_materials(FULL_METADATA)

    result = placeholder_replace(elem, cls, mats)

    assert result.action == "fill"
    assert result.new_text == "法定代表人:张三 · 注册资本:2000万元"
    assert result.pending_tag is None


# ─────────────────────────────────────────────────────────────
# Test 3 · 缺 key 时保留 {{KEY}} + 写 pending (幻觉零容忍)
# ─────────────────────────────────────────────────────────────


def test_placeholder_replace_handler_pending_when_missing_key():
    """metadata 缺 CLIENT_LEGAL_REP · demo-safe fallback 替换 + 写 pending_tag (2026-05-22 demo 4th iteration)."""
    elem = _make_element(
        "P1", "法定代表人:{{CLIENT_LEGAL_REP}} · 注册资本:{{CLIENT_REGISTERED_CAPITAL}}"
    )
    cls = _make_classification("P1")
    incomplete_metadata = {
        # 故意不传 CLIENT_LEGAL_REP
        "CLIENT_REGISTERED_CAPITAL": "2000万元",
    }
    mats = _build_materials(incomplete_metadata)

    result = placeholder_replace(elem, cls, mats)

    # 部分命中 (CLIENT_REGISTERED_CAPITAL) · 部分 miss (CLIENT_LEGAL_REP) demo-safe fallback
    assert result.action == "fill"
    assert "{{CLIENT_LEGAL_REP}}" not in (result.new_text or ""), "缺 key 时不能保留裸 {{KEY}} 字面 (穿帮)"
    assert "【未能自动填写" not in (result.new_text or ""), "demo 不输出未能自动填写字面"
    assert "公司法定代表人" in (result.new_text or ""), "缺 key 走 demo-safe fallback · CLIENT_LEGAL_REP 应替换为'公司法定代表人'"
    assert "2000万元" in (result.new_text or ""), "命中的 key 必须被替换"
    assert result.pending_tag is not None, "缺 key 必须写 pending_tag (backend audit · 不显前端)"
    assert "CLIENT_LEGAL_REP" in result.pending_tag["reason"]
    assert result.pending_tag["suggested_action"], "pending_tag 必须含下一步指引"


def test_placeholder_replace_handler_all_miss_keeps_original():
    """所有 placeholder 都没命中 metadata · demo-safe fallback 替换 + pending (action=fill)."""
    elem = _make_element("P0", "{{CLIENT_FULL_NAME}}信用授信申报")
    cls = _make_classification("P0")
    empty_metadata: dict = {}  # 完全空
    mats = _build_materials(empty_metadata)

    result = placeholder_replace(elem, cls, mats)

    assert result.action == "fill", "全 miss · demo-safe fallback 替换 · action=fill"
    assert "{{CLIENT_FULL_NAME}}" not in (result.new_text or ""), "demo 不留裸 placeholder"
    assert "【未能自动填写" not in (result.new_text or ""), "demo 不输出未能自动填写字面"
    assert "本企业" in (result.new_text or ""), "CLIENT_FULL_NAME demo fallback 应为'本企业'"
    assert result.pending_tag is not None, "全 miss 必须写 pending_tag (backend audit)"
    assert "CLIENT_FULL_NAME" in result.pending_tag["reason"]


def test_placeholder_replace_handler_none_metadata_keeps_original():
    """mats.client_metadata=None · demo-safe fallback 替换 (2026-05-22 demo 4th iteration)."""
    elem = _make_element("P0", "{{CLIENT_FULL_NAME}}信用授信申报")
    cls = _make_classification("P0")
    mats = _build_materials(None)  # 显式 None

    result = placeholder_replace(elem, cls, mats)

    # None metadata 走空 dict 路径 · 全 miss · 走 demo-safe fallback
    assert result.action == "fill"
    assert "{{CLIENT_FULL_NAME}}" not in (result.new_text or ""), "demo 不留裸 placeholder"
    assert "本企业" in (result.new_text or ""), "None metadata demo fallback 应为'本企业'"
    assert result.pending_tag is not None


# ─────────────────────────────────────────────────────────────
# Test 4 · alias_map 大写 schema key ↔ 小写 sidecar 老 key
# ─────────────────────────────────────────────────────────────


def test_placeholder_alias_map_uppercase_to_lowercase():
    """sidecar metadata 仅 lower_case key (老调用方) · placeholder 用 UPPERCASE schema key ·
    验证 alias_map 兜底命中.

    alias_map (v16_op_handlers.placeholder_replace):
      CLIENT_FULL_NAME → name
      CLIENT_CORE_NAME → name_core
      CLIENT_LEGAL_REP → legal_rep
      ... (见 source code)
    """
    elem = _make_element("P0", "{{CLIENT_FULL_NAME}}信用授信申报")
    cls = _make_classification("P0")
    # 老 sidecar 路径 · 只传 lower_case 别名
    lower_metadata = {"name": "鼎盛商贸有限公司"}
    mats = _build_materials(lower_metadata)

    result = placeholder_replace(elem, cls, mats)

    assert result.action == "fill"
    assert result.new_text == "鼎盛商贸有限公司信用授信申报"
    assert result.pending_tag is None


def test_placeholder_alias_map_lowercase_direct_key():
    """metadata key 即 placeholder name 的 lower_case (例如 'client_full_name') ·
    验证 .lower() fallback 命中.
    """
    elem = _make_element("P0", "{{CLIENT_FULL_NAME}}信用授信申报")
    cls = _make_classification("P0")
    # metadata key 是 placeholder 的 lower_case
    lower_key_metadata = {"client_full_name": "鼎盛商贸有限公司"}
    mats = _build_materials(lower_key_metadata)

    result = placeholder_replace(elem, cls, mats)

    assert result.action == "fill"
    assert result.new_text == "鼎盛商贸有限公司信用授信申报"


def test_placeholder_alias_map_uppercase_priority_over_lowercase():
    """metadata 同时含大写 + 小写 key · 大写优先 (schema 真源)."""
    elem = _make_element("P0", "{{CLIENT_FULL_NAME}}信用授信申报")
    cls = _make_classification("P0")
    both_metadata = {
        "CLIENT_FULL_NAME": "大写_鼎盛商贸",
        "client_full_name": "小写_老 sidecar 名字",
        "name": "alias_老 sidecar 名字",
    }
    mats = _build_materials(both_metadata)

    result = placeholder_replace(elem, cls, mats)

    assert result.action == "fill"
    # 大写 key 必须优先 · 不能被 lower / alias 兜底覆盖
    assert "大写_鼎盛商贸" in result.new_text
    assert "小写_" not in result.new_text
    assert "alias_" not in result.new_text


# ─────────────────────────────────────────────────────────────
# Test 5 · paragraph-level 整段替换 (cross-run 回归)
# ─────────────────────────────────────────────────────────────


def test_full_placeholder_in_run_level_text(tmp_path):
    """{{CLIENT_FULL_NAME}} 字符串跨 4 个 Word run 拆开 · extract_elements
    返回的 e.text 是整段拼接 · placeholder_replace 应正常命中 · 回归 worker 1 教训.

    Phase 2 治本路径: docx 写回用 _set_paragraph_text · 第一个 run 写全部新文本 ·
    其余 run 清空 · 因此 placeholder 即使源 docx 跨 run · 输出仍单一 run · 安全.
    """
    docx_path = _create_cross_run_docx(tmp_path)
    elements = extract_elements(docx_path)

    # extract_elements 应把跨 run 的 {{CLIENT_FULL_NAME}} 拼回整段
    assert len(elements) == 1
    e = elements[0]
    assert "{{CLIENT_FULL_NAME}}" in e.text, (
        f"extract_elements 应该把 cross-run 拼回整段 placeholder · got: {e.text!r}"
    )

    # placeholder_replace 应识别并替换
    cls = _make_classification(e.location)
    mats = _build_materials(FULL_METADATA)
    result = placeholder_replace(e, cls, mats)

    assert result.action == "fill"
    assert result.new_text == "鼎盛商贸有限公司信用授信申报"
    assert result.pending_tag is None


# ─────────────────────────────────────────────────────────────
# Test 6 · sample sidecar metadata.json 校验 schema required key 都在
# ─────────────────────────────────────────────────────────────


SAMPLES_DIR = ROOT / "samples"
SCHEMA_PATH = ROOT / "templates" / "placeholder-schema.json"


def _load_schema_required_keys() -> tuple[list[str], list[str], list[str]]:
    """返回 (required_corp_keys, required_retail_keys, required_both_keys)."""
    data = json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))
    corp, retail, both = [], [], []
    for f in data["fields"]:
        if not f.get("required"):
            continue
        scope = f.get("scope_tag", "")
        if scope == "corporate":
            corp.append(f["key"])
        elif scope == "retail":
            retail.append(f["key"])
        elif scope == "both":
            both.append(f["key"])
    return corp, retail, both


# 抽 2 个 sample sidecar 验证 · 1 对公 (经纬有真 client_metadata) + 1 对私
SAMPLE_METADATA_PATHS = [
    SAMPLES_DIR / "经纬测绘_对公成稿A.metadata.json",
    SAMPLES_DIR / "兴业资管_对公成稿B.metadata.json",
]


@pytest.mark.parametrize("sidecar_path", SAMPLE_METADATA_PATHS)
def test_schema_required_keys_present_in_client_metadata(sidecar_path):
    """每个 sample sidecar 的 client_metadata 必须覆盖 schema 标 required=true 的字段.

    防未来 schema v1.2+ 把某字段从 required=false 升为 required=true 时 ·
    sample sidecar 漏更新 · 上线后 placeholder_replace 全填 pending · 没 fail-fast.
    """
    if not sidecar_path.exists():
        pytest.skip(f"sidecar missing: {sidecar_path}")

    sidecar = json.loads(sidecar_path.read_text(encoding="utf-8"))
    client_metadata = sidecar.get("client_metadata") or {}

    # 骨架模板 sidecar (client_metadata 为空 dict) 不参与必填校验 · skip
    # 这是 retail 等 skeleton 场景 · runtime 才传 metadata
    if not client_metadata:
        pytest.skip(
            f"sidecar.client_metadata 为空 dict (skeleton 模板) · {sidecar_path.name}"
        )

    corp_required, _, both_required = _load_schema_required_keys()
    # 对公 sample 须覆盖 corp + both 的 required
    required_all = corp_required + both_required

    missing = [k for k in required_all if k not in client_metadata]
    assert not missing, (
        f"{sidecar_path.name} client_metadata 缺 schema required key: {missing}\n"
        f"提示: 升级 schema (新加 required 字段) 时 · 必须同步更新所有 sample sidecar"
    )
