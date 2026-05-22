# -*- coding: utf-8 -*-
"""Integration · Q6-B CRUD 完整 (agent17 TODO medium-value 3 项) · 2026-05-21.

Why (per "先做完不需 review" 实施单):
  agent17 b9c7fbf 只做 upload + list · 缺 DELETE / 重命名 / 同名 supersede ·
  客户经理上传错版本只能堆叠 · dropdown 越用越脏 · 本期补齐 CRUD.

覆盖 (per 实施单 §Step 5):
  1. test_delete_user_template_archives_to_archived_dir
  2. test_delete_builtin_rejected_400  (CANNOT_DELETE_BUILTIN)
  3. test_delete_nonexistent_404       (TEMPLATE_NOT_FOUND)
  4. test_delete_then_list_excludes_archived
  5. test_rename_user_template
  6. test_rename_builtin_rejected_400  (CANNOT_PATCH_BUILTIN)
  7. test_rename_nonexistent_404
  8. test_rename_invalid_name_400      (长度 < 2 / > 80)
  9. test_rename_no_body_returns_noop  (PATCH 空 body = noop · UX 友好)
 10. test_upload_same_name_supersedes_old (archive 老 · 新 active · superseded_ids 含老 id)
 11. test_upload_different_name_no_supersede  (sanity check · 不同名不归档)
 12. test_archive_template_helper_unit (template_store.archive_template 直测)
"""
from __future__ import annotations

import io
import json
import shutil
import sys
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


@pytest.fixture(autouse=True)
def _autoinject_admin_cookie(monkeypatch):
    """注 admin cookie · pass require_action gate (同 test_template_upload.py)."""
    from auth_service.dependencies import COOKIE_NAME
    from auth_service.jwt_util import issue

    original_init = TestClient.__init__

    def patched_init(self, *args, **kwargs):
        original_init(self, *args, **kwargs)
        self.cookies.set(COOKIE_NAME, issue("u_test", "admin"))

    monkeypatch.setattr(TestClient, "__init__", patched_init)


@pytest.fixture
def isolated_template_root(tmp_path, monkeypatch):
    """每个 test 隔离 TEMPLATE_ROOT · 防互染."""
    from agent_report import template_store
    test_root = tmp_path / "templates"
    test_root.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(template_store, "TEMPLATE_ROOT", test_root)
    yield test_root
    if test_root.exists():
        shutil.rmtree(test_root, ignore_errors=True)


@pytest.fixture
def client():
    from api_server import app
    return TestClient(app)


# ─────────────────────────────────────────────────────────────
# Mini docx 合成 (与 test_template_upload.py 同 helper)
# ─────────────────────────────────────────────────────────────


def _make_clean_template_bytes() -> bytes:
    """合成 mini.docx · 全 placeholder · 0 residue · validation=PASS."""
    doc = Document()
    doc.add_paragraph("{{CLIENT_FULL_NAME}}信用授信申报")
    doc.add_paragraph("法定代表人:{{CLIENT_LEGAL_REP}}")
    doc.add_paragraph("注册资本:{{CLIENT_REGISTERED_CAPITAL}}")
    doc.add_paragraph("申请{{CREDIT_AMOUNT}}综合授信 {{CREDIT_PERIOD}}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload_template(client: TestClient, name: str) -> str:
    """helper · 上传一个 clean 模板 · 返 template_id."""
    content = _make_clean_template_bytes()
    resp = client.post(
        "/api/report/upload-template",
        params={"template_name": name, "business_line": "corporate"},
        files={"file": (f"{name}.docx", content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    return resp.json()["template_id"]


# ─────────────────────────────────────────────────────────────
# DELETE
# ─────────────────────────────────────────────────────────────


def test_delete_user_template_archives_to_archived_dir(client, isolated_template_root):
    """DELETE user template · 200 · 真搬到 .archived/{id}/ · metadata 加 archived_at."""
    template_id = _upload_template(client, "del-test")
    src_dir = isolated_template_root / template_id
    assert src_dir.exists(), "上传后源目录应存在"

    resp = client.delete(f"/api/report/templates/{template_id}")
    assert resp.status_code == 200, resp.text
    j = resp.json()
    assert j["deleted"] == template_id
    assert ".archived" in j["archive_path"], f"archive_path 应含 .archived · 实际 {j['archive_path']!r}"
    assert j["archived_at"], "应返 archived_at 时间戳"
    assert j["archive_reason"] == "user_deleted"

    # 真磁盘验证: 源目录已搬走 · .archived/{id}/ 应存在
    assert not src_dir.exists(), "源目录应被搬走"
    archived_dir = isolated_template_root / ".archived" / template_id
    assert archived_dir.exists(), f"归档目录应存在 · path={archived_dir}"
    # metadata.json 含 archived_at + reason
    meta = json.loads((archived_dir / "metadata.json").read_text(encoding="utf-8"))
    assert meta["archived_at"]
    assert meta["archive_reason"] == "user_deleted"


def test_delete_builtin_rejected_400(client, isolated_template_root):
    """DELETE builtin (template_id 不以 user- 开头) · 400 CANNOT_DELETE_BUILTIN."""
    resp = client.delete("/api/report/templates/samples-builtin-fake")
    assert resp.status_code == 400, resp.text
    assert resp.json()["detail"]["error"]["code"] == "CANNOT_DELETE_BUILTIN"


def test_delete_nonexistent_404(client, isolated_template_root):
    """DELETE 不存在的 user template · 404 TEMPLATE_NOT_FOUND."""
    resp = client.delete("/api/report/templates/user-nonexistent12345")
    assert resp.status_code == 404, resp.text
    assert resp.json()["detail"]["error"]["code"] == "TEMPLATE_NOT_FOUND"


def test_delete_then_list_excludes_archived(client, isolated_template_root):
    """删除后 GET /templates 不再返被删 · .archived/ skip 不染列表."""
    tid = _upload_template(client, "list-after-delete")

    list_before = client.get("/api/report/templates").json()
    user_ids_before = [t["template_id"] for t in list_before["user"]]
    assert tid in user_ids_before

    del_resp = client.delete(f"/api/report/templates/{tid}")
    assert del_resp.status_code == 200

    list_after = client.get("/api/report/templates").json()
    user_ids_after = [t["template_id"] for t in list_after["user"]]
    assert tid not in user_ids_after, f"删除后 list 仍含 · ids={user_ids_after}"


# ─────────────────────────────────────────────────────────────
# PATCH (rename)
# ─────────────────────────────────────────────────────────────


def test_rename_user_template(client, isolated_template_root):
    """PATCH user template · template_name 改成新值 · metadata.renamed_at 加."""
    tid = _upload_template(client, "old-name")
    resp = client.patch(
        f"/api/report/templates/{tid}",
        json={"template_name": "new-name-改后"},
    )
    assert resp.status_code == 200, resp.text
    j = resp.json()
    assert j["updated"] == tid
    assert j["noop"] is False
    assert j["metadata"]["template_name"] == "new-name-改后"
    assert j["metadata"]["renamed_at"]

    # GET 验真 (磁盘 metadata 真改了)
    list_resp = client.get("/api/report/templates").json()
    found = [t for t in list_resp["user"] if t["template_id"] == tid]
    assert len(found) == 1
    assert found[0]["name"] == "new-name-改后"


def test_rename_builtin_rejected_400(client, isolated_template_root):
    """PATCH builtin · 400 CANNOT_PATCH_BUILTIN."""
    resp = client.patch(
        "/api/report/templates/builtin-fake-id",
        json={"template_name": "试图改 builtin"},
    )
    assert resp.status_code == 400, resp.text
    assert resp.json()["detail"]["error"]["code"] == "CANNOT_PATCH_BUILTIN"


def test_rename_nonexistent_404(client, isolated_template_root):
    """PATCH 不存在的 user template · 404 TEMPLATE_NOT_FOUND."""
    resp = client.patch(
        "/api/report/templates/user-nope12345",
        json={"template_name": "新名"},
    )
    assert resp.status_code == 404, resp.text
    assert resp.json()["detail"]["error"]["code"] == "TEMPLATE_NOT_FOUND"


def test_rename_invalid_name_400(client, isolated_template_root):
    """PATCH 名字过短 (<2) · 400 TEMPLATE_NAME_INVALID."""
    tid = _upload_template(client, "invalid-name-test")
    resp = client.patch(
        f"/api/report/templates/{tid}",
        json={"template_name": "a"},
    )
    assert resp.status_code == 400, resp.text
    assert resp.json()["detail"]["error"]["code"] == "TEMPLATE_NAME_INVALID"


def test_rename_no_body_returns_noop(client, isolated_template_root):
    """PATCH 空 body (template_name=None) · 200 noop · 返当前 metadata (UX 友好)."""
    tid = _upload_template(client, "noop-test")
    resp = client.patch(f"/api/report/templates/{tid}", json={})
    assert resp.status_code == 200, resp.text
    j = resp.json()
    assert j["noop"] is True
    assert j["metadata"]["template_id"] == tid


# ─────────────────────────────────────────────────────────────
# Supersede (upload 同名 archive 老的)
# ─────────────────────────────────────────────────────────────


def test_upload_same_name_supersedes_old(client, isolated_template_root):
    """上传同 template_name · 老 user 模板归档 · 新 active · superseded_ids 含老 id."""
    # 上传第一份 · 名 "我的对公 v1"
    tid1 = _upload_template(client, "我的对公 v1")
    # 上传第二份 · 同 template_name
    content = _make_clean_template_bytes()
    resp2 = client.post(
        "/api/report/upload-template",
        params={"template_name": "我的对公 v1", "business_line": "corporate"},
        files={"file": ("v2.docx", content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp2.status_code == 200, resp2.text
    j2 = resp2.json()
    tid2 = j2["template_id"]
    assert tid2 != tid1, "新模板 UUID 应不同"
    # superseded_ids 应含 tid1
    assert tid1 in j2["superseded_ids"], \
        f"superseded_ids 应含老 id={tid1} · 实际 {j2['superseded_ids']}"

    # GET /templates · 只剩新的 active
    list_resp = client.get("/api/report/templates").json()
    user_ids = [t["template_id"] for t in list_resp["user"]]
    assert tid2 in user_ids, f"新模板应 active · ids={user_ids}"
    assert tid1 not in user_ids, f"老模板应被归档 · ids={user_ids}"

    # 磁盘验证: tid1 搬到 .archived/{tid1}/ · metadata reason 含 superseded_by:{tid2}
    archived_dir = isolated_template_root / ".archived" / tid1
    assert archived_dir.exists(), f".archived/{tid1}/ 应存在"
    archived_meta = json.loads((archived_dir / "metadata.json").read_text(encoding="utf-8"))
    assert archived_meta["archive_reason"] == f"superseded_by:{tid2}"


def test_upload_different_name_no_supersede(client, isolated_template_root):
    """上传不同 template_name · 老的不动 · superseded_ids 空."""
    tid1 = _upload_template(client, "模板 A")
    content = _make_clean_template_bytes()
    resp2 = client.post(
        "/api/report/upload-template",
        params={"template_name": "模板 B (不同名)", "business_line": "corporate"},
        files={"file": ("b.docx", content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    j2 = resp2.json()
    assert j2["superseded_ids"] == [], f"不同名不应 supersede · 实际 {j2['superseded_ids']}"
    # 两个都 active
    list_resp = client.get("/api/report/templates").json()
    user_ids = [t["template_id"] for t in list_resp["user"]]
    assert tid1 in user_ids
    assert j2["template_id"] in user_ids


# ─────────────────────────────────────────────────────────────
# Unit · template_store.archive_template 直测 (不走 HTTP)
# ─────────────────────────────────────────────────────────────


def test_archive_template_helper_unit(client, isolated_template_root):
    """template_store.archive_template · soft-delete · 真搬 · 写 archive_reason."""
    from agent_report.template_store import archive_template

    tid = _upload_template(client, "unit-archive-test")
    src = isolated_template_root / tid
    assert src.exists()

    result = archive_template(tid, reason="manual_test", archived_by="u_unit_test")
    assert result["archive_reason"] == "manual_test"
    assert result["archived_by"] == "u_unit_test"
    assert not src.exists(), "源目录应搬走"
    assert (isolated_template_root / ".archived" / tid).exists()


def test_archive_template_builtin_rejected(isolated_template_root):
    """archive_template 拒 builtin · PermissionError."""
    from agent_report.template_store import archive_template
    with pytest.raises(PermissionError):
        archive_template("builtin-something", reason="试图")


def test_archive_template_nonexistent_raises(isolated_template_root):
    """archive_template · 不存在 · FileNotFoundError."""
    from agent_report.template_store import archive_template
    with pytest.raises(FileNotFoundError):
        archive_template("user-nope-xyz", reason="试图")
