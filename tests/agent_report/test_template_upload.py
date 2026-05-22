# -*- coding: utf-8 -*-
"""Integration · Q6-B 用户自定义模板上传 + lint + V16FillRequest 接 user-template (2026-05-21).

覆盖 (per Q6-B 实施单 §Step 6):
  1. test_upload_valid_template_pass · 上传含 {{KEY}} 的 mini.docx · validation=PASS · 0 residue
  2. test_upload_template_with_residue_warn · 含 specific 客户字面 (公司全称) · validation=WARN · residue≥1
  3. test_upload_rejects_non_docx · 上传 .pdf · 400 TEMPLATE_EXT_INVALID
  4. test_upload_rejects_oversized · 11MB · 413 TEMPLATE_PERSIST_REJECTED
  5. test_upload_rejects_empty_template_name · template_name 空 · 400 TEMPLATE_NAME_INVALID
  6. test_list_templates_returns_builtin_5_and_user_uploaded · GET /api/report/templates
  7. test_resolve_user_template_path · resolve_template_path("user-template:{id}") → 真路径
  8. test_resolve_user_template_404 · template_id 不存在 · FileNotFoundError
  9. test_v16_fill_accepts_user_template · POST /api/report/v16/fill source_docx="user-template:{id}" 不 404
"""
from __future__ import annotations

import io
import json
import shutil
import sys
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


@pytest.fixture(autouse=True)
def _autoinject_admin_cookie(monkeypatch):
    """注 admin cookie · 与 tests/agent_channel/conftest.py 同样 pattern · pass require_action gate."""
    from auth_service.dependencies import COOKIE_NAME
    from auth_service.jwt_util import issue

    original_init = TestClient.__init__

    def patched_init(self, *args, **kwargs):
        original_init(self, *args, **kwargs)
        self.cookies.set(COOKIE_NAME, issue("u_test", "admin"))

    monkeypatch.setattr(TestClient, "__init__", patched_init)


@pytest.fixture
def isolated_template_root(tmp_path, monkeypatch):
    """每个 test 隔离 TEMPLATE_ROOT · 防互染."""
    from agent_report import template_store
    test_root = tmp_path / "templates"
    test_root.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(template_store, "TEMPLATE_ROOT", test_root)
    yield test_root
    # cleanup
    if test_root.exists():
        shutil.rmtree(test_root, ignore_errors=True)


@pytest.fixture
def client():
    from api_server import app
    return TestClient(app)


# ─────────────────────────────────────────────────────────────
# Mini docx 合成 · 含 {{KEY}} placeholder + 可选 residue
# ─────────────────────────────────────────────────────────────


def _make_clean_template_bytes() -> bytes:
    """合成 mini.docx · 全 placeholder · 0 residue · validation=PASS."""
    doc = Document()
    doc.add_paragraph("{{CLIENT_FULL_NAME}}信用授信申报")
    doc.add_paragraph("法定代表人:{{CLIENT_LEGAL_REP}}")
    doc.add_paragraph("注册资本:{{CLIENT_REGISTERED_CAPITAL}}")
    doc.add_paragraph("申请{{CREDIT_AMOUNT}}综合授信 {{CREDIT_PERIOD}}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_residue_template_bytes() -> bytes:
    """合成 mini.docx · 含 placeholder + 未替换 specific 字面 (公司全称 / 法人姓名)."""
    doc = Document()
    doc.add_paragraph("{{CLIENT_FULL_NAME}}信用授信申报")
    # 客户经理忘了把 "经纬测绘有限公司" placeholder 化 · 后端应检出 residue (company pattern)
    doc.add_paragraph("申请单位:浙江经纬测绘有限公司")
    # 法定代表人姓名未替换 · person pattern 应命中 (Surname 王 + 中文名)
    doc.add_paragraph("法定代表人:王明华")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────
# Case 1 · 干净模板 · validation=PASS
# ─────────────────────────────────────────────────────────────


def test_upload_valid_template_pass(client, isolated_template_root):
    """上传含 {{KEY}} 的 mini docx · 0 residue · validation=PASS."""
    content = _make_clean_template_bytes()
    resp = client.post(
        "/api/report/upload-template",
        params={"template_name": "测试干净模板", "business_line": "corporate"},
        files={"file": ("clean.docx", content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    j = resp.json()
    assert j["template_id"].startswith("user-"), f"template_id 应 'user-' 前缀: {j['template_id']!r}"
    assert j["template_path"] == f"user-template:{j['template_id']}"
    vr = j["validation_report"]
    assert vr["validation"] == "PASS", f"clean 模板应 PASS · 实际 {vr['validation']}, samples={vr['residue_samples']}"
    assert vr["placeholder_count"] >= 5, f"应≥5 placeholder · 实际 {vr['placeholder_count']}"
    assert vr["residue_count"] == 0, f"residue 应 0 · 实际 {vr['residue_count']}"
    # placeholder_keys 应含 CLIENT_FULL_NAME 等
    assert "CLIENT_FULL_NAME" in vr["placeholder_keys"]
    assert "CREDIT_AMOUNT" in vr["placeholder_keys"]


# ─────────────────────────────────────────────────────────────
# Case 2 · 含 residue · validation=WARN
# ─────────────────────────────────────────────────────────────


def test_upload_template_with_residue_warn(client, isolated_template_root):
    """模板含未 placeholder 化的 specific 字面 · validation=WARN · residue≥1."""
    content = _make_residue_template_bytes()
    resp = client.post(
        "/api/report/upload-template",
        params={"template_name": "含残留模板", "business_line": "corporate"},
        files={"file": ("residue.docx", content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    vr = resp.json()["validation_report"]
    assert vr["validation"] == "WARN", f"含 specific 字面应 WARN · 实际 {vr['validation']}"
    assert vr["residue_count"] >= 1, f"residue_count 应≥1 · 实际 {vr['residue_count']}"
    # residue_samples 含至少 1 个 sample (公司或人名)
    kinds = [s["kind"] for s in vr["residue_samples"]]
    assert any(k in ("company", "person") for k in kinds), \
        f"应检出 company 或 person residue · 实际 kinds={kinds}"


# ─────────────────────────────────────────────────────────────
# Case 3 · 拒非 docx
# ─────────────────────────────────────────────────────────────


def test_upload_rejects_non_docx(client, isolated_template_root):
    """非 .docx 文件 (e.g. .pdf) · 400 TEMPLATE_EXT_INVALID."""
    resp = client.post(
        "/api/report/upload-template",
        params={"template_name": "测试 pdf", "business_line": "corporate"},
        files={"file": ("fake.pdf", b"fake pdf content", "application/pdf")},
    )
    assert resp.status_code == 400, resp.text
    assert resp.json()["detail"]["error"]["code"] == "TEMPLATE_EXT_INVALID"


# ─────────────────────────────────────────────────────────────
# Case 4 · 拒超大文件
# ─────────────────────────────────────────────────────────────


def test_upload_rejects_oversized(client, isolated_template_root):
    """11MB · 超 10MB 上限 · 413 TEMPLATE_PERSIST_REJECTED."""
    # 11MB 假 docx · persist_template 内 ValueError(超限) → 413
    big_content = b"PK" + b"x" * (11 * 1024 * 1024)  # 11MB · 假 zip header
    resp = client.post(
        "/api/report/upload-template",
        params={"template_name": "big template", "business_line": "corporate"},
        files={"file": ("big.docx", big_content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 413, resp.text
    assert resp.json()["detail"]["error"]["code"] == "TEMPLATE_PERSIST_REJECTED"


# ─────────────────────────────────────────────────────────────
# Case 5 · 拒空 template_name
# ─────────────────────────────────────────────────────────────


def test_upload_rejects_empty_template_name(client, isolated_template_root):
    """template_name 空 · 400 TEMPLATE_NAME_INVALID."""
    content = _make_clean_template_bytes()
    resp = client.post(
        "/api/report/upload-template",
        params={"template_name": "", "business_line": "corporate"},
        files={"file": ("clean.docx", content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 400, resp.text
    assert resp.json()["detail"]["error"]["code"] == "TEMPLATE_NAME_INVALID"


# ─────────────────────────────────────────────────────────────
# Case 6 · GET /api/report/templates · 5 builtin + N user
# ─────────────────────────────────────────────────────────────


def test_list_templates_returns_builtin_5_and_user_uploaded(client, isolated_template_root):
    """GET /templates · 5 builtin · user 上传后增量."""
    # 初始: 5 builtin · 0 user
    resp = client.get("/api/report/templates")
    assert resp.status_code == 200, resp.text
    j = resp.json()
    assert len(j["builtin"]) == 5, f"builtin 应 5 · 实际 {len(j['builtin'])}"
    assert j["user"] == [], f"初始 user 应空 · 实际 {j['user']}"
    # builtin paths 都以 samples/ 开头
    for b in j["builtin"]:
        assert b["template_path"].startswith("samples/"), b
        assert b["type"] == "builtin"

    # 上传 1 个 · user 列应含 1
    content = _make_clean_template_bytes()
    up = client.post(
        "/api/report/upload-template",
        params={"template_name": "我的模板 A", "business_line": "corporate"},
        files={"file": ("a.docx", content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert up.status_code == 200, up.text
    template_id = up.json()["template_id"]

    resp2 = client.get("/api/report/templates")
    j2 = resp2.json()
    assert len(j2["user"]) == 1, f"应 1 user template · 实际 {len(j2['user'])}"
    assert j2["user"][0]["template_path"] == f"user-template:{template_id}"
    assert j2["user"][0]["name"] == "我的模板 A"
    assert j2["user"][0]["validation"] == "PASS"
    assert j2["user"][0]["uploader"]  # 非空 · 有审计


# ─────────────────────────────────────────────────────────────
# Case 7 · resolve_template_path · user-template prefix
# ─────────────────────────────────────────────────────────────


def test_resolve_user_template_path(client, isolated_template_root):
    """resolve_template_path('user-template:{id}') 返真路径 · 文件存在."""
    from agent_report.template_store import resolve_template_path

    content = _make_clean_template_bytes()
    up = client.post(
        "/api/report/upload-template",
        params={"template_name": "resolve 测试", "business_line": "corporate"},
        files={"file": ("rt.docx", content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert up.status_code == 200, up.text
    template_id = up.json()["template_id"]

    resolved = resolve_template_path(f"user-template:{template_id}")
    assert resolved.exists(), f"resolved path 不存在: {resolved}"
    assert resolved.suffix == ".docx"
    assert resolved.parent.name == template_id


def test_resolve_user_template_404(isolated_template_root):
    """template_id 不存在 · FileNotFoundError."""
    from agent_report.template_store import resolve_template_path
    with pytest.raises(FileNotFoundError):
        resolve_template_path("user-template:nonexistent12345")


# ─────────────────────────────────────────────────────────────
# Case 8 · V16FillRequest source_docx 接 user-template (不 404)
# ─────────────────────────────────────────────────────────────


def test_v16_fill_accepts_user_template(client, isolated_template_root, monkeypatch):
    """POST /v16/fill source_docx='user-template:{id}' · 不返 404.

    不验真 fill 跑通 (依赖 DEEPSEEK_API_KEY) · 只验 source_docx 解析层不挂.
    用 mock=True + admin role · 触发 mock_v16_stream · 跳过 LLM 真调.
    """
    # 上传一个 template
    content = _make_clean_template_bytes()
    up = client.post(
        "/api/report/upload-template",
        params={"template_name": "v16 接受测试", "business_line": "corporate"},
        files={"file": ("v16t.docx", content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    template_id = up.json()["template_id"]

    # 用 mock=True 走 mock 路径 (避免真 LLM 调用) · admin role 已通过 _autoinject_admin_cookie
    resp = client.post(
        "/api/report/v16/fill",
        json={
            "report_id": "test-rid",
            "source_docx": f"user-template:{template_id}",
            "business_line": "corporate",
            "mock": True,
        },
    )
    # mock 路径 · 应返 200 + SSE stream (不 404)
    # 注意 mock 在 v16_runner 内可能因 SSE 流写 / classifier 路径有别的 error · 这里只断不返 404 TEMPLATE_NOT_FOUND
    assert resp.status_code != 404 or "TEMPLATE_NOT_FOUND" not in resp.text, \
        f"source_docx user-template 解析应成功 · 不返 TEMPLATE_NOT_FOUND · 实际 {resp.status_code}: {resp.text[:200]}"


def test_v16_fill_rejects_unknown_user_template(client, isolated_template_root):
    """V16FillRequest source_docx='user-template:不存在' · 404 TEMPLATE_NOT_FOUND."""
    resp = client.post(
        "/api/report/v16/fill",
        json={
            "report_id": "test-rid",
            "source_docx": "user-template:nonexistent_xyz",
            "business_line": "corporate",
            "mock": True,
        },
    )
    assert resp.status_code == 404, resp.text
    assert resp.json()["detail"]["error"]["code"] == "TEMPLATE_NOT_FOUND"
