# -*- coding: utf-8 -*-
"""V16 e2e-with-LLM 端到端测试 · 红线 fix 防御层 (2026-05-21).

为什么这个测试存在 (agent15 第 3 次 dogfooding 之前必须跑通):
  - test_render_smoke.py 只测 placeholder_replace handler 的确定性路径
  - **不测** REWRITE/REWRITE handler (LLM 路径 · 幻觉重灾区)
  - agent15 实测 REWRITE handler 重度幻觉 + 跨样本污染 + 未解析 placeholder 进 docx:
    * 蓝汀家电报告 P46 含 "兴业国际信托" (其他 sample 客户名 · 跨样本污染)
    * "福州市仓山区XX工业园区 / 240人 / 微型企业" (client_metadata 没给的字段全编造)
    * "授信500万元 / ESG评级A / PD评级4级" (LLM 凭空编造授信条件)
    * `{{CLIENT_INDUSTRY_FULL}}` / `{{CREDIT_PERIOD}}` 字面进 docx (LLM 没解析)
  - 这些是用户看到的最严重渲染 bug · 必须 e2e LLM 跑全链路才能复现 + 锁死防御

测试范围 (2 fixture · agent15 验证当时的 2 个客户):
  - DP002_蓝汀家电 (data/mock/deep-pillar/DP002_蓝汀家电/client_metadata.json)
  - corp-dingsheng-001 (data/mock/workspace/credit/scenarios/corp-dingsheng-001.json)

依赖 (运行时):
  - .env 含 DEEPSEEK_API_KEY (无 key 跳过 · 不阻 CI)
  - samples/经纬测绘_对公成稿A.docx (已 placeholderize · 红线 fix Phase 2 产物)
  - outputs/v16_llm_classified.json (classifier 产物 · 已 commit)
  - samples/ 含 经纬测绘材料 (用于触发跨样本污染场景)

测试断言 (per 红线 fix Step 5):
  1. 跨样本污染: 不出现其他 sample 客户名 (兴业国际信托 / 经纬数字科技 / 鼎盛 等
     与当前客户不同的 sample 客户名)
  2. 编造字段: 不出现 client_metadata 没给但 LLM 可能编造的字段 (福州市仓山区 /
     240人 / ESG评级 / PD评级)
  3. 未解析 placeholder: rendered docx 不含 {{KEY}} 字面 (LLM REWRITE 漏的 + final
     sweep 已兜底)
  4. 新客户名命中: 当前客户的 CLIENT_FULL_NAME 在 rendered 文本中出现 ≥1 次
"""
from __future__ import annotations

import json
import os
import re
import sys
import tempfile
from pathlib import Path

import pytest

# repo 根入 sys.path · 允许 import v16 模块 (tests/integration/ 在两层下)
REPO = Path(__file__).resolve().parents[2]
if str(REPO) not in sys.path:
    sys.path.insert(0, str(REPO))


# ─────────────────────────────────────────────────────────────
# 环境准备 · .env 加载 + LLM cache 关
# ─────────────────────────────────────────────────────────────


def _load_dotenv() -> None:
    """轻量 .env loader · 不依赖第三方包."""
    env_file = REPO / ".env"
    if not env_file.exists():
        return
    for line in env_file.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, _, v = line.partition("=")
        k, v = k.strip(), v.strip()
        if v and ((v[0] == '"' and v[-1] == '"') or (v[0] == "'" and v[-1] == "'")):
            v = v[1:-1]
        os.environ.setdefault(k, v)


_load_dotenv()


def _disable_llm_cache() -> None:
    """关闭 LLMClient cache · 防止读到 agent15 之前生成的污染 cache.

    污染样本 (供 debug 参考):
      .cache/llm/deepseek/1c96da9ef692b9a29fa2364d.json (P25)
      .cache/llm/deepseek/053a67626ba3aab1f2b2087a.json (P46 · 含 "兴业国际信托")
    """
    import llm

    _orig_init = llm.LLMClient.__init__

    def _no_cache_init(self, provider="deepseek", api_key="", cache_enabled=True):
        _orig_init(self, provider, api_key, cache_enabled=False)

    llm.LLMClient.__init__ = _no_cache_init


# ─────────────────────────────────────────────────────────────
# Fixtures (2 case · 红线 fix scope)
# ─────────────────────────────────────────────────────────────

FIXTURES = [
    ("DP002_蓝汀家电", REPO / "data/mock/deep-pillar/DP002_蓝汀家电/client_metadata.json"),
    ("corp-dingsheng-001", REPO / "data/mock/workspace/credit/scenarios/corp-dingsheng-001.json"),
]

TEMPLATE_DOCX = REPO / "samples" / "经纬测绘_对公成稿A.docx"
MATERIAL_DIR = REPO / "samples"
CLASSIFIED_JSON = REPO / "outputs" / "v16_llm_classified.json"

# 跨样本污染清单 (老 sample 客户名 · 当前客户不应包含这些字面)
# 注意: 这些 keyword 是从 5 个老 sample (经纬 / 兴业 / 龙峰精工 / 蓝汀 / 宸星家装 ...)
#       的客户名 / 法人 / 关键标识中拉的 · 当前客户的 CLIENT_FULL_NAME 若含某 keyword,
#       则此 keyword 不视为污染 (例如本身就是经纬时不算).
POLLUTION_KEYWORDS = [
    "兴业国际信托",
    "兴业资产管理",
    "兴业国信资产管理",
    "福建经纬数字科技",
    "经纬数字科技",
    "经纬测绘",
    "测绘地理信息",
    "郑志煌",
    "福建省招标股份",
    "福建省招标采购集团",
    "招标采购集团",
]

# 编造字段 hallucination 清单 (agent15 实测当时 LLM 编造的字段 · 不应再现)
# 这些都是 DP002 蓝汀 / corp-dingsheng 的 client_metadata 没给的字段 ·
# LLM 看到原模板段落含 {{KEY}} 时如果还编造 · 说明 prompt 约束 fail.
HALLUCINATION_KEYWORDS = [
    "福州市仓山区",  # DP002 / corp-dingsheng 都不在福州 · LLM 编造
    "ESG评级",  # client_metadata 没给 · LLM 编造
    "PD评级4级",  # DP002=未给 / corp-dingsheng=未给 (corp 给的是 risk_grade 不是 PD)
    "240人",  # DP002 给的是 180人 · corp-dingsheng 给的是 38人 · 240 是 LLM 编造
]


# ─────────────────────────────────────────────────────────────
# 工具 · 加载 metadata / 抽 docx 文本
# ─────────────────────────────────────────────────────────────


def _load_client_metadata(metadata_path: Path) -> tuple[dict, str]:
    """加载 fixture client_metadata · 返回 (metadata 字典, CLIENT_FULL_NAME).

    支持 3 种格式 (与 test_render_smoke._load_client_metadata 一致):
      a) 3 section nested (新): {client_metadata: {...}, credit_terms: {...}, material_facts: {...}}
      b) 旧 nested (corp scenarios): {client_metadata: {全 45 key}, profile: {...}, ...}
      c) 旧 flat (DP00X): {schema_version: ..., CLIENT_FULL_NAME: ..., ...}
    """
    raw = json.loads(metadata_path.read_text(encoding="utf-8"))

    has_three_section = (
        isinstance(raw.get("client_metadata"), dict)
        and isinstance(raw.get("credit_terms"), dict)
        and isinstance(raw.get("material_facts"), dict)
    )

    if has_three_section:
        merged: dict = {}
        for section in ("material_facts", "credit_terms", "client_metadata"):
            sub = raw.get(section)
            if isinstance(sub, dict):
                merged.update(sub)
        merged["_sections"] = {
            s: dict(raw.get(s, {})) for s in ("client_metadata", "credit_terms", "material_facts")
        }
        client_metadata = merged
        new_client_name = merged.get("CLIENT_FULL_NAME") or raw.get("profile", {}).get("company_name")
    elif "client_metadata" in raw and isinstance(raw["client_metadata"], dict):
        client_metadata = raw["client_metadata"]
        new_client_name = client_metadata.get("CLIENT_FULL_NAME") or raw.get("profile", {}).get("company_name")
    else:
        client_metadata = raw
        new_client_name = client_metadata.get("CLIENT_FULL_NAME")

    return client_metadata, new_client_name or ""


def _extract_docx_text(docx_path: Path) -> str:
    """从 docx 抽全量段落 + 表格 cell 文本."""
    from docx import Document

    doc = Document(str(docx_path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


# ─────────────────────────────────────────────────────────────
# 主测试 (parametrize · 2 fixture)
# ─────────────────────────────────────────────────────────────


@pytest.mark.skipif(
    not os.environ.get("DEEPSEEK_API_KEY"),
    reason="需 DEEPSEEK_API_KEY (从 .env 加载) · 跑真 LLM 调用",
)
@pytest.mark.skipif(
    not CLASSIFIED_JSON.exists(),
    reason=f"classifier 产物缺失: {CLASSIFIED_JSON} (跑 py v16_classifier.py 生成)",
)
@pytest.mark.parametrize("name,metadata_path", FIXTURES, ids=[f[0] for f in FIXTURES])
def test_e2e_llm_no_hallucination_no_cross_sample_pollution(name, metadata_path, tmp_path):
    """跑 v16 全 pipeline (含 LLM REWRITE) · 断言 4 件事:

    1. 无跨样本污染 (其他 sample 客户名不出现)
    2. 无编造字段 (LLM 编造的地址 / 员工数 / 评级 不出现)
    3. 无未解析 placeholder ({{KEY}} 字面不进 rendered docx)
    4. 新客户名命中 (当前客户 CLIENT_FULL_NAME 出现 ≥1 次)
    """
    # 0. 前置条件
    assert metadata_path.exists(), f"fixture missing: {metadata_path}"
    assert TEMPLATE_DOCX.exists(), f"template missing: {TEMPLATE_DOCX}"

    # 1. 加载 metadata
    client_metadata, new_client_name = _load_client_metadata(metadata_path)
    assert client_metadata, f"{name} client_metadata empty"
    assert new_client_name, f"{name} 缺 CLIENT_FULL_NAME / company_name"

    # 2. 关 LLM cache (防读污染 cache)
    _disable_llm_cache()

    # 3. 跑 v16 pipeline (含 LLM REWRITE)
    from v16_pipeline import run_pipeline

    out_dir = tmp_path / f"e2e_{name}"
    out_dir.mkdir(parents=True, exist_ok=True)

    result = run_pipeline(
        source_docx=TEMPLATE_DOCX,
        material_dir=MATERIAL_DIR,
        classified_json=CLASSIFIED_JSON,
        output_dir=out_dir,
        client_metadata=client_metadata,
    )

    output_docx = Path(result["output_docx"])
    assert output_docx.exists(), f"{name} output docx missing: {output_docx}"

    # 4. 抽 rendered docx 文本
    text = _extract_docx_text(output_docx)
    assert text and len(text) > 1000, f"{name} rendered text too short: {len(text)}"

    # ─── 断言 1: 跨样本污染 ───
    leaked_pollution: list[str] = []
    for kw in POLLUTION_KEYWORDS:
        if not kw:
            continue
        # 若当前客户 CLIENT_FULL_NAME 本身含 kw,该 kw 不算污染
        if kw in new_client_name:
            continue
        if kw in text:
            leaked_pollution.append(kw)
    assert not leaked_pollution, (
        f"[{name}] 跨样本污染: 其他 sample 客户名出现在 rendered docx\n"
        f"  pollution_hits: {leaked_pollution}\n"
        f"  current_client: {new_client_name}\n"
        f"  output: {output_docx}\n"
        f"  → 检查 REWRITE handler 是否过滤了跨样本 raw_statements / facts"
    )

    # ─── 断言 2: 编造字段 ───
    leaked_hallucination: list[str] = []
    for kw in HALLUCINATION_KEYWORDS:
        if not kw:
            continue
        if kw in text:
            leaked_hallucination.append(kw)
    assert not leaked_hallucination, (
        f"[{name}] LLM 编造字段: client_metadata 没给的字段被 LLM 凭空编造\n"
        f"  hallucination_hits: {leaked_hallucination}\n"
        f"  output: {output_docx}\n"
        f"  → 检查 REWRITE prompt 是否强约束 '档案没列的字段必须未能自动填写'"
    )

    # ─── 断言 3: 未解析 placeholder ───
    naked_placeholders = re.findall(r"\{\{([A-Z][A-Z0-9_]*)\}\}", text)
    assert not naked_placeholders, (
        f"[{name}] rendered docx 残留 {{{{KEY}}}} 字面 placeholder\n"
        f"  naked: {sorted(set(naked_placeholders))}\n"
        f"  output: {output_docx}\n"
        f"  → 检查 REWRITE handler 是否跑了 _final_placeholder_sweep"
    )

    # ─── 断言 4: 新客户名命中 ───
    assert new_client_name in text, (
        f"[{name}] 新客户名 {new_client_name!r} 不在 rendered docx 中\n"
        f"  output: {output_docx}\n"
        f"  → 检查 placeholder_replace 是否把 {{{{CLIENT_FULL_NAME}}}} 替换为 metadata 实值"
    )


# ─────────────────────────────────────────────────────────────
# 2026-05-22 D 真治本 agent20: 新增结构相似度断言 ·
# 用户 dogfood 4th 暴露根因 — LLM REWRITE 后段落仍与原经纬测绘模板高度相似 (改了公司名但段落结构/数字/比例没动)
# 防御: rendered docx 段落与原 .bak-pre-placeholder 段落 SequenceMatcher 相似度 ≥ 0.6 的"长段"(≥ 50 字符)
#       不应超过 3 段 (允许少量公共法律/格式短句 · 但报告主体不能像模板)
# ─────────────────────────────────────────────────────────────


ORIGINAL_TEMPLATE_BAK = REPO / "samples" / "经纬测绘_对公成稿A.docx.bak-pre-placeholder"


def _extract_paragraphs(docx_path: Path) -> list[str]:
    """从 docx 抽段落 list (paragraph + table cell · 不 join)."""
    from docx import Document

    doc = Document(str(docx_path))
    paras: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            paras.append(t)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    ct = (p.text or "").strip()
                    if ct:
                        paras.append(ct)
    return paras


@pytest.mark.skipif(
    not os.environ.get("DEEPSEEK_API_KEY"),
    reason="需 DEEPSEEK_API_KEY (从 .env 加载) · 跑真 LLM 调用",
)
@pytest.mark.skipif(
    not CLASSIFIED_JSON.exists(),
    reason=f"classifier 产物缺失: {CLASSIFIED_JSON}",
)
@pytest.mark.skipif(
    not ORIGINAL_TEMPLATE_BAK.exists(),
    reason=f"原 docx bak 缺失: {ORIGINAL_TEMPLATE_BAK}",
)
@pytest.mark.parametrize("name,metadata_path", FIXTURES, ids=[f"struct-{f[0]}" for f in FIXTURES])
def test_rendered_not_structurally_similar_to_original(name, metadata_path, tmp_path):
    """rendered docx 不应跟原经纬测绘模板结构高度相似 (即"换皮"穿帮防御).

    判据:
      - 段落长度 ≥ 50 字符的"长段" · 用 difflib.SequenceMatcher 跟原 .bak-pre-placeholder
        所有 ≥ 50 字符长段比相似度.
      - 任一 rendered 长段与原模板长段相似度 ≥ 0.6 视为"换皮段".
      - 总换皮段数 > 3 → fail (允许少量公共短句 / 格式化保留段 · 但报告主体段必须不一样).
    """
    from difflib import SequenceMatcher

    assert metadata_path.exists(), f"fixture missing: {metadata_path}"
    assert TEMPLATE_DOCX.exists(), f"template missing: {TEMPLATE_DOCX}"
    assert ORIGINAL_TEMPLATE_BAK.exists(), f"bak missing: {ORIGINAL_TEMPLATE_BAK}"

    client_metadata, new_client_name = _load_client_metadata(metadata_path)
    assert client_metadata, f"{name} client_metadata empty"

    _disable_llm_cache()

    from v16_pipeline import run_pipeline

    out_dir = tmp_path / f"struct_{name}"
    out_dir.mkdir(parents=True, exist_ok=True)

    result = run_pipeline(
        source_docx=TEMPLATE_DOCX,
        material_dir=MATERIAL_DIR,
        classified_json=CLASSIFIED_JSON,
        output_dir=out_dir,
        client_metadata=client_metadata,
    )
    output_docx = Path(result["output_docx"])
    assert output_docx.exists(), f"{name} output docx missing"

    rendered_paras = _extract_paragraphs(output_docx)
    original_paras = _extract_paragraphs(ORIGINAL_TEMPLATE_BAK)

    rendered_long = [p for p in rendered_paras if len(p) >= 50]
    original_long = [p for p in original_paras if len(p) >= 50]

    assert rendered_long, f"{name} rendered no long paragraphs (≥50 chars)"
    assert original_long, f"{name} original_bak no long paragraphs (≥50 chars · 模板异常)"

    high_sim_pairs: list[tuple[float, str, str]] = []
    for r_p in rendered_long:
        best_sim = 0.0
        best_o = ""
        for o_p in original_long:
            sim = SequenceMatcher(None, r_p, o_p).ratio()
            if sim > best_sim:
                best_sim = sim
                best_o = o_p
        if best_sim >= 0.6:
            high_sim_pairs.append((best_sim, r_p[:100], best_o[:100]))

    high_sim_count = len(high_sim_pairs)
    assert high_sim_count <= 3, (
        f"[{name}] 结构高度相似段落 {high_sim_count} > 3 (经纬模板换皮穿帮)\n"
        f"  阈值: SequenceMatcher ≥ 0.6 视为换皮 · 容忍上限 3 段 (允许公共格式化短段)\n"
        f"  current_client: {new_client_name}\n"
        f"  output: {output_docx}\n"
        f"  high_sim_samples (top 5):\n"
        + "\n".join(
            f"    sim={sim:.2f}\n      r: {r}...\n      o: {o}..."
            for sim, r, o in sorted(high_sim_pairs, key=lambda x: -x[0])[:5]
        )
        + "\n  → 检查 _build_material_summary_for_rewrite 是否注入 financial_metrics 真值块"
        + "\n  → 检查 _REWRITE_SYSTEM_PROMPT clause 1 数据来源优先级是否生效"
    )
