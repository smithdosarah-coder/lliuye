# -*- coding: utf-8 -*-
"""V16 Pipeline — classifier → generator → QC gate 全链路编排.

设计:
  1. classify(source_docx) — 若对应 `outputs/v16_llm_classified.json` 已存在则复用,
     否则调用 v16_classifier 生成(需客户材料齐全的场景才真正跑分类)。
  2. generate(classified_json, source_docx, material_dir) → output_docx + pending.json
  3. qc_gate(output_docx) → QualityReport + md 报告
  4. 汇总产出:outputs/{report_v16.docx, pending_tags.json, qc_report_v16.md}

CLI:
  py v16_pipeline.py --source samples/经纬测绘_对公成稿A.docx --material samples
"""
from __future__ import annotations

import argparse
import json
import sys
from dataclasses import asdict
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).parent))

OUTPUT_DIR = Path(__file__).parent / "outputs"
DEFAULT_CLASSIFIED_JSON = OUTPUT_DIR / "v16_llm_classified.json"
QC_PASS_THRESHOLD = 75.0


def _extract_docx_text(docx_path: Path) -> str:
    """把 docx 所有段落 + 表格文本串成单个字符串,供 QualityScorer 消费."""
    from docx import Document
    doc = Document(str(docx_path))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = p.text
        if t is not None:
            lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text
                    if t is not None:
                        lines.append(t)
                for sub in cell.tables:
                    for sr in sub.rows:
                        for sc in sr.cells:
                            for p in sc.paragraphs:
                                if p.text is not None:
                                    lines.append(p.text)
    return "\n".join(lines)


def _format_qc_markdown(report, docx_name: str) -> str:
    """QualityReport → markdown 审阅报告."""
    lines: list[str] = []
    status = "✅ PASS" if report.passed else "❌ FAIL"
    lines.append(f"# V16 QC 报告 — {docx_name}\n")
    lines.append(f"- 状态: **{status}**")
    lines.append(f"- 总分: {report.total_score:.1f} / 100 "
                 f"(阈值 {report.pass_threshold})")
    lines.append(f"- 去惩罚前: {report.total_score_before_penalty:.1f}")
    if report.fatal_fail:
        lines.append(f"- ⚠ 一票否决: {'; '.join(report.fatal_reasons)}")
    lines.append("")
    lines.append("## 9 维度评分")
    for d in report.dimensions:
        lines.append(f"- {d.name}: {d.raw_score:.1f} (权重 {d.weight}, "
                     f"加权 {d.weighted_score:.1f})")
        if d.missed_items:
            for n in d.missed_items[:3]:
                lines.append(f"    - miss: {n}")
        if d.llm_comment:
            lines.append(f"    - note: {d.llm_comment[:120]}")
    if report.hallucinations:
        lines.append("")
        lines.append(f"## 疑似幻觉 ({len(report.hallucinations)})")
        for h in report.hallucinations[:10]:
            lines.append(f"- [{getattr(h, 'kind', '?')}] "
                         f"{getattr(h, 'reason', '')[:120]}")
    return "\n".join(lines)


def run_pipeline(
    source_docx: Path,
    material_dir: Path,
    classified_json: Path = DEFAULT_CLASSIFIED_JSON,
    output_dir: Path = OUTPUT_DIR,
    client_metadata: dict | None = None,
) -> dict[str, Any]:
    """跑 v16 全链路,返回聚合结果字典.

    2026-05-21 治本 Phase 1: ``client_metadata`` 透传到 ``generate()`` ·
    驱动 placeholder REPLACE handler · None 时 generate 内部 fallback 老路径.
    """
    from v16_generator import generate

    source_docx = Path(source_docx)
    material_dir = Path(material_dir)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    stem = source_docx.stem
    output_docx = output_dir / f"{stem}_v16.docx"
    qc_md_path = output_dir / f"{stem}_v16_qc.md"

    print("=" * 80)
    print(f"V16 Pipeline — {source_docx.name}")
    print("=" * 80)

    # Step 1: classify(预期 classified_json 已经存在 — 实盘运行时若缺则需要先跑
    # v16_classifier。此处仅做校验,不自动触发长耗时分类)
    if not classified_json.exists():
        raise FileNotFoundError(
            f"classifier 产物缺失: {classified_json}\n"
            "请先运行 `py v16_classifier.py` 生成。"
        )

    # Step 2: generate
    gen_result = generate(
        classified_json=classified_json,
        template_docx=source_docx,
        material_dir=material_dir,
        output_docx=output_docx,
        client_metadata=client_metadata,
    )

    # Step 3: QC gate
    qc_result: dict[str, Any] = {"passed": None, "score": None, "md": str(qc_md_path)}
    try:
        from quality_scorer import QualityScorer
        text = _extract_docx_text(output_docx)
        scorer = QualityScorer(borrower_name="")
        report = scorer.score(text, source_file=str(output_docx))
        qc_md = _format_qc_markdown(report, source_docx.name)
        qc_md_path.write_text(qc_md, encoding="utf-8")
        qc_result["passed"] = report.passed
        qc_result["score"] = report.total_score
        qc_result["fatal_fail"] = report.fatal_fail
        qc_result["hallucinations"] = len(report.hallucinations)
        qc_result["fatal_reasons"] = list(report.fatal_reasons)
        qc_result["dimensions"] = report.to_json()["dimensions"]
        qc_result["hallucination_details"] = report.to_json()["hallucinations"]
        print(f"\n[QC] score={report.total_score:.1f} / 100  "
              f"passed={report.passed}  halluc={len(report.hallucinations)}")
        print(f"  └ qc md: {qc_md_path}")
    except Exception as exc:
        qc_md_path.write_text(
            f"# QC 报告生成失败\n\n{type(exc).__name__}: {exc}\n",
            encoding="utf-8",
        )
        print(f"\n[QC][WARN] QualityScorer 异常: {exc}")
        qc_result["error"] = str(exc)

    summary = {
        "source_docx": str(source_docx),
        "output_docx": str(output_docx),
        "pending_json": gen_result.get("pending_json"),
        "handler_stats": gen_result.get("handler_stats"),
        "apply_stats": gen_result.get("apply_stats"),
        "pending_count": gen_result.get("pending_count"),
        "qc": qc_result,
    }
    return summary


def _main(argv: list[str]) -> int:
    ap = argparse.ArgumentParser(description="V16 Pipeline orchestrator")
    ap.add_argument("--source", type=str, help="模板 docx 路径",
                    default="samples/经纬测绘_对公成稿A.docx")
    ap.add_argument("--material", type=str, default="samples",
                    help="客户材料目录")
    ap.add_argument("--classified", type=str,
                    default=str(DEFAULT_CLASSIFIED_JSON))
    ap.add_argument("--output-dir", type=str, default=str(OUTPUT_DIR))
    ap.add_argument("--all", action="store_true",
                    help="跑 samples/ 下全部 docx")
    args = ap.parse_args(argv)

    output_dir = Path(args.output_dir)
    classified = Path(args.classified)

    if args.all:
        sources = sorted(Path("samples").glob("*.docx"))
    else:
        sources = [Path(args.source)]

    all_summaries = []
    for src in sources:
        try:
            s = run_pipeline(
                source_docx=src,
                material_dir=Path(args.material),
                classified_json=classified,
                output_dir=output_dir,
            )
            all_summaries.append(s)
        except Exception as exc:
            print(f"[FAIL] {src.name}: {exc}")
            all_summaries.append({"source_docx": str(src), "error": str(exc)})

    summary_path = output_dir / "v16_pipeline_summary.json"
    summary_path.write_text(
        json.dumps(all_summaries, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"\n[summary] {summary_path}")
    return 0


if __name__ == "__main__":
    sys.exit(_main(sys.argv[1:]))
