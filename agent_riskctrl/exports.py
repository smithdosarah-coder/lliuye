# -*- coding: utf-8 -*-
"""agent_riskctrl.exports — 回测报告三件套导出 (Phase A worker-A4 · 2026-04-29).

三个 builder · 各自返 bytes (PDF / DOCX / XLSX) · 由 api.py 路由层 wrap Response.

Builder 签名统一接受 ctx dict (字段全可选 · 缺则 placeholder · backend 不 silent fail)：
  ctx = {
    "ruleset_id":   str,                 # 文件名后缀 + 报告标题
    "ruleset":      dict | None,         # {rules: [...], description?: str}
    "ks":           dict | None,         # {ksPeak, auc, passRate, badRate, points: [...]}
    "samples":      list[dict] | None,   # [{key, label, count, pct, bad_rate}, ...]
    "rule_stats":   list[dict] | None,   # [{rule_id, hit, fp, tn}, ...]
    "metrics":      dict | None,         # {total_records, approved, rejected, ...}
  }

合规：本地 python-docx / openpyxl / reportlab 渲染 · 不走境外 API.
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _normalize_ctx(ctx: dict[str, Any] | None) -> dict[str, Any]:
    """Fill missing fields with safe placeholders."""
    ctx = dict(ctx or {})
    ctx.setdefault("ruleset_id", "rs_unknown")
    ctx.setdefault("ruleset", {"rules": [], "description": "(未提供 ruleset · 占位)"})
    ctx.setdefault("ks", {})
    ctx.setdefault("samples", [])
    ctx.setdefault("rule_stats", [])
    ctx.setdefault("metrics", {})
    return ctx


def _stamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def build_filename(ruleset_id: str, ext: str) -> str:
    safe = "".join(c for c in ruleset_id if c.isalnum() or c in ("-", "_")) or "rs"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"riskctrl_backtest_{safe}_{ts}.{ext}"


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def build_docx(ctx: dict[str, Any] | None) -> bytes:
    """生成回测报告 Docx · §一 DSL 树 + §二 KS + §三 样本 + §四 规则命中.

    使用 python-docx · 不走境外 API.
    """
    from docx import Document
    from docx.shared import Pt

    ctx = _normalize_ctx(ctx)
    doc = Document()

    title = doc.add_heading(f"风控策略回测报告 · {ctx['ruleset_id']}", level=0)
    for run in title.runs:
        run.font.size = Pt(20)
    doc.add_paragraph(f"生成时间：{_stamp()}").alignment = 2  # right
    doc.add_paragraph(
        "本报告由 Agent2 风控 (Forge) 生成 · AI 初版策略稿 · 未经风险总监审批不得上线"
    )

    # §一 DSL 规则树
    doc.add_heading("§一 DSL 规则树", level=1)
    rules = (ctx["ruleset"] or {}).get("rules", []) or []
    if not rules:
        doc.add_paragraph("(暂无规则 · ruleset.rules 空)")
    else:
        rule_table = doc.add_table(rows=1, cols=4)
        rule_table.style = "Light Grid Accent 1"
        hdr = rule_table.rows[0].cells
        hdr[0].text = "规则 ID"
        hdr[1].text = "名称"
        hdr[2].text = "动作"
        hdr[3].text = "优先级"
        for r in rules:
            row = rule_table.add_row().cells
            row[0].text = str(r.get("rule_id", ""))
            row[1].text = str(r.get("name", ""))
            row[2].text = str(r.get("action", ""))
            row[3].text = str(r.get("priority", ""))

    # §二 KS 双线
    doc.add_heading("§二 KS / AUC / 通过率", level=1)
    ks = ctx["ks"] or {}
    ks_para = doc.add_paragraph()
    ks_para.add_run(f"KS 峰值: {ks.get('ksPeak', 0):.3f} · ").bold = True
    ks_para.add_run(f"AUC: {ks.get('auc', 0):.3f} · ")
    ks_para.add_run(f"通过率: {ks.get('passRate', 0):.1f}% · ")
    ks_para.add_run(f"坏账率: {ks.get('badRate', 0):.1f}%")

    points = ks.get("points", []) or []
    if points:
        doc.add_paragraph("KS 11 点曲线 (TPR / FPR / KS):")
        ks_table = doc.add_table(rows=1, cols=4)
        ks_table.style = "Light Grid Accent 1"
        hdr = ks_table.rows[0].cells
        hdr[0].text = "Bin"
        hdr[1].text = "TPR"
        hdr[2].text = "FPR"
        hdr[3].text = "KS"
        for p in points:
            row = ks_table.add_row().cells
            row[0].text = str(p.get("bin", ""))
            row[1].text = f"{p.get('tpr', 0):.3f}"
            row[2].text = f"{p.get('fpr', 0):.3f}"
            row[3].text = f"{p.get('ks', 0):.3f}"

    # §三 样本分布
    doc.add_heading("§三 样本分布 (pass / review / block)", level=1)
    samples = ctx["samples"] or []
    if not samples:
        doc.add_paragraph("(无样本数据)")
    else:
        s_table = doc.add_table(rows=1, cols=5)
        s_table.style = "Light Grid Accent 1"
        hdr = s_table.rows[0].cells
        hdr[0].text = "档"
        hdr[1].text = "标签"
        hdr[2].text = "数量"
        hdr[3].text = "占比"
        hdr[4].text = "坏账率"
        for s in samples:
            row = s_table.add_row().cells
            row[0].text = str(s.get("key", ""))
            row[1].text = str(s.get("label", ""))
            row[2].text = f"{s.get('count', 0):,}"
            row[3].text = f"{s.get('pct', 0):.1f}%"
            row[4].text = f"{s.get('bad_rate', 0):.1f}%"

    # §四 规则命中明细
    doc.add_heading("§四 规则命中明细 (per-rule FP/TN/hit)", level=1)
    rs = ctx["rule_stats"] or []
    if not rs:
        doc.add_paragraph("(无 rule_stats)")
    else:
        rs_table = doc.add_table(rows=1, cols=4)
        rs_table.style = "Light Grid Accent 1"
        hdr = rs_table.rows[0].cells
        hdr[0].text = "规则 ID"
        hdr[1].text = "命中"
        hdr[2].text = "FP"
        hdr[3].text = "TN"
        for r in rs:
            row = rs_table.add_row().cells
            row[0].text = str(r.get("rule_id", ""))
            row[1].text = str(r.get("hit", 0))
            row[2].text = str(r.get("fp", 0))
            row[3].text = str(r.get("tn", 0))

    # Footer
    foot = doc.add_paragraph()
    foot.add_run("— 以上为 AI 初版策略稿 · 未经风险总监审批不得上线 —").italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def build_xlsx(ctx: dict[str, Any] | None) -> bytes:
    """生成回测报告 Xlsx · 3 sheet (Rules / KS Points / Samples + RuleStats).

    使用 openpyxl · 本地渲染.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font

    ctx = _normalize_ctx(ctx)
    wb = Workbook()

    # Sheet 1: Rules
    ws_rules = wb.active
    ws_rules.title = "Rules"
    headers = ["rule_id", "name", "description", "conditions_json", "action", "priority"]
    ws_rules.append(headers)
    for cell in ws_rules[1]:
        cell.font = Font(bold=True)
    rules = (ctx["ruleset"] or {}).get("rules", []) or []
    for r in rules:
        ws_rules.append([
            r.get("rule_id", ""),
            r.get("name", ""),
            r.get("description", ""),
            str(r.get("conditions", [])),
            r.get("action", ""),
            r.get("priority", ""),
        ])

    # Sheet 2: KS Points
    ws_ks = wb.create_sheet("KS Points")
    ws_ks.append(["bin", "tpr", "fpr", "ks"])
    for cell in ws_ks[1]:
        cell.font = Font(bold=True)
    for p in (ctx["ks"] or {}).get("points", []) or []:
        ws_ks.append([p.get("bin", 0), p.get("tpr", 0), p.get("fpr", 0), p.get("ks", 0)])

    # Sheet 3: Samples
    ws_s = wb.create_sheet("Samples")
    ws_s.append(["key", "label", "count", "pct", "bad_rate"])
    for cell in ws_s[1]:
        cell.font = Font(bold=True)
    for s in ctx["samples"] or []:
        ws_s.append([
            s.get("key", ""),
            s.get("label", ""),
            s.get("count", 0),
            s.get("pct", 0),
            s.get("bad_rate", 0),
        ])

    # Sheet 4: Rule Stats
    ws_rs = wb.create_sheet("RuleStats")
    ws_rs.append(["rule_id", "hit", "fp", "tn"])
    for cell in ws_rs[1]:
        cell.font = Font(bold=True)
    for r in ctx["rule_stats"] or []:
        ws_rs.append([
            r.get("rule_id", ""),
            r.get("hit", 0),
            r.get("fp", 0),
            r.get("tn", 0),
        ])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def build_pdf(ctx: dict[str, Any] | None) -> bytes:
    """生成回测报告 PDF · 送审包 · 含 docx 内容 + 签字栏 (留白).

    使用 reportlab.platypus · 本地渲染.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.lib import colors

    ctx = _normalize_ctx(ctx)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
        title=f"风控策略回测报告 {ctx['ruleset_id']}",
    )

    styles = getSampleStyleSheet()
    story: list = []
    title_style = ParagraphStyle(
        "Title", parent=styles["Heading1"], fontSize=18, spaceAfter=14,
    )
    h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"], fontSize=14, spaceAfter=8,
    )
    body = styles["BodyText"]

    story.append(Paragraph(f"风控策略回测报告 · {ctx['ruleset_id']}", title_style))
    story.append(Paragraph(f"生成时间：{_stamp()}", body))
    story.append(Spacer(1, 8))

    # KS summary
    story.append(Paragraph("§一 KS / AUC / 通过率", h2))
    ks = ctx["ks"] or {}
    story.append(Paragraph(
        f"KS 峰值: <b>{ks.get('ksPeak', 0):.3f}</b> · "
        f"AUC: {ks.get('auc', 0):.3f} · "
        f"通过率: {ks.get('passRate', 0):.1f}% · "
        f"坏账率: {ks.get('badRate', 0):.1f}%",
        body,
    ))
    story.append(Spacer(1, 8))

    # Rules
    story.append(Paragraph("§二 DSL 规则", h2))
    rules = (ctx["ruleset"] or {}).get("rules", []) or []
    if rules:
        rule_data = [["规则 ID", "名称", "动作", "优先级"]]
        for r in rules[:20]:  # 限 20 条避溢页
            rule_data.append([
                str(r.get("rule_id", ""))[:20],
                str(r.get("name", ""))[:30],
                str(r.get("action", "")),
                str(r.get("priority", "")),
            ])
        t = Table(rule_data, colWidths=[30 * mm, 60 * mm, 30 * mm, 25 * mm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ]))
        story.append(t)
    else:
        story.append(Paragraph("(暂无规则)", body))
    story.append(Spacer(1, 8))

    # Samples
    story.append(Paragraph("§三 样本分布", h2))
    samples = ctx["samples"] or []
    if samples:
        s_data = [["档", "数量", "占比", "坏账率"]]
        for s in samples:
            s_data.append([
                str(s.get("label", "")),
                f"{s.get('count', 0):,}",
                f"{s.get('pct', 0):.1f}%",
                f"{s.get('bad_rate', 0):.1f}%",
            ])
        t = Table(s_data, colWidths=[30 * mm, 35 * mm, 35 * mm, 35 * mm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ]))
        story.append(t)
    story.append(Spacer(1, 16))

    # Sign-off block (留白)
    story.append(Paragraph("§四 审批栏", h2))
    sign_data = [
        ["策略经理", "_________________", "日期", "_________________"],
        ["风险总监", "_________________", "日期", "_________________"],
        ["合规复核", "_________________", "日期", "_________________"],
    ]
    sign_table = Table(sign_data, colWidths=[25 * mm, 50 * mm, 20 * mm, 50 * mm])
    sign_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 14),
        ("TOPPADDING", (0, 0), (-1, -1), 14),
    ]))
    story.append(sign_table)
    story.append(Spacer(1, 16))
    story.append(Paragraph(
        "<i>— 以上为 AI 初版策略稿 · 未经风险总监审批不得上线 —</i>",
        body,
    ))

    doc.build(story)
    return buf.getvalue()


__all__ = ["build_docx", "build_xlsx", "build_pdf", "build_filename"]
