"""中锐 ground truth 验证 — v16 输出 vs 真实事实抽取对照."""
import io
import sys
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
sys.path.insert(0, str(Path(__file__).parent))

# 真实中锐 ground truth (来自实际客户材料 OCR)
GT = {
    "company_name": "福建中锐网络股份有限公司",
    "unified_credit_code": "91350105589558110Y",
    "legal_rep": "黄祖海",
    "controller_name": "黄祖海",
    "controller_id": "350124197704201694",
    "controller_birth": "1977年4月20日",
    "controller_address": "福建省福州市仓山区金山新村东区12座104单元",
    "registered_capital": "4100万元",
    "industry": "信息技术服务",
    "main_business": "智慧水利|智慧教育",
    "operating_address": "福州市台江区宁化街道长汀街23号ICC升龙环球中心1310",
    "post_code": "350001",
    "phone": "13860636780",
    "bank_account": "409161646186",
    "auditor": "德赢",
    "audit_years": "2023|2024",
    "is_high_tech": True,
    "lease_area": "2523",
    "lease_location": "福州大学国家科技园",
}


def _collect_text(docx_path):
    from docx import Document
    d = Document(str(docx_path))
    lines = []
    for p in d.paragraphs:
        if p.text: lines.append(p.text)
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if p.text: lines.append(p.text)
    return "\n".join(lines)


def validate(docx_path):
    text = _collect_text(docx_path)
    hit, miss = {}, {}
    for field, expected in GT.items():
        if isinstance(expected, bool):
            continue
        variants = str(expected).split("|") if "|" in str(expected) else [str(expected)]
        if any(v in text for v in variants):
            hit[field] = expected
        else:
            miss[field] = expected
    return hit, miss, text


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", type=str)
    args = ap.parse_args()
    hit, miss, _ = validate(args.docx)
    print(f"Hit: {len(hit)}/{len(GT)-1} ({100*len(hit)/(len(GT)-1):.1f}%)")
    print("\n== HIT ==")
    for k, v in hit.items():
        print(f"  ✓ {k}: {v}")
    print("\n== MISS ==")
    for k, v in miss.items():
        print(f"  ✗ {k}: {v}")
